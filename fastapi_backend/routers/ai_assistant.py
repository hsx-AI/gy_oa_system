# -*- coding: utf-8 -*-
"""
智能制造工艺部 AI 助手
------------------------------------------------------------
整合系统数据库资源，提供带 skills（function calling）的流式问答：
  - 制度/工艺智能检索（向量库 bge-small-zh）
  - 加班 / 请假 / 公出记录查询统计
  - 部门人数 / 人员名单查询
  - 报表查询与下载（生成 Excel 下载链接，复用报表汇聚口径）
  - 数据可视化图表（折线/柱状/饼图 PNG，对话内预览；Word/Excel 可嵌入）

【大模型选择】优先使用 webconfig.deepseek_api_key 指向的 DeepSeek（联网模型，function calling
与流式 token 输出都很稳定）；未配置 deepseek_api_key 时回退到本地 Ollama
（webconfig.llm_base_url / llm_model，默认 qwen3:8b）。注意：本地 qwen 对 OpenAI 风格的
function calling 支持不稳定，建议优先用 DeepSeek。

对话以 SSE (text/event-stream) 形式流式返回，事件类型：
  - {"type": "meta",  "provider": "deepseek|local", "model": "..."}
  - {"type": "tool",  "name": "...", "label": "...", "status": "running|done", "summary": "..."}
  - {"type": "attachment", "label": "...", "url": "...", "filename": "..."}
  - {"type": "chunk", "text": "增量 token"}
  - {"type": "done"}
  - {"type": "error", "message": "..."}
"""
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import StreamingResponse, FileResponse
from typing import Optional, List, Dict, Any, Tuple
from pydantic import BaseModel
from datetime import datetime, timedelta, date
from io import BytesIO
from pathlib import Path
from urllib.parse import urlencode, quote
from collections import defaultdict
import calendar
import asyncio
import os
import re
import json
import time
import uuid
import logging

from database import db

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/ai-assistant", tags=["AI助手"])

# 通用文档（Word/Excel）临时输出目录，供 create_document 生成、download 下载
TEMP_DOC_DIR = Path(__file__).resolve().parent.parent / "temp_docs"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PNG_MIME = "image/png"
TEMP_DOC_TTL = 24 * 3600  # 临时文件保留时长（秒）


def _cleanup_temp_docs():
    """清理过期临时文档，避免无限堆积。"""
    try:
        if not TEMP_DOC_DIR.exists():
            return
        now = time.time()
        for p in TEMP_DOC_DIR.iterdir():
            try:
                if p.is_file() and now - p.stat().st_mtime > TEMP_DOC_TTL:
                    p.unlink()
            except Exception:
                pass
    except Exception:
        pass

# ==================== 通用只读数据库查询（安全沙箱） ====================
# 允许 AI 在受控范围内对业务库执行 SELECT，回答超出预设技能的统计类问题。
# 只读 + 表白名单 + 敏感字段屏蔽 + 强制 LIMIT，杜绝写操作与隐私字段泄露。

# 公开表：全员可查（含明细），无数据权限限制
PUBLIC_TABLES = {
    "yggl",                 # 员工（屏蔽身份证/密码/邮箱授权码字段）
    "tech_problem_manual",  # 工艺问题知识库
    "holiday",              # 节假日
    "dept_policy",          # 部门制度
    "shift_config", "shift_day_lock", "shift_day_plan", "shift_schedule",  # 排班
}

# 隐私表：请假/加班/公出/换休，按角色权限。
#   - 经理/副经理/综合技术室主任副主任 (scope=all)：可查全部门任意成员明细
#   - 室主任/副主任/组长 (scope=dept)、普通员工 (scope=self)：仅可做聚合统计，不返回逐条个人明细
PRIVATE_TABLES = {"qj", "jiaban", "gcsqb", "hxp"}

ALLOWED_TABLES = PUBLIC_TABLES | PRIVATE_TABLES

# 敏感字段：仅身份证、密码、邮箱授权码禁止查询（其余员工信息属公开）
BLOCKED_COLUMNS = {"pass", "password", "sfzh", "email_auth_code"}

# 危险操作（写/DDL/注入/慢查询函数等）
_DANGEROUS_PATTERNS = [
    r"\binsert\b", r"\bupdate\b", r"\bdelete\b", r"\bdrop\b", r"\balter\b",
    r"\btruncate\b", r"\bcreate\b", r"\bgrant\b", r"\brevoke\b", r"\bmerge\b",
    r"\breplace\s+into\b", r"\binto\s+outfile\b", r"\binto\s+dumpfile\b",
    r"\bload_file\b", r"\bload\s+data\b", r"\bcall\b", r"\bset\s",
    r"--", r"/\*", r"\bsleep\s*\(", r"\bbenchmark\s*\(",
]

DEFAULT_QUERY_LIMIT = 500

SHIFT_BUSINESS_DEPARTMENTS = [
    "水轮机工艺室",
    "水发工艺室",
    "汽发工艺室",
    "焊接工艺室",
    "综合技术室",
    "非标技术室",
    "工具技术室",
    "智能制造技术室",
    "数控编程室",
]

DEPARTMENT_ALIASES: Dict[str, str] = {
    "水轮机室": "水轮机工艺室", "水轮机": "水轮机工艺室",
    "水发室": "水发工艺室", "水发": "水发工艺室",
    "汽发室": "汽发工艺室", "汽发": "汽发工艺室",
    "焊艺室": "焊接工艺室", "焊艺": "焊接工艺室", "焊接": "焊接工艺室",
    "综合室": "综合技术室", "综合": "综合技术室",
    "非标室": "非标技术室", "非标": "非标技术室",
    "工具室": "工具技术室", "工具": "工具技术室",
    "智能室": "智能制造技术室", "智能制造室": "智能制造技术室",
    "数控室": "数控编程室", "编程室": "数控编程室", "数控": "数控编程室",
    "部办": "部办",
}


def _normalize_lsys(raw: str) -> str:
    """将用户/模型传入的科室简称映射为数据库正式 lsys。"""
    s = (raw or "").strip()
    if not s:
        return ""
    return DEPARTMENT_ALIASES.get(s, s)

DEPARTMENT_KNOWLEDGE = (
    "【部门组织结构】\n"
    "本部门全称：智能制造工艺部（哈电智能制造工艺部门）。\n"
    "下设 9 大科室（调用工具时 lsys 参数须使用下列正式名称）：\n"
    "· 水轮机工艺室（简称：水轮机室）\n"
    "· 水发工艺室（简称：水发室）\n"
    "· 汽发工艺室（简称：汽发室）\n"
    "· 焊接工艺室（简称：焊艺室）\n"
    "· 综合技术室（简称：综合室）\n"
    "· 非标技术室（简称：非标室）\n"
    "· 工具技术室（简称：工具室）\n"
    "· 智能制造技术室（简称：智能室）\n"
    "· 数控编程室（简称：数控室、编程室）\n"
    "另：部门领导层与总师组成部门办公室「部办」（lsys='部办'，与上述 9 个科室并列，不是某个工艺室）。\n"
    "用户用简称提问时（如「智能室加班情况」「焊艺室人数」），须先映射为正式科室名再查询；"
    "「全部门/整个工艺部」指上述 9 个科室及部办全体在职人员。"
)

# 提供给大模型的安全表结构说明（仅含可查询字段）
DB_SCHEMA_HINT = (
    "可用于 query_database 的数据表（仅只读 SELECT，敏感字段不可访问）：\n"
    "【公开表 · 全员可查】\n"
    "- yggl 员工表：name(姓名), gh(工号), xbie(性别,值为 男/女), enterprise_email(企业邮箱), "
    "lsys(隶属科室), jb(职级,如 主任/副主任/组长/员工/无), zaizhi(在职状态,0=在职,其他=离职/调离)。"
    "统计在职人员须加 WHERE zaizhi=0 且 lsys NOT IN ('其他部门员工','其他部门成员') 且 RIGHT(TRIM(name),1)!='1'。\n"
    "- tech_problem_manual 工艺问题知识库：category(分类), department(部门), title(标题), recorder(记录人), "
    "record_time(记录时间), problem_desc(问题描述), cause_analysis(原因分析), measures(解决措施)。"
    "可据此检索同类工艺问题并给出处理建议（用 LIKE 模糊匹配 title/problem_desc/measures）。\n"
    "- holiday 节假日表：year(年份), date(日期 YYYY-MM-DD), type(类型,如 holiday/workday), festival(节日名称)。\n"
    "- dept_policy 制度表：title(标题), issue_time(发布时间)。\n"
    "- 排班表：shift_schedule 排班明细字段 department(科室), employee_name(姓名), shift_date(日期), "
    "shift_type(班次:白班/夜班/白+夜/休息/空), shift_location(值班位置:准备组/服务组), year, month；"
    "shift_day_plan 每日计划字段 department, plan_date, content；"
    "shift_config 排班配置字段 department, workday_day, workday_night, weekend_day, weekend_night。"
    "查询排班情况优先使用 query_shift_schedule，不要猜测 date/name/start_time/end_time 等不存在字段。\n"
    "【隐私表 · 受权限控制】（非高层仅可聚合统计，不可查逐条个人明细）\n"
    "- jiaban 加班表：xm(姓名), xb(性别), jb(职级), bz(部门), jiabanfs(加班方式), timedate(加班日期文本), "
    "timefrom/timeto(起止datetime), tian1(时长/小时,文本), jbf(其他绩效小时,double), "
    "jiabanzt(审批状态,4=已通过), content(内容)。有效加班须 jiabanzt=4。\n"
    "- qj 请假表：xm(姓名), xb(性别), lsys(科室), qjfs(请假类型), timefrom/timeto(起止datetime), "
    "tian(天数文本), xiaoshi(小时文本), qjzt(审批状态,4=已通过), content(事由)。有效请假须 qjzt=4。\n"
    "- gcsqb 公出表：gcr(公出人), gcdw(公出单位), gcdd(公出地点), gclx(公出类型), "
    "wpsj(外派时间datetime), yjcfsj(预计出发), yjfhsj(预计返回), gcrw(公出任务)。\n"
    "- hxp 换休票表：name(姓名), sl(数量,decimal), sj(时间), ly(来源)。\n"
    "注意：时长/天数等为文本字段，做数值统计时用 CAST(... AS DECIMAL(10,2))；性别字段值为中文“男”“女”。"
)

# ==================== 大模型配置 ====================

# 本地 Ollama（OpenAI 兼容）默认值
DEFAULT_LOCAL_BASE_URL = "http://10.42.60.250:11434/v1"
DEFAULT_LOCAL_MODEL = "qwen3:8b"

# DeepSeek 联网模型
DEEPSEEK_BASE_URL = "https://api.deepseek.com/v1"
DEEPSEEK_MODEL = "deepseek-v4-flash"          # 工具调用（function calling）用
DEEPSEEK_REASONER_MODEL = "deepseek-v4-flash"  # 最终回答用，可输出思维链 reasoning_content

# 工具调用最多迭代轮数，防止无限循环
MAX_TOOL_ROUNDS = 4

# ==================== 大模型请求参数（在此调整） ====================
# 说明：400「上下文超出」通常是【输入】过长，不是 max_tokens（输出上限）本身。
# 本地网关/ Ollama 实际可用窗口往往远小于模型宣传的 1M；工具返回的大段 JSON 也会迅速撑爆上下文。
LLM_MAX_OUTPUT_TOKENS = 4096          # 单次回答 max_tokens（输出上限）
LLM_MAX_OUTPUT_TOKENS_DEEPSEEK = 8192   # 联网 DeepSeek 可略大
LLM_TEMPERATURE_TOOL = 0.2              # 阶段1 工具规划
LLM_TEMPERATURE_ANSWER = 0.4            # 阶段2 最终回答（本地模型）
LLM_MAX_HISTORY_TURNS = 16              # 送入模型的最近消息条数（不含 system，约 8 轮问答）
LLM_MAX_MESSAGE_CHARS = 6000            # 单条 user/assistant 正文最大字符（超出截断）
LLM_MAX_TOOL_RESULT_CHARS = 12000       # 单个 tool 摘要 JSON 上限（摘要模式通常远小于此）
LLM_TOOL_SAMPLE_ROWS = 10               # 给模型看的明细样本行数
LLM_TOOL_SAMPLE_ROSTER = 12             # 给模型看的名单样本人数
LLM_FULL_DETAIL_THRESHOLD = 25          # 低于此条数仍返回完整明细


def _clean_text(s: Any) -> str:
    """清理文本中的非法控制字符，避免网关 JSON 解析失败。"""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", s)


def _sanitize_json_value(v: Any) -> Any:
    """递归清理工具结果，确保可被 strict JSON 序列化。"""
    import math
    from decimal import Decimal

    if v is None or isinstance(v, (bool, int)):
        return v
    if isinstance(v, float):
        if math.isnan(v) or math.isinf(v):
            return None
        return v
    if isinstance(v, Decimal):
        try:
            f = float(v)
            if math.isnan(f) or math.isinf(f):
                return None
            return f
        except Exception:
            return str(v)
    if isinstance(v, (datetime,)):
        return v.isoformat(sep=" ", timespec="seconds")
    if isinstance(v, bytes):
        try:
            return _clean_text(v.decode("utf-8", errors="replace"))
        except Exception:
            return ""
    if isinstance(v, str):
        return _clean_text(v)
    if isinstance(v, dict):
        return {str(k): _sanitize_json_value(val) for k, val in v.items()}
    if isinstance(v, (list, tuple)):
        return [_sanitize_json_value(x) for x in v]
    return _clean_text(str(v))


def _aggregate_rows(rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    """对查询结果中的数值列做简单聚合（供模型理解全貌，无需逐行明细）。"""
    if not rows:
        return {}
    agg: Dict[str, Any] = {}
    keys = list(rows[0].keys()) if isinstance(rows[0], dict) else []
    for col in keys:
        nums: List[float] = []
        for r in rows:
            if not isinstance(r, dict):
                continue
            v = r.get(col)
            if v is None or v == "":
                continue
            try:
                nums.append(float(v))
            except (TypeError, ValueError):
                continue
        if len(nums) >= 2:
            agg[col] = {
                "sum": round(sum(nums), 2),
                "avg": round(sum(nums) / len(nums), 2),
                "min": round(min(nums), 2),
                "max": round(max(nums), 2),
                "count": len(nums),
            }
    return agg


def _monthly_breakdown(records: List[Dict[str, Any]], date_keys: Tuple[str, ...] = ("date", "start", "timedate")) -> Dict[str, int]:
    """按 YYYY-MM 统计记录条数。"""
    counts: Dict[str, int] = {}
    for r in records:
        if not isinstance(r, dict):
            continue
        raw = ""
        for k in date_keys:
            if r.get(k):
                raw = str(r[k])
                break
        key = raw[:7] if len(raw) >= 7 else raw
        if key:
            counts[key] = counts.get(key, 0) + 1
    return dict(sorted(counts.items()))


def _count_by_field(items: List[Dict[str, Any]], field: str) -> Dict[str, int]:
    out: Dict[str, int] = {}
    for it in items:
        if not isinstance(it, dict):
            continue
        k = (it.get(field) or "未填").strip() or "未填"
        out[k] = out.get(k, 0) + 1
    return dict(sorted(out.items(), key=lambda x: (-x[1], x[0])))


def _summarize_tool_result_for_llm(fname: str, result: Any, fargs: Optional[Dict[str, Any]] = None) -> Any:
    """把工具原始结果转为「模型友好摘要」：保留统计/分布/样本，去掉大段明细。
    完整数据仍可通过 generate_report / create_document / create_chart 重新查库导出，不影响业务质量。"""
    if not isinstance(result, dict):
        return result
    if not result.get("ok", True):
        return result

    fargs = fargs or {}
    out: Dict[str, Any] = {"ok": True, "_view": "summary"}

    # 小结果 / 无明细：原样返回（对话质量无损）
    def _small_enough(n: int) -> bool:
        return n <= LLM_FULL_DETAIL_THRESHOLD

    if fname == "query_department":
        roster = result.get("roster") or []
        total = int(result.get("total_count") or len(roster))
        out.update({"lsys": result.get("lsys"), "total_count": total})
        out["by_department"] = _count_by_field(roster, "lsys")
        out["by_jb"] = _count_by_field(roster, "jb")
        if any(isinstance(x, dict) and x.get("xbie") for x in roster):
            out["by_gender"] = _count_by_field(roster, "xbie")
        if _small_enough(total):
            out["roster"] = roster
            out["members"] = result.get("members") or [x.get("name") for x in roster]
        else:
            out["roster_sample"] = roster[:LLM_TOOL_SAMPLE_ROSTER]
            out["_hint"] = (
                f"共 {total} 人，已附各科室/职级汇总与 {len(out['roster_sample'])} 人样本。"
                "若需完整名单：让用户指定科室(lsys)查询，或调用 generate_report/create_document 导出。"
            )
        return out

    if fname == "query_database":
        rows = result.get("rows") or []
        total = int(result.get("row_count") or len(rows))
        out.update({"sql": result.get("sql"), "row_count": total, "message": result.get("message")})
        if _small_enough(total):
            out["rows"] = rows
        else:
            out["columns"] = list(rows[0].keys()) if rows and isinstance(rows[0], dict) else []
            out["sample_rows"] = rows[:LLM_TOOL_SAMPLE_ROWS]
            out["aggregates"] = _aggregate_rows(rows)
            out["_hint"] = (
                f"共 {total} 行，已附 {len(out['sample_rows'])} 行样本与数值聚合。"
                "回答统计/趋势/对比类问题请优先用 aggregates；"
                "若需完整明细表格请 create_document 或缩小 SQL 范围。"
            )
        return out

    if fname in ("query_dept_comparison", "query_person_ranking", "query_monthly_summary",
                 "query_full_attendance", "query_work_intensity", "query_discipline_clock",
                 "query_leader_overtime"):
        for k in ("scope", "start_date", "end_date", "count", "message", "metric", "metric_label", "unit",
                  "totals", "chart_hint", "year", "month", "lsys", "workdays", "totalPeople", "fullCount",
                  "rate", "totalOvertimeHours", "totalLeaveHours", "totalActualHours", "overallIntensity",
                  "rankInfo", "clockInTotal", "clockOutTotal", "clockInRange", "clockOutRange",
                  "totalHours", "personCount", "range", "periodMonths", "rankBaselineYear", "rankTotal"):
            if k in result:
                out[k] = result[k]
        for key in ("rows", "byDept", "byPerson", "fullNames_sample", "clockIn", "clockOut", "list", "details_sample"):
            arr = result.get(key) or []
            if not isinstance(arr, list):
                continue
            if _small_enough(len(arr)):
                out[key] = arr
            else:
                out[f"{key}_sample"] = arr[:LLM_TOOL_SAMPLE_ROWS]
                out[f"{key}_count"] = len(arr)
        return out

    if fname in ("query_overtime", "query_leave", "query_business_trip"):
        for k in ("name", "year", "month", "total_hours", "total_days", "total_count", "by_type", "message"):
            if k in result:
                out[k] = result[k]
        records = result.get("records") or []
        n = len(records)
        if _small_enough(n):
            out["records"] = records
        else:
            out["records_sample"] = records[:LLM_TOOL_SAMPLE_ROWS]
            out["monthly_breakdown"] = _monthly_breakdown(records)
            out["_hint"] = f"共 {n} 条记录，已附样本与按月分布；导出完整明细请 generate_report。"
        return out

    if fname == "query_business_trip_summary":
        for k in ("scope", "trip_type", "start_date", "end_date", "raw_record_count", "person_count",
                  "total_days", "by_department", "message"):
            if k in result:
                out[k] = result[k]
        people = result.get("by_person") or []
        if _small_enough(len(people)):
            out["by_person"] = people
        else:
            out["by_person_sample"] = people[:LLM_TOOL_SAMPLE_ROWS]
            out["_hint"] = f"共 {len(people)} 人，已附前 {len(out['by_person_sample'])} 人样本；总天数以 total_days 为准。"
        return out

    if fname == "search_policy":
        out["count"] = result.get("count", 0)
        out["results"] = (result.get("results") or [])[:5]
        return out

    # 下载/图表类：只保留状态，URL 已在 attachment 事件里
    if fname in ("generate_report", "export_report_center", "create_document", "create_chart"):
        out.update({k: result[k] for k in ("message", "label", "filename", "chart_type") if k in result})
        if result.get("download_url"):
            out["download_url"] = "(已生成，见对话附件)"
        return out

    # 其他技能：保留主要字段，去掉可能很大的嵌套列表
    for k, v in result.items():
        if k in ("rows", "roster", "members", "records", "results") and isinstance(v, list) and len(v) > LLM_FULL_DETAIL_THRESHOLD:
            out[f"{k}_sample"] = v[:LLM_TOOL_SAMPLE_ROWS]
            out[f"{k}_count"] = len(v)
        elif k not in ("preview_url",):
            out[k] = v
    return out


def _compact_tool_result(result: Any) -> Any:
    """兜底压缩（智能摘要后仍超长时使用）。"""
    if not isinstance(result, dict):
        return result
    out = dict(result)
    for key, cap in (("rows", LLM_TOOL_SAMPLE_ROWS), ("roster", LLM_TOOL_SAMPLE_ROSTER),
                     ("records", LLM_TOOL_SAMPLE_ROWS), ("members", 30)):
        arr = out.get(key)
        if isinstance(arr, list) and len(arr) > cap:
            out[key] = arr[:cap]
            out[f"_{key}_note"] = f"共 {len(arr)} 条，已截为 {cap} 条"
    if out.get("download_url"):
        out["download_url"] = "(已生成，见对话附件)"
    out.pop("preview_url", None)
    return out


def _tool_result_json_for_llm(fname: str, result: Any, fargs: Optional[Dict[str, Any]] = None,
                               max_chars: int = LLM_MAX_TOOL_RESULT_CHARS) -> str:
    """先智能摘要 → 再 strict JSON 序列化（保证网关可解析且信息密度高）。"""
    summarized = _summarize_tool_result_for_llm(fname, result, fargs)
    cleaned = _sanitize_json_value(_compact_tool_result(summarized))
    try:
        text = json.dumps(cleaned, ensure_ascii=False, allow_nan=False)
    except (TypeError, ValueError) as e:
        logger.warning("工具结果 JSON 序列化失败: %s", e)
        text = json.dumps({"ok": False, "message": "工具结果序列化失败"}, ensure_ascii=False)
    if len(text) > max_chars:
        summary: Dict[str, Any] = {
            "_truncated": True,
            "_note": "摘要仍过长，已保留关键统计；请缩小查询或用导出工具获取完整数据",
        }
        if isinstance(cleaned, dict):
            for k in ("ok", "message", "total_count", "row_count", "count", "sql", "lsys",
                      "by_department", "by_jb", "by_gender", "aggregates", "monthly_breakdown", "_hint"):
                if k in cleaned:
                    summary[k] = cleaned[k]
        text = json.dumps(summary, ensure_ascii=False, allow_nan=False)
    return text


def _trim_chat_history(history: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """保留 system + 最近若干轮对话；用户问题尽量完整保留，仅压缩过长的 assistant 回复。"""
    if not history:
        return history
    system = history[0] if history[0].get("role") == "system" else None
    rest = history[1:] if system else list(history)
    if len(rest) > LLM_MAX_HISTORY_TURNS:
        rest = rest[-LLM_MAX_HISTORY_TURNS:]
    out = ([system] if system else []) + rest
    for msg in out:
        if not isinstance(msg.get("content"), str):
            continue
        c = _clean_text(msg["content"])
        role = msg.get("role")
        limit = LLM_MAX_MESSAGE_CHARS if role == "assistant" else LLM_MAX_MESSAGE_CHARS + 2000
        if len(c) > limit:
            if role == "assistant":
                # 保留开头结论 + 结尾，中间省略
                head, tail = limit // 2 - 30, limit // 2 - 30
                msg["content"] = c[:head] + "\n…（中间回复已压缩，关键数据以本轮工具结果为准）…\n" + c[-tail:]
            else:
                msg["content"] = c[: limit - 20] + "…"
        else:
            msg["content"] = c
    return out


def _safe_tool_arguments(raw: str) -> str:
    """本地模型有时返回不合法 JSON arguments，需清洗后再回传。"""
    s = _clean_text(raw or "").strip()
    if not s:
        return "{}"
    try:
        return json.dumps(json.loads(s), ensure_ascii=False, allow_nan=False)
    except Exception:
        return "{}"


def _has_tool_context(messages: List[Dict[str, Any]]) -> bool:
    return any(
        isinstance(m, dict) and (m.get("role") == "tool" or m.get("tool_calls"))
        for m in messages
    )


def _tool_content_as_plain_text(content: str) -> str:
    text = _clean_text(content or "").strip()
    if not text:
        return ""
    try:
        data = json.loads(text)
        text = json.dumps(data, ensure_ascii=False, allow_nan=False)
    except Exception:
        pass
    return text.replace('"', "'")


def _strip_external_chart_links(text: str) -> str:
    """模型偶尔会编造外部图表/占位图链接；图表必须由 create_chart 本地生成。"""
    if not text:
        return text
    s = text
    s = re.sub(r"!\[[^\]]*\]\(\s*https?://[^)\s]*(?:placeholder|chart|quickchart|image-charts)[^)]*\)", "", s, flags=re.I)
    s = re.sub(r"\[[^\]]*(?:图表|折线图|下载|PNG|图片)[^\]]*\]\(\s*https?://[^)]*\)", "（图表请点击下方本地生成的附件下载）", s, flags=re.I)
    s = re.sub(r"https?://(?:via\.placeholder\.com|placehold\.co|quickchart\.io|image-charts\.com|chart\.googleapis\.com)[^\s)）]+", "", s, flags=re.I)
    return s


def _user_declines_chart(text: str) -> bool:
    s = (text or "").strip()
    if not s:
        return False
    return bool(re.search(r"(不要|不用|无需|别|不必).{0,8}(图|图表|可视化)|只要(文字|数据|表格|列表)", s))


def _is_work_intensity_or_discipline_question(text: str) -> bool:
    s = (text or "").strip()
    if not s:
        return False
    # 权限类意图识别必须偏保守：本地模型不一定遵循提示词，进入模型前先挡住。
    if re.search(r"(工作|劳动|人员|科室|部门|个人).{0,8}(强度|负荷|负载|饱和|压力|忙碌|忙不忙)", s, flags=re.I):
        return True
    if re.search(r"(谁|哪个人|哪位|哪个科室|哪个部门|本部门|全部门).{0,8}(最忙|更忙|比较忙|忙碌)", s, flags=re.I):
        return True
    if re.search(r"(强度|负荷|负载|饱和度|忙碌度).{0,8}(统计|排行|排名|对比|分析|报表|表|图|趋势)", s, flags=re.I):
        return True
    if re.search(r"(加班|请假|公出|出勤|考勤).{0,12}(在岗|强度|负荷|负载|饱和|压力)", s, flags=re.I):
        return True
    return bool(re.search(
        r"(工作强度|强度排行|强度统计|强度表|负荷统计|负荷排行|负载统计|劳动强度|"
        r"考勤纪律|打卡纪律|早到|晚走|踩点|卡点|实际在岗|在岗小时|出勤纪律|"
        r"口径\s*[ABCＡＢＣ]|口径[一二三]|公式\s*[ABCＡＢＣ]|强度口径)",
        s,
        flags=re.I,
    ))


def _chart_field_from_text(text: str, default_field: str, default_name: str) -> Tuple[str, str, str]:
    s = text or ""
    if "人均" in s and ("加班" in s or "工时" in s):
        return "overtime_per_capita", "人均加班", "小时/人"
    if "人均" in s and "请假" in s:
        return "leave_per_capita", "人均请假", "天/人"
    if "人均" in s and ("公出" in s or "出差" in s):
        return "trip_per_capita", "人均公出", "天/人"
    if "请假" in s:
        return "leave_days", "请假天数", "天"
    if "公出" in s or "出差" in s:
        return "trip_days", "公出天数", "天"
    if "满勤率" in s:
        return "rate", "满勤率", "%"
    if "满勤" in s:
        return "fullCount", "满勤人数", "人"
    return default_field, default_name, ""


def _row_label(row: Dict[str, Any]) -> str:
    for k in ("name", "lsys", "department", "key", "month"):
        v = row.get(k)
        if v not in (None, ""):
            return str(v)
    return "项目"


def _to_chart_value(value: Any, percent: bool = False) -> float:
    try:
        num = float(value or 0)
    except (TypeError, ValueError):
        num = 0.0
    return round(num * 100, 2) if percent else round(num, 2)


def _auto_chart_spec_from_tool_result(fname: str, result: Dict[str, Any], fargs: Dict[str, Any], user_text: str) -> Optional[Dict[str, Any]]:
    if not isinstance(result, dict) or not result.get("ok", True) or _user_declines_chart(user_text):
        return None
    if fname in ("query_database", "query_overtime", "query_leave", "query_business_trip", "query_shift_schedule"):
        return None

    scope = result.get("scope") or result.get("lsys") or fargs.get("department") or fargs.get("lsys") or "统计"
    start = result.get("start_date") or ""
    end = result.get("end_date") or ""
    period = f"{start}至{end}" if start and end else str(result.get("year") or datetime.now().year)
    text = user_text or ""

    if fname == "query_monthly_summary":
        hint = result.get("chart_hint") or {}
        labels = hint.get("labels") or []
        series = hint.get("series") or []
        if len(labels) >= 2 and series:
            chart_type = "area" if "面积" in text else ("bar" if "柱" in text else "line")
            return {
                "chart_type": chart_type,
                "title": f"{scope}{period}月度趋势",
                "x_label": "月份",
                "y_label": "数值",
                "labels": labels,
                "series": series,
            }

    if fname == "query_dept_comparison":
        rows = (result.get("rows") or [])[:12]
        if len(rows) >= 2:
            labels = [_row_label(r) for r in rows]
            if not any(k in text for k in ("加班", "请假", "公出", "出差", "人均")):
                series = [
                    {"name": "加班小时", "values": [_to_chart_value(r.get("overtime_hours")) for r in rows]},
                    {"name": "请假天数", "values": [_to_chart_value(r.get("leave_days")) for r in rows]},
                    {"name": "公出天数", "values": [_to_chart_value(r.get("trip_days")) for r in rows]},
                ]
                return {"chart_type": "bar", "title": f"{period}科室横向对比", "x_label": "科室", "y_label": "数值", "labels": labels, "series": series}
            field, name, unit = _chart_field_from_text(text, "overtime_hours", "加班小时")
            return {"chart_type": "bar", "title": f"{period}{name}科室对比", "x_label": "科室", "y_label": unit or name, "labels": labels, "series": [{"name": name, "values": [_to_chart_value(r.get(field), field == "rate") for r in rows]}]}

    if fname == "query_person_ranking":
        rows = (result.get("rows") or [])[:15]
        if len(rows) >= 2:
            return {
                "chart_type": "horizontal_bar",
                "title": f"{period}{result.get('metric_label') or '人员'}排行",
                "x_label": result.get("unit") or "数值",
                "y_label": "人员",
                "labels": [_row_label(r) for r in rows],
                "series": [{"name": result.get("metric_label") or "数值", "values": [_to_chart_value(r.get("value")) for r in rows]}],
            }

    if fname == "query_full_attendance":
        rows = result.get("byDept") or []
        if len(rows) >= 2:
            field, name, unit = _chart_field_from_text(text, "rate", "满勤率")
            return {
                "chart_type": "bar",
                "title": f"{scope}{period}{name}",
                "x_label": "科室",
                "y_label": unit or name,
                "labels": [_row_label(r) for r in rows],
                "series": [{"name": name, "values": [_to_chart_value(r.get(field), field == "rate") for r in rows]}],
            }
        total = int(result.get("totalPeople") or 0)
        full = int(result.get("fullCount") or 0)
        if total > 0:
            return {"chart_type": "pie", "title": f"{scope}{period}满勤构成", "labels": ["满勤", "未满勤"], "series": [{"name": "人数", "values": [full, max(total - full, 0)]}]}

    if fname == "query_work_intensity":
        rows = result.get("byDept") or result.get("byPerson") or []
        rows = rows[:15]
        if len(rows) >= 2:
            return {
                "chart_type": "horizontal_bar",
                "title": f"{scope}{period}工作强度排行",
                "x_label": "工作强度(%)",
                "y_label": "对象",
                "labels": [_row_label(r) for r in rows],
                "series": [{"name": "工作强度", "values": [_to_chart_value(r.get("intensity"), True) for r in rows]}],
            }

    if fname == "query_discipline_clock":
        rows = result.get("clockOut") if any(k in text for k in ("下班", "晚走", "退勤")) else result.get("clockIn")
        name = "晚走打卡" if rows is result.get("clockOut") else "早到打卡"
        rows = (rows or [])[:15]
        if len(rows) >= 2:
            return {"chart_type": "horizontal_bar", "title": f"{scope}{period}{name}排行", "x_label": "次数", "y_label": "对象", "labels": [_row_label(r) for r in rows], "series": [{"name": "次数", "values": [_to_chart_value(r.get("count")) for r in rows]}]}

    if fname == "query_leader_overtime":
        rows = (result.get("list") or [])[:15]
        if len(rows) >= 2:
            return {"chart_type": "horizontal_bar", "title": f"{period}领导加班排行", "x_label": "小时", "y_label": "人员", "labels": [_row_label(r) for r in rows], "series": [{"name": "加班小时", "values": [_to_chart_value(r.get("hours")) for r in rows]}]}

    return None


def _plain_messages_for_answer(messages: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    """把 tool-calling 历史压成普通对话消息，兼容不完整支持 OpenAI tool roles 的网关。"""
    out: List[Dict[str, str]] = []
    pending_tool_names: Dict[str, str] = {}
    for m in messages:
        if not isinstance(m, dict):
            continue
        role = m.get("role")
        if role in ("system", "user"):
            content = _clean_text(m.get("content") or "").strip()
            if content:
                out.append({"role": role, "content": content})
            continue
        if role == "assistant":
            for tc in (m.get("tool_calls") or []):
                if not isinstance(tc, dict):
                    continue
                fid = str(tc.get("id") or "")
                fn = ((tc.get("function") or {}).get("name") or "tool").strip()
                if fid:
                    pending_tool_names[fid] = fn
            content = _clean_text(m.get("content") or "").strip()
            if content:
                out.append({"role": "assistant", "content": content})
            continue
        if role == "tool":
            tool_id = str(m.get("tool_call_id") or "")
            tool_name = pending_tool_names.get(tool_id, "工具")
            content = _tool_content_as_plain_text(m.get("content") or "")
            if content:
                out.append({"role": "user", "content": f"【{tool_name} 查询结果】\n{content}"})
    return out


def _llm_request_kwargs(cfg: dict, **extra) -> dict:
    """统一大模型请求参数（修改 LLM_* 常量即可）。"""
    kw = dict(extra)
    if cfg.get("provider") == "deepseek":
        kw.setdefault("max_tokens", LLM_MAX_OUTPUT_TOKENS_DEEPSEEK)
    else:
        kw.setdefault("max_tokens", LLM_MAX_OUTPUT_TOKENS)
    return kw


def _normalize_llm_base_url(url: str) -> str:
    u = (url or "").strip()
    if not u:
        return u
    u = u.rstrip("/")
    if u.endswith("/chat/completions"):
        u = u[: -len("/chat/completions")].rstrip("/")
    return u


_last_llm_log = ""  # 仅在所选模型变化时打印一次日志，避免每次请求刷屏


def _resolve_llm() -> dict:
    """
    解析对话使用的大模型。返回:
      {"provider","base_url","model","api_key","use_extra"}
    规则：以 webconfig.deepseek_api_key 为开关——非空则使用联网 DeepSeek；
    为空则跳过联网 API，使用本地模型（webconfig.llm_base_url / llm_model，默认 qwen3:8b）。
    use_extra: 是否附带 Ollama/qwen 专用的 enable_thinking 参数（DeepSeek 不能带）。
    """
    global _last_llm_log

    deepseek_key = ""
    local_base_url = DEFAULT_LOCAL_BASE_URL
    local_model = DEFAULT_LOCAL_MODEL
    local_api_key = ""        # 本地模型鉴权 token（如 DeepSeek-V4 网关的 JWT）；空则用占位 "ollama"
    local_use_extra = True    # True=Ollama 本地（带 enable_thinking）；False=OpenAI 兼容网关

    # deepseek_api_key 以 webconfig 为准（数据库为唯一开关）
    try:
        rows = db.execute_query("SELECT deepseek_api_key FROM webconfig WHERE id=%s LIMIT 1", ("1",))
        if rows:
            deepseek_key = (rows[0].get("deepseek_api_key") or "").strip()
    except Exception as e:
        logger.debug(f"读取 webconfig.deepseek_api_key 失败: {e}")

    try:
        rows = db.execute_query("SELECT llm_base_url, llm_model FROM webconfig WHERE id=%s LIMIT 1", ("1",))
        if rows:
            r = rows[0]
            if (r.get("llm_base_url") or "").strip():
                local_base_url = (r["llm_base_url"] or "").strip()
            if (r.get("llm_model") or "").strip():
                local_model = (r["llm_model"] or "").strip()
    except Exception as e:
        logger.debug(f"读取 webconfig 本地大模型配置失败: {e}")

    # 本地模型鉴权 token 与接口类型（新列，旧库可能尚未创建，单独 try 容错）
    try:
        rows = db.execute_query("SELECT llm_api_key, llm_use_extra FROM webconfig WHERE id=%s LIMIT 1", ("1",))
        if rows:
            local_api_key = (rows[0].get("llm_api_key") or "").strip()
            v = rows[0].get("llm_use_extra")
            if v is not None:
                local_use_extra = bool(int(v))
    except Exception as e:
        logger.debug(f"读取 webconfig 本地模型鉴权/接口类型失败（可能为旧库无此列）: {e}")

    if deepseek_key:
        cfg = {
            "provider": "deepseek",
            "base_url": DEEPSEEK_BASE_URL,
            "model": DEEPSEEK_MODEL,
            "api_key": deepseek_key,
            "use_extra": False,
        }
    else:
        cfg = {
            "provider": "local",
            "base_url": _normalize_llm_base_url(local_base_url),
            "model": local_model,
            "api_key": local_api_key or "ollama",
            "use_extra": local_use_extra,
        }

    log_line = f"{cfg['provider']} | {cfg['model']} | {cfg['base_url']}"
    if log_line != _last_llm_log:
        _last_llm_log = log_line
        if cfg["provider"] == "deepseek":
            logger.info(f"AI 助手大模型：联网 DeepSeek（{cfg['model']}）")
        else:
            logger.info(f"AI 助手大模型：本地模型（{cfg['model']} @ {cfg['base_url']}）— webconfig.deepseek_api_key 为空，已跳过联网 API")
    return cfg


def _model_label(cfg: dict) -> str:
    """面向用户的模型展示名（不暴露 base_url / api_key 等敏感信息）。"""
    if cfg.get("provider") == "deepseek":
        return "DeepSeek（联网）"
    return cfg.get("model") or "本地模型"


# ==================== 请求模型 ====================


class ChatMessage(BaseModel):
    role: str
    content: str = ""


class ChatRequest(BaseModel):
    messages: List[ChatMessage] = []
    current_user: str = ""


# ==================== 工具函数 ====================


def _fmt_dt(d) -> str:
    if d is None:
        return ""
    if hasattr(d, "strftime"):
        try:
            return d.strftime("%Y-%m-%d %H:%M")
        except Exception:
            return str(d)
    return str(d)[:16]


def _to_float(v) -> float:
    try:
        return float(v)
    except Exception:
        return 0.0


def _get_user_dept(name: str) -> Dict[str, str]:
    """从 yggl 取某人的隶属科室 lsys 与级别 jb。"""
    try:
        rows = db.execute_query("SELECT jb, lsys FROM yggl WHERE name=%s LIMIT 1", (name,))
        if rows:
            return {
                "jb": (rows[0].get("jb") or "").strip(),
                "lsys": (rows[0].get("lsys") or "").strip(),
            }
    except Exception as e:
        logger.debug(f"读取用户科室失败: {e}")
    return {"jb": "", "lsys": ""}


# ==================== Skills（工具）实现 ====================


def skill_search_policy(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """部门制度 / 工艺制度智能检索（向量库）。"""
    query = (args.get("query") or "").strip()
    top_k = int(args.get("top_k") or 5)
    top_k = max(1, min(10, top_k))
    if not query:
        return {"ok": False, "message": "缺少检索关键词 query"}
    try:
        from services.policy_vector import search
        hits = search(query, top_k=top_k)
    except Exception as e:
        logger.warning(f"制度向量检索失败: {e}")
        return {"ok": False, "message": f"制度检索服务不可用：{e}"}

    if not hits:
        return {"ok": True, "count": 0, "results": [], "message": "未检索到相关制度"}

    ids = [h[0] for h in hits]
    placeholders = ",".join(["%s"] * len(ids))
    meta = {}
    try:
        rows = db.execute_query(
            f"SELECT id, title, issue_time, remark FROM dept_policy WHERE id IN ({placeholders})",
            tuple(ids),
        )
        for r in rows:
            meta[str(r["id"])] = r
    except Exception as e:
        logger.debug(f"读取制度标题失败: {e}")

    results = []
    for pid, score, snippet in hits:
        m = meta.get(str(pid), {})
        results.append({
            "title": (m.get("title") or "（未命名制度）"),
            "issue_time": str(m.get("issue_time") or ""),
            "score": round(float(score) * 100, 1),
            "snippet": (snippet or "")[:300],
        })
    return {"ok": True, "count": len(results), "results": results}


def skill_query_overtime(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """查询某员工已审批通过(jiabanzt=4)的加班记录。"""
    name = (args.get("name") or current_user or "").strip()
    if not name:
        return {"ok": False, "message": "缺少员工姓名"}
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None

    if month:
        month_str = f"{year}-{month:02d}"
        rows = db.execute_query(
            """SELECT timedate, jiabanfs, jbf, tian1, content
               FROM jiaban WHERE xm=%s AND jiabanzt=4
               AND (timedate LIKE %s OR substr(timedate,1,7)=%s)
               ORDER BY timedate DESC""",
            (name, f"{year}-{month:02d}%", month_str),
        )
    else:
        rows = db.execute_query(
            """SELECT timedate, jiabanfs, jbf, tian1, content
               FROM jiaban WHERE xm=%s AND jiabanzt=4
               AND (timedate LIKE %s OR substr(timedate,1,4)=%s)
               ORDER BY timedate DESC""",
            (name, f"{year}%", str(year)),
        )
    total_hours = 0.0
    records = []
    for r in rows:
        h = _to_float(r.get("jbf")) or _to_float(r.get("tian1"))
        total_hours += h
        records.append({
            "date": str(r.get("timedate") or ""),
            "type": r.get("jiabanfs") or "",
            "hours": h,
            "content": (r.get("content") or "")[:80],
        })
    return {
        "ok": True, "name": name, "year": year, "month": month,
        "total_count": len(records), "total_hours": round(total_hours, 2),
        "records": records[:30],
    }


def skill_query_leave(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """查询某员工已审批通过(qjzt=4)的请假记录。"""
    name = (args.get("name") or current_user or "").strip()
    if not name:
        return {"ok": False, "message": "缺少员工姓名"}
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None

    if month:
        month_str = f"{year}-{month:02d}"
        rows = db.execute_query(
            """SELECT qjfs, timefrom, timeto, tian, xiaoshi, content
               FROM qj WHERE xm=%s AND qjzt=4
               AND (timefrom LIKE %s OR substr(timefrom,1,7)=%s
                    OR timefromdate LIKE %s OR substr(timefromdate,1,7)=%s)
               ORDER BY timefrom DESC""",
            (name, f"{year}-{month:02d}%", month_str, f"{year}-{month:02d}%", month_str),
        )
    else:
        rows = db.execute_query(
            """SELECT qjfs, timefrom, timeto, tian, xiaoshi, content
               FROM qj WHERE xm=%s AND qjzt=4
               AND (timefrom LIKE %s OR substr(timefrom,1,4)=%s
                    OR timefromdate LIKE %s OR substr(timefromdate,1,4)=%s)
               ORDER BY timefrom DESC""",
            (name, f"{year}%", str(year), f"{year}%", str(year)),
        )
    total_days = 0.0
    total_hours = 0.0
    records = []
    by_type: Dict[str, float] = {}
    for r in rows:
        d = _to_float(r.get("tian"))
        h = _to_float(r.get("xiaoshi"))
        total_days += d
        total_hours += h
        t = r.get("qjfs") or "其他"
        by_type[t] = round(by_type.get(t, 0.0) + (d if d else h / 8.0), 2)
        records.append({
            "type": t,
            "from": _fmt_dt(r.get("timefrom")),
            "to": _fmt_dt(r.get("timeto")),
            "days": d,
            "hours": h,
            "content": (r.get("content") or "")[:80],
        })
    return {
        "ok": True, "name": name, "year": year, "month": month,
        "total_count": len(records), "total_days": round(total_days, 2),
        "total_hours": round(total_hours, 2), "by_type": by_type,
        "records": records[:30],
    }


def skill_query_business_trip(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """查询某员工公出记录（gcsqb 表）。"""
    name = (args.get("name") or current_user or "").strip()
    if not name:
        return {"ok": False, "message": "缺少员工姓名"}
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None

    if month:
        month_str = f"{year}-{month:02d}"
        rows = db.execute_query(
            """SELECT wpdw, gcdd, gcsj, yjfhsj, sjfhtime, gcrw
               FROM gcsqb WHERE gcr=%s AND (gcsj LIKE %s OR DATE_FORMAT(gcsj,'%%Y-%%m')=%s)
               ORDER BY gcsj DESC""",
            (name, f"{month_str}%", month_str),
        )
    else:
        rows = db.execute_query(
            """SELECT wpdw, gcdd, gcsj, yjfhsj, sjfhtime, gcrw
               FROM gcsqb WHERE gcr=%s AND (gcsj LIKE %s OR YEAR(gcsj)=%s)
               ORDER BY gcsj DESC""",
            (name, f"{year}%", year),
        )
    total_days = 0.0
    records = []
    for r in rows:
        gcsj = r.get("gcsj")
        end_dt = r.get("sjfhtime") or r.get("yjfhsj") or gcsj
        days = 1.0
        if gcsj and end_dt:
            try:
                d1 = gcsj if hasattr(gcsj, "day") else datetime.strptime(str(gcsj)[:10], "%Y-%m-%d")
                d2 = end_dt if hasattr(end_dt, "day") else datetime.strptime(str(end_dt)[:10], "%Y-%m-%d")
                days = max(1, (d2 - d1).days + 1)
            except Exception:
                pass
        total_days += days
        records.append({
            "unit": r.get("wpdw") or "",
            "place": r.get("gcdd") or "",
            "start": _fmt_dt(gcsj),
            "expected_return": _fmt_dt(r.get("yjfhsj")),
            "actual_return": _fmt_dt(r.get("sjfhtime")),
            "task": (r.get("gcrw") or "")[:80],
            "days": round(days, 2),
        })
    return {
        "ok": True, "name": name, "year": year, "month": month,
        "total_count": len(records), "total_days": round(total_days, 2),
        "records": records[:30],
    }


def _parse_db_date(v) -> Optional[date]:
    if v is None:
        return None
    if hasattr(v, "date"):
        return v.date()
    s = str(v).strip()[:10]
    if len(s) == 10 and s[4] == "-" and s[7] == "-":
        try:
            return datetime.strptime(s, "%Y-%m-%d").date()
        except Exception:
            return None
    return None


def _merge_date_intervals(intervals: List[Tuple[date, date]]) -> List[Tuple[date, date]]:
    items = sorted((s, e) for s, e in intervals if s and e and e >= s)
    if not items:
        return []
    merged: List[Tuple[date, date]] = []
    cur_s, cur_e = items[0]
    for s, e in items[1:]:
        if s <= cur_e + timedelta(days=1):
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e))
    return merged


def _merged_interval_days(intervals: List[Tuple[date, date]]) -> int:
    return sum((e - s).days + 1 for s, e in _merge_date_intervals(intervals))


def _trip_period_from_args(args: Dict[str, Any]) -> Tuple[date, date]:
    today = datetime.now().date()
    start_raw = (args.get("start_date") or "").strip()
    end_raw = (args.get("end_date") or "").strip()
    try:
        end_day = datetime.strptime(end_raw, "%Y-%m-%d").date() if end_raw else today
    except Exception:
        end_day = today
    try:
        start_day = datetime.strptime(start_raw, "%Y-%m-%d").date() if start_raw else None
    except Exception:
        start_day = None
    if start_day is None:
        months = int(args.get("months") or args.get("months_back") or 3)
        months = max(1, min(months, 24))
        start_day = end_day - timedelta(days=months * 31 - 1)
    if end_day < start_day:
        start_day, end_day = end_day, start_day
    return start_day, end_day


def skill_query_business_trip_summary(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """按统计页口径汇总部门/科室公出天数，后端完成区间合并与裁剪。"""
    start_day, end_day = _trip_period_from_args(args)
    dept = _resolve_shift_department(args.get("department") or args.get("lsys") or "", current_user)
    scope_info = _get_data_scope(current_user)
    person_filter = ""
    scope = (args.get("scope") or "").strip().lower()
    if dept == "__ALL__" or scope == "all":
        dept = ""
    elif not dept:
        if scope_info.get("scope") == "all":
            dept = ""
        else:
            dept = scope_info.get("lsys") or ""
    if scope_info.get("scope") == "dept":
        own_dept = scope_info.get("lsys") or ""
        if dept and dept != own_dept:
            return {"ok": False, "message": f"权限不足：您只能查询本科室（{own_dept}）公出汇总。"}
        dept = own_dept
    elif scope_info.get("scope") != "all":
        person_filter = (current_user or "").strip()
        dept = scope_info.get("lsys") or dept
    trip_type = (args.get("trip_type") or args.get("gclx") or "").strip()
    if trip_type and trip_type not in ("市内公出", "境内公出", "境外公出"):
        trip_type = ""

    ds_lo = start_day.strftime("%Y-%m-%d")
    ds_hi = end_day.strftime("%Y-%m-%d")
    where = [
        "RIGHT(TRIM(g.gcr),1) != '1'",
        "g.bldzt = 2",
        "g.szrzt = 2",
        "COALESCE(g.gcsj, g.yjcfsj) <= %s",
        "COALESCE(g.sjfhtime, g.yjfhsj) >= %s",
    ]
    params: List[Any] = [ds_hi, ds_lo]
    if dept:
        where.append("y.lsys = %s")
        params.append(dept)
    if person_filter:
        where.append("g.gcr = %s")
        params.append(person_filter)
    if trip_type:
        where.append("COALESCE(NULLIF(TRIM(g.gclx), ''), '境内公出') = %s")
        params.append(trip_type)

    sql = f"""
        SELECT g.gcr, g.gclx, g.gcsj, g.sjfhtime, g.yjfhsj, g.yjcfsj, y.lsys
        FROM gcsqb g
        INNER JOIN yggl y ON g.gcr = y.name
          AND y.lsys IS NOT NULL AND y.lsys != ''
          AND RIGHT(TRIM(y.name),1) != '1'
          AND RIGHT(TRIM(y.lsys),1) != '1'
          AND TRIM(y.lsys) NOT IN ('其他部门员工','其他部门成员')
          AND COALESCE(y.zaizhi,0) = 0
        WHERE {' AND '.join(where)}
    """
    try:
        rows = db.execute_query(sql, tuple(params)) or []
    except Exception as e:
        logger.warning("公出汇总查询失败: %s", e)
        return {"ok": False, "message": f"公出汇总查询失败：{e}"}

    by_person_intervals: Dict[Tuple[str, str], List[Tuple[date, date]]] = defaultdict(list)
    raw_records = 0
    for r in rows:
        name = (r.get("gcr") or "").strip()
        lsys = (r.get("lsys") or "").strip()
        start = _parse_db_date(r.get("gcsj") or r.get("yjcfsj"))
        end = _parse_db_date(r.get("sjfhtime") or r.get("yjfhsj"))
        if not name or not lsys or not start or not end or end < start:
            continue
        clipped_start = max(start, start_day)
        clipped_end = min(end, end_day)
        if clipped_end >= clipped_start:
            raw_records += 1
            by_person_intervals[(name, lsys)].append((clipped_start, clipped_end))

    by_department_acc: Dict[str, float] = defaultdict(float)
    by_person = []
    for (name, lsys), intervals in by_person_intervals.items():
        days = _merged_interval_days(intervals)
        if days <= 0:
            continue
        by_department_acc[lsys] += days
        by_person.append({
            "name": name,
            "lsys": lsys,
            "days": round(days, 2),
            "merged_intervals": [
                {"start": s.strftime("%Y-%m-%d"), "end": e.strftime("%Y-%m-%d")}
                for s, e in _merge_date_intervals(intervals)
            ],
        })
    by_person.sort(key=lambda x: (-x["days"], x["lsys"], x["name"]))
    by_department = [
        {"lsys": k, "days": round(v, 2)}
        for k, v in sorted(by_department_acc.items(), key=lambda x: (-x[1], x[0]))
    ]
    total_days = round(sum(x["days"] for x in by_department), 2)
    return {
        "ok": True,
        "scope": dept or "全部门",
        "trip_type": trip_type or "全部公出类型",
        "start_date": ds_lo,
        "end_date": ds_hi,
        "raw_record_count": raw_records,
        "person_count": len(by_person),
        "total_days": total_days,
        "by_department": by_department,
        "by_person": by_person[:50],
            "message": (
            f"{dept or '全部门'} {ds_lo} 至 {ds_hi} 已审批通过公出："
            f"{raw_records} 条，{len(by_person)} 人，合计 {total_days} 天。"
            "天数已按人合并重叠区间并裁剪到查询区间。"
        ),
    }


def _parse_period_args(args: Dict[str, Any]) -> Tuple[date, date]:
    today = datetime.now().date()
    year = int(args.get("year") or today.year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    start_raw = (args.get("start_date") or "").strip()
    end_raw = (args.get("end_date") or "").strip()
    if start_raw or end_raw:
        try:
            start_day = datetime.strptime(start_raw, "%Y-%m-%d").date() if start_raw else date(year, 1, 1)
        except Exception:
            start_day = date(year, 1, 1)
        try:
            end_day = datetime.strptime(end_raw, "%Y-%m-%d").date() if end_raw else today
        except Exception:
            end_day = today
    elif month:
        start_day = date(year, month, 1)
        end_day = date(year, month, calendar.monthrange(year, month)[1])
    else:
        start_day = date(year, 1, 1)
        end_day = date(year, 12, 31)
    if end_day < start_day:
        start_day, end_day = end_day, start_day
    return start_day, end_day


def _months_between(start_day: date, end_day: date) -> List[Tuple[int, int, date, date]]:
    out: List[Tuple[int, int, date, date]] = []
    cur = date(start_day.year, start_day.month, 1)
    while cur <= end_day:
        last = date(cur.year, cur.month, calendar.monthrange(cur.year, cur.month)[1])
        out.append((cur.year, cur.month, max(start_day, cur), min(end_day, last)))
        if cur.month == 12:
            cur = date(cur.year + 1, 1, 1)
        else:
            cur = date(cur.year, cur.month + 1, 1)
    return out


def _leave_bounds(row: Dict[str, Any]) -> Tuple[Optional[date], Optional[date]]:
    start = _parse_db_date(row.get("timefrom") or row.get("timefromdate"))
    end = _parse_db_date(row.get("timeto") or row.get("timefrom") or row.get("timefromdate"))
    return start, end


def _allocate_days_to_period(total_days: float, start: date, end: date, period_start: date, period_end: date) -> float:
    if not start or not end or end < start or total_days <= 0:
        return 0.0
    ov_s = max(start, period_start)
    ov_e = min(end, period_end)
    if ov_e < ov_s:
        return 0.0
    total_span = (end - start).days + 1
    ov_span = (ov_e - ov_s).days + 1
    if total_span <= 0:
        return 0.0
    return total_days * ov_span / total_span


def _visible_scope_filter(current_user: str, requested_dept: str = "", requested_name: str = "") -> Tuple[Optional[str], str, str]:
    scope_info = _get_data_scope(current_user)
    scope = scope_info.get("scope")
    own_dept = scope_info.get("lsys") or ""
    req_dept = _resolve_shift_department(requested_dept, current_user) if requested_dept else ""
    if req_dept == "__ALL__":
        req_dept = ""
    req_name = (requested_name or "").strip()
    if scope == "all":
        return None, req_dept, req_name
    if scope == "dept":
        if req_dept and req_dept != own_dept:
            return f"权限不足：您只能查询本科室（{own_dept}）统计。", "", ""
        if req_name and not _can_access_person(scope_info, current_user, req_name):
            return f"权限不足：您只能查询本科室（{own_dept}）人员统计。", "", ""
        return None, own_dept, req_name
    if req_name and req_name != current_user:
        return "权限不足：您只能查询本人统计。", "", ""
    return None, own_dept, current_user


def _staff_rows_for_scope(dept: str = "", name: str = "") -> List[Dict[str, Any]]:
    where = [
        "name IS NOT NULL AND TRIM(name) != ''",
        "RIGHT(TRIM(name),1) != '1'",
        "RIGHT(TRIM(lsys),1) != '1'",
        "TRIM(lsys) NOT IN ('其他部门员工','其他部门成员')",
        "COALESCE(zaizhi,0) = 0",
    ]
    params: List[Any] = []
    if dept:
        where.append("lsys = %s")
        params.append(dept)
    if name:
        where.append("name = %s")
        params.append(name)
    rows = db.execute_query(
        "SELECT name, lsys, jb FROM yggl WHERE " + " AND ".join(where) + " ORDER BY lsys, name",
        tuple(params),
    ) or []
    return [
        {"name": (r.get("name") or "").strip(), "lsys": (r.get("lsys") or "").strip(), "jb": (r.get("jb") or "").strip()}
        for r in rows if (r.get("name") or "").strip()
    ]


def _overtime_hours_by_person(start_day: date, end_day: date, names: Optional[List[str]] = None) -> Dict[str, float]:
    where = [
        "jiabanzt = 4",
        "RIGHT(TRIM(xm),1) != '1'",
        "SUBSTRING(timedate,1,10) >= %s",
        "SUBSTRING(timedate,1,10) <= %s",
    ]
    params: List[Any] = [start_day.strftime("%Y-%m-%d"), end_day.strftime("%Y-%m-%d")]
    if names is not None:
        if not names:
            return {}
        ph = ",".join(["%s"] * len(names))
        where.append(f"xm IN ({ph})")
        params.extend(names)
    rows = db.execute_query(
        "SELECT xm AS name, SUM(CAST(COALESCE(NULLIF(tian1,''), jbf, 0) AS DECIMAL(10,2))) AS hours "
        "FROM jiaban WHERE " + " AND ".join(where) + " GROUP BY xm",
        tuple(params),
    ) or []
    return {(r.get("name") or "").strip(): round(float(r.get("hours") or 0), 2) for r in rows if r.get("name")}


def _leave_days_by_person(start_day: date, end_day: date, names: Optional[List[str]] = None) -> Dict[str, float]:
    where = [
        "qjzt = 4",
        "RIGHT(TRIM(xm),1) != '1'",
        "COALESCE(DATE(timeto), DATE(timefrom), DATE(timefromdate)) >= %s",
        "COALESCE(DATE(timefrom), DATE(timefromdate)) <= %s",
    ]
    params: List[Any] = [start_day.strftime("%Y-%m-%d"), end_day.strftime("%Y-%m-%d")]
    if names is not None:
        if not names:
            return {}
        ph = ",".join(["%s"] * len(names))
        where.append(f"xm IN ({ph})")
        params.extend(names)
    rows = db.execute_query(
        "SELECT xm AS name, timefrom, timeto, timefromdate, CAST(COALESCE(NULLIF(tian,''),0) AS DECIMAL(10,2)) AS days "
        "FROM qj WHERE " + " AND ".join(where),
        tuple(params),
    ) or []
    acc: Dict[str, float] = defaultdict(float)
    for r in rows:
        name = (r.get("name") or "").strip()
        s, e = _leave_bounds(r)
        if not name or not s or not e:
            continue
        acc[name] += _allocate_days_to_period(float(r.get("days") or 0), s, e, start_day, end_day)
    return {k: round(v, 2) for k, v in acc.items()}


def _trip_days_by_person(start_day: date, end_day: date, names: Optional[List[str]] = None) -> Dict[str, float]:
    where = [
        "bldzt = 2",
        "szrzt = 2",
        "RIGHT(TRIM(gcr),1) != '1'",
        "COALESCE(gcsj, yjcfsj) <= %s",
        "COALESCE(sjfhtime, yjfhsj) >= %s",
    ]
    params: List[Any] = [end_day.strftime("%Y-%m-%d"), start_day.strftime("%Y-%m-%d")]
    if names is not None:
        if not names:
            return {}
        ph = ",".join(["%s"] * len(names))
        where.append(f"gcr IN ({ph})")
        params.extend(names)
    rows = db.execute_query(
        "SELECT gcr AS name, gcsj, sjfhtime, yjfhsj, yjcfsj FROM gcsqb WHERE " + " AND ".join(where),
        tuple(params),
    ) or []
    intervals: Dict[str, List[Tuple[date, date]]] = defaultdict(list)
    for r in rows:
        name = (r.get("name") or "").strip()
        s = _parse_db_date(r.get("gcsj") or r.get("yjcfsj"))
        e = _parse_db_date(r.get("sjfhtime") or r.get("yjfhsj"))
        if name and s and e and e >= s:
            cs, ce = max(s, start_day), min(e, end_day)
            if ce >= cs:
                intervals[name].append((cs, ce))
    return {name: round(_merged_interval_days(items), 2) for name, items in intervals.items()}


def skill_query_dept_comparison(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """科室横向对比：加班、请假、公出总量与人均。"""
    start_day, end_day = _parse_period_args(args)
    err, dept, name_filter = _visible_scope_filter(current_user, args.get("department") or args.get("lsys") or "", "")
    if err:
        return {"ok": False, "message": err}
    staff = _staff_rows_for_scope(dept=dept, name=name_filter)
    names = [r["name"] for r in staff]
    by_name_dept = {r["name"]: r["lsys"] for r in staff}
    person_count: Dict[str, int] = defaultdict(int)
    for r in staff:
        person_count[r["lsys"]] += 1
    ot = _overtime_hours_by_person(start_day, end_day, names)
    lv = _leave_days_by_person(start_day, end_day, names)
    tr = _trip_days_by_person(start_day, end_day, names)
    acc: Dict[str, Dict[str, float]] = defaultdict(lambda: {"overtime_hours": 0.0, "leave_days": 0.0, "trip_days": 0.0})
    for n, v in ot.items():
        if n in by_name_dept:
            acc[by_name_dept[n]]["overtime_hours"] += v
    for n, v in lv.items():
        if n in by_name_dept:
            acc[by_name_dept[n]]["leave_days"] += v
    for n, v in tr.items():
        if n in by_name_dept:
            acc[by_name_dept[n]]["trip_days"] += v
    rows = []
    for lsys in sorted(person_count.keys()):
        pc = person_count[lsys]
        data = acc[lsys]
        rows.append({
            "lsys": lsys,
            "person_count": pc,
            "overtime_hours": round(data["overtime_hours"], 2),
            "leave_days": round(data["leave_days"], 2),
            "trip_days": round(data["trip_days"], 2),
            "overtime_per_capita": round(data["overtime_hours"] / pc, 2) if pc else 0,
            "leave_per_capita": round(data["leave_days"] / pc, 2) if pc else 0,
            "trip_per_capita": round(data["trip_days"] / pc, 2) if pc else 0,
        })
    return {
        "ok": True,
        "scope": dept or "全部门",
        "start_date": start_day.strftime("%Y-%m-%d"),
        "end_date": end_day.strftime("%Y-%m-%d"),
        "count": len(rows),
        "rows": rows,
        "message": "科室横向对比已由后端按统一统计口径计算完成。",
    }


def skill_query_person_ranking(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """人员排行：加班小时、请假天数、公出天数。"""
    start_day, end_day = _parse_period_args(args)
    metric = (args.get("metric") or args.get("type") or "overtime").strip().lower()
    metric_map = {
        "overtime": ("加班小时", "小时", _overtime_hours_by_person),
        "leave": ("请假天数", "天", _leave_days_by_person),
        "trip": ("公出天数", "天", _trip_days_by_person),
        "business_trip": ("公出天数", "天", _trip_days_by_person),
    }
    if metric not in metric_map:
        metric = "overtime"
    err, dept, name_filter = _visible_scope_filter(current_user, args.get("department") or args.get("lsys") or "", args.get("name") or "")
    if err:
        return {"ok": False, "message": err}
    staff = _staff_rows_for_scope(dept=dept, name=name_filter)
    names = [r["name"] for r in staff]
    by_name_dept = {r["name"]: r["lsys"] for r in staff}
    label, unit, fn = metric_map[metric]
    vals = fn(start_day, end_day, names)
    limit = int(args.get("limit") or 20)
    limit = max(1, min(limit, 100))
    ranked = [
        {"rank": i + 1, "name": n, "lsys": by_name_dept.get(n, ""), "value": round(v, 2), "unit": unit}
        for i, (n, v) in enumerate(sorted(vals.items(), key=lambda x: (-x[1], x[0]))[:limit])
        if v > 0
    ]
    return {
        "ok": True,
        "metric": metric,
        "metric_label": label,
        "unit": unit,
        "scope": dept or ("本人" if name_filter else "全部门"),
        "start_date": start_day.strftime("%Y-%m-%d"),
        "end_date": end_day.strftime("%Y-%m-%d"),
        "count": len(ranked),
        "rows": ranked,
        "message": f"{label}排行已由后端按统一统计口径计算完成。",
    }


def skill_query_monthly_summary(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """月度趋势汇总：按月返回加班、请假、公出。"""
    start_day, end_day = _parse_period_args(args)
    err, dept, name_filter = _visible_scope_filter(current_user, args.get("department") or args.get("lsys") or "", args.get("name") or "")
    if err:
        return {"ok": False, "message": err}
    staff = _staff_rows_for_scope(dept=dept, name=name_filter)
    names = [r["name"] for r in staff]
    rows = []
    for y, m, ms, me in _months_between(start_day, end_day):
        ot = sum(_overtime_hours_by_person(ms, me, names).values())
        lv = sum(_leave_days_by_person(ms, me, names).values())
        tr = sum(_trip_days_by_person(ms, me, names).values())
        rows.append({
            "month": f"{y}-{m:02d}",
            "overtime_hours": round(ot, 2),
            "leave_days": round(lv, 2),
            "trip_days": round(tr, 2),
        })
    totals = {
        "overtime_hours": round(sum(r["overtime_hours"] for r in rows), 2),
        "leave_days": round(sum(r["leave_days"] for r in rows), 2),
        "trip_days": round(sum(r["trip_days"] for r in rows), 2),
    }
    return {
        "ok": True,
        "scope": name_filter or dept or "全部门",
        "start_date": start_day.strftime("%Y-%m-%d"),
        "end_date": end_day.strftime("%Y-%m-%d"),
        "totals": totals,
        "rows": rows,
        "chart_hint": {
            "labels": [r["month"] for r in rows],
            "series": [
                {"name": "加班小时", "values": [r["overtime_hours"] for r in rows]},
                {"name": "请假天数", "values": [r["leave_days"] for r in rows]},
                {"name": "公出天数", "values": [r["trip_days"] for r in rows]},
            ],
        },
        "message": "月度趋势已由后端按统一统计口径计算完成，可直接用于折线图、柱状图或面积图。",
    }


def _statistics_module():
    try:
        from routers import statistics as stats
    except Exception:
        from fastapi_backend.routers import statistics as stats
    return stats


def _run_async_tool(coro):
    try:
        asyncio.get_running_loop()
    except RuntimeError:
        return asyncio.run(coro)

    import concurrent.futures

    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        return executor.submit(asyncio.run, coro).result()


def skill_query_full_attendance(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """管理驾驶舱满勤统计：月度/年度满勤率、满勤人数、科室分布。"""
    stats = _statistics_module()
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    lsys = _normalize_lsys(args.get("lsys") or args.get("department") or "") or None
    try:
        if month:
            data = _run_async_tool(stats.get_leader_full_attendance(year=year, month=month, lsys=lsys))
        else:
            data = _run_async_tool(stats.get_leader_full_attendance_year(year=year, lsys=lsys))
    except HTTPException as e:
        return {"ok": False, "message": str(e.detail)}
    if not isinstance(data, dict) or not data.get("success"):
        return {"ok": False, "message": "满勤统计查询失败。"}
    by_dept = data.get("byDept") or []
    return {
        "ok": True,
        "year": year,
        "month": month,
        "lsys": lsys or "全部门",
        "workdays": data.get("workdays"),
        "totalPeople": data.get("totalPeople"),
        "fullCount": data.get("fullCount"),
        "rate": data.get("rate"),
        "fullNames_sample": (data.get("fullNames") or [])[:50],
        "byDept": [
            {
                "lsys": d.get("lsys"),
                "totalPeople": d.get("totalPeople"),
                "fullCount": d.get("fullCount"),
                "rate": d.get("rate"),
                "fullNames_sample": (d.get("fullNames") or [])[:20],
            }
            for d in by_dept[:20]
        ],
        "message": "满勤统计已由管理驾驶舱后端口径计算完成。",
    }


def skill_query_work_intensity(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """管理驾驶舱工作强度统计：口径 A/B/C、按科室和个人排行。"""
    stats = _statistics_module()
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    lsys = _normalize_lsys(args.get("lsys") or args.get("department") or "") or None
    name = (args.get("name") or "").strip() or None
    date_from = (args.get("start_date") or args.get("date_from") or "").strip() or None
    date_to = (args.get("end_date") or args.get("date_to") or "").strip() or None
    formula = (args.get("formula") or args.get("intensity_formula") or "a").strip().lower()
    limit = max(1, min(int(args.get("limit") or 20), 100))
    try:
        data = _run_async_tool(stats.get_work_intensity(
            year=year, month=month, lsys=lsys, current_user=current_user, name=name,
            date_from=date_from, date_to=date_to, intensity_formula=formula,
        ))
    except HTTPException as e:
        return {"ok": False, "message": str(e.detail)}
    if not isinstance(data, dict) or not data.get("success"):
        return {"ok": False, "message": "工作强度统计查询失败。"}
    return {
        "ok": True,
        "year": year,
        "month": month,
        "lsys": lsys or "全部门",
        "formula": data.get("intensityFormula") or formula,
        "workdays": data.get("workdays"),
        "expectedHoursPerPerson": data.get("expectedHoursPerPerson"),
        "totalPeople": data.get("totalPeople"),
        "totalOvertimeHours": data.get("totalOvertimeHours"),
        "totalLeaveHours": data.get("totalLeaveHours"),
        "totalActualHours": data.get("totalActualHours"),
        "overallIntensity": data.get("overallIntensity"),
        "rankInfo": data.get("rankInfo"),
        "byDept": (data.get("byDept") or [])[:20],
        "byPerson": (data.get("rankList") or data.get("byPerson") or [])[:limit],
        "message": "工作强度统计已由管理驾驶舱后端口径计算完成。",
    }


def skill_query_discipline_clock(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """考勤纪律审查：早到、晚走打卡次数统计，可按人员/科室/月聚合。"""
    stats = _statistics_module()
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    lsys = _normalize_lsys(args.get("lsys") or args.get("department") or "") or None
    dimension = (args.get("dimension") or "person").strip()
    limit = max(1, min(int(args.get("limit") or 20), 100))
    try:
        data = _run_async_tool(stats.get_clock_in_discipline_stats(
            year=year,
            month=month,
            lsys=lsys,
            dimension=dimension,
            clock_in_minutes=int(args.get("clock_in_minutes") or 2),
            clock_out_minutes=int(args.get("clock_out_minutes") or 2),
        ))
    except HTTPException as e:
        return {"ok": False, "message": str(e.detail)}
    if not isinstance(data, dict) or not data.get("success"):
        return {"ok": False, "message": "考勤纪律统计查询失败。"}
    return {
        "ok": True,
        "year": year,
        "month": month,
        "lsys": lsys or "全部门",
        "dimension": dimension,
        "clockInTotal": data.get("clockInTotal"),
        "clockOutTotal": data.get("clockOutTotal"),
        "clockInRange": data.get("clockInRange"),
        "clockOutRange": data.get("clockOutRange"),
        "clockIn": (data.get("clockIn") or [])[:limit],
        "clockOut": (data.get("clockOut") or [])[:limit],
        "message": "考勤纪律审查统计已由管理驾驶舱后端口径计算完成。",
    }


def skill_query_leader_overtime(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """领导加班统计：部办人员打卡加班识别、领导岗位月均天数和估算排名。"""
    stats = _statistics_module()
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    date_from = (args.get("start_date") or args.get("date_from") or "").strip() or None
    date_to = (args.get("end_date") or args.get("date_to") or "").strip() or None
    limit = max(1, min(int(args.get("limit") or 20), 100))
    try:
        data = _run_async_tool(stats.get_leader_overtime_from_attendance(
            year=year,
            month=month,
            date_from=date_from,
            date_to=date_to,
            current_user=current_user,
        ))
    except HTTPException as e:
        return {"ok": False, "message": str(e.detail)}
    if not isinstance(data, dict) or not data.get("success"):
        return {"ok": False, "message": "领导加班统计查询失败。"}
    return {
        "ok": True,
        "year": data.get("year"),
        "month": data.get("month"),
        "range": data.get("range"),
        "periodMonths": data.get("periodMonths"),
        "rankBaselineYear": data.get("rankBaselineYear"),
        "rankTotal": data.get("rankTotal"),
        "totalHours": data.get("totalHours"),
        "personCount": data.get("personCount"),
        "list": (data.get("list") or [])[:limit],
        "details_sample": (data.get("details") or [])[:30],
        "message": "领导加班统计已由管理驾驶舱后端口径计算完成。",
    }


def skill_query_department(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """查询科室 / 全部门的在职人数与人员名单。
    - scope='all'：全部门全体人员；
    - lsys=科室名：该科室人员；
    - 都不填：当前用户所在科室。
    部门人员名单/人数属公开信息，不做权限限制。"""
    scope = (args.get("scope") or "").strip().lower()
    lsys = (args.get("lsys") or "").strip()

    base = (
        "SELECT name, jb, lsys, xbie FROM yggl WHERE name IS NOT NULL AND name != '' "
        "AND RIGHT(TRIM(name),1) != '1' AND RIGHT(TRIM(lsys),1) != '1' "
        "AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0)"
    )
    try:
        if scope == "all":
            rows = db.execute_query(base + " ORDER BY lsys, name")
            lsys_label = "全部门"
        else:
            if not lsys:
                lsys = _get_user_dept(current_user).get("lsys", "")
            else:
                lsys = _normalize_lsys(lsys)
            if not lsys:
                return {"ok": False, "message": "未指定科室，且无法识别当前用户所属科室"}
            rows = db.execute_query(base + " AND lsys = %s ORDER BY name", (lsys,))
            lsys_label = lsys
    except Exception as e:
        logger.error(f"查询科室人员失败: {e}")
        return {"ok": False, "message": f"查询失败：{e}"}

    roster = [
        {"name": (r.get("name") or "").strip(),
         "jb": (r.get("jb") or "").strip(),
         "lsys": (r.get("lsys") or "").strip(),
         "xbie": (r.get("xbie") or "").strip()}
        for r in (rows or []) if r.get("name")
    ]
    names = [x["name"] for x in roster]
    return {
        "ok": True, "lsys": lsys_label, "total_count": len(names),
        "members": names, "roster": roster,
    }


def skill_generate_report(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """生成报表下载链接（Excel）。report_type: overtime|leave|business_trip。"""
    report_type = (args.get("report_type") or "").strip()
    type_label = {"overtime": "加班", "leave": "请假", "business_trip": "公出"}
    if report_type not in type_label:
        return {"ok": False, "message": "report_type 仅支持 overtime / leave / business_trip"}
    scope = (args.get("scope") or "").strip().lower()
    lsys = (args.get("lsys") or "").strip()
    name = (args.get("name") or "").strip()

    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    period = f"{year}年" + (f"{month}月" if month else "全年")

    params: Dict[str, Any] = {"report_type": report_type, "year": year}
    if month:
        params["month"] = month
    if current_user:
        params["requester"] = current_user  # 下载时按请求者权限二次校验

    if scope == "all":
        params["scope"] = "all"
        title = "全部门"
        filename = f"全部门-{type_label[report_type]}报表-{period}.xlsx"
    elif lsys:
        params["lsys"] = lsys
        title = lsys
        filename = f"{lsys}-{type_label[report_type]}报表-{period}.xlsx"
    else:
        name = name or current_user
        if not name:
            return {"ok": False, "message": "缺少员工姓名"}
        params["name"] = name
        title = name
        filename = f"{name}-{type_label[report_type]}报表-{period}.xlsx"

    url = "/api/ai-assistant/export?" + urlencode(params)
    return {
        "ok": True,
        "download_url": url,
        "filename": filename,
        "label": f"{title} {period}{type_label[report_type]}报表",
        "message": f"已生成《{filename}》下载链接，请提示用户点击下方按钮下载。",
    }


def _report_center_period(args: Dict[str, Any]) -> Tuple[int, Optional[int], str]:
    year = int(args.get("year") or datetime.now().year)
    month = args.get("month")
    month = int(month) if month not in (None, "", 0) else None
    label = f"{year}年{month}月" if month else f"{year}年全年"
    return year, month, label


def skill_export_report_center(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """导出报表中心中已有后端文件流报表，返回真实系统下载附件。"""
    raw_type = (args.get("report_type") or args.get("report_id") or "").strip().lower().replace("-", "_")
    aliases = {
        "attendance_exception": "attendance_exceptions",
        "attendance_exceptions": "attendance_exceptions",
        "exceptions": "attendance_exceptions",
        "leave_handler": "leave_handler",
        "suggestion_attendance": "suggestion_attendance",
        "attendance_day": "suggestion_attendance",
        "daily_attendance": "suggestion_attendance",
        "attendance_report": "attendance_report",
        "attendance_word": "attendance_report",
        "holiday_duty": "holiday_duty",
        "duty": "holiday_duty",
        "shift": "shift_schedule",
        "shift_schedule": "shift_schedule",
        "employees": "employees",
        "employee": "employees",
        "file_numbering": "file_numbering",
        "numbering": "file_numbering",
        "confidentiality_ledger": "confidentiality_ledger",
        "confidentiality": "confidentiality_ledger",
    }
    report_type = aliases.get(raw_type, raw_type)
    year, month, period_label = _report_center_period(args)
    params: Dict[str, Any] = {}
    path = ""
    filename = "report"
    label = "报表"

    if report_type in ("attendance_exceptions", "leave_handler"):
        if not month:
            return {"ok": False, "message": "该报表必须指定 month。"}
        params = {"year": year, "month": month}
        if current_user:
            params["current_user"] = current_user
        if report_type == "attendance_exceptions":
            path = "/api/attendance/exceptions/export"
            filename = f"考勤异常_{year}{month:02d}.xlsx"
            label = f"{period_label}考勤异常明细"
        else:
            path = "/api/attendance/leave-handler-export"
            filename = f"异常处理表_{year}{month:02d}.xlsx"
            label = f"{period_label}异常处理表"
    elif report_type == "suggestion_attendance":
        start_date = (args.get("start_date") or "").strip()
        end_date = (args.get("end_date") or "").strip()
        if not start_date or not end_date:
            if not month:
                return {"ok": False, "message": "全员考勤日表必须指定 start_date/end_date，或指定 year+month。"}
            start_date = f"{year}-{month:02d}-01"
            end_date = f"{year}-{month:02d}-{calendar.monthrange(year, month)[1]:02d}"
        params = {"start_date": start_date, "end_date": end_date}
        if current_user:
            params["current_user"] = current_user
        path = "/api/attendance/suggestion-attendance-report/export"
        filename = f"全员考勤日表_{start_date}_{end_date}.xlsx"
        label = f"{start_date} 至 {end_date} 全员考勤日表"
    elif report_type == "attendance_report":
        if not month:
            return {"ok": False, "message": "考勤表 Word 必须指定 month。"}
        params = {"year": year, "month": month}
        lsys = _normalize_lsys(args.get("lsys") or args.get("department") or "")
        if lsys:
            params["lsys"] = lsys
            filename = f"考勤表_{year}年{month}月_{lsys}.docx"
            label = f"{lsys} {period_label}考勤表"
        else:
            filename = f"考勤表_{year}年{month}月_各科室.zip"
            label = f"全部门 {period_label}考勤表"
        path = "/api/leader/attendance-report-export"
    elif report_type == "holiday_duty":
        start_date = (args.get("start_date") or "").strip()
        end_date = (args.get("end_date") or "").strip()
        if not start_date or not end_date:
            return {"ok": False, "message": "假期值班出勤核查必须指定 start_date 和 end_date。"}
        params = {"start_date": start_date, "end_date": end_date}
        lsys = _normalize_lsys(args.get("lsys") or args.get("department") or "")
        if lsys:
            params["lsys"] = lsys
        if current_user:
            params["name"] = current_user
        path = "/api/discipline/holiday-duty-attendance/export"
        filename = f"{lsys or '全部科室'}_{start_date}_{end_date}_假期值班出勤核查.xlsx"
        label = f"{start_date} 至 {end_date} 假期值班出勤核查"
    elif report_type == "shift_schedule":
        fmt = (args.get("format") or args.get("shift_format") or "month").strip()
        department = _normalize_lsys(args.get("department") or args.get("lsys") or "")
        if not department:
            department = "__ALL__" if fmt == "holiday" else _get_user_dept(current_user).get("lsys", "")
        params = {"department": department, "year": year, "format": fmt}
        if fmt == "month":
            if not month:
                return {"ok": False, "message": "月排班表必须指定 month。"}
            params["month"] = month
        elif fmt == "week":
            week_date = (args.get("week_date") or args.get("date") or "").strip()
            if not week_date:
                return {"ok": False, "message": "周排班表必须指定 week_date。"}
            params["week_date"] = week_date
        elif fmt == "holiday":
            holiday = (args.get("holiday") or "").strip()
            if not holiday:
                return {"ok": False, "message": "假期排班表必须指定 holiday。"}
            params["holiday"] = holiday
        path = "/api/shift/export-excel"
        filename = f"{department}_{year}年排班表.xlsx"
        label = f"{department}排班表"
    elif report_type == "employees":
        if current_user:
            params["current_user"] = current_user
        path = "/api/admin/export-employees"
        filename = f"在职员工表_{datetime.now().strftime('%Y-%m-%d')}.xlsx"
        label = "在职员工表"
    elif report_type == "file_numbering":
        table = (args.get("table") or "tech").strip()
        params = {"table": table}
        if current_user:
            params["name"] = current_user
        path = "/api/file-numbering/export/excel"
        filename = f"文件编号_{table}.xlsx"
        label = "文件编号台账"
    elif report_type == "confidentiality_ledger":
        keyword = (args.get("keyword") or "").strip()
        if keyword:
            params["keyword"] = keyword
        path = "/api/confidentiality-ledger/export"
        filename = "智能制造工艺部论文保密审批台账.xlsx"
        label = "论文保密审批台账"
    else:
        return {"ok": False, "message": f"报表中心暂不支持该报表类型：{raw_type}"}

    url = path + ("?" + urlencode(params) if params else "")
    return {
        "ok": True,
        "download_url": url,
        "filename": filename,
        "label": label,
        "message": f"已生成《{filename}》下载附件，请提示用户点击下方按钮下载。",
    }


def _safe_select(sql: str) -> Tuple[Optional[str], Optional[str]]:
    """校验并规整只读 SQL。返回 (安全SQL, None) 或 (None, 拒绝原因)。"""
    s = (sql or "").strip().rstrip(";").strip()
    if not s:
        return None, "SQL 为空"
    low = s.lower()
    if not (low.startswith("select") or low.startswith("with")):
        return None, "仅允许 SELECT 只读查询"
    if ";" in s:
        return None, "不允许一次执行多条语句"
    for pat in _DANGEROUS_PATTERNS:
        if re.search(pat, low):
            return None, "检测到不允许的关键字/操作（仅支持只读查询）"
    if re.search(r"select\s+\*", low):
        return None, "请显式列出需要的字段，不允许 SELECT *（避免带出敏感字段）"
    for col in BLOCKED_COLUMNS:
        if re.search(r"\b" + re.escape(col) + r"\b", low):
            return None, f"字段「{col}」为敏感字段，禁止查询"
    tables = re.findall(r"(?:from|join)\s+`?([a-zA-Z_][\w]*)`?", low)
    if not tables:
        return None, "未识别到查询的数据表"
    for t in tables:
        if t not in ALLOWED_TABLES:
            return None, f"数据表「{t}」不在允许查询的范围内"
    if not re.search(r"\blimit\s+\d+", low):
        s = f"{s} LIMIT {DEFAULT_QUERY_LIMIT}"
    return s, None


def skill_query_database(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """通用只读数据库查询：在安全沙箱内执行模型生成的 SELECT，回答统计类问题。"""
    sql = (args.get("sql") or "").strip()
    if not sql:
        return {"ok": False, "message": "缺少查询语句 sql"}
    safe_sql, err = _safe_select(sql)
    if err:
        return {"ok": False, "message": f"查询被拒绝：{err}"}
    try:
        rows = db.execute_query(safe_sql) or []
    except Exception as e:
        logger.warning(f"通用查询失败: {e} | SQL={safe_sql}")
        return {"ok": False, "message": f"查询执行失败：{e}"}

    capped = rows[:DEFAULT_QUERY_LIMIT]
    note = f"查询返回 {len(rows)} 行" + ("（仅展示前 %d 行）" % DEFAULT_QUERY_LIMIT if len(rows) > DEFAULT_QUERY_LIMIT else "")
    return {
        "ok": True,
        "sql": safe_sql,
        "row_count": len(rows),
        "rows": capped,
        "message": note,
    }


def _resolve_shift_department(raw: str, current_user: str) -> str:
    dept = (raw or "").strip()
    aliases = {
        "全部": "__ALL__",
        "全部门": "__ALL__",
        "全体": "__ALL__",
        "所有科室": "__ALL__",
        "各科室": "__ALL__",
        "部办": "__ALL__",
        "智能室": "智能制造技术室",
        "智能制造室": "智能制造技术室",
        "智能制造": "智能制造技术室",
        "综合室": "综合技术室",
        "综合技术": "综合技术室",
        "水轮机": "水轮机工艺室",
        "水轮机室": "水轮机工艺室",
        "水发": "水发工艺室",
        "水发室": "水发工艺室",
        "汽发": "汽发工艺室",
        "汽发室": "汽发工艺室",
        "焊接": "焊接工艺室",
        "焊艺": "焊接工艺室",
        "焊艺室": "焊接工艺室",
        "工具": "工具技术室",
        "工具室": "工具技术室",
        "非标": "非标技术室",
        "非标室": "非标技术室",
        "数控": "数控编程室",
        "数控室": "数控编程室",
        "编程室": "数控编程室",
    }
    if dept in aliases:
        return aliases[dept]
    if not dept:
        user_dept = _get_user_dept(current_user).get("lsys", "")
        scope_info = _get_data_scope(current_user)
        if user_dept == "部办" or scope_info.get("scope") == "all":
            return "__ALL__"
        return user_dept
    try:
        rows = db.execute_query(
            "SELECT DISTINCT lsys FROM yggl WHERE lsys IS NOT NULL AND TRIM(lsys) != '' "
            "AND COALESCE(zaizhi,0)=0 AND (lsys=%s OR lsys LIKE %s) LIMIT 5",
            (dept, f"%{dept}%"),
        )
        candidates = [(r.get("lsys") or "").strip() for r in rows or [] if (r.get("lsys") or "").strip()]
        if dept in candidates:
            return dept
        if len(candidates) == 1:
            return candidates[0]
    except Exception as e:
        logger.debug("解析排班科室失败: %s", e)
    return dept


def skill_query_shift_schedule(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """查询科室近期排班和每日计划。"""
    scope = (args.get("scope") or "").strip().lower()
    raw_department = args.get("department") or args.get("lsys") or ("__ALL__" if scope == "all" else "")
    department = _resolve_shift_department(raw_department, current_user)
    if not department:
        return {"ok": False, "message": "缺少科室名称，且无法识别当前用户所属科室"}

    today = datetime.now().date()
    start_raw = (args.get("start_date") or "").strip()
    end_raw = (args.get("end_date") or "").strip()
    try:
        start_day = datetime.strptime(start_raw, "%Y-%m-%d").date() if start_raw else today
    except Exception:
        start_day = today
    days = int(args.get("days") or 14)
    days = max(1, min(days, 31))
    try:
        end_day = datetime.strptime(end_raw, "%Y-%m-%d").date() if end_raw else start_day + timedelta(days=days - 1)
    except Exception:
        end_day = start_day + timedelta(days=days - 1)
    if end_day < start_day:
        start_day, end_day = end_day, start_day
    if (end_day - start_day).days > 60:
        end_day = start_day + timedelta(days=60)

    ds_lo = start_day.strftime("%Y-%m-%d")
    ds_hi = end_day.strftime("%Y-%m-%d")
    target_departments = SHIFT_BUSINESS_DEPARTMENTS if department == "__ALL__" else [department]
    placeholders = ",".join(["%s"] * len(target_departments))
    try:
        rows = db.execute_query(
            "SELECT department, employee_name, shift_date, shift_type, shift_location FROM shift_schedule "
            f"WHERE department IN ({placeholders}) AND shift_date >= %s AND shift_date <= %s "
            "AND shift_type IS NOT NULL AND TRIM(shift_type) != '' AND shift_type != '休息' "
            "ORDER BY department ASC, shift_date ASC, employee_name ASC",
            tuple(target_departments) + (ds_lo, ds_hi),
        ) or []
        plan_rows = db.execute_query(
            "SELECT department, plan_date, content FROM shift_day_plan "
            f"WHERE department IN ({placeholders}) AND plan_date >= %s AND plan_date <= %s "
            "AND content IS NOT NULL AND TRIM(content) != '' ORDER BY department ASC, plan_date ASC",
            tuple(target_departments) + (ds_lo, ds_hi),
        ) or []
    except Exception as e:
        logger.warning("查询排班失败: %s", e)
        return {"ok": False, "message": f"查询排班失败：{e}"}

    by_date: Dict[str, Dict[str, Any]] = {}
    for r in rows:
        sd = r.get("shift_date")
        date_text = sd.strftime("%Y-%m-%d") if hasattr(sd, "strftime") else str(sd)[:10]
        row_dept = (r.get("department") or "").strip()
        key = f"{row_dept}|{date_text}" if department == "__ALL__" else date_text
        item = by_date.setdefault(key, {"date": date_text, "department": row_dept, "day": [], "night": [], "other": []})
        name = (r.get("employee_name") or "").strip()
        shift_type = (r.get("shift_type") or "").strip()
        location = (r.get("shift_location") or "").strip()
        label = name + (f"({location})" if location else "")
        if shift_type == "白班":
            item["day"].append(label)
        elif shift_type == "夜班":
            item["night"].append(label)
        elif shift_type == "白+夜":
            item["day"].append(label)
            item["night"].append(label)
        else:
            item["other"].append(f"{label}:{shift_type}")

    plans = []
    for r in plan_rows:
        pd = r.get("plan_date")
        plans.append({
            "department": (r.get("department") or "").strip(),
            "date": pd.strftime("%Y-%m-%d") if hasattr(pd, "strftime") else str(pd)[:10],
            "content": (r.get("content") or "").strip()[:300],
        })

    schedule = list(by_date.values())
    department_label = "全部业务科室" if department == "__ALL__" else department
    return {
        "ok": True,
        "department": department_label,
        "departments": target_departments,
        "start_date": ds_lo,
        "end_date": ds_hi,
        "row_count": len(rows),
        "schedule": schedule,
        "plans": plans,
        "message": f"查询到 {department_label} {ds_lo} 至 {ds_hi} 的排班记录 {len(rows)} 条、每日计划 {len(plans)} 条",
    }


def _parse_chart_spec(args: Dict[str, Any]) -> Tuple[str, List[str], List[Dict[str, Any]], str, str, str]:
    """解析并校验图表参数，返回 (chart_type, labels, series, title, x_label, y_label)。"""
    raw_chart_type = (args.get("chart_type") or "line").strip().lower()
    chart_aliases = {
        "折线": "line", "折线图": "line", "趋势图": "line", "line_chart": "line",
        "柱状": "bar", "柱状图": "bar", "柱形图": "bar", "bar_chart": "bar", "column": "bar",
        "条形": "horizontal_bar", "条形图": "horizontal_bar", "横向柱状图": "horizontal_bar",
        "horizontal_bar_chart": "horizontal_bar", "barh": "horizontal_bar",
        "堆叠柱状图": "stacked_bar", "堆叠柱形图": "stacked_bar", "stacked": "stacked_bar",
        "面积": "area", "面积图": "area", "area_chart": "area",
        "饼": "pie", "饼图": "pie", "占比图": "pie", "pie_chart": "pie",
    }
    chart_type = chart_aliases.get(raw_chart_type, raw_chart_type)
    if chart_type not in ("line", "bar", "horizontal_bar", "stacked_bar", "area", "pie"):
        chart_type = "line"
    title = (args.get("title") or "图表").strip()
    x_label = (args.get("x_label") or "").strip()
    y_label = (args.get("y_label") or "").strip()

    labels = args.get("labels") or []
    if not isinstance(labels, list):
        labels = []
    labels = [str(x) for x in labels]

    series_raw = args.get("series") or []
    if not isinstance(series_raw, list):
        series_raw = []
    series: List[Dict[str, Any]] = []
    for item in series_raw:
        if not isinstance(item, dict):
            continue
        name = (item.get("name") or "数据").strip()
        vals = item.get("values") or []
        if not isinstance(vals, list):
            continue
        nums = []
        for v in vals:
            try:
                nums.append(float(v))
            except (TypeError, ValueError):
                nums.append(0.0)
        series.append({"name": name, "values": nums})

    if not series and args.get("values") is not None:
        vals = args.get("values")
        if isinstance(vals, list):
            name = (args.get("series_name") or "数据").strip()
            nums = []
            for v in vals:
                try:
                    nums.append(float(v))
                except (TypeError, ValueError):
                    nums.append(0.0)
            series = [{"name": name, "values": nums}]

    if not series:
        raise ValueError("缺少数据：请提供 series（含 name/values）或 values")
    if not labels:
        raise ValueError("缺少横轴/分类 labels")
    if chart_type == "pie":
        if len(series) != 1:
            raise ValueError("饼图仅支持单组数据（一个 series）")
        if len(labels) != len(series[0]["values"]):
            raise ValueError("饼图 labels 数量须与 values 一致")
    elif chart_type in ("stacked_bar", "area") and len(series) < 1:
        raise ValueError("缺少数据系列")
    else:
        n = len(labels)
        for s in series:
            if len(s["values"]) != n:
                raise ValueError(f"系列「{s['name']}」的数据点数量须与 labels 一致（{n} 个）")
    return chart_type, labels, series, title, x_label, y_label


def _render_chart_png(fp: Path, spec: Dict[str, Any]) -> None:
    """用 matplotlib 渲染折线/柱状/饼图并保存为 PNG。"""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from utils.chart_font import configure_matplotlib_cjk, get_cjk_font_properties

    configure_matplotlib_cjk()
    cjk_font = get_cjk_font_properties()
    text_kw = {"fontproperties": cjk_font} if cjk_font else {}

    chart_type, labels, series, title, x_label, y_label = _parse_chart_spec(spec)
    fig, ax = plt.subplots(figsize=(8.5, 4.8), dpi=120)

    if chart_type == "pie":
        vals = series[0]["values"]
        pie_text = {**text_kw, "fontsize": 10}
        ax.pie(vals, labels=labels, autopct="%1.1f%%", startangle=90, textprops=pie_text)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
    elif chart_type == "bar":
        x_idx = list(range(len(labels)))
        n_series = len(series)
        width = min(0.8 / max(n_series, 1), 0.35)
        for i, s in enumerate(series):
            offset = (i - (n_series - 1) / 2) * width
            xs = [xi + offset for xi in x_idx]
            ax.bar(xs, s["values"], width=width, label=s["name"])
        ax.set_xticks(x_idx)
        ax.set_xticklabels(labels, rotation=25 if len(labels) > 6 else 0, ha="right" if len(labels) > 6 else "center", **text_kw)
        if n_series > 1:
            ax.legend(loc="best", fontsize=9, prop=cjk_font)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
        if x_label:
            ax.set_xlabel(x_label, **text_kw)
        if y_label:
            ax.set_ylabel(y_label, **text_kw)
        ax.grid(True, axis="y", alpha=0.3)
    elif chart_type == "horizontal_bar":
        y_idx = list(range(len(labels)))
        n_series = len(series)
        height = min(0.8 / max(n_series, 1), 0.35)
        for i, s in enumerate(series):
            offset = (i - (n_series - 1) / 2) * height
            ys = [yi + offset for yi in y_idx]
            ax.barh(ys, s["values"], height=height, label=s["name"])
        ax.set_yticks(y_idx)
        ax.set_yticklabels(labels, **text_kw)
        ax.invert_yaxis()
        if n_series > 1:
            ax.legend(loc="best", fontsize=9, prop=cjk_font)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
        if x_label:
            ax.set_xlabel(x_label, **text_kw)
        if y_label:
            ax.set_ylabel(y_label, **text_kw)
        ax.grid(True, axis="x", alpha=0.3)
    elif chart_type == "stacked_bar":
        x_idx = list(range(len(labels)))
        bottoms = [0.0 for _ in labels]
        for s in series:
            ax.bar(x_idx, s["values"], bottom=bottoms, label=s["name"])
            bottoms = [b + float(v) for b, v in zip(bottoms, s["values"])]
        ax.set_xticks(x_idx)
        ax.set_xticklabels(labels, rotation=25 if len(labels) > 6 else 0, ha="right" if len(labels) > 6 else "center", **text_kw)
        if len(series) > 1:
            ax.legend(loc="best", fontsize=9, prop=cjk_font)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
        if x_label:
            ax.set_xlabel(x_label, **text_kw)
        if y_label:
            ax.set_ylabel(y_label, **text_kw)
        ax.grid(True, axis="y", alpha=0.3)
    elif chart_type == "area":
        x_idx = list(range(len(labels)))
        values = [s["values"] for s in series]
        names = [s["name"] for s in series]
        ax.stackplot(x_idx, values, labels=names, alpha=0.75)
        ax.set_xticks(x_idx)
        ax.set_xticklabels(labels, rotation=25 if len(labels) > 6 else 0, ha="right" if len(labels) > 6 else "center", **text_kw)
        if len(series) > 1:
            ax.legend(loc="best", fontsize=9, prop=cjk_font)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
        if x_label:
            ax.set_xlabel(x_label, **text_kw)
        if y_label:
            ax.set_ylabel(y_label, **text_kw)
        ax.grid(True, alpha=0.25)
    else:  # line
        for s in series:
            ax.plot(labels, s["values"], marker="o", linewidth=2, markersize=5, label=s["name"])
        if len(series) > 1:
            ax.legend(loc="best", fontsize=9, prop=cjk_font)
        ax.set_title(title, fontsize=14, fontweight="bold", **text_kw)
        if x_label:
            ax.set_xlabel(x_label, **text_kw)
        if y_label:
            ax.set_ylabel(y_label, **text_kw)
        ax.grid(True, alpha=0.3)
        if len(labels) > 6:
            plt.setp(ax.get_xticklabels(), rotation=25, ha="right", **text_kw)
        elif cjk_font:
            for lbl in ax.get_xticklabels() + ax.get_yticklabels():
                lbl.set_fontproperties(cjk_font)

    fig.tight_layout()
    fig.savefig(fp, format="png", bbox_inches="tight")
    plt.close(fig)


def _add_xlsx_chart(ws, chart_spec: dict, data_start_row: int, n_cols: int, n_data_rows: int) -> None:
    """在已有表格数据下方插入 openpyxl 原生折线/柱状图（仅 Excel）。"""
    from openpyxl.chart import LineChart, BarChart, Reference

    chart_type, _, _, _, _, _ = _parse_chart_spec(chart_spec)
    title = (chart_spec.get("title") or "图表").strip()
    x_label = (chart_spec.get("x_label") or "").strip()
    y_label = (chart_spec.get("y_label") or "").strip()

    if n_cols < 2 or n_data_rows < 1:
        return
    ChartCls = BarChart if chart_type in ("bar", "horizontal_bar", "stacked_bar") else LineChart
    chart = ChartCls()
    if chart_type == "horizontal_bar":
        chart.type = "bar"
    elif chart_type == "stacked_bar":
        chart.type = "col"
        chart.grouping = "stacked"
        chart.overlap = 100
    chart.title = title
    if x_label:
        chart.x_axis.title = x_label
    if y_label:
        chart.y_axis.title = y_label

    # 第一列为分类轴，其余列为数据系列
    data = Reference(ws, min_col=2, min_row=data_start_row, max_col=n_cols, max_row=data_start_row + n_data_rows)
    cats = Reference(ws, min_col=1, min_row=data_start_row + 1, max_row=data_start_row + n_data_rows)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.width = 16
    chart.height = 9
    anchor_row = data_start_row + n_data_rows + 2
    ws.add_chart(chart, f"A{anchor_row}")


def _render_xlsx(fp: Path, title: str, body: str, columns: list, rows: list, charts: Optional[list] = None) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = (title or "数据")[:28] or "数据"
    span = max(1, len(columns) if columns else 1)

    r = 1
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=span)
    tc = ws.cell(row=r, column=1, value=title or "文档")
    tc.font = Font(size=14, bold=True)
    tc.alignment = Alignment(horizontal="center", vertical="center")
    r += 1

    if body:
        for line in body.splitlines():
            line = line.strip()
            if not line:
                continue
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=span)
            ws.cell(row=r, column=1, value=line).alignment = Alignment(wrap_text=True, vertical="top")
            r += 1
        r += 1

    table_header_row = 0
    n_data_rows = 0
    if columns:
        header_fill = PatternFill(start_color="1890FF", end_color="1890FF", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)
        table_header_row = r
        for ci, h in enumerate(columns, start=1):
            c = ws.cell(row=r, column=ci, value=str(h))
            c.fill = header_fill
            c.font = header_font
            c.alignment = Alignment(horizontal="center", vertical="center")
        r += 1
        for row in (rows or []):
            cells = row if isinstance(row, (list, tuple)) else [row]
            for ci, val in enumerate(cells, start=1):
                ws.cell(row=r, column=ci, value=val if val is not None else "")
            r += 1
            n_data_rows += 1
        for ci, h in enumerate(columns, start=1):
            ws.column_dimensions[get_column_letter(ci)].width = max(14, len(str(h)) + 8)

        # Excel 内置折线/柱状图（引用上方表格数据，第一列为分类轴）
        if charts and table_header_row and n_data_rows:
            for spec in charts:
                if isinstance(spec, dict):
                    try:
                        _add_xlsx_chart(ws, spec, table_header_row, len(columns), n_data_rows)
                    except Exception as e:
                        logger.warning("Excel 图表插入失败: %s", e)

    wb.save(fp)


def _render_docx(fp: Path, title: str, body: str, columns: list, rows: list, charts: Optional[list] = None) -> None:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    h = doc.add_heading(title or "文档", level=0)
    try:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # body 支持简单 Markdown：# / ## 标题，- 列表，其余为段落
    for raw in (body or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line.lstrip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.lstrip()), style="List Number")
        else:
            doc.add_paragraph(line)

    if columns:
        table = doc.add_table(rows=1, cols=len(columns))
        try:
            table.style = "Light Grid Accent 1"
        except Exception:
            pass
        hdr = table.rows[0].cells
        for ci, c in enumerate(columns):
            hdr[ci].text = str(c)
            for p in hdr[ci].paragraphs:
                for run in p.runs:
                    run.bold = True
        for row in (rows or []):
            cells = row if isinstance(row, (list, tuple)) else [row]
            tr = table.add_row().cells
            for ci in range(len(columns)):
                tr[ci].text = "" if ci >= len(cells) or cells[ci] is None else str(cells[ci])

    # Word 文档嵌入图表（PNG 图片）
    if charts:
        for spec in charts:
            if not isinstance(spec, dict):
                continue
            try:
                _cleanup_temp_docs()
                TEMP_DOC_DIR.mkdir(parents=True, exist_ok=True)
                tmp_token = uuid.uuid4().hex
                chart_fp = TEMP_DOC_DIR / f"{tmp_token}.png"
                _render_chart_png(chart_fp, spec)
                ct = (spec.get("title") or "图表").strip()
                doc.add_heading(ct, level=2)
                doc.add_picture(str(chart_fp), width=Inches(6.2))
                doc.add_paragraph("")
            except Exception as e:
                logger.warning("Word 图表嵌入失败: %s", e)
                doc.add_paragraph(f"（图表生成失败：{e}）")

    doc.save(fp)


def skill_create_chart(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """生成折线/柱状/饼图 PNG，供对话内预览与下载；也可供 create_document 引用相同数据结构。"""
    try:
        chart_type, _, _, _, _, _ = _parse_chart_spec(args)
    except ValueError as e:
        return {"ok": False, "message": str(e)}

    _cleanup_temp_docs()
    TEMP_DOC_DIR.mkdir(parents=True, exist_ok=True)
    token = uuid.uuid4().hex
    fp = TEMP_DOC_DIR / f"{token}.png"
    try:
        _render_chart_png(fp, args)
    except ImportError:
        return {"ok": False, "message": "服务端未安装 matplotlib，无法生成图表。请联系管理员执行 pip install matplotlib"}
    except Exception as e:
        logger.error("图表生成失败: %s", e)
        return {"ok": False, "message": f"图表生成失败：{e}"}

    title = (args.get("title") or "图表").strip()
    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title).strip() or "图表"
    filename = f"{safe_title}.png"
    url = "/api/ai-assistant/download?" + urlencode({"token": token, "filename": filename})
    type_label = {
        "line": "折线图",
        "bar": "柱状图",
        "horizontal_bar": "条形图",
        "stacked_bar": "堆叠柱状图",
        "area": "面积图",
        "pie": "饼图",
    }.get(chart_type, "图表")
    return {
        "ok": True,
        "download_url": url,
        "preview_url": url,
        "filename": filename,
        "label": f"{title}（{type_label}）",
        "chart_type": chart_type,
        "message": f"已生成{type_label}《{filename}》，对话中会直接展示预览，用户也可点击下载。",
    }


def skill_create_document(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """通用文档生成：把大模型整理好的内容渲染为可下载的 Word(docx)/Excel(xlsx)。
    适用于系统没有现成模板的自定义报表/指导文档/统计汇总等。"""
    fmt = (args.get("format") or "xlsx").strip().lower()
    if fmt not in ("xlsx", "docx"):
        return {"ok": False, "message": "format 仅支持 xlsx（Excel）/ docx（Word）"}
    title = (args.get("title") or "文档").strip()
    body = (args.get("body") or "").strip()
    columns = args.get("columns") or []
    rows = args.get("rows") or []
    charts = args.get("charts") or []
    if not isinstance(columns, list):
        columns = []
    if not isinstance(rows, list):
        rows = []
    if not isinstance(charts, list):
        charts = []
    if not body and not columns and not charts:
        return {"ok": False, "message": "缺少内容：请提供正文 body、表格 columns/rows 或 charts 图表数据"}

    _cleanup_temp_docs()
    TEMP_DOC_DIR.mkdir(parents=True, exist_ok=True)
    token = uuid.uuid4().hex
    fp = TEMP_DOC_DIR / f"{token}.{fmt}"
    try:
        if fmt == "xlsx":
            _render_xlsx(fp, title, body, columns, rows, charts=charts if charts else None)
        else:
            _render_docx(fp, title, body, columns, rows, charts=charts if charts else None)
    except Exception as e:
        logger.error(f"文档生成失败: {e}")
        return {"ok": False, "message": f"文档生成失败：{e}"}

    safe_title = re.sub(r'[\\/:*?"<>|]', "_", title).strip() or "文档"
    filename = f"{safe_title}.{fmt}"
    url = "/api/ai-assistant/download?" + urlencode({"token": token, "filename": filename})
    return {
        "ok": True,
        "download_url": url,
        "filename": filename,
        "label": title,
        "message": f"已生成《{filename}》，请提示用户点击下方按钮下载。",
    }


def skill_get_info_feed(args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    """获取中转推送的实时天气 / 新闻信息（全员可查）。"""
    kind = (args.get("kind") or "").strip().lower()
    if kind not in ("weather", "news"):
        return {"ok": False, "message": "kind 仅支持 weather / news"}
    try:
        from routers.info_feed import _load_store
        store = _load_store() or {}
    except Exception as e:
        logger.warning(f"读取信息源失败: {e}")
        return {"ok": False, "message": f"信息源暂不可用：{e}"}

    prefix = "weather:" if kind == "weather" else "news:"
    items = [(k, v) for k, v in store.items() if k.startswith(prefix)]
    if not items:
        return {"ok": True, "kind": kind, "count": 0, "items": [],
                "message": ("暂无天气数据" if kind == "weather" else "暂无新闻数据")}

    out = []
    for k, v in items[:20]:
        out.append({
            "key": k,
            "updated_at": (v or {}).get("updated_at", ""),
            "data": (v or {}).get("data"),
        })
    return {"ok": True, "kind": kind, "count": len(items), "items": out}


SKILL_FUNCS = {
    "search_policy": skill_search_policy,
    "query_overtime": skill_query_overtime,
    "query_leave": skill_query_leave,
    "query_business_trip": skill_query_business_trip,
    "query_business_trip_summary": skill_query_business_trip_summary,
    "query_dept_comparison": skill_query_dept_comparison,
    "query_person_ranking": skill_query_person_ranking,
    "query_monthly_summary": skill_query_monthly_summary,
    "query_full_attendance": skill_query_full_attendance,
    "query_work_intensity": skill_query_work_intensity,
    "query_discipline_clock": skill_query_discipline_clock,
    "query_leader_overtime": skill_query_leader_overtime,
    "query_department": skill_query_department,
    "generate_report": skill_generate_report,
    "export_report_center": skill_export_report_center,
    "query_database": skill_query_database,
    "query_shift_schedule": skill_query_shift_schedule,
    "get_info_feed": skill_get_info_feed,
    "create_document": skill_create_document,
    "create_chart": skill_create_chart,
}

SKILL_LABELS = {
    "search_policy": "制度智能检索",
    "query_overtime": "加班记录查询",
    "query_leave": "请假记录查询",
    "query_business_trip": "公出记录查询",
    "query_business_trip_summary": "公出汇总统计",
    "query_dept_comparison": "科室横向对比",
    "query_person_ranking": "人员排行统计",
    "query_monthly_summary": "月度趋势汇总",
    "query_department": "部门人员查询",
    "generate_report": "生成报表下载",
    "query_database": "数据库智能查询",
    "query_shift_schedule": "排班情况查询",
    "get_info_feed": "天气/新闻",
    "create_document": "生成文档",
    "create_chart": "生成图表",
}

SKILL_LABELS["export_report_center"] = "报表中心导出"
SKILL_LABELS["query_full_attendance"] = "满勤统计"
SKILL_LABELS["query_work_intensity"] = "工作强度统计"
SKILL_LABELS["query_discipline_clock"] = "考勤纪律审查"
SKILL_LABELS["query_leader_overtime"] = "领导加班统计"

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "search_policy",
            "description": "对部门制度、工艺规范、管理规定等文档进行语义智能检索，返回最相关的制度片段。当用户询问规章制度、流程规定、工艺要求、报销规定等知识性问题时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "检索关键词或自然语言问题"},
                    "top_k": {"type": "integer", "description": "返回条数，默认5"},
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_overtime",
            "description": "查询某员工已审批通过的加班记录与累计工时。用户问“我的加班/某人加班了多少”时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "员工姓名，不填则查询当前登录用户"},
                    "year": {"type": "integer", "description": "年份，默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12，不填为全年"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_leave",
            "description": "查询某员工已审批通过的请假记录、累计请假天数/小时及按类型分布。",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "员工姓名，不填则查询当前登录用户"},
                    "year": {"type": "integer", "description": "年份，默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12，不填为全年"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_business_trip",
            "description": "查询某员工公出（出差）记录与累计公出天数。",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "员工姓名，不填则查询当前登录用户"},
                    "year": {"type": "integer", "description": "年份，默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12，不填为全年"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_business_trip_summary",
            "description": "汇总部门/科室公出情况与公出总天数。用户问“部门近三个月公出情况”“各科室公出天数”“某科室本月公出统计”等聚合问题时优先使用本工具，不要用 query_database 自己算天数。本工具按已审批通过记录、实际/预计时间兜底、查询区间裁剪、同一人重叠区间合并去重计算。",
            "parameters": {
                "type": "object",
                "properties": {
                    "department": {"type": "string", "description": "科室名称或简称，如“智能室”“综合室”；不填时领导查全部门，普通员工查本人科室"},
                    "scope": {"type": "string", "enum": ["all"], "description": "传 all 查询全部门公出汇总"},
                    "months": {"type": "integer", "description": "最近几个月，默认3；例如近三个月传3"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；可替代 months"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；不填默认今天"},
                    "trip_type": {"type": "string", "enum": ["市内公出", "境内公出", "境外公出"], "description": "可选，指定公出类型"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_dept_comparison",
            "description": "科室横向对比统计。用户问各科室加班/请假/公出对比、人均对比、哪个科室最多、部门整体横向比较时使用。后端按统一口径计算，不要用 query_database 自行聚合。",
            "parameters": {
                "type": "object",
                "properties": {
                    "department": {"type": "string", "description": "可选，科室名称或简称；不填时按当前用户权限范围"},
                    "year": {"type": "integer", "description": "年份；默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12；不填为全年"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；可替代 year/month"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；可替代 year/month"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_person_ranking",
            "description": "人员排行统计。用户问谁加班最多、谁请假最多、谁公出最多、前十名/排行榜时使用。后端按统一口径计算，不要用 query_database 自行排序。",
            "parameters": {
                "type": "object",
                "properties": {
                    "metric": {"type": "string", "enum": ["overtime", "leave", "trip", "business_trip"], "description": "排行指标：overtime=加班小时，leave=请假天数，trip/business_trip=公出天数"},
                    "department": {"type": "string", "description": "可选，科室名称或简称；不填时按当前用户权限范围"},
                    "name": {"type": "string", "description": "可选，限定某个人；普通员工只能查本人"},
                    "year": {"type": "integer", "description": "年份；默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12；不填为全年"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；可替代 year/month"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；可替代 year/month"},
                    "limit": {"type": "integer", "description": "返回前 N 名，默认20，最多100"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_monthly_summary",
            "description": "月度趋势汇总。用户问某人/某科室/全部门每月加班、请假、公出趋势，或要求画月度折线图时使用。后端按月计算并返回可直接用于 create_chart 的 chart_hint。",
            "parameters": {
                "type": "object",
                "properties": {
                    "department": {"type": "string", "description": "可选，科室名称或简称；不填时按当前用户权限范围"},
                    "name": {"type": "string", "description": "可选，员工姓名；普通员工只能查本人"},
                    "year": {"type": "integer", "description": "年份；默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12；通常趋势不传 month"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；可替代 year/month"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；可替代 year/month"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_full_attendance",
            "description": "管理驾驶舱满勤统计。用于查询满勤率、满勤人数、满勤名单样本、各科室满勤对比；后端口径与管理驾驶舱一致。",
            "parameters": {
                "type": "object",
                "properties": {
                    "year": {"type": "integer", "description": "年份，默认当前年。"},
                    "month": {"type": "integer", "description": "月份 1-12；不传则查询全年满勤。"},
                    "department": {"type": "string", "description": "科室名称，可选；不传为全部门。"},
                    "lsys": {"type": "string", "description": "科室名称，同 department。"}
                }
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_work_intensity",
            "description": "管理驾驶舱工作强度统计。用于查询工作强度、科室强度排行、人员强度排行。后端使用驾驶舱统一口径计算，避免模型自行计算；具体口径说明和统计结果均受权限控制。",
            "parameters": {
                "type": "object",
                "properties": {
                    "year": {"type": "integer", "description": "年份，默认当前年。"},
                    "month": {"type": "integer", "description": "月份 1-12；不传为全年。"},
                    "department": {"type": "string", "description": "科室名称，可选。"},
                    "lsys": {"type": "string", "description": "科室名称，同 department。"},
                    "name": {"type": "string", "description": "员工姓名，可选。"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD。"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD。"},
                    "formula": {"type": "string", "enum": ["a", "b", "c"], "description": "工作强度口径编号，必须在权限通过后由后端解释和计算，模型不得自行展开公式。"},
                    "limit": {"type": "integer", "description": "返回个人排行条数，默认 20。"}
                }
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_discipline_clock",
            "description": "考勤纪律审查统计。用于查询早到、晚走打卡次数，支持按 person/dept/month 聚合排行。",
            "parameters": {
                "type": "object",
                "properties": {
                    "year": {"type": "integer", "description": "年份，默认当前年。"},
                    "month": {"type": "integer", "description": "月份 1-12；不传为全年。"},
                    "department": {"type": "string", "description": "科室名称，可选。"},
                    "lsys": {"type": "string", "description": "科室名称，同 department。"},
                    "dimension": {"type": "string", "enum": ["person", "dept", "month"], "description": "聚合维度，默认 person。"},
                    "clock_in_minutes": {"type": "integer", "description": "上班前阈值分钟，默认 2。"},
                    "clock_out_minutes": {"type": "integer", "description": "下班后阈值分钟，默认 2。"},
                    "limit": {"type": "integer", "description": "返回排行条数，默认 20。"}
                }
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_leader_overtime",
            "description": "领导加班统计。用于查询部办人员/领导岗位打卡识别加班、月均加班天数、估算排名。后端口径与管理驾驶舱领导加班统计一致。",
            "parameters": {
                "type": "object",
                "properties": {
                    "year": {"type": "integer", "description": "年份，默认当前年。"},
                    "month": {"type": "integer", "description": "月份 1-12；不传为全年。"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD。"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD。"},
                    "limit": {"type": "integer", "description": "返回排行条数，默认 20。"}
                }
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_department",
            "description": "查询科室或全部门的在职人数与人员名单。用户问“我们科室多少人”时按当前科室查；问“全部门/全体/整个部门多少人”时必须传 scope='all'；问“某科室有哪些人”时传 lsys。该信息属公开信息，所有用户均可查询。",
            "parameters": {
                "type": "object",
                "properties": {
                    "lsys": {"type": "string", "description": "科室名称，如“综合技术室”。不填且未传 scope 时查询当前用户所在科室"},
                    "scope": {"type": "string", "enum": ["all"], "description": "传 'all' 查询全部门全体人员（人数/名单）"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_report",
            "description": "仅用于导出【加班/请假/公出】这三类标准考勤报表（数据直接来自考勤库）。支持三种范围：单人(传 name)、某科室全员(传 lsys)、全部门全员(传 scope='all')。注意：仅当用户明确要导出加班/请假/公出考勤报表时才用本工具；若用户要的是工艺指导、知识整理、统计汇总、或任何非加班/请假/公出的自定义文档，不要用本工具，改用 create_document。",
            "parameters": {
                "type": "object",
                "properties": {
                    "report_type": {
                        "type": "string",
                        "enum": ["overtime", "leave", "business_trip"],
                        "description": "报表类型：overtime=加班，leave=请假，business_trip=公出",
                    },
                    "name": {"type": "string", "description": "员工姓名（单人报表）。不填且未指定 lsys/scope 时为当前登录用户"},
                    "lsys": {"type": "string", "description": "科室名称（科室全员汇总报表），如“综合技术室”"},
                    "scope": {"type": "string", "enum": ["all"], "description": "传 'all' 表示导出全部门全体人员汇总报表"},
                    "year": {"type": "integer", "description": "年份，默认当前年"},
                    "month": {"type": "integer", "description": "月份1-12，不填为全年"},
                },
                "required": ["report_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "export_report_center",
            "description": "导出“报表汇聚/导出中心”里已有的后端文件流报表。适用于考勤异常明细、异常处理表、全员考勤日表、考勤表Word、假期值班出勤核查、排班表、在职员工表、文件编号台账、论文保密审批台账。生成后只返回系统附件，不要在正文编造下载链接。",
            "parameters": {
                "type": "object",
                "properties": {
                    "report_type": {
                        "type": "string",
                        "enum": [
                            "attendance_exceptions",
                            "leave_handler",
                            "suggestion_attendance",
                            "attendance_report",
                            "holiday_duty",
                            "shift_schedule",
                            "employees",
                            "file_numbering",
                            "confidentiality_ledger"
                        ],
                        "description": "报表类型。attendance_exceptions=考勤异常明细；leave_handler=异常处理表；suggestion_attendance=全员考勤日表；attendance_report=考勤表Word；holiday_duty=假期值班出勤核查；shift_schedule=排班表；employees=在职员工表；file_numbering=文件编号台账；confidentiality_ledger=论文保密审批台账。"
                    },
                    "year": {"type": "integer", "description": "年份。"},
                    "month": {"type": "integer", "description": "月份 1-12；月报、考勤异常、异常处理表、考勤表Word、月排班表需要。"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；全员考勤日表、假期值班核查需要。"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；全员考勤日表、假期值班核查需要。"},
                    "department": {"type": "string", "description": "科室名称；考勤表、假期值班、排班表可用。"},
                    "lsys": {"type": "string", "description": "科室名称，同 department。"},
                    "format": {"type": "string", "enum": ["month", "week", "holiday"], "description": "排班表格式：month=月排班，week=周排班，holiday=假期值班表。"},
                    "week_date": {"type": "string", "description": "周排班所处日期 YYYY-MM-DD。"},
                    "holiday": {"type": "string", "description": "假期排班表的假期名称。"},
                    "table": {"type": "string", "enum": ["tech", "jsgl", "manage", "gygch", "scszh"], "description": "文件编号台账类型。"},
                    "keyword": {"type": "string", "description": "论文保密审批台账可选关键词。"}
                },
                "required": ["report_type"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_shift_schedule",
            "description": "查询某科室近期排班情况和值班每日计划。用户问排班、值班、白班、夜班、最近/本周/下周排班时优先使用本工具，不要用 query_database 猜排班表字段。科室简称如“智能室”应传 department='智能室'，工具会映射到真实科室。",
            "parameters": {
                "type": "object",
                "properties": {
                    "department": {"type": "string", "description": "科室名称或简称，如“智能室”“智能制造技术室”“综合室”。不填时普通员工查本人科室，领导/部办查全部业务科室"},
                    "scope": {"type": "string", "enum": ["all"], "description": "传 all 查询全部业务科室排班"},
                    "start_date": {"type": "string", "description": "开始日期 YYYY-MM-DD；不填默认今天"},
                    "end_date": {"type": "string", "description": "结束日期 YYYY-MM-DD；不填按 days 计算"},
                    "days": {"type": "integer", "description": "查询天数，默认14，最多31"},
                },
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "query_database",
            "description": (
                "通用只读数据库查询。当用户的问题超出其他专用工具的能力范围时使用——"
                "例如统计男女比例、各科室人数分布、按性别/职级/科室分组聚合、跨表统计等。"
                "查询排班/值班/白班/夜班情况时优先使用 query_shift_schedule，不要用本工具猜排班字段。"
                "查询部门/科室公出汇总和公出总天数时优先使用 query_business_trip_summary，不要用本工具自行计算公出天数。"
                "查询科室横向对比、人员排行、月度趋势时优先使用 query_dept_comparison / query_person_ranking / query_monthly_summary。"
                "你需要根据系统提供的表结构编写一条标准 MySQL SELECT 语句。"
                "仅支持只读查询，不能修改数据；敏感字段（密码/身份证/邮箱授权码等）不可访问。"
                "涉及具体个人的加班/请假/公出明细应优先使用对应的专用查询工具（受权限控制），"
                "本工具主要用于聚合统计与公开信息查询。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "sql": {
                        "type": "string",
                        "description": "完整的 MySQL SELECT 语句（只读）。可用表：yggl, jiaban, qj, gcsqb, dept_policy, tech_problem_manual, holiday, shift_schedule, shift_day_plan, shift_config。请严格按系统给出的表结构编写。",
                    },
                    "purpose": {"type": "string", "description": "本次查询要回答的问题简述"},
                },
                "required": ["sql"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_info_feed",
            "description": "获取系统中转的实时天气或新闻信息（全员可查）。当用户询问天气、气温、最近天气情况、新闻、资讯等时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "kind": {"type": "string", "enum": ["weather", "news"], "description": "weather=天气，news=新闻"},
                },
                "required": ["kind"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_document",
            "description": (
                "通用文档生成工具：把你整理/分析/检索得到的内容生成可下载的 Word 或 Excel 文件。"
                "适用于系统没有现成模板的任意自定义文档，例如：工艺指导报表、知识库整理汇总、"
                "统计分析报告、自定义清单等。先通过查询/检索工具获取真实数据，再用本工具把内容整理成文档。"
                "注意：加班/请假/公出三类标准考勤报表请用 generate_report，不要用本工具。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "format": {"type": "string", "enum": ["xlsx", "docx"], "description": "xlsx=Excel表格，docx=Word文档。以表格数据为主用 xlsx；以图文说明为主用 docx"},
                    "title": {"type": "string", "description": "文档标题（同时作为下载文件名）"},
                    "body": {"type": "string", "description": "文档正文，支持简单 Markdown（# 标题、- 列表、数字列表、段落）。Word 作为正文，Excel 作为表格上方说明，可选"},
                    "columns": {"type": "array", "items": {"type": "string"}, "description": "表格表头数组，可选"},
                    "rows": {"type": "array", "items": {"type": "array"}, "description": "表格数据，二维数组，每行与 columns 对应，可选"},
                    "charts": {
                        "type": "array",
                        "description": "可选，嵌入文档的图表列表（Word 嵌入 PNG 图片；Excel 在表格下方插入原生折线/柱状图）。每项结构与 create_chart 相同",
                        "items": {
                            "type": "object",
                            "properties": {
                                "chart_type": {"type": "string", "enum": ["line", "bar", "horizontal_bar", "stacked_bar", "area", "pie"], "description": "line=折线图，bar=柱状图，horizontal_bar=条形图，stacked_bar=堆叠柱状图，area=面积图，pie=饼图"},
                                "title": {"type": "string", "description": "图表标题"},
                                "x_label": {"type": "string", "description": "X 轴标题（折线/柱状/条形/面积图）"},
                                "y_label": {"type": "string", "description": "Y 轴标题（折线/柱状/条形/面积图）"},
                                "labels": {"type": "array", "items": {"type": "string"}, "description": "X 轴分类或饼图扇区名称"},
                                "series": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "name": {"type": "string", "description": "系列名称"},
                                            "values": {"type": "array", "items": {"type": "number"}, "description": "数值数组，长度须与 labels 一致"},
                                        },
                                        "required": ["values"],
                                    },
                                    "description": "数据系列，可多条（折线/柱状/条形/堆叠柱状/面积图支持多系列）",
                                },
                                "values": {"type": "array", "items": {"type": "number"}, "description": "单系列简写：仅一组数据时可只传 values + series_name"},
                                "series_name": {"type": "string", "description": "单系列名称（与 values 配合）"},
                            },
                            "required": ["chart_type", "title", "labels"],
                        },
                    },
                },
                "required": ["format", "title"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_chart",
            "description": (
                "生成数据可视化图表（折线图/柱状图/条形图/堆叠柱状图/面积图/饼图），输出 PNG 图片。"
                "当用户要求画图、做趋势分析、数据对比、可视化统计结果时使用。"
                "先通过 query_database / 查询类工具获取真实数据，再调用本工具绘图。"
                "图表会在对话中直接展示预览，并支持下载；下载地址只能由本工具返回的本地 /api/ai-assistant/download 生成。"
                "禁止在正文里编造 via.placeholder.com、quickchart.io、image-charts.com 或任何外部图表/占位图链接。"
                "若用户还要 Word/Excel 文档，可单独 create_document 并在 charts 参数中传入相同图表数据。"
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "chart_type": {"type": "string", "enum": ["line", "bar", "horizontal_bar", "stacked_bar", "area", "pie"], "description": "line=折线图（趋势），bar=柱状图（分类对比），horizontal_bar=条形图（人员/科室排行），stacked_bar=堆叠柱状图（构成对比），area=面积图（累计/趋势），pie=饼图（占比）"},
                    "title": {"type": "string", "description": "图表标题"},
                    "x_label": {"type": "string", "description": "X 轴标题（折线/柱状/条形/面积图，可选）"},
                    "y_label": {"type": "string", "description": "Y 轴标题（折线/柱状/条形/面积图，可选）"},
                    "labels": {"type": "array", "items": {"type": "string"}, "description": "横轴分类标签，如 ['1月','2月','3月'] 或 ['男','女']"},
                    "series": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string", "description": "系列名称，如 '加班时长'"},
                                "values": {"type": "array", "items": {"type": "number"}, "description": "数值，与 labels 等长"},
                            },
                            "required": ["values"],
                        },
                        "description": "数据系列；多系列折线/柱状/条形/堆叠柱状/面积图可传多条",
                    },
                    "values": {"type": "array", "items": {"type": "number"}, "description": "单系列简写：只传一组数据时使用"},
                    "series_name": {"type": "string", "description": "单系列名称（与 values 配合）"},
                },
                "required": ["chart_type", "title", "labels"],
            },
        },
    },
]


def _tools_for_scope(scope_info: Dict[str, str]) -> List[Dict[str, Any]]:
    """按服务端权限裁剪下发给模型的工具 schema，避免本地模型从描述里学习受限口径。"""
    if _is_manager_jb(scope_info.get("jb") or ""):
        return TOOLS
    blocked = {"query_work_intensity", "query_discipline_clock"}
    out: List[Dict[str, Any]] = []
    for tool in TOOLS:
        fn = (tool.get("function") or {}).get("name")
        if fn in blocked:
            continue
        out.append(tool)
    return out


# ==================== 数据权限（与报表/统计页 level 口径一致） ====================
# scope: all=全部门任意人员；dept=本人及本科室人员；self=仅本人


def _get_data_scope(current_user: str) -> Dict[str, str]:
    """返回 {'scope': 'all'|'dept'|'self', 'lsys': str, 'jb': str}。
    判定与 /report/statistics-permission 一致：
      - 部长/经理/副部长、综合技术室主任副主任、系统管理员(admin1)、人事管理员(admin2)、部办 → all
      - 主任/副主任/组长 → dept（本科室）
      - 其他 → self
    """
    name = (current_user or "").strip()
    if not name:
        return {"scope": "self", "lsys": "", "jb": ""}
    try:
        from routers.approvers import _get_user_info, _jb_match, can_access_leader_dashboard
    except Exception as e:
        logger.debug(f"导入权限判定失败: {e}")
        return {"scope": "self", "lsys": "", "jb": ""}

    user = _get_user_info(name) or {}
    jb = (user.get("jb") or "").strip()
    lsys = (user.get("lsys") or "").strip()

    try:
        if can_access_leader_dashboard(name):
            return {"scope": "all", "lsys": lsys, "jb": jb}
    except Exception:
        pass
    if _jb_match_safe(jb, "部长") or _jb_match_safe(jb, "副部长"):
        return {"scope": "all", "lsys": lsys, "jb": jb}
    if lsys == "部办":
        return {"scope": "all", "lsys": lsys, "jb": jb}
    if _is_manager_jb(jb):
        return {"scope": "dept", "lsys": lsys, "jb": jb}
    if _jb_match_safe(jb, "主任") or _jb_match_safe(jb, "组长") or (jb and "副主任" in jb):
        return {"scope": "dept", "lsys": lsys, "jb": jb}
    return {"scope": "self", "lsys": lsys, "jb": jb}


def _jb_match_safe(jb: str, target: str) -> bool:
    try:
        from routers.approvers import _jb_match
        return _jb_match(jb, target)
    except Exception:
        return False


def _is_manager_jb(jb: str) -> bool:
    """AI 统计白名单：yggl.jb 为经理/副经理/经理助理。"""
    s = (jb or "").strip()
    if not s:
        return False
    return bool("经理助理" in s or "副经理" in s or _jb_match_safe(s, "经理") or s == "经理")


def _is_manager_stats_global(scope_info: Dict[str, str]) -> bool:
    """工作强度/考勤纪律白名单中的部领导按部门级范围处理。"""
    return _is_manager_jb(scope_info.get("jb") or "")


def _can_export_all_attendance_reports(scope_info: Dict[str, str]) -> bool:
    """标准加班/请假/公出报表的全部门导出权限，按 yggl.jb/lsys 判断，不因 admin 身份自动放大。"""
    jb = scope_info.get("jb") or ""
    lsys = _normalize_lsys(scope_info.get("lsys") or "")
    return bool(lsys == "部办" or _jb_match_safe(jb, "部长") or _jb_match_safe(jb, "副部长"))


def _requested_lsys_from_args(fargs: Dict[str, Any]) -> str:
    raw = fargs.get("lsys") or fargs.get("department") or ""
    lsys = _normalize_lsys(raw)
    if not lsys and (fargs.get("scope") or "").strip().lower() == "all":
        return "__ALL__"
    return lsys


def _scope_desc(scope_info: Dict[str, str]) -> str:
    scope = scope_info.get("scope")
    lsys = scope_info.get("lsys") or "本科室"
    if scope == "all":
        return "全部门所有人员（可查询/导出任意员工、任意科室的数据）"
    if scope == "dept":
        return f"本人及本科室（{lsys}）全体人员"
    return "仅限本人"


def _can_access_person(scope_info: Dict[str, str], current_user: str, target_name: str) -> bool:
    scope = scope_info.get("scope")
    cu = (current_user or "").strip()
    tn = (target_name or "").strip()
    if scope == "all":
        return True
    if tn and tn == cu:
        return True
    if scope == "dept":
        cur_lsys = scope_info.get("lsys") or ""
        if not cur_lsys or not tn:
            return False
        return _get_user_dept(tn).get("lsys") == cur_lsys
    return False


def _can_access_dept(scope_info: Dict[str, str], target_lsys: str) -> bool:
    scope = scope_info.get("scope")
    target_lsys = _normalize_lsys(target_lsys)
    if scope == "all":
        return True
    cur_lsys = _normalize_lsys(scope_info.get("lsys") or "")
    return bool(target_lsys) and target_lsys == cur_lsys


def _check_skill_permission(fname: str, fargs: Dict[str, Any], current_user: str,
                            scope_info: Dict[str, str]) -> Optional[str]:
    """越权时返回拒绝原因文本；允许时返回 None。"""
    scope = scope_info.get("scope")
    own_lsys = _normalize_lsys(scope_info.get("lsys") or "")
    if fname in ("query_overtime", "query_leave", "query_business_trip"):
        target = (fargs.get("name") or current_user or "").strip()
        if not _can_access_person(scope_info, current_user, target):
            extra = "及本科室人员" if scope == "dept" else ""
            return f"权限不足：您当前只能查询本人{extra}的数据，无权查询【{target}】的记录。"
    elif fname == "generate_report":
        req_scope = (fargs.get("scope") or "").strip().lower()
        req_lsys = _normalize_lsys(fargs.get("lsys") or fargs.get("department") or "")
        if req_scope == "all":
            if not _can_export_all_attendance_reports(scope_info):
                return "权限不足：您无权导出全部门报表，仅能导出本人或权限范围内的数据。"
        elif req_lsys:
            if not _can_access_dept(scope_info, req_lsys):
                return f"权限不足：您无权导出【{req_lsys}】科室的报表。"
            fargs["lsys"] = req_lsys
            fargs["department"] = req_lsys
        else:
            target = (fargs.get("name") or current_user or "").strip()
            if not _can_access_person(scope_info, current_user, target):
                extra = "及本科室人员" if scope == "dept" else ""
                return f"权限不足：您当前只能导出本人{extra}的报表，无权导出【{target}】的报表。"
    elif fname in ("query_work_intensity", "query_discipline_clock"):
        if not _is_manager_jb(scope_info.get("jb") or ""):
            return "权限不足：工作强度和考勤纪律统计仅限 yggl.jb 为经理、副经理、经理助理的用户查询。"
        req_lsys = _requested_lsys_from_args(fargs)
        if req_lsys == "__ALL__":
            if not _is_manager_stats_global(scope_info):
                return f"权限不足：您只能查询本科室（{own_lsys}）的统计，不能查询全部门。"
        elif req_lsys:
            if not _can_access_dept(scope_info, req_lsys):
                return f"权限不足：您无权查询【{req_lsys}】的统计。"
        elif not _is_manager_stats_global(scope_info):
            fargs["lsys"] = own_lsys
            fargs["department"] = own_lsys
    elif fname in ("query_dept_comparison", "query_person_ranking", "query_monthly_summary", "query_full_attendance"):
        req_lsys = _requested_lsys_from_args(fargs)
        if scope == "self":
            target = (fargs.get("name") or "").strip()
            if not target or target != (current_user or "").strip() or fname == "query_full_attendance":
                return "权限不足：该统计工具仅支持经理/主任等管理范围查询；普通员工请查询本人明细。"
        if req_lsys == "__ALL__":
            if scope != "all":
                return f"权限不足：您只能查询本科室（{own_lsys}）统计，不能查询全部门。"
        elif req_lsys:
            if not _can_access_dept(scope_info, req_lsys):
                return f"权限不足：您无权查询【{req_lsys}】的统计。"
        elif scope == "dept":
            fargs["lsys"] = own_lsys
            fargs["department"] = own_lsys
    elif fname == "export_report_center":
        raw_type = (fargs.get("report_type") or fargs.get("report_id") or "").strip().lower().replace("-", "_")
        aliases = {
            "attendance_exception": "attendance_exceptions",
            "exceptions": "attendance_exceptions",
            "attendance_day": "suggestion_attendance",
            "daily_attendance": "suggestion_attendance",
            "attendance_word": "attendance_report",
            "duty": "holiday_duty",
            "shift": "shift_schedule",
            "employee": "employees",
            "numbering": "file_numbering",
            "confidentiality": "confidentiality_ledger",
        }
        report_type = aliases.get(raw_type, raw_type)
        req_lsys = _requested_lsys_from_args(fargs)
        scoped_report_types = {"attendance_report", "holiday_duty", "shift_schedule"}
        all_only_report_types = {
            "attendance_exceptions", "leave_handler", "suggestion_attendance",
            "employees", "file_numbering", "confidentiality_ledger",
        }
        if scope == "self":
            return "权限不足：报表汇聚导出仅限具备科室或部门管理权限的用户使用。"
        if report_type in all_only_report_types:
            if scope != "all":
                return "权限不足：该报表为全部门报表，您当前无权导出。"
        elif report_type in scoped_report_types:
            if req_lsys == "__ALL__":
                if scope != "all":
                    return f"权限不足：您只能导出本科室（{own_lsys}）报表，不能导出全部门。"
            elif req_lsys:
                if not _can_access_dept(scope_info, req_lsys):
                    return f"权限不足：您无权导出【{req_lsys}】报表。"
                normalized = _normalize_lsys(req_lsys)
                fargs["lsys"] = normalized
                fargs["department"] = normalized
            elif scope == "dept":
                fargs["lsys"] = own_lsys
                fargs["department"] = own_lsys
    elif fname == "query_database":
        # 公开表（员工/知识库/排班/节假日/制度）全员可查；隐私表（请假/加班/公出/换休）受权限控制：
        # 仅 scope=all（经理/副经理/综合技术室主任副主任）可查逐条个人明细，
        # scope=dept/self 只能对隐私表做聚合统计，避免用自由 SQL 绕过隐私明细权限。
        sql_low = (fargs.get("sql") or "").lower()
        purpose = (fargs.get("purpose") or "").strip()
        user_text = (fargs.get("_user_text") or "").strip()
        used_tables = set(re.findall(r"(?:from|join)\s+`?([a-zA-Z_]\w*)`?", sql_low))
        discipline_like = bool(
            used_tables & {"attendance_records"}
            or _is_work_intensity_or_discipline_question(user_text)
            or re.search(r"(工作强度|考勤纪律|早到|晚走|打卡纪律|踩点|在岗小时|实际在岗)", purpose)
            or re.search(r"(工作强度|考勤纪律|早到|晚走|打卡纪律|踩点|在岗小时|实际在岗)", fargs.get("sql") or "")
        )
        if discipline_like and not _is_manager_jb(scope_info.get("jb") or ""):
            return "权限不足：工作强度和考勤纪律相关数据仅限 yggl.jb 为经理、副经理、经理助理的用户查询。"
        if (used_tables & PRIVATE_TABLES) and scope != "all":
            is_agg = bool(
                re.search(r"\b(count|sum|avg|min|max)\s*\(", sql_low)
                or re.search(r"\bgroup\s+by\b", sql_low)
            )
            if not is_agg:
                return ("权限不足：请假/加班/公出/换休属个人隐私数据，您当前权限仅可对其做聚合统计"
                        "（计数/求和/分组占比等），不能查询逐条个人明细。如需查询个人记录，"
                        "请使用对应的加班/请假/公出查询工具（按您的权限范围）。")
    # query_department：部门人数/名单等属公开信息，不做权限限制
    return None


def _execute_skill(name: str, args: Dict[str, Any], current_user: str) -> Dict[str, Any]:
    func = SKILL_FUNCS.get(name)
    if not func:
        return {"ok": False, "message": f"未知技能：{name}"}
    try:
        return func(args, current_user)
    except Exception as e:
        logger.exception(f"技能 {name} 执行失败")
        return {"ok": False, "message": f"技能执行失败：{e}"}


# ==================== System Prompt ====================


def _build_system_prompt(current_user: str, scope_info: Dict[str, str]) -> str:
    today = datetime.now().strftime("%Y-%m-%d")
    dept = scope_info.get("lsys") or "未知"
    scope_text = _scope_desc(scope_info)
    return (
        "你是「智能制造工艺部 AI 助手」，服务于哈电智能制造工艺部门的集成办公平台。"
        "你可以整合系统数据库资源，帮助员工查询考勤、加班、请假、公出数据，统计部门人员，检索部门制度与工艺规范，并生成可下载的报表。\n\n"
        + DEPARTMENT_KNOWLEDGE
        + "\n\n"
        f"当前登录用户：{current_user or '未知'}；所在科室：{dept}；今天日期：{today}。\n"
        f"该用户的数据查询权限范围：{scope_text}。\n\n"
        "工作要求：\n"
        "1. 始终使用简体中文回答，语气专业、简洁、友好。\n"
        "2. 凡是涉及系统数据的问题（加班/请假/公出记录与统计、部门人数、人员职级/科室、制度规定、报表导出），"
        "必须调用相应的工具获取真实数据。答案中出现的每一个姓名、职级、科室、数字等事实，"
        "都必须严格来自工具返回的结果，绝不能凭空编造、猜测或用常识推断；"
        "工具结果中没有的信息，要明确说明“未查询到”，不要自行补全。\n"
        "3. 用户未指明姓名时，默认查询当前登录用户；未指明科室时，默认当前用户所在科室。"
        "用户提到科室简称（水轮机室/水发室/汽发室/焊艺室/综合室/非标室/工具室/智能室/数控室/部办等）时，"
        "须先按【部门组织结构】映射为正式 lsys 名称再调用工具。\n"
        "4. 用户要求导出/下载报表时，调用 generate_report 生成下载链接，并明确告诉用户报表已生成、可点击下方按钮下载。"
        "注意区分导出范围：只有部门级权限用户导出“全部门/全体/所有人员”报表时才传 scope='all'；"
        "科室级用户说“我们部门/我们科室/本部门”时默认传当前用户所在 lsys；只导出单个人时传 name。"
        "若用户无全部门权限，不得主动传 scope='all' 或拼全部门下载链接。\n"
        "   注意：生成下载文件后，正文中只需提示“点击下方按钮下载”，切勿把下载网址（URL/链接地址）直接写进回答文本里——"
        "下载按钮会由系统自动渲染在消息下方。\n"
        "5. 引用制度检索结果时注明制度标题，并基于检索到的片段作答，避免脱离原文。\n"
        "6. 回答使用 Markdown（标题、列表、表格、加粗）让结构清晰美观。\n"
        "7. 工具返回无数据时，如实说明未查询到相关记录，不要编造。\n"
        "7a. 工具返回可能是「摘要视图」（含 total_count、by_department、by_jb、by_gender、aggregates、"
        "sample_rows 等汇总字段，而非完整明细）。回答统计/分布/对比类问题时，优先依据这些汇总数字作答；"
        "若用户需要完整名单或明细表格，应调用 generate_report / create_document / create_chart 重新导出，"
        "不要因摘要中没有逐条明细就编造或猜测。\n"
        "8. 严格遵守该用户的数据权限范围：若其权限不足以查询他人/其他科室数据，工具会返回"
        "“权限不足”，此时应礼貌告知用户其权限范围，不要尝试绕过。\n"
        "8a. 用户询问部门/科室/近几个月的公出汇总、公出总天数、各科室公出对比时，"
        "必须调用 query_business_trip_summary，并直接使用工具返回的 total_days、by_department、by_person。"
        "不要用 query_database 返回明细后自行计算公出天数。\n"
        "8b. 用户询问科室横向对比、人员排行、月度趋势、按月折线图、最多/最少/Top N 等统计问题时，"
        "必须优先调用 query_dept_comparison、query_person_ranking 或 query_monthly_summary。"
        "不要用 query_database 拉明细后让模型自己计算、排序或画趋势。\n"
        "9. 当问题超出专用工具的能力（如男女比例、各科室人数分布、按职级/性别分组统计、"
        "员工职级/邮箱/排班/节假日、工艺问题知识库检索等），调用 query_database 编写只读 SELECT 查询，"
        "严格依据下方表结构编写，禁止查询敏感字段；写不出合规 SQL 时如实说明。\n"
        "9a. 查询排班/值班/白班/夜班/每日计划时，优先调用 query_shift_schedule；"
        "不要用 query_database 猜测排班表字段，也不要在最终回答中输出 <｜DSML｜tool_calls> 等工具调用标记。\n"
        "10. 当用户咨询工艺/技术问题时，可检索 tech_problem_manual 知识库中的同类问题"
        "（problem_desc/cause_analysis/measures），并结合检索到的处理措施给出专业建议。\n"
        "11. 当用户询问天气或新闻时，调用 get_info_feed 获取系统中转的实时信息后作答。\n"
        "12. 关于“报表/文档导出”要正确区分：\n"
        "    - 仅当用户要导出【加班/请假/公出】标准考勤报表时，才用 generate_report；\n"
        "    - 其他任何需要生成可下载文件的场景（如工艺指导报表、知识库整理、统计分析报告、"
        "自定义清单等），先用查询/检索工具拿到真实数据，再用 create_document 生成 Word 或 Excel。\n"
        "    - 不要一看到“报表”二字就盲目导出加班报表；要先理解用户真正想要的内容主题。\n"
        "13. 当用户要求画图，或问题本身适合用统计图表达时（趋势、排行、横向对比、占比、构成、驾驶舱统计等）：\n"
        "    - 先用查询工具拿到真实数据，再调用 create_chart 生成 PNG 图表（对话中会直接展示预览）。\n"
        "    - 图表文件必须由本系统本地生成，下载地址只能来自 create_chart 返回的 /api/ai-assistant/download 附件；"
        "正文不要写任何外部图表链接、占位图链接或项目官网链接，尤其禁止 via.placeholder.com、quickchart.io、image-charts.com。\n"
        "    - 选择图表类型：趋势默认 line，分类对比默认 bar，人员/科室排行优先 horizontal_bar，构成占比用 pie，多系列构成对比用 stacked_bar，累计或连续趋势可用 area。\n"
        "    - 若用户还要 Word/Excel 文档且需要图表，create_document 的 charts 参数传入相同图表数据"
        "（Word 嵌入图片，Excel 在表格下方插入原生图表；Excel 需同时提供 columns/rows 表格数据）。\n"
        "    - 不要用 Markdown/ASCII 字符画折线代替真实图表。\n\n"
        + "\n报表工具规则：用户要导出“报表汇聚/导出中心”中的考勤异常明细、异常处理表、全员考勤日表、考勤表Word、假期值班出勤核查、排班表、在职员工表、文件编号台账、论文保密审批台账时，必须调用 export_report_center。只有导出加班/请假/公出三类标准明细报表时才调用 generate_report。工资、总时长、科室对比、人员排行、月度趋势等统计类问题优先调用 query_dept_comparison、query_person_ranking、query_monthly_summary 等统计工具获取系统计算结果，不要先拉 SQL 明细后自行计算。\n"
        + "管理驾驶舱统计规则：满勤/满勤率/满勤名单必须调用 query_full_attendance；工作强度/负荷/强度排行/实际在岗小时必须调用 query_work_intensity；考勤纪律审查/早到晚走/打卡纪律排行必须调用 query_discipline_clock；领导加班统计/部办加班/月均加班天数/领导估算排名必须调用 query_leader_overtime。工作强度和考勤纪律仅限 yggl.jb 为经理、副经理、经理助理的用户查询；该白名单用户可查询全体人员或指定科室。\n"
        + DB_SCHEMA_HINT
    )


# ==================== 流式对话接口 ====================


def _sse(payload: dict) -> str:
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


@router.get("/model-info")
def model_info(current_user: str = Query("", description="当前登录用户（可选）")):
    """返回当前生效的大模型展示信息（仅模型名/类型，不含 base_url、api_key 等敏感信息），供前端只读展示。"""
    try:
        cfg = _resolve_llm()
        return {"success": True, "provider": cfg.get("provider"), "model": cfg.get("model"), "label": _model_label(cfg)}
    except Exception as e:
        logger.debug(f"获取模型信息失败: {e}")
        return {"success": False, "label": ""}


@router.post("/chat-stream")
async def chat_stream(req: ChatRequest):
    """带工具调用的流式对话（SSE）。真·流式：边生成边逐字输出。"""
    try:
        from openai import OpenAI
    except ImportError:
        raise HTTPException(status_code=500, detail="服务端未安装 openai SDK，无法调用大模型")

    cfg = _resolve_llm()
    if not cfg.get("base_url") or not cfg.get("model"):
        raise HTTPException(status_code=500, detail="未配置可用大模型：请在 webconfig 配置 deepseek_api_key，或 llm_base_url/llm_model")

    current_user = (req.current_user or "").strip()
    provider = cfg["provider"]
    model = cfg["model"]
    scope_info = _get_data_scope(current_user)

    history: List[Dict[str, Any]] = [{"role": "system", "content": _build_system_prompt(current_user, scope_info)}]
    for m in (req.messages or []):
        role = (m.role or "").strip()
        if role not in ("user", "assistant"):
            continue
        content = (m.content or "").strip()
        if not content:
            continue
        history.append({"role": role, "content": content})

    if len(history) <= 1:
        raise HTTPException(status_code=400, detail="对话内容为空")

    latest_user_text = ""
    for hm in reversed(history):
        if hm.get("role") == "user":
            latest_user_text = hm.get("content") or ""
            break

    if _is_work_intensity_or_discipline_question(latest_user_text) and not _is_manager_jb(scope_info.get("jb") or ""):
        def denied_gen():
            yield _sse({"type": "meta", "provider": provider, "model": model, "label": model_label})
            msg = (
                "权限不足：工作强度、考勤纪律及其统计口径说明仅限 yggl.jb 为"
                "经理、副经理、经理助理的用户查询。"
            )
            yield _sse({"type": "chunk", "text": msg})
            yield _sse({"type": "done"})

        model_label = _model_label(cfg)
        return StreamingResponse(
            denied_gen(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",
                "Connection": "keep-alive",
            },
        )

    history = _trim_chat_history(history)

    extra_body = {"chat_template_kwargs": {"enable_thinking": False}} if cfg.get("use_extra") else None

    model_label = _model_label(cfg)
    available_tools = _tools_for_scope(scope_info)

    def gen():
        yield _sse({"type": "meta", "provider": provider, "model": model, "label": model_label})
        messages = list(history)
        latest_user_text = ""
        for hm in reversed(history):
            if hm.get("role") == "user":
                latest_user_text = hm.get("content") or ""
                break
        try:
            client = OpenAI(base_url=cfg["base_url"], api_key=cfg["api_key"], timeout=120.0)
        except Exception as e:
            yield _sse({"type": "error", "message": f"初始化大模型客户端失败：{e}"})
            return

        # ---------- 阶段 1：工具调用循环（非流式，DeepSeek 工具调用很快） ----------
        # 注意：不在带 tools 的请求里做流式输出——很多 OpenAI 兼容网关在 tools+stream
        # 模式下会缓冲后一次性返回文本，导致前端看不到逐字效果。最终答案放到阶段 2 纯流式生成。
        try:
            for _round in range(MAX_TOOL_ROUNDS):
                yield _sse({"type": "status", "text": "正在理解问题并规划任务…" if _round == 0 else "正在结合查询结果继续分析…"})
                kwargs = _llm_request_kwargs(cfg,
                    model=model,
                    messages=messages,
                    tools=available_tools,
                    tool_choice="auto",
                    temperature=LLM_TEMPERATURE_TOOL,
                    stream=False,
                )
                if extra_body:
                    kwargs["extra_body"] = extra_body
                try:
                    resp = client.chat.completions.create(**kwargs)
                except Exception as api_err:
                    err_s = str(api_err)
                    if "json_invalid" in err_s or "JSON decode" in err_s or "400" in err_s:
                        if _has_tool_context(messages):
                            logger.warning(
                                "工具结果后的继续规划请求被模型网关拒绝，改为直接生成最终回答(%s, user=%s, msgs=%d): %s",
                                provider, current_user, len(messages), err_s,
                            )
                            break
                        logger.error("大模型请求体异常(%s, user=%s, msgs=%d): %s",
                                       provider, current_user, len(messages), err_s)
                        yield _sse({"type": "error", "message":
                            "模型网关拒绝请求（请求体过大或格式异常）。"
                            "领导账号查询全部门数据时较常见，已启用截断；请重试或缩小查询范围（如指定科室/月份）。"})
                        return
                    raise
                msg = resp.choices[0].message
                tool_calls = getattr(msg, "tool_calls", None) or []
                if not tool_calls:
                    break

                messages.append({
                    "role": "assistant",
                    "content": msg.content or "",
                    "tool_calls": [
                        {
                            "id": tc.id or f"call_{i}",
                            "type": "function",
                            "function": {
                                "name": tc.function.name,
                                "arguments": _safe_tool_arguments(tc.function.arguments or "{}"),
                            },
                        }
                        for i, tc in enumerate(tool_calls)
                    ],
                })

                tool_names_this_round = {tc.function.name for tc in tool_calls if getattr(tc, "function", None)}
                auto_chart_notes: List[str] = []

                for i, tc in enumerate(tool_calls):
                    fname = tc.function.name
                    label = SKILL_LABELS.get(fname, fname)
                    try:
                        fargs = json.loads(tc.function.arguments or "{}")
                    except Exception:
                        fargs = {}
                    fargs["_user_text"] = latest_user_text
                    yield _sse({"type": "tool", "name": fname, "label": label, "status": "running"})
                    yield _sse({"type": "status", "text": f"正在{label}…"})

                    # 数据权限校验：越权直接拒绝，不执行查询/不生成下载链接
                    denied = _check_skill_permission(fname, fargs, current_user, scope_info)
                    if denied:
                        result = {"ok": False, "message": denied}
                    else:
                        result = _execute_skill(fname, fargs, current_user)

                    if fname in ("generate_report", "export_report_center", "create_document", "create_chart") and result.get("ok") and result.get("download_url"):
                        att_evt = {
                            "type": "attachment",
                            "label": result.get("label") or "文件下载",
                            "url": result.get("download_url"),
                            "filename": result.get("filename") or "document",
                        }
                        if fname == "create_chart" or (result.get("filename") or "").lower().endswith(".png"):
                            att_evt["kind"] = "image"
                        yield _sse(att_evt)

                    if fname != "create_chart" and "create_chart" not in tool_names_this_round and isinstance(result, dict):
                        chart_spec = _auto_chart_spec_from_tool_result(fname, result, fargs, latest_user_text)
                        if chart_spec:
                            yield _sse({"type": "tool", "name": "create_chart", "label": SKILL_LABELS.get("create_chart", "生成图表"), "status": "running"})
                            yield _sse({"type": "status", "text": "正在生成统计图…"})
                            chart_result = _execute_skill("create_chart", chart_spec, current_user)
                            if chart_result.get("ok") and chart_result.get("download_url"):
                                att_evt = {
                                    "type": "attachment",
                                    "label": chart_result.get("label") or "统计图",
                                    "url": chart_result.get("download_url"),
                                    "filename": chart_result.get("filename") or "chart.png",
                                    "kind": "image",
                                }
                                yield _sse(att_evt)
                                auto_chart_notes.append(f"已自动生成本地统计图附件：{chart_result.get('label') or chart_result.get('filename') or '统计图'}。")
                                yield _sse({"type": "tool", "name": "create_chart", "label": SKILL_LABELS.get("create_chart", "生成图表"), "status": "done", "summary": "已生成图表"})
                            else:
                                yield _sse({"type": "tool", "name": "create_chart", "label": SKILL_LABELS.get("create_chart", "生成图表"), "status": "done", "summary": chart_result.get("message", "图表生成失败")})

                    summary = ""
                    if isinstance(result, dict):
                        if not result.get("ok", True):
                            summary = result.get("message", "执行失败")
                        elif "total_count" in result:
                            summary = f"共 {result.get('total_count', 0)} 条"
                        elif "row_count" in result:
                            summary = f"返回 {result.get('row_count', 0)} 行"
                        elif "count" in result:
                            summary = f"命中 {result.get('count', 0)} 条"
                    yield _sse({"type": "tool", "name": fname, "label": label, "status": "done", "summary": summary})

                    messages.append({
                        "role": "tool",
                        "tool_call_id": tc.id or f"call_{i}",
                        "content": _tool_result_json_for_llm(fname, result, fargs),
                    })

                if auto_chart_notes:
                    messages.append({
                        "role": "assistant",
                        "content": "\n".join(auto_chart_notes) + "\n最终回答中不要再编写 Markdown 图片链接；只需说明统计图已在对话中展示，可点击附件查看或下载。",
                    })
        except Exception as e:
            logger.error(f"AI 助手工具调用失败({provider}): {e}")
            yield _sse({"type": "error", "message": f"模型调用失败：{e}"})
            return

        # ---------- 阶段 2：最终答案（纯流式，不带 tools，逐 token 输出） ----------
        # DeepSeek 优先用 reasoner 模型：先流式输出思维链(reasoning_content)，再输出正文，
        # 让前端能看到"思考过程"。reasoner 不可用时回退普通对话模型。本地模型直接流式。
        def _emit_final(model2):
            final_messages = _plain_messages_for_answer(messages) if _has_tool_context(messages) else messages
            kwargs = _llm_request_kwargs(cfg, model=model2, messages=final_messages, stream=True)
            if provider != "deepseek":
                kwargs["temperature"] = LLM_TEMPERATURE_ANSWER
                if extra_body:
                    kwargs["extra_body"] = extra_body
            stream = client.chat.completions.create(**kwargs)
            produced = False
            suppress_tool_markup = False
            for event in stream:
                if not getattr(event, "choices", None):
                    continue
                delta = event.choices[0].delta
                # 思维链（DeepSeek reasoner 扩展字段）
                rc = getattr(delta, "reasoning_content", None)
                if rc is None:
                    me = getattr(delta, "model_extra", None)
                    if me:
                        rc = me.get("reasoning_content")
                if rc:
                    yield _sse({"type": "reasoning", "text": rc})
                piece = getattr(delta, "content", None) or ""
                if piece and ("<｜DSML" in piece or "<|DSML" in piece):
                    piece = re.split(r"<[｜|]DSML", piece, maxsplit=1)[0]
                    suppress_tool_markup = True
                if suppress_tool_markup:
                    if piece:
                        piece = _strip_external_chart_links(piece)
                        produced = True
                        yield _sse({"type": "chunk", "text": piece})
                    continue
                if piece:
                    piece = _strip_external_chart_links(piece)
                if piece:
                    produced = True
                    yield _sse({"type": "chunk", "text": piece})
            return produced

        # 关键：deepseek-reasoner 不支持 function calling，无法消费对话历史里的
        # tool_calls / tool 结果消息。若本轮调用过工具仍用 reasoner，会丢失真实数据导致“胡编乱造”。
        # 因此：本轮调用过工具时，最终答案必须用 deepseek-chat（能正确读取工具返回结果）；
        # 纯对话（未调用工具）时才用 reasoner 展示思维链。
        used_tools = any(isinstance(m, dict) and m.get("role") == "tool" for m in messages)
        yield _sse({"type": "status", "text": "正在生成回答…"})
        try:
            got = False
            if provider == "deepseek" and not used_tools:
                try:
                    got = yield from _emit_final(DEEPSEEK_REASONER_MODEL)
                except Exception as e_reason:
                    logger.warning(f"DeepSeek reasoner 不可用，回退 {model}: {e_reason}")
                    got = yield from _emit_final(model)
            else:
                got = yield from _emit_final(model)
            if not got:
                yield _sse({"type": "chunk", "text": "（未生成回复，请重试或换一种问法）"})
            yield _sse({"type": "done"})
        except Exception as e:
            logger.error(f"AI 助手流式输出失败({provider}): {e}")
            yield _sse({"type": "error", "message": f"流式输出失败：{e}"})

    return StreamingResponse(
        gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
            "Connection": "keep-alive",
        },
    )


# ==================== 报表导出（Excel 下载） ====================


def _list_members(target_lsys: Optional[str]) -> List[Tuple[str, str]]:
    """返回 [(姓名, 科室)] 列表。target_lsys 为空时返回全部门在职人员，否则返回该科室人员。
    过滤规则与 /report/statistics-employees 一致（排除带尾缀'1'、其他部门、离职）。"""
    base = (
        "SELECT name, lsys FROM yggl WHERE name IS NOT NULL AND name != '' "
        "AND RIGHT(TRIM(name),1) != '1' AND RIGHT(TRIM(lsys),1) != '1' "
        "AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0)"
    )
    try:
        if target_lsys:
            rows = db.execute_query(base + " AND lsys = %s ORDER BY name", (target_lsys,))
        else:
            rows = db.execute_query(base + " ORDER BY lsys, name")
    except Exception as e:
        logger.error(f"获取人员名单失败: {e}")
        return []
    return [((r.get("name") or "").strip(), (r.get("lsys") or "").strip()) for r in rows if r.get("name")]


@router.get("/export")
def export_report(
    report_type: str = Query(..., description="overtime|leave|business_trip"),
    name: Optional[str] = Query(None, description="员工姓名（单人报表）"),
    lsys: Optional[str] = Query(None, description="科室名称（科室全员汇总）"),
    scope: Optional[str] = Query(None, description="传 all 表示全部门全员汇总"),
    year: Optional[int] = Query(None),
    month: Optional[int] = Query(None, ge=1, le=12),
    requester: Optional[str] = Query(None, description="发起下载的当前用户，用于权限校验"),
):
    """生成并下载加班 / 请假 / 公出 Excel 报表。
    支持单人(name) / 科室全员(lsys) / 全部门(scope=all) 三种范围。"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        raise HTTPException(status_code=500, detail="服务端未安装 openpyxl，无法导出 Excel")

    headers_map = {
        "overtime": (["日期", "加班方式", "时长(小时)", "内容"], "加班报表", skill_query_overtime,
                     lambda r: [r["date"], r["type"], r["hours"], r["content"]]),
        "leave": (["请假类型", "开始时间", "结束时间", "天数", "小时", "事由"], "请假报表", skill_query_leave,
                  lambda r: [r["type"], r["from"], r["to"], r["days"], r["hours"], r["content"]]),
        "business_trip": (["外派单位", "公出地点", "公出时间", "预计返回", "实际返回", "公出任务", "天数"], "公出报表",
                          skill_query_business_trip,
                          lambda r: [r["unit"], r["place"], r["start"], r["expected_return"], r["actual_return"], r["task"], r["days"]]),
    }
    if report_type not in headers_map:
        raise HTTPException(status_code=400, detail="report_type 仅支持 overtime / leave / business_trip")

    scope_v = (scope or "").strip().lower()
    lsys_v = (lsys or "").strip()
    name_v = (name or "").strip()
    req = (requester or "").strip()

    # 数据权限校验：按发起者(requester)权限范围核对导出范围
    if req:
        scope_info = _get_data_scope(req)
        if scope_v == "all":
            if not _can_export_all_attendance_reports(scope_info):
                raise HTTPException(status_code=403, detail="权限不足：您无权导出全部门报表")
        elif lsys_v:
            if not _can_access_dept(scope_info, lsys_v):
                raise HTTPException(status_code=403, detail=f"权限不足：您无权导出【{lsys_v}】科室报表")
        else:
            tgt = name_v or req
            if not _can_access_person(scope_info, req, tgt):
                raise HTTPException(status_code=403, detail=f"权限不足：您无权下载【{tgt}】的报表")

    year = year or datetime.now().year
    period = f"{year}年" + (f"{month}月" if month else "全年")
    cols, sheet_title, skill_func, row_mapper = headers_map[report_type]

    is_batch = scope_v == "all" or bool(lsys_v)
    if scope_v == "all":
        members = _list_members(None)
        batch_title = "全部门"
    elif lsys_v:
        members = _list_members(lsys_v)
        batch_title = lsys_v
    else:
        nm = name_v or req
        if not nm:
            raise HTTPException(status_code=400, detail="缺少导出对象：请指定 name / lsys / scope")
        members = [(nm, _get_user_dept(nm).get("lsys", ""))]
        batch_title = nm

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_title

    header_fill = PatternFill(start_color="1890FF", end_color="1890FF", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    out_cols = (["科室", "姓名"] + cols) if is_batch else cols

    title_text = f"{batch_title} - {sheet_title}（{period}）"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(out_cols))
    tcell = ws.cell(row=1, column=1, value=title_text)
    tcell.font = Font(size=14, bold=True)
    tcell.alignment = Alignment(horizontal="center", vertical="center")

    for ci, h in enumerate(out_cols, start=1):
        c = ws.cell(row=2, column=ci, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center", vertical="center")

    r_idx = 3
    g_count = 0
    g_hours = 0.0
    g_days = 0.0
    for mname, mlsys in members:
        data = skill_func({"name": mname, "year": year, "month": month}, mname)
        records = data.get("records", []) if isinstance(data, dict) else []
        g_count += int(data.get("total_count", 0) or 0)
        g_hours += float(data.get("total_hours", 0) or 0)
        g_days += float(data.get("total_days", 0) or 0)
        for rec in records:
            mapped = list(row_mapper(rec))
            row_vals = ([mlsys, mname] + mapped) if is_batch else mapped
            for ci, val in enumerate(row_vals, start=1):
                ws.cell(row=r_idx, column=ci, value=val)
            r_idx += 1

    if report_type == "overtime":
        total_text = f"合计：{g_count} 条，{round(g_hours, 2)} 小时"
    elif report_type == "leave":
        total_text = f"合计：{g_count} 条，{round(g_days, 2)} 天 / {round(g_hours, 2)} 小时"
    else:
        total_text = f"合计：{g_count} 条，{round(g_days, 2)} 天"
    if is_batch:
        total_text = f"共 {len(members)} 人，" + total_text
    ws.cell(row=r_idx, column=1, value=total_text).font = Font(bold=True)

    for ci, h in enumerate(out_cols, start=1):
        ws.column_dimensions[ws.cell(row=2, column=ci).column_letter].width = max(12, len(str(h)) + 8)

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    filename = f"{batch_title}-{sheet_title}-{period}.xlsx"
    disposition = f"attachment; filename*=UTF-8''{quote(filename)}"
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": disposition},
    )


# ==================== 通用文档下载（create_document 生成物） ====================


@router.get("/download")
def download_document(
    token: str = Query(..., description="文档标识"),
    filename: Optional[str] = Query(None, description="下载文件名"),
):
    """下载 create_document 生成的临时 Word/Excel 文件。"""
    if not re.fullmatch(r"[0-9a-fA-F]{8,64}", token or ""):
        raise HTTPException(status_code=400, detail="无效的下载标识")
    for ext, mime in (("xlsx", XLSX_MIME), ("docx", DOCX_MIME), ("png", PNG_MIME)):
        fp = TEMP_DOC_DIR / f"{token}.{ext}"
        if fp.exists():
            dl_name = (filename or fp.name).strip() or fp.name
            return FileResponse(path=str(fp), media_type=mime, filename=dl_name)
    raise HTTPException(status_code=404, detail="文件不存在或已过期，请重新生成")
