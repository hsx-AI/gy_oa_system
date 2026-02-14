# -*- coding: utf-8 -*-
"""
部门制度向量检索服务 - bge-small-zh-v1.5 + Chroma
"""
import os
import logging
from typing import Optional, List, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

_BASE_DIR = Path(__file__).resolve().parent.parent
CHROMA_DIR = _BASE_DIR / "data" / "policy_chroma"
COLLECTION_NAME = "dept_policy_cos"  # 使用余弦距离，需重新回填

_model = None
_chroma_client = None
_collection = None


def _get_model_path():
    """模型路径：优先从配置读取，否则使用 HuggingFace 模型名（首次自动下载）"""
    from config import settings
    path = getattr(settings, "EMBEDDING_MODEL_PATH", "") or ""
    path = str(path).strip()
    if path and os.path.isdir(path):
        return path
    return "BAAI/bge-small-zh-v1.5"


def _get_model():
    """懒加载 embedding 模型"""
    global _model
    if _model is not None:
        return _model
    try:
        from sentence_transformers import SentenceTransformer
        path = _get_model_path()
        logger.info(f"加载向量模型: {path}")
        _model = SentenceTransformer(path)
        return _model
    except Exception as e:
        logger.error(f"加载向量模型失败: {e}")
        raise


def _get_collection():
    """获取 Chroma 集合（使用余弦距离，便于相关性分数直观 0-100%）"""
    global _chroma_client, _collection
    if _collection is not None:
        return _collection
    try:
        import chromadb
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        _collection = _chroma_client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={"hnsw:space": "cosine", "description": "部门制度向量库"}
        )
        return _collection
    except Exception as e:
        logger.error(f"Chroma 初始化失败: {e}")
        raise


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """从 PDF/Word/Excel 提取正文文本"""
    path = Path(file_path)
    if not path.is_file():
        return ""
    ft = (file_type or "").lower()
    text_parts = []
    try:
        if ft == "pdf":
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
        elif ft in ("doc", "docx"):
            from docx import Document
            doc = Document(str(path))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
        elif ft in ("xls", "xlsx"):
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    vals = [str(v).strip() for v in (row or []) if v is not None and str(v).strip()]
                    if vals:
                        text_parts.append(" ".join(vals))
            wb.close()
    except Exception as e:
        logger.warning(f"提取文本失败 {path}: {e}")
    text = "\n".join(text_parts).strip()
    if len(text) > 50000:
        text = text[:50000]
    return text


def _split_chunks(text: str, chunk_size: int, overlap: int) -> List[str]:
    """按字符切分，带重叠"""
    if not text or chunk_size <= 0:
        return [text] if text else []
    chunks = []
    stride = max(1, chunk_size - overlap)
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += stride
        if start >= len(text):
            break
    return chunks if chunks else [text[:chunk_size]]


def add_to_index(policy_id: str, title: str, issue_time: str, remark: str, file_path: str, file_type: str) -> bool:
    """将制度加入向量库（按切片存储，便于展示匹配切片）"""
    try:
        from config import settings
        chunk_size = getattr(settings, "VECTOR_CHUNK_SIZE", 400) or 400
        overlap = getattr(settings, "VECTOR_CHUNK_OVERLAP", 80) or 80
        chunk_size = max(100, min(2000, int(chunk_size)))
        overlap = max(0, min(chunk_size - 50, int(overlap)))

        coll = _get_collection()
        model = _get_model()
        file_full = _BASE_DIR / "data" / file_path.replace("/", os.sep)
        doc_text = extract_text_from_file(str(file_full), file_type)
        prefix = f"{title or ''}\n{issue_time or ''}\n{remark or ''}\n".strip()
        combined = f"{prefix}\n{doc_text}".strip()
        if not combined:
            combined = title or policy_id

        chunks = _split_chunks(combined, chunk_size, overlap)
        if not chunks:
            chunks = [combined[:chunk_size]]

        ids = [f"{policy_id}_c{i}" for i in range(len(chunks))]
        metadatas = [{"policy_id": policy_id, "chunk_index": i} for i in range(len(chunks))]
        emb = model.encode(chunks, normalize_embeddings=True)
        coll.upsert(
            ids=ids,
            embeddings=emb.tolist(),
            documents=chunks,
            metadatas=metadatas,
        )
        coll.delete(ids=[policy_id])  # 删除旧的全文档记录（如有）
        logger.info(f"向量入库成功: {policy_id}, {len(chunks)} 个切片")
        return True
    except Exception as e:
        logger.error(f"向量入库失败 {policy_id}: {e}")
        return False


def remove_from_index(policy_id: str) -> bool:
    """从向量库删除（删除该 policy 的所有切片）"""
    try:
        coll = _get_collection()
        try:
            existing = coll.get(where={"policy_id": policy_id}, include=[])
            if existing and existing.get("ids"):
                coll.delete(ids=existing["ids"])
        except Exception:
            pass
        try:
            coll.delete(ids=[policy_id])
        except Exception:
            pass
        return True
    except Exception as e:
        logger.warning(f"向量删除失败 {policy_id}: {e}")
        return False


def search(query: str, top_k: int = 20) -> List[Tuple[str, float, str]]:
    """向量检索，返回 [(policy_id, score, snippet), ...]，snippet 为匹配到的切片原文"""
    if not (query or "").strip():
        return []
    try:
        coll = _get_collection()
        model = _get_model()
        q_emb = model.encode([query.strip()], normalize_embeddings=True)
        n_results = min(top_k * 3, 150)  # 多取一些，便于按 policy 去重后保留 top_k
        results = coll.query(
            query_embeddings=q_emb.tolist(),
            n_results=n_results,
            include=["metadatas", "distances", "documents"],
        )
        ids = results.get("ids", [[]])[0] or []
        dists = results.get("distances", [[]])[0] or []
        docs = results.get("documents", [[]])[0] or []
        metadatas = results.get("metadatas", [[]])[0] or []
        seen = {}
        for i, chunk_id in enumerate(ids):
            meta = metadatas[i] if i < len(metadatas) else {}
            policy_id = meta.get("policy_id") or (chunk_id.split("_c")[0] if "_c" in str(chunk_id) else chunk_id)
            d = float(dists[i] if i < len(dists) else 1.0)
            # 余弦距离：distance = 1 - cos_sim，故 cos_sim = 1 - distance，直接作为相关性
            score = max(0.0, min(1.0, 1.0 - d))
            doc_text = docs[i] if i < len(docs) else ""
            snippet = (doc_text or "").strip()
            if policy_id not in seen or score > seen[policy_id][1]:
                seen[policy_id] = (snippet, score)
        out = [(pid, sc, sn) for pid, (sn, sc) in seen.items()]
        out.sort(key=lambda x: -x[1])
        return out[:top_k]
    except Exception as e:
        logger.error(f"向量检索失败: {e}")
        raise
