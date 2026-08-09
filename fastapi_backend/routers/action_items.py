# -*- coding: utf-8 -*-
"""行动项督办：会议纪要、AI 提取、执行、审批、提醒与统计。"""
from __future__ import annotations

import csv
import calendar
import io
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, List, Literal, Optional

from fastapi import APIRouter, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, ConfigDict, Field, ValidationError, field_validator

from database import db
from routers.ai_assistant import (
    DEPARTMENT_ALIASES, DEPARTMENT_KNOWLEDGE, SHIFT_BUSINESS_DEPARTMENTS,
)

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/action-items", tags=["行动项督办"])

BASE_DIR = Path(__file__).resolve().parents[1]
UPLOAD_DIR = BASE_DIR / "data" / "action_items"
MINUTES_DIR = UPLOAD_DIR / "minutes"
EVIDENCE_DIR = UPLOAD_DIR / "evidence"
ALLOWED_MINUTES_EXT = {".doc", ".docx", ".pdf", ".txt"}
ALLOWED_EVIDENCE_EXT = {
    ".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".csv",
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".zip", ".rar", ".7z",
}
MAX_FILE_SIZE = 30 * 1024 * 1024

STATUS_DRAFT = "草稿"
STATUS_PENDING_PUBLISH = "待发布"
STATUS_PENDING_RECEIVE = "待接收"
STATUS_IN_PROGRESS = "进行中"
STATUS_PENDING_COMPLETION = "待完工审批"
STATUS_RETURNED = "退回整改"
STATUS_COMPLETED = "已完成"
STATUS_CANCELLED = "已取消"
ACTIVE_STATUSES = (
    STATUS_PENDING_RECEIVE, STATUS_IN_PROGRESS, STATUS_PENDING_COMPLETION, STATUS_RETURNED
)


DDL_STATEMENTS = [
    """
    CREATE TABLE IF NOT EXISTS meeting_minutes (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      meeting_name VARCHAR(255) NOT NULL,
      meeting_type VARCHAR(100) NOT NULL DEFAULT '',
      meeting_date DATE NULL,
      minutes_number VARCHAR(100) NOT NULL DEFAULT '',
      meeting_subject VARCHAR(500) NOT NULL DEFAULT '',
      publish_date DATE NULL,
      minutes_text LONGTEXT NULL,
      original_attachment VARCHAR(1000) NOT NULL DEFAULT '',
      original_file_name VARCHAR(255) NOT NULL DEFAULT '',
      creator VARCHAR(100) NOT NULL,
      created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      status VARCHAR(50) NOT NULL DEFAULT '草稿',
      ai_raw_json LONGTEXT NULL,
      ai_batch_id VARCHAR(64) NOT NULL DEFAULT '',
      ai_extracted_at DATETIME NULL,
      KEY idx_minutes_date (meeting_date),
      KEY idx_minutes_number (minutes_number),
      KEY idx_minutes_status (status)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_item (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_number VARCHAR(80) NOT NULL DEFAULT '',
      title VARCHAR(500) NOT NULL,
      content LONGTEXT NOT NULL,
      source_type VARCHAR(50) NOT NULL DEFAULT '会议纪要',
      source_meeting_id BIGINT NULL,
      minutes_number VARCHAR(100) NOT NULL DEFAULT '',
      source_quote LONGTEXT NULL,
      responsible_department_id VARCHAR(100) NULL,
      responsible_person_id VARCHAR(100) NULL,
      responsible_department_ids TEXT NULL,
      responsible_person_ids TEXT NULL,
      collaborating_departments TEXT NULL,
      collaborating_people TEXT NULL,
      supervisor_id VARCHAR(100) NULL,
      required_completion_date DATE NULL,
      actual_completion_date DATE NULL,
      priority VARCHAR(20) NOT NULL DEFAULT '中',
      current_progress INT NOT NULL DEFAULT 0,
      current_status VARCHAR(50) NOT NULL DEFAULT '草稿',
      risk_status VARCHAR(100) NOT NULL DEFAULT '',
      published_at DATETIME NULL,
      created_by VARCHAR(100) NOT NULL,
      created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
      last_progress_at DATETIME NULL,
      uncertain_fields TEXT NULL,
      ai_batch_id VARCHAR(64) NOT NULL DEFAULT '',
      KEY idx_action_meeting (source_meeting_id),
      KEY idx_action_status (current_status),
      KEY idx_action_dept (responsible_department_id),
      KEY idx_action_person (responsible_person_id),
      KEY idx_action_supervisor (supervisor_id),
      KEY idx_action_deadline (required_completion_date)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_progress (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NOT NULL,
      progress_percent INT NOT NULL,
      current_progress LONGTEXT NOT NULL,
      completed_work LONGTEXT NULL,
      existing_problems LONGTEXT NULL,
      next_plan LONGTEXT NULL,
      expected_completion_date DATE NULL,
      delay_risk TINYINT NOT NULL DEFAULT 0,
      attachments TEXT NULL,
      reporter VARCHAR(100) NOT NULL,
      reported_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      KEY idx_progress_action (action_item_id, reported_at)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS completion_application (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NOT NULL,
      completion_description LONGTEXT NOT NULL,
      actual_completion_date DATE NULL,
      completion_results LONGTEXT NULL,
      evidence_materials TEXT NULL,
      remaining_issues LONGTEXT NULL,
      applicant VARCHAR(100) NOT NULL,
      applied_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      approval_status VARCHAR(50) NOT NULL DEFAULT '待审批',
      KEY idx_completion_action (action_item_id, applied_at),
      KEY idx_completion_status (approval_status)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_department_execution (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NOT NULL,
      department VARCHAR(100) NOT NULL,
      responsible_person VARCHAR(100) NULL,
      execution_status VARCHAR(50) NOT NULL DEFAULT '待接收',
      progress_percent INT NOT NULL DEFAULT 0,
      received_by VARCHAR(100) NULL,
      received_at DATETIME NULL,
      completed_at DATETIME NULL,
      updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
      UNIQUE KEY uk_action_department (action_item_id,department),
      KEY idx_department_execution_status (execution_status,department),
      KEY idx_department_execution_person (responsible_person)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_approval (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      business_type VARCHAR(50) NOT NULL,
      business_id BIGINT NOT NULL,
      action_item_id BIGINT NOT NULL,
      approver VARCHAR(100) NOT NULL,
      approval_result VARCHAR(50) NOT NULL,
      approval_opinion LONGTEXT NULL,
      approved_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      KEY idx_approval_action (action_item_id, approved_at),
      KEY idx_approval_business (business_type, business_id)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_change (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NOT NULL,
      change_type VARCHAR(50) NOT NULL,
      before_content LONGTEXT NULL,
      after_content LONGTEXT NULL,
      change_reason LONGTEXT NOT NULL,
      applicant VARCHAR(100) NOT NULL,
      applied_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      approval_status VARCHAR(50) NOT NULL DEFAULT '待审批',
      approver VARCHAR(100) NULL,
      approval_opinion LONGTEXT NULL,
      approved_at DATETIME NULL,
      KEY idx_change_action (action_item_id, applied_at),
      KEY idx_change_status (approval_status)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_reminder (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NOT NULL,
      reminder_type VARCHAR(50) NOT NULL,
      reminder_target VARCHAR(100) NOT NULL,
      reminder_channel VARCHAR(50) NOT NULL DEFAULT '站内',
      reminder_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      send_status VARCHAR(50) NOT NULL DEFAULT '待发送',
      reminder_note LONGTEXT NULL,
      reminder_key VARCHAR(255) NOT NULL,
      read_at DATETIME NULL,
      UNIQUE KEY uk_action_reminder_key (reminder_key),
      KEY idx_reminder_target (reminder_target, read_at, reminder_time)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_attachment (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NULL,
      business_type VARCHAR(50) NOT NULL,
      business_id BIGINT NOT NULL,
      original_name VARCHAR(255) NOT NULL,
      stored_path VARCHAR(1000) NOT NULL,
      content_type VARCHAR(255) NOT NULL DEFAULT '',
      file_size BIGINT NOT NULL DEFAULT 0,
      uploader VARCHAR(100) NOT NULL,
      uploaded_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      KEY idx_attachment_business (business_type, business_id),
      KEY idx_attachment_action (action_item_id)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_event_log (
      id BIGINT AUTO_INCREMENT PRIMARY KEY,
      action_item_id BIGINT NULL,
      meeting_minutes_id BIGINT NULL,
      event_type VARCHAR(50) NOT NULL,
      operator VARCHAR(100) NOT NULL,
      event_content LONGTEXT NOT NULL,
      event_data LONGTEXT NULL,
      created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,
      KEY idx_event_action (action_item_id, created_at),
      KEY idx_event_minutes (meeting_minutes_id, created_at)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    CREATE TABLE IF NOT EXISTS action_supervision_config (
      leader_name VARCHAR(100) PRIMARY KEY,
      departments TEXT NOT NULL,
      work_division LONGTEXT NULL,
      enabled TINYINT NOT NULL DEFAULT 1,
      updated_by VARCHAR(100) NOT NULL DEFAULT '',
      updated_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
      KEY idx_supervision_enabled (enabled)
    ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    """,
    """
    INSERT IGNORE INTO action_supervision_config
      (leader_name,departments,work_division,enabled,updated_by)
    SELECT name,'["焊接工艺室","智能制造技术室"]','负责技术准备和降本增效工作',1,'系统初始化'
    FROM yggl
    WHERE name='房金彪' AND COALESCE(zaizhi,0)=0
    LIMIT 1
    """,
]


def ensure_tables() -> None:
    for statement in DDL_STATEMENTS:
        if db.execute_update(statement) < 0:
            raise RuntimeError("行动项督办数据表初始化失败")
    for column, definition in (
        ("responsible_department_ids", "TEXT NULL"),
        ("responsible_person_ids", "TEXT NULL"),
    ):
        existing = db.execute_query(f"SHOW COLUMNS FROM action_item LIKE '{column}'")
        if not existing and db.execute_update(
            f"ALTER TABLE action_item ADD COLUMN {column} {definition} "
            "AFTER responsible_person_id"
        ) < 0:
            raise RuntimeError(f"行动项督办字段 {column} 初始化失败")
    for table, column in (
        ("action_progress", "responsible_department"),
        ("completion_application", "responsible_department"),
    ):
        existing = db.execute_query(f"SHOW COLUMNS FROM {table} LIKE '{column}'")
        if not existing and db.execute_update(
            f"ALTER TABLE {table} ADD COLUMN {column} VARCHAR(100) NULL AFTER action_item_id"
        ) < 0:
            raise RuntimeError(f"行动项督办字段 {table}.{column} 初始化失败")
    for directory in (MINUTES_DIR, EVIDENCE_DIR):
        directory.mkdir(parents=True, exist_ok=True)


def _json(value: Any, default: Any = None) -> Any:
    if value in (None, ""):
        return [] if default is None else default
    if isinstance(value, (dict, list)):
        return value
    try:
        return json.loads(value)
    except Exception:
        return [] if default is None else default


def _responsible_departments(action: dict) -> list[str]:
    values = _json(action.get("responsible_department_ids"))
    scalar = str(action.get("responsible_department_id") or "").strip()
    return list(dict.fromkeys(
        str(value or "").strip() for value in (values or [scalar])
        if str(value or "").strip()
    ))


def _responsible_people(action: dict) -> list[str]:
    values = _json(action.get("responsible_person_ids"))
    scalar = str(action.get("responsible_person_id") or "").strip()
    return list(dict.fromkeys(
        str(value or "").strip() for value in (values or [scalar])
        if str(value or "").strip()
    ))


def _ensure_department_executions(
    action: dict,
    people: Optional[dict[str, str]] = None,
) -> list[dict]:
    """按责任科室补齐执行单元，兼容升级前已存在的行动项。"""
    departments = _responsible_departments(action)
    if not departments:
        return []
    if people is None:
        _departments, people, _jobs, _supervisors, _defaults = _action_directory_maps()
    responsible_people = _responsible_people(action)
    for department in departments:
        person = next(
            (name for name in responsible_people if people.get(name) == department),
            None,
        )
        old_status = action.get("current_status")
        execution_status = (
            STATUS_PENDING_RECEIVE if old_status == STATUS_PENDING_RECEIVE
            else STATUS_COMPLETED if old_status == STATUS_COMPLETED
            else STATUS_PENDING_COMPLETION if old_status == STATUS_PENDING_COMPLETION
            else STATUS_RETURNED if old_status == STATUS_RETURNED
            else STATUS_IN_PROGRESS
        )
        db.execute_update(
            "INSERT IGNORE INTO action_department_execution "
            "(action_item_id,department,responsible_person,execution_status,progress_percent) "
            "VALUES (%s,%s,%s,%s,%s)",
            (
                action["id"], department, person, execution_status,
                int(action.get("current_progress") or 0),
            ),
        )
    return db.execute_query(
        "SELECT * FROM action_department_execution "
        "WHERE action_item_id=%s ORDER BY id",
        (action["id"],),
    )


def _refresh_action_rollup(action_id: int) -> dict:
    """按科室执行单元等权汇总行动项总进度和主状态。"""
    rows = db.execute_query(
        "SELECT * FROM action_department_execution WHERE action_item_id=%s ORDER BY id",
        (action_id,),
    )
    if not rows:
        return {}
    progress = round(sum(int(row.get("progress_percent") or 0) for row in rows) / len(rows))
    statuses = [row.get("execution_status") or STATUS_PENDING_RECEIVE for row in rows]
    if all(status == STATUS_COMPLETED for status in statuses):
        status = STATUS_COMPLETED
    elif any(status == STATUS_PENDING_RECEIVE for status in statuses):
        status = STATUS_PENDING_RECEIVE
    elif any(status == STATUS_RETURNED for status in statuses):
        status = STATUS_RETURNED
    elif all(status in (STATUS_PENDING_COMPLETION, STATUS_COMPLETED) for status in statuses):
        status = STATUS_PENDING_COMPLETION
    else:
        status = STATUS_IN_PROGRESS
    db.execute_update(
        "UPDATE action_item SET current_progress=%s,current_status=%s,"
        "actual_completion_date=CASE WHEN %s=%s THEN CURDATE() ELSE actual_completion_date END "
        "WHERE id=%s",
        (progress, status, status, STATUS_COMPLETED, action_id),
    )
    return {"current_progress": progress, "current_status": status, "departments": rows}


def _is_responsible_department(action: dict, department: str) -> bool:
    return bool(department and department in _responsible_departments(action))


def _is_responsible_person(action: dict, person: str) -> bool:
    return bool(person and person in _responsible_people(action))


def _load_supervision_configs(enabled_only: bool = True) -> list[dict]:
    where = "WHERE c.enabled=1" if enabled_only else ""
    rows = db.execute_query(
        "SELECT c.leader_name,c.departments,c.work_division,c.enabled,c.updated_by,"
        "c.updated_at,y.jb FROM action_supervision_config c "
        "LEFT JOIN yggl y ON y.name=c.leader_name AND COALESCE(y.zaizhi,0)=0 "
        f"{where} ORDER BY c.leader_name"
    ) or []
    result = []
    for row in rows:
        result.append({
            "leader_name": (row.get("leader_name") or "").strip(),
            "departments": list(dict.fromkeys(
                _normalize_department_name(item)
                for item in _json(row.get("departments"))
                if str(item or "").strip()
            )),
            "work_division": (row.get("work_division") or "").strip(),
            "enabled": bool(row.get("enabled")),
            "job": (row.get("jb") or "").strip(),
            "updated_by": (row.get("updated_by") or "").strip(),
            "updated_at": _dt_text(row.get("updated_at")),
        })
    return result


def _date_text(value: Any) -> str:
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d")
    return str(value)[:10]


def _dt_text(value: Any) -> str:
    if not value:
        return ""
    if hasattr(value, "strftime"):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    return str(value)[:19]


def _minutes_reference_date(minutes: dict) -> Any:
    """截止日默认值优先依据会议日期，其次发布日期和纪要创建日期。"""
    return (
        minutes.get("meeting_date")
        or minutes.get("publish_date")
        or minutes.get("created_at")
    )


def _user(name: str) -> dict:
    rows = db.execute_query(
        "SELECT name, gh, lsys, jb, enterprise_email FROM yggl "
        "WHERE name=%s AND COALESCE(zaizhi,0)=0 LIMIT 1",
        ((name or "").strip(),),
    )
    if not rows:
        raise HTTPException(status_code=401, detail="当前用户不存在或已离职")
    return rows[0]


def _role_context(name: str) -> dict:
    user = _user(name)
    n = (user.get("name") or "").strip()
    dept = (user.get("lsys") or "").strip()
    jb = (user.get("jb") or "").strip()
    try:
        from routers.approvers import _jb_match, is_admin1_user, is_admin2_user
        admin = is_admin1_user(n)
        admin2 = is_admin2_user(n)
        department_leader = _jb_match(jb, "部长") or _jb_match(jb, "副部长")
        dept_manager = (
            _jb_match(jb, "主任") or _jb_match(jb, "副主任")
            or _jb_match(jb, "组长") or _jb_match(jb, "经理")
            or _jb_match(jb, "副经理")
        )
        company_manager = _jb_match(jb, "经理") or _jb_match(jb, "副经理")
    except Exception:
        admin = admin2 = department_leader = dept_manager = company_manager = False
    minutes_admin = admin or dept.startswith("综合") or dept == "综合技术室"
    minutes_uploader = minutes_admin or company_manager
    action_creator = minutes_admin or department_leader or company_manager
    return {
        "name": n, "dept": dept, "jb": jb, "admin": admin, "admin2": admin2,
        "department_leader": department_leader,
        "dept_manager": dept_manager,
        "dept_director": _job_director_priority(jb) is not None,
        "minutes_admin": minutes_admin, "minutes_uploader": minutes_uploader,
        "action_creator": action_creator,
        "view_all": admin or department_leader or minutes_admin,
    }


def _permissions(ctx: dict) -> dict:
    return {
        "systemConfig": bool(ctx["admin"]),
        "minutesManage": bool(ctx["minutes_admin"]),
        "minutesUpload": bool(ctx["minutes_uploader"]),
        "actionCreate": bool(ctx["action_creator"]),
        "extract": bool(ctx["minutes_admin"]),
        "publish": bool(ctx["minutes_admin"]),
        "supervise": bool(ctx["minutes_admin"] or ctx["department_leader"]),
        "export": bool(ctx["minutes_admin"] or ctx["department_leader"]),
        "departmentAssign": bool(ctx["dept_manager"] or ctx["view_all"]),
        "progressUpdate": True,
        "completionApply": True,
        "completionApprove": True,
        "changeApply": True,
        "changeApprove": True,
        "dashboardAll": bool(ctx["view_all"]),
    }


def _visible_action_sql(ctx: dict, alias: str = "a") -> tuple[str, list]:
    if ctx["view_all"]:
        return "1=1", []
    terms = [
        f"{alias}.responsible_person_id=%s",
        f"JSON_CONTAINS(COALESCE({alias}.responsible_person_ids, '[]'), JSON_QUOTE(%s))",
        f"{alias}.supervisor_id=%s",
        f"JSON_CONTAINS(COALESCE({alias}.collaborating_people, '[]'), JSON_QUOTE(%s))",
    ]
    params: list[Any] = [ctx["name"], ctx["name"], ctx["name"], ctx["name"]]
    if ctx["action_creator"]:
        terms.append(f"{alias}.created_by=%s")
        params.append(ctx["name"])
    if ctx["dept_manager"] and ctx["dept"]:
        terms.extend([
            f"{alias}.responsible_department_id=%s",
            f"JSON_CONTAINS(COALESCE({alias}.responsible_department_ids, '[]'), JSON_QUOTE(%s))",
            f"JSON_CONTAINS(COALESCE({alias}.collaborating_departments, '[]'), JSON_QUOTE(%s))",
        ])
        params.extend([ctx["dept"], ctx["dept"], ctx["dept"]])
    return "(" + " OR ".join(terms) + ")", params


def _action_row(action_id: int) -> dict:
    rows = db.execute_query("SELECT * FROM action_item WHERE id=%s LIMIT 1", (action_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="行动项不存在")
    return rows[0]


def _assert_action_visible(ctx: dict, action: dict) -> None:
    if (
        action.get("current_status") in (STATUS_DRAFT, STATUS_PENDING_PUBLISH)
        and not ctx["minutes_admin"]
        and action.get("created_by") != ctx["name"]
    ):
        raise HTTPException(status_code=403, detail="仅创建人、综合室或系统管理员可查看未发布行动项")
    if ctx["view_all"]:
        return
    if _is_responsible_person(action, ctx["name"]):
        return
    if (action.get("supervisor_id") or "") == ctx["name"]:
        return
    if ctx["name"] in _json(action.get("collaborating_people")):
        return
    if ctx["dept_manager"] and ctx["dept"]:
        if _is_responsible_department(action, ctx["dept"]):
            return
        if ctx["dept"] in _json(action.get("collaborating_departments")):
            return
    raise HTTPException(status_code=403, detail="无权查看该行动项")


def _event(
    operator: str, event_type: str, content: str, action_id: Optional[int] = None,
    meeting_id: Optional[int] = None, data: Any = None,
) -> None:
    db.execute_update(
        "INSERT INTO action_event_log "
        "(action_item_id, meeting_minutes_id, event_type, operator, event_content, event_data) "
        "VALUES (%s,%s,%s,%s,%s,%s)",
        (action_id, meeting_id, event_type, operator, content,
         json.dumps(data, ensure_ascii=False, default=str) if data is not None else None),
    )


_REMINDER_EVENT_CONTENT_RE = re.compile(
    r"^(?P<reminder_type>[^：:]+)[：:](?P<target>[^；;]+)[；;](?P<note>.*)$"
)


def _group_timeline_events(rows: list[dict]) -> list[dict]:
    """合并同一秒、同一类型和内容的多人提醒，仅影响详情页展示。"""
    result: list[dict] = []
    reminder_groups: dict[tuple[str, str, str, str, str], dict] = {}
    for source in rows:
        row = dict(source)
        row["event_data"] = _json(row.get("event_data"), {})
        row["created_at"] = _dt_text(row.get("created_at"))
        content = str(row.get("event_content") or "").strip()
        match = (
            _REMINDER_EVENT_CONTENT_RE.match(content)
            if row.get("event_type") == "提醒"
            else None
        )
        if not match:
            result.append(row)
            continue
        reminder_type = match.group("reminder_type").strip()
        target = match.group("target").strip()
        note = match.group("note").strip()
        key = (
            row["created_at"],
            str(row.get("operator") or ""),
            str(row.get("event_type") or ""),
            reminder_type,
            note,
        )
        grouped = reminder_groups.get(key)
        if grouped is None:
            grouped = row
            grouped["_reminder_type"] = reminder_type
            grouped["_reminder_note"] = note
            grouped["_reminder_targets"] = []
            grouped["_grouped_event_ids"] = []
            reminder_groups[key] = grouped
            result.append(grouped)
        if target and target not in grouped["_reminder_targets"]:
            grouped["_reminder_targets"].append(target)
        if row.get("id") is not None:
            grouped["_grouped_event_ids"].append(row["id"])
    for row in result:
        targets = row.pop("_reminder_targets", None)
        if targets is None:
            continue
        reminder_type = row.pop("_reminder_type")
        note = row.pop("_reminder_note")
        event_ids = row.pop("_grouped_event_ids")
        row["event_content"] = (
            f"{reminder_type}：{'、'.join(targets)}"
            f"{f'；{note}' if note else ''}"
        )
        row["group_count"] = len(event_ids)
        row["event_data"] = {
            **(row.get("event_data") or {}),
            "groupedEventIds": event_ids,
            "groupedTargets": targets,
        }
    return result


async def _save_upload(
    file: UploadFile, directory: Path, business_type: str, business_id: int,
    uploader: str, action_id: Optional[int] = None, allowed: Optional[set[str]] = None,
) -> dict:
    name = Path(file.filename or "").name
    ext = Path(name).suffix.lower()
    if not name or (allowed is not None and ext not in allowed):
        raise HTTPException(status_code=400, detail=f"不支持的附件格式：{ext or '未知'}")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail=f"附件 {name} 为空")
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"附件 {name} 超过 30MB")
    directory.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}{ext}"
    path = directory / stored
    path.write_bytes(content)
    attachment_id = db.execute_insert(
        "INSERT INTO action_attachment "
        "(action_item_id,business_type,business_id,original_name,stored_path,content_type,file_size,uploader) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
        (action_id, business_type, business_id, name, str(path.relative_to(BASE_DIR)).replace("\\", "/"),
         file.content_type or "", len(content), uploader),
    )
    return {"id": attachment_id, "name": name, "size": len(content)}


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if ext == ".doc":
        from config import settings
        command = (
            str(getattr(settings, "LIBREOFFICE_CMD", "") or "").strip()
            or shutil.which("libreoffice") or shutil.which("soffice")
        )
        if not command:
            raise ValueError("解析 .doc 需要服务器安装 LibreOffice，也可另存为 .docx 后上传")
        with tempfile.TemporaryDirectory(prefix="action_doc_") as output_dir:
            profile_dir = Path(output_dir) / "profile"
            profile_dir.mkdir()
            result = subprocess.run(
                [
                    command, "--headless", "--norestore",
                    f"-env:UserInstallation=file:///{str(profile_dir).replace(chr(92), '/')}",
                    "--convert-to", "pdf", "--outdir", output_dir, str(path),
                ],
                capture_output=True, timeout=120,
            )
            pdf_path = Path(output_dir) / f"{path.stem}.pdf"
            if result.returncode != 0 or not pdf_path.is_file():
                detail = result.stderr.decode("utf-8", errors="ignore").strip()
                raise ValueError(f"LibreOffice 转换 .doc 失败：{detail or '未生成 PDF'}")
            import fitz
            doc = fitz.open(str(pdf_path))
            return "\n".join(page.get_text("text") for page in doc).strip()
    if ext == ".pdf":
        import fitz
        doc = fitz.open(str(path))
        return "\n".join(page.get_text("text") for page in doc).strip()
    return path.read_text(encoding="utf-8", errors="ignore").strip()


class ExtractedAction(BaseModel):
    model_config = ConfigDict(extra="forbid")

    title: str = Field(min_length=1)
    content: str = Field(min_length=1)
    source_quote: str = Field(min_length=1)
    responsible_department: Optional[str] = None
    responsible_person: Optional[str] = None
    collaborating_departments: List[str] = Field(default_factory=list)
    supervisor: Optional[str] = None
    deadline: Optional[str] = Field(
        default=None,
        description="仅允许 YYYY-MM-DD；无法从原文确定到具体日期时必须为 null",
    )
    priority: Literal["高", "中", "低"] = "中"
    uncertain_fields: List[Literal[
        "responsible_department", "responsible_person",
        "collaborating_departments", "supervisor", "deadline",
    ]] = Field(default_factory=list)

    @field_validator("title", "content", "source_quote")
    @classmethod
    def non_blank_text(cls, value: str) -> str:
        if not value.strip():
            raise ValueError("标题、内容和原文依据不得为空")
        return value.strip()

    @field_validator("deadline")
    @classmethod
    def valid_deadline(cls, value: Optional[str]) -> Optional[str]:
        if not value:
            return None
        try:
            date.fromisoformat(value)
        except ValueError as exc:
            raise ValueError("deadline 必须为 YYYY-MM-DD") from exc
        return value


class ExtractedPayload(BaseModel):
    model_config = ConfigDict(extra="forbid")
    items: List[ExtractedAction] = Field(min_length=1)


AI_JSON_SCHEMA = ExtractedPayload.model_json_schema()


class ExtractedMinutesMetadata(BaseModel):
    """会议纪要基础信息；召集领导字段已废弃，不再提取。"""

    model_config = ConfigDict(extra="forbid")

    meeting_name: Optional[str] = None
    meeting_type: Optional[str] = None
    meeting_date: Optional[str] = None
    minutes_number: Optional[str] = None
    meeting_subject: Optional[str] = None
    publish_date: Optional[str] = None

    @field_validator(
        "meeting_name", "meeting_type", "minutes_number", "meeting_subject",
        mode="before",
    )
    @classmethod
    def blank_metadata_text(cls, value: Any) -> Optional[str]:
        text = str(value or "").strip()
        return text or None

    @field_validator("meeting_date", "publish_date")
    @classmethod
    def valid_metadata_date(cls, value: Optional[str]) -> Optional[str]:
        if not value:
            return None
        try:
            return date.fromisoformat(value).isoformat()
        except ValueError as exc:
            raise ValueError("日期必须为 YYYY-MM-DD") from exc


MINUTES_METADATA_JSON_SCHEMA = ExtractedMinutesMetadata.model_json_schema()
MINUTES_METADATA_PROMPT = """你是会议纪要基础信息识别助手。
请仅根据文件名和会议纪要正文提取以下字段：
meeting_name、meeting_type、meeting_date、minutes_number、meeting_subject、publish_date。

规则：
1. meeting_type 只允许：综合管理例会、质量例会、产品提升专题会、其他。
2. meeting_date 和 publish_date 只能返回 YYYY-MM-DD；原文未明确时返回 null。
3. minutes_number、meeting_subject 未明确时返回 null，不得虚构。
4. meeting_name 优先采用正文中的正式会议名称；无法确定时返回 null。
5. 召集领导字段已废弃，禁止输出召集领导或任何额外字段。
6. 只返回符合 JSON Schema 的 JSON，不得输出 Markdown 或解释。"""

AI_SYSTEM_PROMPT = f"""你是会议行动项提取助手，服务于智能制造工艺部的行动项督办工作。

{DEPARTMENT_KNOWLEDGE}

科室简称必须映射为上述数据库正式名称后再写入 responsible_department 和 collaborating_departments；
例如“焊艺室”输出“焊接工艺室”，“智能室”输出“智能制造技术室”，“编程室”输出“数控编程室”。
原文出现“各室、各科室、所有科室、全部科室、各专业室”等全覆盖表述时，
responsible_department 原样返回“各科室”作为多科室标记；后端只保留一条行动项，
同时关联 9 个正式业务科室及各室主任、主任责和副主任，禁止拆成多条或只选择其中某一个科室。
原文未明确责任科室、科室名称无法识别，或返回的名称不能映射到有效科室时，responsible_department 返回 null
并标记 uncertain_fields；后端会将该事项作为一条多科室行动项关联全部 9 个正式业务科室，不得因此放弃行动项或中断提取。
只提取具有执行、落实、完成、解决、制定、推进、协调、检查等明确行动含义的事项。
一个段落包含多个独立任务时必须拆分；不同责任主体承担不同任务时必须拆分。
每条必须保留纪要原文 source_quote。原文没有明确责任科室、责任人、协同科室或截止日期时返回 null/空数组，
并把字段名加入 uncertain_fields。主管领导可依据后附的系统管理员配置判断；配置也无法唯一确定时返回 null，
并把 supervisor 加入 uncertain_fields。deadline 只能返回 YYYY-MM-DD，严禁返回“年底”“月底”“尽快”等自然语言；
原文只明确到“某年年底/某月月底”时可换算为该年/该月最后一天，其他无法确定到具体日期的表述必须返回 null。
原文未指定完成时间时，deadline 必须取会议日期后第 7 天；后端还会再次强制补全该默认值。
原文明确责任人但未明确责任科室时，不得猜测科室；原文只明确责任科室时，不得猜测具体责任人。
系统会根据在职人员目录自动校正责任人的所属科室，并在仅有责任科室时补充该科室负责人。
责任人明确为经理或副经理时，supervisor 应与 responsible_person 返回同一个人；
但“各室/全部科室”共同负责的行动项优先选择在职经理作为 supervisor。
每条行动项最终必须具备责任科室、责任人、主管领导、完成时间四项信息；
不得根据常识、组织习惯或上下文虚构信息。只返回符合 JSON Schema 的 JSON，不得输出 Markdown 或解释。"""


def _supervision_prompt() -> str:
    configs = _load_supervision_configs()
    if not configs:
        return (
            "【主管领导配置】系统管理员尚未配置科室分管领导。"
            "全部科室共同负责时 supervisor 选择在职经理；"
            "其他情况除责任人本身为经理/副经理外，supervisor 返回 null 并标记待确认。"
        )
    lines = [
        "【主管领导配置】以下内容来自系统管理员当前生效的配置，允许用于判断 supervisor："
    ]
    for item in configs:
        departments = "、".join(item["departments"]) or "未指定科室"
        division = item["work_division"] or "未填写工作分工"
        job = f"（{item['job']}）" if item.get("job") else ""
        lines.append(
            f"- {item['leader_name']}{job}；分管科室：{departments}；工作分工：{division}"
        )
    lines.extend([
        "主管领导判断顺序：",
        "1. “各室/各科室/全部科室”共同负责时，supervisor 必须选择在职经理；经理主管全部科室；",
        "2. 仅部分科室负责时，结合责任科室和行动项内容，按上述分管科室及工作分工选择；",
        "3. 责任人是经理或副经理时，supervisor 与 responsible_person 相同；",
        "4. 多名领导均可能匹配或无法可靠判断时返回 null，并将 supervisor 加入 uncertain_fields；",
        "5. supervisor 只能返回上述已配置领导或在职经理/副经理姓名，不得虚构。",
    ])
    return "\n".join(lines)


def _parse_json_response(content: str) -> dict:
    text = (content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?", "", text, flags=re.I).strip()
        text = re.sub(r"```$", "", text).strip()
    try:
        return json.loads(text)
    except Exception:
        start, end = text.find("{"), text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start:end + 1])
        raise


_UNCERTAIN_FIELDS = {
    "responsible_department", "responsible_person",
    "collaborating_departments", "supervisor", "deadline",
}
_UNCERTAIN_ALIASES = {
    "责任科室": "responsible_department",
    "责任部门": "responsible_department",
    "责任人": "responsible_person",
    "协同科室": "collaborating_departments",
    "协同部门": "collaborating_departments",
    "主管领导": "supervisor",
    "完成时间": "deadline",
    "截止时间": "deadline",
    "完成日期": "deadline",
}


def _as_date(value: Any) -> Optional[date]:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if value:
        try:
            return date.fromisoformat(str(value)[:10])
        except ValueError:
            return None
    return None


def _normalize_ai_deadline(value: Any, meeting_date: Any = None) -> Optional[str]:
    """把可无歧义确定的中文期限换算成日期；其余返回 None 交人工确认。"""
    if value in (None, ""):
        return None
    text = str(value).strip()
    if not text:
        return None

    # 已经是 ISO 日期，或常见中文/斜杠日期。
    match = re.search(r"(?<!\d)(\d{4})[-/.年](\d{1,2})[-/.月](\d{1,2})(?:日|号)?", text)
    if match:
        try:
            return date(int(match.group(1)), int(match.group(2)), int(match.group(3))).isoformat()
        except ValueError:
            return None

    # “2026 年底/年末/年内”具有明确的期间终点。
    match = re.search(r"(?<!\d)(\d{4})\s*年\s*(?:底|末|内)", text)
    if match:
        return date(int(match.group(1)), 12, 31).isoformat()

    # “2026 年 6 月底/末”换算为该月最后一天。
    match = re.search(r"(?<!\d)(\d{4})\s*年\s*(\d{1,2})\s*月\s*(?:底|末)", text)
    if match:
        year, month = int(match.group(1)), int(match.group(2))
        if 1 <= month <= 12:
            return date(year, month, calendar.monthrange(year, month)[1]).isoformat()
        return None

    # 明确季度末。
    match = re.search(r"(?<!\d)(\d{4})\s*年?\s*(?:第?\s*)?([一二三四1-4])\s*季度", text)
    if match:
        quarter_map = {"一": 1, "二": 2, "三": 3, "四": 4}
        quarter = quarter_map.get(match.group(2), int(match.group(2)) if match.group(2).isdigit() else 0)
        month = quarter * 3
        return date(int(match.group(1)), month, calendar.monthrange(int(match.group(1)), month)[1]).isoformat()

    context = _as_date(meeting_date)
    if context:
        if re.fullmatch(r"(?:今年|本年)?(?:年)?(?:底|末)(?:前|之前)?", text):
            return date(context.year, 12, 31).isoformat()
        if re.fullmatch(r"(?:本月|这个月|当月)?(?:月)?(?:底|末)(?:前|之前)?", text):
            return date(
                context.year, context.month,
                calendar.monthrange(context.year, context.month)[1],
            ).isoformat()
        if re.fullmatch(r"(?:本周|这周)(?:内|末|前|之前)?", text):
            return (context + timedelta(days=6 - context.weekday())).isoformat()
        if re.fullmatch(r"(?:下周)(?:内|末|前|之前)?", text):
            return (context + timedelta(days=13 - context.weekday())).isoformat()
        match = re.fullmatch(r"(\d{1,3})\s*天(?:内|前|之内)", text)
        if match:
            return (context + timedelta(days=int(match.group(1)))).isoformat()
    return None


def _normalize_ai_payload(data: dict, meeting_date: Any = None) -> dict:
    """在严格 Schema 校验前修复模型常见的、可安全归一化的格式偏差。"""
    if not isinstance(data, dict) or not isinstance(data.get("items"), list):
        return data
    meeting_day = _as_date(meeting_date)
    default_deadline = (
        (meeting_day + timedelta(days=7)).isoformat() if meeting_day else None
    )
    for item in data["items"]:
        if not isinstance(item, dict):
            continue
        raw_deadline = item.get("deadline")
        raw_collaborators = item.get("collaborating_departments")
        normalized_deadline = _normalize_ai_deadline(raw_deadline, meeting_date)
        if raw_deadline in (None, "") and default_deadline:
            normalized_deadline = default_deadline
        item["deadline"] = normalized_deadline
        if raw_collaborators is None:
            item["collaborating_departments"] = []

        raw_uncertain = item.get("uncertain_fields")
        if not isinstance(raw_uncertain, list):
            raw_uncertain = []
        uncertain = []
        for field in raw_uncertain:
            normalized = _UNCERTAIN_ALIASES.get(str(field).strip(), str(field).strip())
            if normalized in _UNCERTAIN_FIELDS and normalized not in uncertain:
                uncertain.append(normalized)
        if raw_deadline not in (None, "") and normalized_deadline is None and "deadline" not in uncertain:
            uncertain.append("deadline")
        if normalized_deadline:
            uncertain = [field for field in uncertain if field != "deadline"]
        if raw_collaborators is None and "collaborating_departments" not in uncertain:
            uncertain.append("collaborating_departments")
        if item.get("collaborating_departments") == [] and "collaborating_departments" in [
            _UNCERTAIN_ALIASES.get(str(x).strip(), str(x).strip()) for x in raw_uncertain
        ] and "collaborating_departments" not in uncertain:
            uncertain.append("collaborating_departments")
        item["uncertain_fields"] = uncertain

        for key in ("responsible_department", "responsible_person", "supervisor"):
            if item.get(key) == "":
                item[key] = None
        priority = str(item.get("priority") or "中").strip()
        item["priority"] = {
            "高优先级": "高", "紧急": "高", "重要": "高",
            "中等": "中", "一般": "中", "普通": "中",
            "低优先级": "低",
        }.get(priority, priority)
    return data


def _resolve_action_llm() -> dict:
    try:
        scene_rows = db.execute_query(
            "SELECT model_id FROM llm_scene_config WHERE scene='action_items' LIMIT 1"
        )
        scene_model_id = int(scene_rows[0].get("model_id") or 0) if scene_rows else 0
        if scene_model_id == -1:
            from routers.ai_assistant import _resolve_llm
            online = _resolve_llm()
            if online.get("provider") != "deepseek":
                raise ValueError("行动项提取已指定联网 DeepSeek，但管理员尚未配置有效 API Key")
            return online
        rows = db.execute_query(
            "SELECT m.base_url,m.model,m.api_key,m.use_extra "
            "FROM llm_scene_config s JOIN llm_models m ON m.id=s.model_id "
            "WHERE s.scene='action_items' LIMIT 1"
        )
        if rows:
            r = rows[0]
            return {
                "provider": "scene", "base_url": (r.get("base_url") or "").strip(),
                "model": (r.get("model") or "").strip(),
                "api_key": (r.get("api_key") or "").strip() or "ollama",
                "use_extra": bool(r.get("use_extra")),
            }
    except ValueError:
        raise
    except Exception:
        pass
    from routers.ai_assistant import _resolve_llm
    return _resolve_llm()


def _normalize_minutes_type(value: Any, text: str = "") -> str:
    raw = str(value or "").strip()
    allowed = {"综合管理例会", "质量例会", "产品提升专题会", "其他"}
    if raw in allowed:
        return raw
    source = f"{raw}\n{text[:3000]}"
    if re.search(r"质量(?:工作)?(?:例会|会议)", source):
        return "质量例会"
    if re.search(r"产品提升(?:工作)?(?:专题会|专题会议|会议)", source):
        return "产品提升专题会"
    if re.search(r"综合管理(?:工作)?(?:例会|会议)", source):
        return "综合管理例会"
    return "其他"


def _fallback_minutes_metadata(minutes_text: str, file_name: str) -> dict:
    """AI 不可用时保留可上传性，只采用正文明确值和文件名，不虚构信息。"""
    text = (minutes_text or "").strip()
    stem = Path(file_name or "会议纪要").stem.strip() or "会议纪要"
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if line.strip()]
    meeting_name = ""
    for line in lines[:30]:
        if len(line) <= 180 and any(word in line for word in ("会议", "例会", "纪要")):
            meeting_name = line
            break
    meeting_name = (meeting_name or stem)[:255]

    meeting_date = _normalize_ai_deadline(text[:5000])
    publish_match = re.search(
        r"(?:发布日期|印发日期|发布于|印发于)\s*[：:]?\s*"
        r"((?:19|20)\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}(?:日|号)?)",
        text[:10000],
    )
    publish_date = _normalize_ai_deadline(publish_match.group(1)) if publish_match else None
    number_match = re.search(
        r"(?:纪要编号|编号|文号)\s*[：:]?\s*([^\s，,；;。]{2,100})",
        text[:10000],
    )
    subject_match = re.search(
        r"(?:会议主题|主题|议题)\s*[：:]\s*([^\n\r]{2,500})",
        text[:10000],
    )
    return {
        "meeting_name": meeting_name,
        "meeting_type": _normalize_minutes_type("", f"{stem}\n{text}"),
        "meeting_date": meeting_date,
        "minutes_number": (number_match.group(1).strip()[:100] if number_match else None),
        "meeting_subject": (subject_match.group(1).strip()[:500] if subject_match else None),
        "publish_date": publish_date,
    }


def _extract_minutes_metadata(minutes_text: str, file_name: str) -> tuple[dict, str]:
    fallback = _fallback_minutes_metadata(minutes_text, file_name)
    try:
        from openai import OpenAI

        cfg = _resolve_action_llm()
        client = OpenAI(
            base_url=cfg["base_url"], api_key=cfg.get("api_key") or "ollama",
            timeout=90.0, max_retries=1,
        )
        messages = [
            {"role": "system", "content": MINUTES_METADATA_PROMPT},
            {
                "role": "user",
                "content": (
                    "JSON Schema:\n"
                    + json.dumps(MINUTES_METADATA_JSON_SCHEMA, ensure_ascii=False)
                    + f"\n\n文件名：{file_name}"
                    + "\n\n会议纪要正文：\n"
                    + minutes_text[:60000]
                ),
            },
        ]
        kwargs: dict[str, Any] = {
            "model": cfg["model"], "messages": messages, "temperature": 0,
        }
        if cfg.get("use_extra"):
            kwargs["extra_body"] = {"chat_template_kwargs": {"enable_thinking": False}}
        try:
            completion = client.chat.completions.create(
                **kwargs, response_format={"type": "json_object"},
            )
        except Exception:
            completion = client.chat.completions.create(**kwargs)
        raw = (completion.choices[0].message.content or "").strip()
        parsed = _parse_json_response(raw)
        allowed = set(ExtractedMinutesMetadata.model_fields)
        parsed = {key: value for key, value in parsed.items() if key in allowed}
        for key in ("meeting_date", "publish_date"):
            parsed[key] = _normalize_ai_deadline(parsed.get(key))
        metadata = ExtractedMinutesMetadata.model_validate(parsed).model_dump()
        merged = {
            key: metadata.get(key) or fallback.get(key)
            for key in ExtractedMinutesMetadata.model_fields
        }
        merged["meeting_type"] = _normalize_minutes_type(
            metadata.get("meeting_type"), f"{file_name}\n{minutes_text}",
        )
        merged["meeting_name"] = str(merged.get("meeting_name") or fallback["meeting_name"])[:255]
        if merged.get("minutes_number"):
            merged["minutes_number"] = str(merged["minutes_number"])[:100]
        if merged.get("meeting_subject"):
            merged["meeting_subject"] = str(merged["meeting_subject"])[:500]
        return merged, "AI"
    except Exception as exc:
        logger.warning("会议纪要基础信息 AI 识别失败，已使用文件内容兜底：%s", exc)
        return fallback, "正文兜底"


def _build_extraction_request(
    minutes_text: str, meeting_date: Any = None, stream: bool = False,
) -> tuple[Any, dict, dict]:
    from openai import OpenAI
    cfg = _resolve_action_llm()
    client = OpenAI(
        base_url=cfg["base_url"], api_key=cfg.get("api_key") or "ollama",
        timeout=120.0, max_retries=1,
    )
    messages = [
        {"role": "system", "content": AI_SYSTEM_PROMPT + "\n\n" + _supervision_prompt()},
        {"role": "user", "content": "JSON Schema:\n" + json.dumps(AI_JSON_SCHEMA, ensure_ascii=False)
         + f"\n\n会议日期：{_date_text(meeting_date) or '未提供'}"
         + "\n会议纪要正文：\n" + minutes_text[:60000]},
    ]
    kwargs: dict[str, Any] = {
        "model": cfg["model"], "messages": messages, "temperature": 0, "stream": stream,
    }
    if cfg.get("use_extra"):
        kwargs["extra_body"] = {"chat_template_kwargs": {"enable_thinking": False}}
    return client, cfg, kwargs


def _call_extraction_llm(
    minutes_text: str, meeting_date: Any = None,
) -> tuple[ExtractedPayload, str]:
    client, _, kwargs = _build_extraction_request(minutes_text, meeting_date, stream=False)
    try:
        completion = client.chat.completions.create(
            **kwargs, response_format={"type": "json_object"}
        )
    except Exception:
        completion = client.chat.completions.create(**kwargs)
    raw = (completion.choices[0].message.content or "").strip()
    parsed = _normalize_ai_payload(_parse_json_response(raw), meeting_date)
    payload = ExtractedPayload.model_validate(parsed)
    if not payload.items:
        raise ValueError("大模型未提取到可执行行动项")
    return payload, raw


_DEFAULT_RESPONSIBLE_JOB_PRIORITY = {
    "主任": 0,
    "主任责": 1,
    "副主任": 2,
}
_DEPARTMENT_LEADER_JOBS = {"经理", "副经理"}
_ALL_DEPARTMENT_MARKERS = {
    "各室", "各科室", "各专业室", "所有科室", "全部科室", "全体科室",
    "所有专业室", "全部专业室", "各业务科室",
}


def _job_director_priority(job: str) -> Optional[int]:
    """科室主任类职务优先级：主任 > 主任责 > 副主任；无法识别返回 None。"""
    j = (job or "").strip()
    if not j:
        return None
    if "副主任" in j or j.startswith("副主任"):
        return 2
    if "主任责" in j:
        return 1
    if j == "主任" or j.startswith("主任"):
        return 0
    return None


def _normalize_department_name(value: Any) -> str:
    """复用 AI 助手的部门简称词典，统一输出数据库正式科室名。"""
    name = str(value or "").strip()
    return DEPARTMENT_ALIASES.get(name, name)


def _default_responsibles(
    user_rows: list[dict], people: dict[str, str],
) -> dict[str, str]:
    """为每个科室选出稳定的默认负责人，避免同一批数据重复处理时结果漂移。"""
    candidates: dict[str, list[tuple[int, str]]] = {}
    for row in user_rows:
        name = (row.get("name") or "").strip()
        department = people.get(name, "")
        priority = _job_director_priority(row.get("jb") or "")
        if not name or not department or priority is None:
            continue
        candidates.setdefault(department, []).append((priority, name))
    return {
        department: sorted(items, key=lambda item: (item[0], item[1]))[0][1]
        for department, items in candidates.items()
    }


def _department_responsible_people(
    department: str, people: dict[str, str], jobs: dict[str, str],
) -> list[str]:
    candidates = [
        (_job_director_priority(jobs.get(name, "")), name)
        for name, person_department in people.items()
        if person_department == department
        and _job_director_priority(jobs.get(name, "")) is not None
    ]
    return [name for _, name in sorted(candidates, key=lambda item: (item[0], item[1]))]


def _publish_receive_targets(
    action: dict,
    people: Optional[dict[str, str]] = None,
    jobs: Optional[dict[str, str]] = None,
) -> list[str]:
    """发布后应收到「待接收」提醒的对象：责任人 + 责任科室主任/主任责/副主任。"""
    targets = list(_responsible_people(action))
    if people is None or jobs is None:
        _departments, people, jobs, _supervisors, _defaults = _action_directory_maps()
    for department in _responsible_departments(action):
        for name in _department_responsible_people(department, people, jobs):
            if name not in targets:
                targets.append(name)
    return targets


def _close_department_receiver_reminders(
    action_id: int,
    department: str,
    people: Optional[dict[str, str]] = None,
    jobs: Optional[dict[str, str]] = None,
) -> int:
    """同一科室任一负责人作出接收/分配选择后，关闭该科室负责人组的未读提醒。"""
    department = str(department or "").strip()
    if not department:
        return 0
    if people is None or jobs is None:
        _departments, people, jobs, _supervisors, _defaults = _action_directory_maps()
    targets = _department_responsible_people(department, people, jobs)
    if not targets:
        return 0
    placeholders = ",".join(["%s"] * len(targets))
    return db.execute_update(
        "UPDATE action_reminder SET read_at=COALESCE(read_at,NOW()) "
        f"WHERE action_item_id=%s AND reminder_target IN ({placeholders}) "
        "AND read_at IS NULL",
        (action_id, *targets),
    )


def _can_receive_action(ctx: dict, action: dict) -> bool:
    """责任人，或责任科室主任/主任责/副主任，可接收待接收行动项。"""
    if action.get("current_status") != STATUS_PENDING_RECEIVE:
        return False
    if _is_responsible_person(action, ctx["name"]):
        return True
    if (
        ctx.get("dept")
        and _is_responsible_department(action, ctx["dept"])
        and _job_director_priority(ctx.get("jb") or "") is not None
    ):
        return True
    return False


def _assignment_departments(ctx: dict, action: dict) -> list[str]:
    """按当前角色返回发布后可分配责任人的科室范围。"""
    responsible_departments = _responsible_departments(action)
    if ctx["minutes_admin"] or ctx["department_leader"]:
        return responsible_departments
    if (
        ctx.get("dept_director")
        and ctx.get("dept")
        and ctx["dept"] in responsible_departments
    ):
        return [ctx["dept"]]
    return []


def _is_all_departments_action(
    item: ExtractedAction, people: Optional[dict[str, str]] = None,
) -> bool:
    raw_department = (item.responsible_department or "").strip()
    if raw_department in _ALL_DEPARTMENT_MARKERS:
        return True
    text = "\n".join((item.title, item.content, item.source_quote))
    if re.search(
        r"(?:各|所有|全部|全体)(?:业务|专业|责任)?(?:科室|室)(?:主任|负责人)?",
        text,
    ):
        return True
    normalized_department = _normalize_department_name(raw_department)
    valid_departments = {*SHIFT_BUSINESS_DEPARTMENTS, "部办"}
    if normalized_department in valid_departments:
        return False
    responsible_person = (item.responsible_person or "").strip()
    if responsible_person and (people is None or responsible_person in people):
        return False
    return True


def _work_division_terms(value: str) -> list[str]:
    parts = re.split(
        r"[，,、；;。：“”\"'（）()【】/\\\n]|以及|并负责|负责|分管|牵头|和|及|与",
        value or "",
    )
    terms = []
    for part in parts:
        term = re.sub(r"(相关)?(工作|业务|事项|方面|领域)$", "", part.strip())
        if len(term) >= 2 and term not in terms:
            terms.append(term)
    return terms


def _work_division_score(action_text: str, work_division: str) -> int:
    text = re.sub(r"\s+", "", action_text or "")
    if not text or not work_division:
        return 0
    return sum(len(term) ** 2 for term in _work_division_terms(work_division) if term in text)


def _configured_supervisor(
    department: Optional[str], action_text: str, supervisors: set[str],
    supervision_configs: list[dict],
) -> Optional[str]:
    valid = [
        item for item in supervision_configs
        if item.get("enabled", True) and item.get("leader_name") in supervisors
    ]
    if not valid:
        return None
    department_candidates = [
        item for item in valid
        if department and department in (item.get("departments") or [])
    ]
    candidates = department_candidates or valid
    scored = [
        (_work_division_score(action_text, item.get("work_division") or ""), item)
        for item in candidates
    ]
    best_score = max((score for score, _ in scored), default=0)
    if best_score > 0:
        winners = [item for score, item in scored if score == best_score]
        if len(winners) == 1:
            return winners[0]["leader_name"]
    if len(department_candidates) == 1:
        return department_candidates[0]["leader_name"]
    return None


def _department_manager(jobs: dict[str, str], supervisors: set[str]) -> Optional[str]:
    exact = sorted(
        name for name, job in jobs.items()
        if name in supervisors and str(job or "").strip() == "经理"
    )
    if exact:
        return exact[0]
    candidates = sorted(
        name for name, job in jobs.items()
        if name in supervisors and "经理" in str(job or "") and "副经理" not in str(job or "")
    )
    return candidates[0] if candidates else None


def _configured_supervisor_for_departments(
    responsibility_departments: list[str], action_text: str,
    supervisors: set[str], supervision_configs: list[dict],
) -> Optional[str]:
    selected = set(responsibility_departments)
    valid = [
        item for item in supervision_configs
        if item.get("enabled", True)
        and item.get("leader_name") in supervisors
        and selected.intersection(item.get("departments") or [])
    ]
    if not valid:
        return None
    scored = [
        (_work_division_score(action_text, item.get("work_division") or ""), item)
        for item in valid
    ]
    best_score = max((score for score, _ in scored), default=0)
    if best_score > 0:
        winners = [item for score, item in scored if score == best_score]
        if len(winners) == 1:
            return winners[0]["leader_name"]
    full_coverage = [
        item for item in valid
        if selected.issubset(set(item.get("departments") or []))
    ]
    if len(full_coverage) == 1:
        return full_coverage[0]["leader_name"]
    overlap_counts = [
        (len(selected.intersection(item.get("departments") or [])), item)
        for item in valid
    ]
    best_overlap = max((count for count, _ in overlap_counts), default=0)
    winners = [item for count, item in overlap_counts if count == best_overlap]
    return winners[0]["leader_name"] if best_overlap and len(winners) == 1 else None


def _resolve_scope_supervisor(
    responsibility_departments: list[str], responsibility_people: list[str],
    requested_supervisor: Any, jobs: dict[str, str], supervisors: set[str],
    supervision_configs: Optional[list[dict]], action_text: str,
) -> Optional[str]:
    selected = set(responsibility_departments)
    if set(SHIFT_BUSINESS_DEPARTMENTS).issubset(selected):
        return _department_manager(jobs, supervisors)
    for person in responsibility_people:
        if person in supervisors and jobs.get(person, "") in _DEPARTMENT_LEADER_JOBS:
            return person
    configs = (
        supervision_configs if supervision_configs is not None
        else _load_supervision_configs()
    )
    configured = _configured_supervisor_for_departments(
        responsibility_departments, action_text, supervisors, configs,
    )
    if configured:
        return configured
    requested = str(requested_supervisor or "").strip()
    return requested if requested in supervisors else None


def _resolve_assignment_rules(
    department: Any, person: Any, supervisor: Any,
    departments: set[str], people: dict[str, str], jobs: dict[str, str],
    supervisors: set[str], default_responsibles: dict[str, str],
    supervision_configs: Optional[list[dict]] = None, action_text: str = "",
) -> tuple[Optional[str], Optional[str], Optional[str]]:
    """统一执行责任人优先、科室负责人兜底和部门领导自主管规则。"""
    resolved_department = _normalize_department_name(department)
    resolved_person = str(person or "").strip()
    resolved_supervisor = str(supervisor or "").strip()

    if resolved_department not in departments:
        resolved_department = ""
    if resolved_person not in people:
        resolved_person = ""

    # 责任人是人员主数据，所属科室必须以其当前在职信息为准。
    if resolved_person:
        resolved_department = people[resolved_person]
    elif resolved_department:
        # 只识别到科室时，按职务优先级补齐该科室负责人。
        resolved_person = default_responsibles.get(resolved_department, "")

    # 部门经理/副经理亲自负责时，不再另行指定主管领导。
    if resolved_person and jobs.get(resolved_person, "") in _DEPARTMENT_LEADER_JOBS:
        resolved_supervisor = resolved_person
    elif resolved_supervisor not in supervisors:
        resolved_supervisor = _configured_supervisor(
            resolved_department or None, action_text, supervisors,
            supervision_configs if supervision_configs is not None
            else _load_supervision_configs(),
        ) or ""

    return (
        resolved_department or None,
        resolved_person or None,
        resolved_supervisor or None,
    )


def _resolve_responsibility_lists(
    department_values: Any, person_values: Any,
    departments: set[str], people: dict[str, str], jobs: dict[str, str],
) -> tuple[list[str], list[str]]:
    raw_departments = (
        department_values if isinstance(department_values, list)
        else [department_values]
    )
    raw_people = person_values if isinstance(person_values, list) else [person_values]
    resolved_people = list(dict.fromkeys(
        str(name or "").strip() for name in raw_people
        if str(name or "").strip() in people
    ))
    resolved_departments = list(dict.fromkeys(
        _normalize_department_name(value) for value in raw_departments
        if _normalize_department_name(value) in departments
    ))
    for name in resolved_people:
        department = people[name]
        if department and department not in resolved_departments:
            resolved_departments.append(department)
    if not resolved_people:
        resolved_people = list(dict.fromkeys(
            name
            for department in resolved_departments
            for name in _department_responsible_people(department, people, jobs)
        ))
    return resolved_departments, resolved_people


def _normalize_ai_item(
    item: ExtractedAction, departments: set[str],
    people: dict[str, str], jobs: dict[str, str], supervisors: set[str],
    default_responsibles: dict[str, str],
    supervision_configs: Optional[list[dict]] = None,
) -> dict:
    uncertain = set(item.uncertain_fields or [])
    all_departments = _is_all_departments_action(item, people)
    dept = _normalize_department_name(item.responsible_department)
    requested_person = (item.responsible_person or "").strip()
    person = requested_person
    supervisor = (item.supervisor or "").strip()
    requested_collaborators = [
        _normalize_department_name(d)
        for d in item.collaborating_departments if str(d or "").strip()
    ]
    collaborators = list(dict.fromkeys(
        d for d in requested_collaborators if d in departments
    ))
    if any(d not in departments for d in requested_collaborators):
        uncertain.add("collaborating_departments")
    action_text = "\n".join((item.title, item.content, item.source_quote))
    if all_departments:
        responsible_departments = [
            department for department in SHIFT_BUSINESS_DEPARTMENTS
            if department in departments
        ]
        responsible_people = list(dict.fromkeys(
            name
            for department in responsible_departments
            for name in _department_responsible_people(department, people, jobs)
        ))
        dept = responsible_departments[0] if responsible_departments else None
        person = responsible_people[0] if responsible_people else None
        uncertain.discard("responsible_department")
        if responsible_people:
            uncertain.discard("responsible_person")
    else:
        dept, person, supervisor = _resolve_assignment_rules(
            dept, person, supervisor, departments, people, jobs, supervisors,
            default_responsibles, supervision_configs, action_text,
        )
        responsible_departments = [dept] if dept else []
        if dept and requested_person not in people:
            responsible_people = _department_responsible_people(dept, people, jobs)
        else:
            responsible_people = [person] if person else []
        if person and person not in responsible_people:
            responsible_people.insert(0, person)
    supervisor = _resolve_scope_supervisor(
        responsible_departments, responsible_people, item.supervisor,
        jobs, supervisors, supervision_configs, action_text,
    )
    collaborating_people = []
    for field, value in (
        ("responsible_department", responsible_departments),
        ("responsible_person", responsible_people),
        ("supervisor", supervisor),
    ):
        if value:
            uncertain.discard(field)
        else:
            uncertain.add(field)
    if not item.deadline:
        uncertain.add("deadline")
    return {
        "title": item.title.strip()[:500],
        "content": item.content.strip(),
        "source_quote": item.source_quote.strip(),
        "responsible_department_id": dept,
        "responsible_person_id": person,
        "responsible_department_ids": responsible_departments,
        "responsible_person_ids": responsible_people,
        "collaborating_departments": collaborators,
        "collaborating_people": collaborating_people,
        "supervisor_id": supervisor,
        "required_completion_date": item.deadline,
        "priority": item.priority,
        "uncertain_fields": sorted(uncertain),
    }


def _action_directory_maps() -> tuple[
    set[str], dict[str, str], dict[str, str], set[str], dict[str, str],
]:
    user_rows = db.execute_query(
        "SELECT name,lsys,jb FROM yggl WHERE COALESCE(zaizhi,0)=0 AND name IS NOT NULL"
    )
    people = {
        (r.get("name") or "").strip(): (r.get("lsys") or "").strip()
        for r in user_rows if (r.get("name") or "").strip()
    }
    jobs = {
        (r.get("name") or "").strip(): (r.get("jb") or "").strip()
        for r in user_rows if (r.get("name") or "").strip()
    }
    departments = {
        (r.get("lsys") or "").strip()
        for r in user_rows if (r.get("lsys") or "").strip()
    }
    supervisors = {
        (r.get("name") or "").strip() for r in user_rows
        if any(k in (r.get("jb") or "") for k in ("部长", "副部长", "经理", "副经理"))
    }
    return departments, people, jobs, supervisors, _default_responsibles(user_rows, people)


def _archive_old_ai_drafts(minutes_id: int, operator: str) -> int:
    """重新提取时逻辑取消旧 AI 草稿，人工新增草稿不受影响。"""
    old = db.execute_query(
        "SELECT id FROM action_item WHERE source_meeting_id=%s "
        "AND current_status IN ('草稿','待发布') AND COALESCE(ai_batch_id,'')<>''",
        (minutes_id,),
    )
    for row in old:
        db.execute_update(
            "UPDATE action_item SET current_status=%s WHERE id=%s",
            (STATUS_CANCELLED, row["id"]),
        )
        _event(
            operator, "AI重新提取", "新批次提取时保留并取消旧草稿",
            action_id=row["id"], meeting_id=minutes_id,
        )
    return len(old)


def _save_extracted_drafts(
    minutes: dict, payload: ExtractedPayload, raw: str, operator: str,
    departments: set[str], people: dict[str, str], jobs: dict[str, str],
    supervisors: set[str], default_responsibles: dict[str, str],
) -> dict:
    minutes_id = int(minutes["id"])
    batch_id = uuid.uuid4().hex
    ids = []
    supervision_configs = _load_supervision_configs()
    normalized_items = []
    default_day = _as_date(_minutes_reference_date(minutes))
    default_deadline = (
        (default_day + timedelta(days=7)).isoformat() if default_day else None
    )
    for extracted in payload.items:
        item = _normalize_ai_item(
            extracted, departments, people, jobs, supervisors, default_responsibles,
            supervision_configs,
        )
        if not item["required_completion_date"]:
            item["required_completion_date"] = default_deadline
            if default_deadline:
                item["uncertain_fields"] = [
                    field for field in item["uncertain_fields"] if field != "deadline"
                ]
        normalized_items.append(item)

    missing = []
    required = (
        ("责任科室", "responsible_department_ids"),
        ("责任人", "responsible_person_ids"),
        ("主管领导", "supervisor_id"),
        ("完成时间", "required_completion_date"),
    )
    for index, item in enumerate(normalized_items, start=1):
        absent = [label for label, key in required if not item.get(key)]
        if absent:
            missing.append(f"第{index}项《{item['title']}》缺少{'、'.join(absent)}")
    if missing:
        raise HTTPException(
            status_code=422,
            detail=(
                "AI 提取结果未能补齐四项必填信息："
                + "；".join(missing[:12])
                + "。请完善科室分管领导配置或在纪要中明确责任主体后重新提取。"
            ),
        )

    archived = _archive_old_ai_drafts(minutes_id, operator)
    for item in normalized_items:
        action_id = db.execute_insert(
            "INSERT INTO action_item "
            "(title,content,source_type,source_meeting_id,minutes_number,source_quote,"
            "responsible_department_id,responsible_person_id,responsible_department_ids,"
            "responsible_person_ids,collaborating_departments,"
            "collaborating_people,"
            "supervisor_id,required_completion_date,priority,current_status,created_by,"
            "uncertain_fields,ai_batch_id) VALUES (%s,%s,'会议纪要',%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
            (
                item["title"], item["content"], minutes_id,
                minutes.get("minutes_number") or "", item["source_quote"],
                item["responsible_department_id"], item["responsible_person_id"],
                json.dumps(item["responsible_department_ids"], ensure_ascii=False),
                json.dumps(item["responsible_person_ids"], ensure_ascii=False),
                json.dumps(item["collaborating_departments"], ensure_ascii=False),
                json.dumps(item["collaborating_people"], ensure_ascii=False),
                item["supervisor_id"], item["required_completion_date"], item["priority"],
                STATUS_DRAFT, operator,
                json.dumps(item["uncertain_fields"], ensure_ascii=False), batch_id,
            ),
        )
        ids.append(action_id)
        _event(
            operator, "AI提取", "AI 生成行动项草稿，待人工确认",
            action_id=action_id, meeting_id=minutes_id,
        )
    db.execute_update(
        "UPDATE meeting_minutes SET status='待确认',ai_raw_json=%s,ai_batch_id=%s,"
        "ai_extracted_at=NOW() WHERE id=%s",
        (raw, batch_id, minutes_id),
    )
    _event(
        operator, "AI提取完成", f"提取 {len(ids)} 条行动项草稿",
        meeting_id=minutes_id, data={"batchId": batch_id, "count": len(ids)},
    )
    return {
        "success": True, "count": len(ids), "ids": ids, "batchId": batch_id,
        "archived": archived,
        "message": f"AI 已提取 {len(ids)} 条草稿，请人工核对后发布",
    }


def _serialize_action(row: dict) -> dict:
    result = dict(row)
    for key in (
        "responsible_department_ids", "responsible_person_ids",
        "collaborating_departments", "collaborating_people", "uncertain_fields",
    ):
        result[key] = _json(result.get(key))
    if not result["responsible_department_ids"] and result.get("responsible_department_id"):
        result["responsible_department_ids"] = [result["responsible_department_id"]]
    if not result["responsible_person_ids"] and result.get("responsible_person_id"):
        result["responsible_person_ids"] = [result["responsible_person_id"]]
    for key in ("required_completion_date", "actual_completion_date"):
        result[key] = _date_text(result.get(key))
    for key in ("published_at", "created_at", "updated_at", "last_progress_at"):
        result[key] = _dt_text(result.get(key))
    result["risk_tags"] = _risk_tags(row)
    return result


def _risk_tags(row: dict) -> list[str]:
    if (row.get("current_status") or "") in (
        STATUS_COMPLETED, STATUS_CANCELLED, STATUS_DRAFT, STATUS_PENDING_PUBLISH
    ):
        return []
    tags: list[str] = []
    deadline = row.get("required_completion_date")
    if deadline:
        if isinstance(deadline, datetime):
            deadline = deadline.date()
        elif not isinstance(deadline, date):
            try:
                deadline = date.fromisoformat(str(deadline)[:10])
            except Exception:
                deadline = None
    if deadline:
        delta = (deadline - date.today()).days
        if delta < 0:
            tags.append("已逾期")
        elif delta <= 7:
            tags.append("临期")
    if (row.get("risk_status") or "") == "存在延期风险":
        tags.append("存在延期风险")
    latest = row.get("last_progress_at") or row.get("published_at")
    if latest:
        if isinstance(latest, str):
            try:
                latest = datetime.fromisoformat(latest[:19])
            except Exception:
                latest = None
        if latest and datetime.now() - latest > timedelta(days=7):
            tags.append("长期未更新")
    return tags


@router.get("/permissions")
def get_permissions(current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    return {"success": True, "user": ctx, "permissions": _permissions(ctx)}


@router.get("/directory")
def get_directory(current_user: str = Query(...)):
    ensure_tables()
    _user(current_user)
    rows = db.execute_query(
        "SELECT name,gh,lsys,jb FROM yggl WHERE COALESCE(zaizhi,0)=0 "
        "AND name IS NOT NULL AND TRIM(name)<>'' ORDER BY lsys,name"
    )
    departments = sorted({(r.get("lsys") or "").strip() for r in rows if (r.get("lsys") or "").strip()})
    people = [{
        "name": (r.get("name") or "").strip(), "gh": (r.get("gh") or "").strip(),
        "department": (r.get("lsys") or "").strip(), "job": (r.get("jb") or "").strip(),
    } for r in rows]
    supervisors = [p for p in people if any(k in p["job"] for k in ("部长", "副部长", "经理", "副经理"))]
    people_map = {p["name"]: p["department"] for p in people}
    default_responsibles = _default_responsibles(rows, people_map)
    return {
        "success": True,
        "departments": departments,
        "people": people,
        "supervisors": supervisors,
        "defaultResponsibles": default_responsibles,
    }


@router.post("/minutes")
async def create_minutes(
    current_user: str = Form(...), file: UploadFile = File(...),
):
    ensure_tables()
    ctx = _role_context(current_user)
    if not ctx["minutes_uploader"]:
        raise HTTPException(status_code=403, detail="仅经理、副经理、综合室或系统管理员可上传会议纪要")
    original_name = Path(file.filename or "会议纪要").name
    initial_name = (Path(original_name).stem.strip() or "会议纪要")[:255]
    minutes_id = db.execute_insert(
        "INSERT INTO meeting_minutes "
        "(meeting_name,meeting_type,minutes_text,creator,status) VALUES (%s,%s,%s,%s,%s)",
        (initial_name, "其他", "", ctx["name"], STATUS_DRAFT),
    )
    if not minutes_id:
        raise HTTPException(status_code=500, detail="会议纪要保存失败")
    attachment = await _save_upload(
        file, MINUTES_DIR, "meeting_minutes", minutes_id, ctx["name"],
        allowed=ALLOWED_MINUTES_EXT,
    )
    att_rows = db.execute_query(
        "SELECT stored_path,original_name FROM action_attachment WHERE id=%s",
        (attachment["id"],),
    )
    if not att_rows:
        db.execute_update("UPDATE meeting_minutes SET status='解析失败' WHERE id=%s", (minutes_id,))
        raise HTTPException(status_code=500, detail="会议纪要附件保存失败")
    stored_path = att_rows[0]["stored_path"]
    original_name = att_rows[0]["original_name"]
    path = BASE_DIR / stored_path
    try:
        parsed_text = _extract_text(path)
    except Exception as exc:
        db.execute_update(
            "UPDATE meeting_minutes SET status='解析失败',original_attachment=%s,"
            "original_file_name=%s WHERE id=%s",
            (stored_path, original_name, minutes_id),
        )
        _event(
            ctx["name"], "纪要解析失败", str(exc), meeting_id=minutes_id,
            data={"attachmentId": attachment["id"]},
        )
        raise HTTPException(status_code=400, detail=f"纪要正文解析失败：{exc}") from exc
    if not parsed_text:
        db.execute_update(
            "UPDATE meeting_minutes SET status='解析失败',original_attachment=%s,"
            "original_file_name=%s WHERE id=%s",
            (stored_path, original_name, minutes_id),
        )
        _event(ctx["name"], "纪要解析失败", "未解析到会议纪要正文", meeting_id=minutes_id)
        raise HTTPException(status_code=400, detail="未解析到会议纪要正文，请检查文件内容")

    metadata, metadata_source = _extract_minutes_metadata(parsed_text, original_name)
    db.execute_update(
        "UPDATE meeting_minutes SET meeting_name=%s,meeting_type=%s,meeting_date=%s,"
        "minutes_number=%s,meeting_subject=%s,publish_date=%s,minutes_text=%s,"
        "original_attachment=%s,original_file_name=%s WHERE id=%s",
        (
            metadata["meeting_name"], metadata["meeting_type"], metadata["meeting_date"],
            metadata["minutes_number"] or "", metadata["meeting_subject"] or "",
            metadata["publish_date"], parsed_text, stored_path, original_name, minutes_id,
        ),
    )
    _event(
        ctx["name"], "纪要上传",
        f"上传会议纪要《{metadata['meeting_name']}》，基础信息由{metadata_source}自动识别",
        meeting_id=minutes_id, data={"metadataSource": metadata_source},
    )
    return {
        "success": True, "id": minutes_id, "metadata": metadata,
        "metadataSource": metadata_source,
        "message": f"会议纪要已保存，基础信息由{metadata_source}自动识别",
    }


@router.get("/minutes")
def list_minutes(
    current_user: str = Query(...), keyword: str = Query(""),
    meeting_type: str = Query(""), page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
):
    ensure_tables()
    ctx = _role_context(current_user)
    where = ["1=1"]
    params: list[Any] = []
    if keyword.strip():
        where.append("CONCAT_WS('',meeting_name,minutes_number,meeting_subject) LIKE %s")
        params.append(f"%{keyword.strip()}%")
    if meeting_type.strip():
        where.append("meeting_type=%s")
        params.append(meeting_type.strip())
    clause = " AND ".join(where)
    total = db.execute_scalar(f"SELECT COUNT(*) FROM meeting_minutes WHERE {clause}", tuple(params)) or 0
    count_scope = (
        "a.current_status<>'已取消'" if ctx["minutes_admin"]
        else "a.current_status NOT IN ('草稿','待发布','已取消')"
    )
    rows = db.execute_query(
        f"SELECT m.*,(SELECT COUNT(*) FROM action_item a WHERE a.source_meeting_id=m.id "
        f"AND {count_scope}) action_count FROM meeting_minutes m WHERE {clause} "
        "ORDER BY COALESCE(meeting_date,DATE(created_at)) DESC,id DESC LIMIT %s OFFSET %s",
        tuple(params + [page_size, (page - 1) * page_size]),
    )
    items = []
    for row in rows:
        item = dict(row)
        item.pop("convening_leader", None)
        for key in ("meeting_date", "publish_date"):
            item[key] = _date_text(item.get(key))
        for key in ("created_at", "ai_extracted_at"):
            item[key] = _dt_text(item.get(key))
        item.pop("ai_raw_json", None)
        item["minutes_text_preview"] = (item.pop("minutes_text", "") or "")[:240]
        items.append(item)
    return {"success": True, "items": items, "total": int(total), "page": page, "pageSize": page_size}


@router.delete("/minutes/{minutes_id}")
def delete_draft_minutes(
    minutes_id: int, current_user: str = Query(...),
):
    """删除尚未发布的会议纪要及其草稿行动项、附件和过程记录。"""
    ensure_tables()
    ctx = _role_context(current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可删除未发布会议纪要")
    rows = db.execute_query(
        "SELECT id,meeting_name,status FROM meeting_minutes WHERE id=%s LIMIT 1",
        (minutes_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="会议纪要不存在")
    minutes = rows[0]
    if minutes.get("status") == "已发布":
        raise HTTPException(status_code=400, detail="已发布会议纪要不允许删除")
    published_count = db.execute_scalar(
        "SELECT COUNT(*) FROM action_item WHERE source_meeting_id=%s "
        "AND current_status NOT IN ('草稿','待发布','已取消')",
        (minutes_id,),
    ) or 0
    if published_count:
        raise HTTPException(
            status_code=400,
            detail="该会议纪要存在已发布行动项，不允许删除",
        )

    action_rows = db.execute_query(
        "SELECT id FROM action_item WHERE source_meeting_id=%s",
        (minutes_id,),
    )
    action_ids = [int(row["id"]) for row in action_rows]
    attachment_where = ["(business_type='meeting_minutes' AND business_id=%s)"]
    attachment_params: list[Any] = [minutes_id]
    if action_ids:
        placeholders = ",".join(["%s"] * len(action_ids))
        attachment_where.append(f"action_item_id IN ({placeholders})")
        attachment_params.extend(action_ids)
    attachment_rows = db.execute_query(
        "SELECT stored_path FROM action_attachment WHERE "
        + " OR ".join(attachment_where),
        tuple(attachment_params),
    )

    if action_ids:
        placeholders = ",".join(["%s"] * len(action_ids))
        action_params = tuple(action_ids)
        for table in (
            "action_progress", "completion_application", "action_approval",
            "action_change", "action_reminder", "action_department_execution",
        ):
            db.execute_update(
                f"DELETE FROM {table} WHERE action_item_id IN ({placeholders})",
                action_params,
            )
        db.execute_update(
            f"DELETE FROM action_attachment WHERE action_item_id IN ({placeholders})",
            action_params,
        )
        db.execute_update(
            f"DELETE FROM action_event_log WHERE action_item_id IN ({placeholders})",
            action_params,
        )
        db.execute_update(
            f"DELETE FROM action_item WHERE id IN ({placeholders})",
            action_params,
        )
    db.execute_update(
        "DELETE FROM action_attachment "
        "WHERE business_type='meeting_minutes' AND business_id=%s",
        (minutes_id,),
    )
    db.execute_update(
        "DELETE FROM action_event_log WHERE meeting_minutes_id=%s",
        (minutes_id,),
    )
    db.execute_update("DELETE FROM meeting_minutes WHERE id=%s", (minutes_id,))

    removed_files = 0
    for attachment in attachment_rows:
        try:
            path = (BASE_DIR / str(attachment.get("stored_path") or "")).resolve()
            if BASE_DIR.resolve() in path.parents and path.is_file():
                path.unlink()
                removed_files += 1
        except Exception as exc:
            logger.warning(
                "会议纪要记录已删除，但附件文件清理失败 minutes=%s path=%s: %s",
                minutes_id, attachment.get("stored_path"), exc,
            )
    return {
        "success": True,
        "message": f"未发布会议纪要《{minutes.get('meeting_name') or minutes_id}》已删除",
        "deletedActions": len(action_ids),
        "deletedFiles": removed_files,
    }


@router.get("/minutes/{minutes_id}")
def get_minutes(minutes_id: int, current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    rows = db.execute_query("SELECT * FROM meeting_minutes WHERE id=%s LIMIT 1", (minutes_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="会议纪要不存在")
    item = dict(rows[0])
    item.pop("convening_leader", None)
    item["meeting_date"] = _date_text(item.get("meeting_date"))
    item["publish_date"] = _date_text(item.get("publish_date"))
    item["created_at"] = _dt_text(item.get("created_at"))
    item["ai_extracted_at"] = _dt_text(item.get("ai_extracted_at"))
    action_scope = "" if ctx["minutes_admin"] else " AND current_status NOT IN ('草稿','待发布','已取消')"
    actions = db.execute_query(
        f"SELECT * FROM action_item WHERE source_meeting_id=%s{action_scope} ORDER BY id", (minutes_id,)
    )
    item["actions"] = [_serialize_action(row) for row in actions]
    if not ctx["minutes_admin"]:
        item.pop("ai_raw_json", None)
    attachments = db.execute_query(
        "SELECT id,original_name,file_size,uploaded_at FROM action_attachment "
        "WHERE business_type='meeting_minutes' AND business_id=%s ORDER BY id", (minutes_id,)
    )
    item["attachments"] = [{
        "id": a["id"], "name": a["original_name"], "size": a["file_size"],
        "uploadedAt": _dt_text(a["uploaded_at"]),
    } for a in attachments]
    return {"success": True, "item": item}


@router.post("/minutes/{minutes_id}/extract")
def extract_minutes(minutes_id: int, current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可执行 AI 提取")
    rows = db.execute_query("SELECT * FROM meeting_minutes WHERE id=%s LIMIT 1", (minutes_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="会议纪要不存在")
    minutes = rows[0]
    text = (minutes.get("minutes_text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="会议纪要正文为空")
    try:
        payload, raw = _call_extraction_llm(text, _minutes_reference_date(minutes))
    except ValidationError as exc:
        problems = []
        for error in exc.errors()[:8]:
            location = ".".join(str(part) for part in error.get("loc", []))
            message = error.get("msg") or "格式不正确"
            problems.append(f"{location}：{message}")
        detail = "；".join(problems) or "返回结构不符合约定"
        raise HTTPException(
            status_code=422,
            detail=f"AI 返回内容格式不符合要求，可重新提取：{detail}",
        ) from exc
    except (ValueError, json.JSONDecodeError) as exc:
        raise HTTPException(
            status_code=422,
            detail=f"AI 返回内容无法解析，可重新提取：{exc}",
        ) from exc
    except Exception as exc:
        logger.exception("行动项 AI 提取失败: %s", exc)
        raise HTTPException(status_code=502, detail=f"大模型调用失败：{exc}") from exc

    departments, people, jobs, supervisors, default_responsibles = _action_directory_maps()
    return _save_extracted_drafts(
        minutes, payload, raw, ctx["name"], departments, people, jobs, supervisors,
        default_responsibles,
    )


def _stream_event(event_type: str, **data: Any) -> str:
    return json.dumps(
        {
            "type": event_type,
            "time": datetime.now().strftime("%H:%M:%S"),
            **data,
        },
        ensure_ascii=False,
        default=str,
    ) + "\n"


@router.post("/minutes/{minutes_id}/extract/stream")
async def stream_extract_minutes(minutes_id: int, current_user: str = Query(...)):
    """以 NDJSON 流输出模型文本和提取工作流，并在严格校验后保存草稿。"""
    ensure_tables()
    ctx = _role_context(current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可执行 AI 提取")
    rows = db.execute_query("SELECT * FROM meeting_minutes WHERE id=%s LIMIT 1", (minutes_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="会议纪要不存在")
    minutes = rows[0]
    text = (minutes.get("minutes_text") or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="会议纪要正文为空")

    def generate():
        stream_response = None
        try:
            yield _stream_event(
                "workflow", step="读取会议纪要", status="done",
                message=f"已读取《{minutes.get('meeting_name') or ''}》，正文 {len(text)} 字",
            )
            yield _stream_event(
                "workflow", step="加载提取规则", status="done",
                message="已加载 JSON Schema、部门组织结构、科室简称、分管领导及防虚构规则",
            )

            client, cfg, kwargs = _build_extraction_request(
                text, _minutes_reference_date(minutes), stream=True,
            )
            provider = "联网 DeepSeek" if cfg.get("provider") == "deepseek" else (
                "本地/兼容模型" if cfg.get("provider") in ("local", "scene") else str(cfg.get("provider") or "模型")
            )
            yield _stream_event(
                "workflow", step="选择大模型", status="done",
                message=f"{provider} · {cfg.get('model') or '未命名模型'}",
            )
            yield _stream_event(
                "workflow", step="模型流式生成", status="running",
                message="正在生成结构化行动项 JSON",
            )
            try:
                stream_response = client.chat.completions.create(
                    **kwargs, response_format={"type": "json_object"},
                )
            except Exception:
                yield _stream_event(
                    "workflow", step="模型兼容处理", status="info",
                    message="当前模型不支持 JSON response_format，已切换普通流式输出",
                )
                stream_response = client.chat.completions.create(**kwargs)

            raw_parts: list[str] = []
            for chunk in stream_response:
                if not getattr(chunk, "choices", None):
                    continue
                delta = chunk.choices[0].delta
                reasoning = getattr(delta, "reasoning_content", None)
                if reasoning:
                    yield _stream_event("reasoning", content=str(reasoning))
                content = getattr(delta, "content", None)
                if content:
                    content = str(content)
                    raw_parts.append(content)
                    yield _stream_event("token", content=content)
            raw = "".join(raw_parts).strip()
            if not raw:
                raise ValueError("大模型未返回可解析内容")
            yield _stream_event(
                "workflow", step="模型流式生成", status="done",
                message=f"输出完成，共 {len(raw)} 个字符",
            )

            yield _stream_event(
                "workflow", step="期限归一化", status="running",
                message="正在处理“年底、月底、本周内”等自然语言期限",
            )
            parsed = _normalize_ai_payload(
                _parse_json_response(raw), _minutes_reference_date(minutes),
            )
            yield _stream_event(
                "workflow", step="期限归一化", status="done",
                message="未指定完成时间的行动项已按会议日期后7天补全",
            )

            yield _stream_event(
                "workflow", step="JSON 严格校验", status="running",
                message="正在校验字段、类型、优先级和不确定项",
            )
            payload = ExtractedPayload.model_validate(parsed)
            if not payload.items:
                raise ValueError("大模型未提取到可执行行动项")
            yield _stream_event(
                "workflow", step="JSON 严格校验", status="done",
                message=f"校验通过，识别到 {len(payload.items)} 条行动项",
            )
            yield _stream_event(
                "workflow", step="组织信息匹配", status="running",
                message="正在按责任人、责任科室及领导工作分工匹配主管领导",
            )
            departments, people, jobs, supervisors, default_responsibles = (
                _action_directory_maps()
            )
            yield _stream_event(
                "workflow", step="组织信息匹配", status="done",
                message=f"已加载 {len(departments)} 个科室、{len(people)} 名在职人员",
            )
            multi_department_count = sum(
                1 for item in payload.items
                if _is_all_departments_action(item, people)
            )
            if multi_department_count:
                yield _stream_event(
                    "workflow", step="多科室责任归并", status="done",
                    message=(
                        f"识别到 {multi_department_count} 条全体科室任务；"
                        "每项保留为一条行动项，并关联全部责任科室及其负责人"
                    ),
                )

            yield _stream_event(
                "workflow", step="草稿历史处理", status="running",
                message="四项必填信息校验通过后，将保留历史并逻辑取消旧 AI 草稿",
            )

            yield _stream_event(
                "workflow", step="保存行动项草稿", status="running",
                message="正在校验责任科室、责任人、主管领导、完成时间并写入草稿",
            )
            result = _save_extracted_drafts(
                minutes, payload, raw, ctx["name"],
                departments, people, jobs, supervisors, default_responsibles,
            )
            yield _stream_event(
                "workflow", step="草稿历史处理", status="done",
                message=f"已归档 {result.get('archived', 0)} 条旧 AI 草稿",
            )
            yield _stream_event(
                "workflow", step="保存行动项草稿", status="done",
                message=f"已保存 {result['count']} 条草稿，等待人工核对",
            )
            yield _stream_event("complete", result=result)
        except GeneratorExit:
            return
        except ValidationError as exc:
            problems = []
            for error in exc.errors()[:8]:
                location = ".".join(str(part) for part in error.get("loc", []))
                problems.append(f"{location}：{error.get('msg') or '格式不正确'}")
            message = "；".join(problems) or "返回结构不符合约定"
            _event(ctx["name"], "AI提取失败", message, meeting_id=minutes_id)
            yield _stream_event("error", message=f"JSON 校验失败：{message}")
        except HTTPException as exc:
            message = str(exc.detail)
            _event(ctx["name"], "AI提取失败", message, meeting_id=minutes_id)
            yield _stream_event("error", message=message)
        except Exception as exc:
            logger.exception("行动项流式 AI 提取失败: %s", exc)
            message = str(exc) or "未知错误"
            _event(ctx["name"], "AI提取失败", message, meeting_id=minutes_id)
            yield _stream_event("error", message=message)
        finally:
            close = getattr(stream_response, "close", None)
            if callable(close):
                try:
                    close()
                except Exception:
                    pass

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson; charset=utf-8",
        headers={
            "Cache-Control": "no-cache, no-transform",
            "X-Accel-Buffering": "no",
        },
    )


class ActionEditRequest(BaseModel):
    current_user: str
    title: Optional[str] = None
    content: Optional[str] = None
    source_quote: Optional[str] = None
    responsible_department_id: Optional[str] = None
    responsible_person_id: Optional[str] = None
    responsible_department_ids: Optional[List[str]] = None
    responsible_person_ids: Optional[List[str]] = None
    collaborating_departments: Optional[List[str]] = None
    collaborating_people: Optional[List[str]] = None
    supervisor_id: Optional[str] = None
    required_completion_date: Optional[str] = None
    priority: Optional[Literal["高", "中", "低"]] = None


class DraftCreateRequest(ActionEditRequest):
    source_meeting_id: Optional[int] = None


class PublishRequest(BaseModel):
    current_user: str
    ids: List[int]


class MergeRequest(BaseModel):
    current_user: str
    ids: List[int]
    title: str = ""
    content: str = ""


class SplitPart(BaseModel):
    title: str
    content: str
    source_quote: str = ""


class SplitRequest(BaseModel):
    current_user: str
    parts: List[SplitPart]


class ReceiveRequest(BaseModel):
    current_user: str


class AssignRequest(BaseModel):
    current_user: str
    responsible_person_id: str


class ApprovalRequest(BaseModel):
    current_user: str
    result: Literal["通过", "退回整改", "要求补充材料"]
    opinion: str = ""


class ChangeRequest(BaseModel):
    current_user: str
    change_type: Literal["延期", "负责人变更", "责任科室变更", "内容调整", "取消"]
    after_content: dict = Field(default_factory=dict)
    reason: str


def _clear_resolved_uncertain(action: dict) -> list[str]:
    uncertain = set(_json(action.get("uncertain_fields")))
    checks = {
        "responsible_department": action.get("responsible_department_id"),
        "responsible_person": action.get("responsible_person_id"),
        "supervisor": action.get("supervisor_id"),
        "deadline": action.get("required_completion_date"),
    }
    for key, value in checks.items():
        if value:
            uncertain.discard(key)
    return sorted(uncertain)


def _create_reminder(
    action: dict, target: str, reminder_type: str, note: str, reminder_key: str,
) -> bool:
    target = str(target or "").strip()
    if not target:
        return False
    try:
        users = db.execute_query(
            "SELECT name FROM yggl "
            "WHERE name=%s AND COALESCE(zaizhi,0)=0 LIMIT 1",
            (target,),
        )
        if not users:
            logger.info(
                "跳过行动项提醒：未识别到在职提醒对象 action=%s target=%s",
                action.get("id"), target,
            )
            return False
        inserted = db.execute_insert(
            "INSERT IGNORE INTO action_reminder "
            "(action_item_id,reminder_type,reminder_target,reminder_channel,send_status,reminder_note,reminder_key) "
            "VALUES (%s,%s,%s,%s,'待发送',%s,%s)",
            (action["id"], reminder_type, target, "站内", note, reminder_key),
        )
        if not inserted:
            return False
    except Exception as exc:
        logger.warning(
            "行动项提醒创建失败，已跳过 action=%s target=%s: %s",
            action.get("id"), target, exc,
        )
        return False
    try:
        db.execute_update("UPDATE action_reminder SET send_status=%s WHERE id=%s", ("站内已送达", inserted))
        _event("系统", "提醒", f"{reminder_type}：{target}；{note}", action_id=action["id"])
    except Exception as exc:
        logger.warning(
            "行动项提醒状态记录失败，但不影响业务流程 action=%s target=%s: %s",
            action.get("id"), target, exc,
        )
    return True


def _create_responsible_reminders(
    action: dict, reminder_type: str, note: str, reminder_key: str,
) -> int:
    return sum(
        int(_create_reminder(
            action, target, reminder_type, note,
            f"{reminder_key}:{target}",
        ))
        for target in _responsible_people(action)
    )


@router.post("/drafts")
def create_draft(req: DraftCreateRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["action_creator"]:
        raise HTTPException(status_code=403, detail="仅部门领导、经理、副经理、综合室或系统管理员可新增行动项")
    minutes = []
    if req.source_meeting_id:
        minutes = db.execute_query(
            "SELECT minutes_number FROM meeting_minutes WHERE id=%s LIMIT 1",
            (req.source_meeting_id,),
        )
        if not minutes:
            raise HTTPException(status_code=404, detail="来源会议纪要不存在")
    if not (req.title or "").strip() or not (req.content or "").strip():
        raise HTTPException(status_code=400, detail="标题和内容不能为空")
    departments, people, jobs, supervisors, default_responsibles = _action_directory_maps()
    responsible_departments, responsible_people = _resolve_responsibility_lists(
        req.responsible_department_ids or req.responsible_department_id,
        req.responsible_person_ids or req.responsible_person_id,
        departments, people, jobs,
    )
    responsible_department, responsible_person, supervisor = _resolve_assignment_rules(
        responsible_departments[0] if responsible_departments else None,
        responsible_people[0] if responsible_people else None, req.supervisor_id,
        departments, people, jobs, supervisors, default_responsibles,
        action_text="\n".join((
            (req.title or "").strip(), (req.content or "").strip(),
            (req.source_quote or "").strip(),
        )),
    )
    supervisor = _resolve_scope_supervisor(
        responsible_departments, responsible_people, req.supervisor_id,
        jobs, supervisors, None,
        "\n".join((
            (req.title or "").strip(), (req.content or "").strip(),
            (req.source_quote or "").strip(),
        )),
    )
    action_id = db.execute_insert(
        "INSERT INTO action_item "
        "(title,content,source_meeting_id,minutes_number,source_quote,responsible_department_id,"
        "responsible_person_id,responsible_department_ids,responsible_person_ids,"
        "collaborating_departments,collaborating_people,supervisor_id,"
        "required_completion_date,priority,current_status,created_by,uncertain_fields) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'[]')",
        ((req.title or "").strip(), (req.content or "").strip(), req.source_meeting_id,
         (minutes[0].get("minutes_number") or "") if minutes else "",
         (req.source_quote or "").strip(),
         responsible_department, responsible_person,
         json.dumps(responsible_departments, ensure_ascii=False),
         json.dumps(responsible_people, ensure_ascii=False),
         json.dumps(req.collaborating_departments or [], ensure_ascii=False),
         json.dumps(req.collaborating_people or [], ensure_ascii=False),
         supervisor, req.required_completion_date or None,
         req.priority or "中", STATUS_DRAFT, ctx["name"]),
    )
    _event(ctx["name"], "人工新增", "人工新增行动项草稿", action_id=action_id,
           meeting_id=req.source_meeting_id)
    return {"success": True, "id": action_id}


@router.put("/{action_id}")
def edit_action(action_id: int, req: ActionEditRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    action = _action_row(action_id)
    _assert_action_visible(ctx, action)
    draft_edit = action["current_status"] in (STATUS_DRAFT, STATUS_PENDING_PUBLISH)
    published_adjust = action["current_status"] not in (
        STATUS_DRAFT, STATUS_PENDING_PUBLISH, STATUS_CANCELLED,
    )
    if draft_edit:
        if not (
            ctx["minutes_admin"]
            or (ctx["action_creator"] and action.get("created_by") == ctx["name"])
        ):
            raise HTTPException(status_code=403, detail="仅创建人、综合室或系统管理员可编辑未发布行动项")
    elif published_adjust:
        if not ctx["minutes_admin"]:
            raise HTTPException(status_code=403, detail="仅综合室或系统管理员可调整已发布行动项")
    else:
        raise HTTPException(status_code=400, detail="当前状态不可调整")
    before = _serialize_action(action)
    fields = []
    params: list[Any] = []
    mapping = {
        "title": req.title, "content": req.content, "source_quote": req.source_quote,
        "required_completion_date": req.required_completion_date,
        "priority": req.priority,
    }
    for column, value in mapping.items():
        if value is not None:
            fields.append(f"{column}=%s")
            params.append(value.strip() if isinstance(value, str) else value)
    for column, value in (
        ("collaborating_departments", req.collaborating_departments),
        ("collaborating_people", req.collaborating_people),
    ):
        if value is not None:
            fields.append(f"{column}=%s")
            params.append(json.dumps(value, ensure_ascii=False))

    assignment_fields = {
        "responsible_department_id", "responsible_person_id",
        "responsible_department_ids", "responsible_person_ids", "supervisor_id",
    }
    if assignment_fields.intersection(req.model_fields_set):
        departments, people, jobs, supervisors, default_responsibles = (
            _action_directory_maps()
        )
        department_values = (
            req.responsible_department_ids
            if "responsible_department_ids" in req.model_fields_set
            else (
                req.responsible_department_id
                if "responsible_department_id" in req.model_fields_set
                else _responsible_departments(action)
            )
        )
        person_values = (
            req.responsible_person_ids
            if "responsible_person_ids" in req.model_fields_set
            else (
                req.responsible_person_id
                if "responsible_person_id" in req.model_fields_set
                else _responsible_people(action)
            )
        )
        supervisor = (
            req.supervisor_id
            if "supervisor_id" in req.model_fields_set
            else action.get("supervisor_id")
        )
        # 批量改科室时请求中不会携带责任人，旧责任人不能反向覆盖新科室。
        if (
            (
                "responsible_department_id" in req.model_fields_set
                or "responsible_department_ids" in req.model_fields_set
            )
            and not {
                "responsible_person_id", "responsible_person_ids",
            }.intersection(req.model_fields_set)
        ):
            person_values = []
        responsible_departments, responsible_people = _resolve_responsibility_lists(
            department_values, person_values, departments, people, jobs,
        )
        action_text = "\n".join((
            (req.title if req.title is not None else action.get("title") or ""),
            (req.content if req.content is not None else action.get("content") or ""),
            (
                req.source_quote if req.source_quote is not None
                else action.get("source_quote") or ""
            ),
        ))
        department, person, supervisor = _resolve_assignment_rules(
            responsible_departments[0] if responsible_departments else None,
            responsible_people[0] if responsible_people else None,
            supervisor, departments, people, jobs, supervisors,
            default_responsibles,
            action_text=action_text,
        )
        supervisor = _resolve_scope_supervisor(
            responsible_departments, responsible_people, supervisor,
            jobs, supervisors, None, action_text,
        )
        for column, value in (
            ("responsible_department_id", department),
            ("responsible_person_id", person),
            (
                "responsible_department_ids",
                json.dumps(responsible_departments, ensure_ascii=False),
            ),
            (
                "responsible_person_ids",
                json.dumps(responsible_people, ensure_ascii=False),
            ),
            ("supervisor_id", supervisor),
        ):
            fields.append(f"{column}=%s")
            params.append(value)
    if not fields:
        return {"success": True, "message": "没有需要更新的字段"}
    params.append(action_id)
    db.execute_update(f"UPDATE action_item SET {','.join(fields)} WHERE id=%s", tuple(params))
    updated = _action_row(action_id)
    uncertain = _clear_resolved_uncertain(updated)
    next_status = STATUS_PENDING_PUBLISH if draft_edit else updated["current_status"]
    db.execute_update(
        "UPDATE action_item SET uncertain_fields=%s,current_status=%s WHERE id=%s",
        (json.dumps(uncertain, ensure_ascii=False), next_status, action_id),
    )
    _event(
        ctx["name"],
        "人工编辑" if draft_edit else "台账调整",
        "编辑行动项草稿" if draft_edit else "综合室在台账中调整已发布行动项",
        action_id=action_id,
        meeting_id=action.get("source_meeting_id"),
        data={"before": before, "after": _serialize_action(_action_row(action_id))},
    )
    return {"success": True, "message": "行动项已更新"}


@router.post("/merge")
def merge_actions(req: MergeRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可合并草稿")
    ids = list(dict.fromkeys(req.ids))
    if len(ids) < 2:
        raise HTTPException(status_code=400, detail="至少选择两条行动项")
    rows = db.execute_query(
        f"SELECT * FROM action_item WHERE id IN ({','.join(['%s'] * len(ids))}) ORDER BY id",
        tuple(ids),
    )
    if len(rows) != len(ids) or any(r["current_status"] not in (STATUS_DRAFT, STATUS_PENDING_PUBLISH) for r in rows):
        raise HTTPException(status_code=400, detail="只能合并同批次的未发布草稿")
    meeting_ids = {r.get("source_meeting_id") for r in rows}
    if len(meeting_ids) != 1:
        raise HTTPException(status_code=400, detail="只能合并同一会议纪要的草稿")
    target = rows[0]
    title = req.title.strip() or "；".join(r["title"] for r in rows)
    content = req.content.strip() or "\n".join(r["content"] for r in rows)
    quotes = "\n".join(r.get("source_quote") or "" for r in rows if r.get("source_quote"))
    responsible_departments = list(dict.fromkeys(
        department for row in rows for department in _responsible_departments(row)
    ))
    responsible_people = list(dict.fromkeys(
        person for row in rows for person in _responsible_people(row)
    ))
    db.execute_update(
        "UPDATE action_item SET title=%s,content=%s,source_quote=%s,"
        "responsible_department_id=%s,responsible_person_id=%s,"
        "responsible_department_ids=%s,responsible_person_ids=%s,"
        "current_status=%s WHERE id=%s",
        (
            title[:500], content, quotes,
            responsible_departments[0] if responsible_departments else None,
            responsible_people[0] if responsible_people else None,
            json.dumps(responsible_departments, ensure_ascii=False),
            json.dumps(responsible_people, ensure_ascii=False),
            STATUS_PENDING_PUBLISH, target["id"],
        ),
    )
    for row in rows[1:]:
        db.execute_update("UPDATE action_item SET current_status=%s WHERE id=%s", (STATUS_CANCELLED, row["id"]))
        _event(ctx["name"], "合并", f"已合并至行动项 {target['id']}", action_id=row["id"],
               meeting_id=row.get("source_meeting_id"))
    _event(ctx["name"], "合并", f"合并 {len(rows)} 条草稿", action_id=target["id"],
           meeting_id=target.get("source_meeting_id"), data={"mergedIds": ids})
    return {"success": True, "id": target["id"]}


@router.post("/{action_id}/split")
def split_action(action_id: int, req: SplitRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可拆分草稿")
    action = _action_row(action_id)
    if action["current_status"] not in (STATUS_DRAFT, STATUS_PENDING_PUBLISH):
        raise HTTPException(status_code=400, detail="只能拆分未发布草稿")
    if len(req.parts) < 2:
        raise HTTPException(status_code=400, detail="拆分后至少需要两条行动项")
    new_ids = []
    for part in req.parts:
        if not part.title.strip() or not part.content.strip():
            raise HTTPException(status_code=400, detail="拆分项标题和内容不能为空")
        new_id = db.execute_insert(
            "INSERT INTO action_item "
            "(title,content,source_type,source_meeting_id,minutes_number,source_quote,"
            "responsible_department_id,responsible_person_id,responsible_department_ids,"
            "responsible_person_ids,collaborating_departments,"
            "collaborating_people,supervisor_id,required_completion_date,priority,current_status,"
            "created_by,uncertain_fields,ai_batch_id) "
            "SELECT %s,%s,source_type,source_meeting_id,minutes_number,%s,responsible_department_id,"
            "responsible_person_id,responsible_department_ids,responsible_person_ids,"
            "collaborating_departments,collaborating_people,supervisor_id,"
            "required_completion_date,priority,%s,%s,uncertain_fields,ai_batch_id "
            "FROM action_item WHERE id=%s",
            (part.title.strip(), part.content.strip(), part.source_quote.strip(),
             STATUS_PENDING_PUBLISH, ctx["name"], action_id),
        )
        new_ids.append(new_id)
        _event(ctx["name"], "拆分", f"由行动项 {action_id} 拆分生成", action_id=new_id,
               meeting_id=action.get("source_meeting_id"))
    db.execute_update("UPDATE action_item SET current_status=%s WHERE id=%s", (STATUS_CANCELLED, action_id))
    _event(ctx["name"], "拆分", f"拆分为 {len(new_ids)} 条草稿", action_id=action_id,
           meeting_id=action.get("source_meeting_id"), data={"newIds": new_ids})
    return {"success": True, "ids": new_ids}


@router.post("/{action_id}/cancel-draft")
def cancel_action_draft(action_id: int, req: ReceiveRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可删除草稿")
    action = _action_row(action_id)
    if action["current_status"] not in (STATUS_DRAFT, STATUS_PENDING_PUBLISH):
        raise HTTPException(status_code=400, detail="只能删除未发布草稿")
    db.execute_update("UPDATE action_item SET current_status=%s WHERE id=%s", (STATUS_CANCELLED, action_id))
    _event(ctx["name"], "删除草稿", "人工确认页删除行动项草稿（保留审计记录）",
           action_id=action_id, meeting_id=action.get("source_meeting_id"))
    return {"success": True, "message": "草稿已删除并保留历史记录"}


class CancelPublishedRequest(BaseModel):
    current_user: str
    reason: str = ""


@router.post("/{action_id}/cancel")
def cancel_published_action(action_id: int, req: CancelPublishedRequest):
    """综合室/系统管理员在台账中取消已发布行动项（软删除，保留审计）。"""
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可删除已发布行动项")
    action = _action_row(action_id)
    status = action.get("current_status") or ""
    if status in (STATUS_DRAFT, STATUS_PENDING_PUBLISH):
        raise HTTPException(status_code=400, detail="未发布草稿请使用删除草稿接口")
    if status == STATUS_CANCELLED:
        return {"success": True, "message": "行动项已是取消状态"}
    reason = (req.reason or "").strip() or "台账管理删除"
    db.execute_update(
        "UPDATE action_item SET current_status=%s,last_progress_at=NOW() WHERE id=%s",
        (STATUS_CANCELLED, action_id),
    )
    _event(
        ctx["name"], "台账删除", f"已发布行动项被取消：{reason}",
        action_id=action_id, meeting_id=action.get("source_meeting_id"),
        data={"reason": reason, "previousStatus": status},
    )
    return {"success": True, "message": "行动项已删除（取消）并保留历史记录"}


@router.post("/publish")
def publish_actions(req: PublishRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    if not ctx["minutes_admin"]:
        raise HTTPException(status_code=403, detail="仅综合室或系统管理员可发布行动项")
    ids = list(dict.fromkeys(req.ids))
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要发布的行动项")
    rows = db.execute_query(
        f"SELECT * FROM action_item WHERE id IN ({','.join(['%s'] * len(ids))})", tuple(ids)
    )
    if len(rows) != len(ids):
        raise HTTPException(status_code=404, detail="部分行动项不存在")
    departments, people, jobs, supervisors, default_responsibles = _action_directory_maps()
    supervision_configs = _load_supervision_configs()
    missing = []
    for row in rows:
        action_text = "\n".join((
            row.get("title") or "", row.get("content") or "",
            row.get("source_quote") or "",
        ))
        responsible_departments, responsible_people = _resolve_responsibility_lists(
            _responsible_departments(row), _responsible_people(row),
            departments, people, jobs,
        )
        department, person, supervisor = _resolve_assignment_rules(
            responsible_departments[0] if responsible_departments else None,
            responsible_people[0] if responsible_people else None,
            row.get("supervisor_id"), departments, people, jobs, supervisors,
            default_responsibles, supervision_configs, action_text,
        )
        supervisor = _resolve_scope_supervisor(
            responsible_departments, responsible_people, row.get("supervisor_id"),
            jobs, supervisors, supervision_configs, action_text,
        )
        if (
            department != row.get("responsible_department_id")
            or person != row.get("responsible_person_id")
            or responsible_departments != _responsible_departments(row)
            or responsible_people != _responsible_people(row)
            or supervisor != row.get("supervisor_id")
        ):
            db.execute_update(
                "UPDATE action_item SET responsible_department_id=%s,"
                "responsible_person_id=%s,responsible_department_ids=%s,"
                "responsible_person_ids=%s,supervisor_id=%s WHERE id=%s",
                (
                    department, person,
                    json.dumps(responsible_departments, ensure_ascii=False),
                    json.dumps(responsible_people, ensure_ascii=False),
                    supervisor, row["id"],
                ),
            )
            row["responsible_department_id"] = department
            row["responsible_person_id"] = person
            row["responsible_department_ids"] = responsible_departments
            row["responsible_person_ids"] = responsible_people
            row["supervisor_id"] = supervisor
        if row["current_status"] not in (STATUS_DRAFT, STATUS_PENDING_PUBLISH):
            missing.append(f"{row['id']} 状态不可发布")
        for label, key in (
            ("责任科室", "responsible_department_ids"), ("责任人", "responsible_person_ids"),
            ("主管领导", "supervisor_id"), ("完成时间", "required_completion_date"),
        ):
            if not row.get(key):
                missing.append(f"{row['id']} 缺少{label}")
        if not (row.get("title") or "").strip() or not (row.get("content") or "").strip():
            missing.append(f"{row['id']} 缺少标题或内容")
        supervisor = (row.get("supervisor_id") or "").strip()
        invalid_responsible_depts = [
            value for value in responsible_departments if value not in departments
        ]
        invalid_responsible_people = [
            value for value in responsible_people if value not in people
        ]
        if invalid_responsible_depts:
            missing.append(f"{row['id']} 责任科室无效")
        if invalid_responsible_people:
            missing.append(f"{row['id']} 责任人无效或已离职")
        elif any(people[value] not in responsible_departments for value in responsible_people):
            missing.append(f"{row['id']} 责任人不属于责任科室")
        if supervisor and supervisor not in supervisors:
            missing.append(f"{row['id']} 主管领导无效、已离职或不具备主管职务")
        invalid_depts = [x for x in _json(row.get("collaborating_departments")) if x not in departments]
        invalid_people = [x for x in _json(row.get("collaborating_people")) if x not in people]
        if invalid_depts:
            missing.append(f"{row['id']} 存在无效协同科室")
        if invalid_people:
            missing.append(f"{row['id']} 存在无效协同责任人")
    if missing:
        raise HTTPException(status_code=400, detail="；".join(missing[:20]))
    meeting_ids = set()
    for row in rows:
        number = row.get("action_number") or f"XD-{datetime.now():%Y}-{int(row['id']):05d}"
        db.execute_update(
            "UPDATE action_item SET action_number=%s,current_status=%s,published_at=NOW(),"
            "uncertain_fields='[]' WHERE id=%s",
            (number, STATUS_PENDING_RECEIVE, row["id"]),
        )
        meeting_ids.add(row.get("source_meeting_id"))
        updated = _action_row(row["id"])
        _ensure_department_executions(updated, people)
        _event(ctx["name"], "发布", "行动项经人工确认后发布", action_id=row["id"],
               meeting_id=row.get("source_meeting_id"))
        for target in _publish_receive_targets(updated, people, jobs):
            _create_reminder(
                updated, target, "发布", "您有一条新的行动项待接收",
                f"publish:{row['id']}:{target}",
            )
    for meeting_id in meeting_ids:
        if meeting_id:
            remain = db.execute_scalar(
                "SELECT COUNT(*) FROM action_item WHERE source_meeting_id=%s "
                "AND current_status IN ('草稿','待发布')", (meeting_id,)
            ) or 0
            if not remain:
                db.execute_update("UPDATE meeting_minutes SET status='已发布' WHERE id=%s", (meeting_id,))
    return {"success": True, "count": len(rows), "message": f"已发布 {len(rows)} 条行动项"}


def _build_action_list(
    ctx: dict, keyword: str = "", meeting_id: Optional[int] = None,
    minutes_number: str = "", department: str = "", responsible_person: str = "",
    supervisor: str = "", status: str = "", deadline_from: str = "", deadline_to: str = "",
    overdue: Optional[bool] = None, mine: bool = False, metric: str = "",
    sort_by: str = "updated_at", sort_order: str = "desc",
    page: int = 1, page_size: int = 20,
) -> dict:
    visible_sql, params = _visible_action_sql(ctx, "a")
    where = [visible_sql]
    if not ctx["minutes_admin"]:
        if ctx["action_creator"]:
            where.append(
                "(a.current_status NOT IN ('草稿','待发布') OR a.created_by=%s)"
            )
            params.append(ctx["name"])
        else:
            where.append("a.current_status NOT IN ('草稿','待发布')")
    if mine:
        where.append(
            "("
            "a.responsible_person_id=%s "
            "OR JSON_CONTAINS(COALESCE(a.responsible_person_ids,'[]'),JSON_QUOTE(%s)) "
            "OR JSON_CONTAINS(COALESCE(a.collaborating_people,'[]'),JSON_QUOTE(%s)) "
            "OR ("
            "  a.current_status='待接收' AND %s<>'' AND ("
            "    a.responsible_department_id=%s "
            "    OR JSON_CONTAINS(COALESCE(a.responsible_department_ids,'[]'),JSON_QUOTE(%s))"
            "  )"
            ")"
            ")"
        )
        params.extend([
            ctx["name"], ctx["name"], ctx["name"],
            ctx["dept"] if _job_director_priority(ctx.get("jb") or "") is not None else "",
            ctx["dept"] or "",
            ctx["dept"] or "",
        ])
    if keyword.strip():
        where.append("CONCAT_WS('',a.action_number,a.title,a.content,a.source_quote) LIKE %s")
        params.append(f"%{keyword.strip()}%")
    for value, sql in (
        (meeting_id, "a.source_meeting_id=%s"), (minutes_number.strip(), "a.minutes_number=%s"),
        (
            department.strip(),
            "(a.responsible_department_id=%s OR "
            "JSON_CONTAINS(COALESCE(a.responsible_department_ids,'[]'),JSON_QUOTE(%s)))",
        ),
        (
            responsible_person.strip(),
            "(a.responsible_person_id=%s OR "
            "JSON_CONTAINS(COALESCE(a.responsible_person_ids,'[]'),JSON_QUOTE(%s)))",
        ),
        (supervisor.strip(), "a.supervisor_id=%s"), (status.strip(), "a.current_status=%s"),
        (deadline_from, "a.required_completion_date>=%s"), (deadline_to, "a.required_completion_date<=%s"),
    ):
        if value:
            where.append(sql)
            params.append(value)
            if "JSON_CONTAINS" in sql:
                params.append(value)
    if not status.strip():
        where.append("a.current_status<>'已取消'")
    if mine and status.strip() == STATUS_PENDING_RECEIVE:
        where.append(
            "EXISTS (SELECT 1 FROM action_department_execution ade "
            "WHERE ade.action_item_id=a.id AND ade.execution_status='待接收' AND ("
            "ade.responsible_person=%s OR (%s<>'' AND ade.department=%s)"
            "))"
        )
        director_department = (
            ctx["dept"] if _job_director_priority(ctx.get("jb") or "") is not None else ""
        )
        params.extend([ctx["name"], director_department, director_department])
    if overdue is True:
        where.append("a.required_completion_date<CURDATE() AND a.current_status NOT IN ('已完成','已取消')")
    elif overdue is False:
        where.append("(a.required_completion_date>=CURDATE() OR a.required_completion_date IS NULL)")
    metric_sql = {
        "completed": "a.current_status='已完成'",
        "active": "a.current_status IN ('待接收','进行中','退回整改')",
        "overdue": "a.required_completion_date<CURDATE() AND a.current_status NOT IN ('已完成','已取消')",
        "dueSoon": (
            "a.required_completion_date BETWEEN CURDATE() AND DATE_ADD(CURDATE(),INTERVAL 7 DAY) "
            "AND a.current_status NOT IN ('已完成','已取消')"
        ),
        "pendingApproval": "a.current_status='待完工审批'",
        "stale": (
            "COALESCE(a.last_progress_at,a.published_at)<DATE_SUB(NOW(),INTERVAL 7 DAY) "
            "AND a.current_status NOT IN ('已完成','已取消')"
        ),
        "newThisWeek": (
            "a.published_at>=DATE_SUB(CURDATE(),INTERVAL WEEKDAY(CURDATE()) DAY)"
        ),
        "completedThisWeek": (
            "a.current_status='已完成' AND "
            "a.actual_completion_date>=DATE_SUB(CURDATE(),INTERVAL WEEKDAY(CURDATE()) DAY)"
        ),
    }
    if metric in metric_sql:
        where.append(metric_sql[metric])
    clause = " AND ".join(where)
    total = db.execute_scalar(f"SELECT COUNT(*) FROM action_item a WHERE {clause}", tuple(params)) or 0
    allowed_sort = {
        "updated_at": "a.updated_at", "deadline": "a.required_completion_date",
        "progress": "a.current_progress", "created_at": "a.created_at",
        "priority": "FIELD(a.priority,'高','中','低')",
    }
    order_column = allowed_sort.get(sort_by, "a.updated_at")
    order = "ASC" if sort_order.lower() == "asc" else "DESC"
    rows = db.execute_query(
        f"SELECT a.*,m.meeting_name,m.meeting_date FROM action_item a "
        f"LEFT JOIN meeting_minutes m ON m.id=a.source_meeting_id WHERE {clause} "
        f"ORDER BY {order_column} {order},a.id DESC LIMIT %s OFFSET %s",
        tuple(params + [page_size, (page - 1) * page_size]),
    )
    items = [_serialize_action(r) for r in rows]
    if mine:
        for item in items:
            execution_rows = db.execute_query(
                "SELECT execution_status FROM action_department_execution "
                "WHERE action_item_id=%s AND (responsible_person=%s OR department=%s) "
                "ORDER BY CASE WHEN responsible_person=%s THEN 0 ELSE 1 END LIMIT 1",
                (item["id"], ctx["name"], ctx.get("dept") or "", ctx["name"]),
            )
            item["my_execution_status"] = (
                execution_rows[0].get("execution_status") if execution_rows else ""
            )
    return {
        "items": items, "total": int(total),
        "page": page, "pageSize": page_size,
    }


@router.get("")
def list_actions(
    current_user: str = Query(...), keyword: str = Query(""),
    meeting_id: Optional[int] = Query(None), minutes_number: str = Query(""),
    department: str = Query(""), responsible_person: str = Query(""),
    supervisor: str = Query(""), status: str = Query(""),
    deadline_from: str = Query(""), deadline_to: str = Query(""),
    overdue: Optional[bool] = Query(None), mine: bool = Query(False),
    metric: str = Query(""),
    sort_by: str = Query("updated_at"), sort_order: str = Query("desc"),
    page: int = Query(1, ge=1), page_size: int = Query(20, ge=1, le=200),
):
    ensure_tables()
    ctx = _role_context(current_user)
    if mine and status.strip() == STATUS_PENDING_RECEIVE:
        pending_rows = db.execute_query(
            "SELECT * FROM action_item WHERE current_status=%s",
            (STATUS_PENDING_RECEIVE,),
        )
        if pending_rows:
            _departments, people, _jobs, _supervisors, _defaults = _action_directory_maps()
            for pending_action in pending_rows:
                _ensure_department_executions(pending_action, people)
    data = _build_action_list(
        ctx, keyword, meeting_id, minutes_number, department, responsible_person,
        supervisor, status, deadline_from, deadline_to, overdue, mine,
        metric, sort_by, sort_order, page, page_size,
    )
    return {"success": True, **data}


@router.get("/dashboard")
def dashboard(current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    visible, params = _visible_action_sql(ctx, "a")
    rows = db.execute_query(
        f"SELECT a.* FROM action_item a WHERE {visible} AND a.current_status<>'已取消'", tuple(params)
    )
    items = [_serialize_action(r) for r in rows]
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    total = len(items)
    completed = [x for x in items if x["current_status"] == STATUS_COMPLETED]
    def has_tag(item: dict, tag: str) -> bool:
        return tag in item["risk_tags"]
    summary = {
        "total": total,
        "completionRate": round(len(completed) / total * 100, 1) if total else 0,
        "inProgress": sum(x["current_status"] in (STATUS_PENDING_RECEIVE, STATUS_IN_PROGRESS, STATUS_RETURNED) for x in items),
        "overdue": sum(has_tag(x, "已逾期") for x in items),
        "dueSoon": sum(has_tag(x, "临期") for x in items),
        "pendingApproval": sum(x["current_status"] == STATUS_PENDING_COMPLETION for x in items),
        "stale": sum(has_tag(x, "长期未更新") for x in items),
        "newThisWeek": sum((x.get("published_at") or "")[:10] >= week_start.isoformat() for x in items),
        "completedThisWeek": sum(
            x["current_status"] == STATUS_COMPLETED
            and (x.get("actual_completion_date") or "") >= week_start.isoformat() for x in items
        ),
    }
    depts: dict[str, dict] = {}
    for item in items:
        names = item.get("responsible_department_ids") or ["待确认"]
        for name in names:
            bucket = depts.setdefault(
                name, {"department": name, "total": 0, "completed": 0, "overdue": 0},
            )
            bucket["total"] += 1
            bucket["completed"] += item["current_status"] == STATUS_COMPLETED
            bucket["overdue"] += has_tag(item, "已逾期")
    department_stats = []
    for bucket in depts.values():
        bucket["completionRate"] = round(bucket["completed"] / bucket["total"] * 100, 1)
        bucket["overdueRate"] = round(bucket["overdue"] / bucket["total"] * 100, 1)
        department_stats.append(bucket)
    department_stats.sort(key=lambda x: (-x["total"], x["department"]))
    return {"success": True, "summary": summary, "departments": department_stats}


@router.get("/export")
def export_actions(
    current_user: str = Query(...), keyword: str = Query(""),
    department: str = Query(""), status: str = Query(""),
):
    ensure_tables()
    ctx = _role_context(current_user)
    if not (ctx["minutes_admin"] or ctx["department_leader"]):
        raise HTTPException(status_code=403, detail="仅综合室或部门领导可导出台账")
    data = _build_action_list(
        ctx, keyword=keyword, department=department, status=status, page=1, page_size=10000
    )
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([
        "行动项编号", "标题", "内容", "来源会议", "纪要编号", "责任科室", "责任人",
        "主管领导", "要求完成时间", "实际完成时间", "优先级", "进度", "状态", "风险标签",
    ])
    for item in data["items"]:
        writer.writerow([
            item.get("action_number"), item.get("title"), item.get("content"),
            item.get("meeting_name"), item.get("minutes_number"),
            "、".join(item.get("responsible_department_ids") or []),
            "、".join(item.get("responsible_person_ids") or []),
            item.get("supervisor_id"), item.get("required_completion_date"),
            item.get("actual_completion_date"), item.get("priority"),
            item.get("current_progress"), item.get("current_status"),
            "、".join(item.get("risk_tags") or []),
        ])
    payload = ("\ufeff" + output.getvalue()).encode("utf-8")
    return StreamingResponse(
        io.BytesIO(payload), media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="action_items_{date.today():%Y%m%d}.csv"'},
    )


@router.post("/{action_id}/receive")
def receive_action(action_id: int, req: ReceiveRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    action = _action_row(action_id)
    if not _can_receive_action(ctx, action):
        raise HTTPException(status_code=403, detail="仅行动项责任人或责任科室主任/副主任可接收")
    if action["current_status"] != STATUS_PENDING_RECEIVE:
        raise HTTPException(status_code=400, detail="当前状态不可接收")
    departments, people, jobs, _supervisors, _defaults = _action_directory_maps()
    receiver_department = str(people.get(ctx["name"]) or ctx.get("dept") or "").strip()
    executions = _ensure_department_executions(action, people)
    execution = next(
        (row for row in executions if row.get("department") == receiver_department),
        None,
    )
    if not execution:
        raise HTTPException(status_code=403, detail="当前用户不属于该行动项的责任科室")
    if execution.get("execution_status") != STATUS_PENDING_RECEIVE:
        raise HTTPException(status_code=400, detail="本科室已接收该行动项，无需重复接收")
    db.execute_update(
        "UPDATE action_department_execution SET responsible_person=%s,"
        "execution_status=%s,received_by=%s,received_at=NOW() "
        "WHERE action_item_id=%s AND department=%s",
        (ctx["name"], STATUS_IN_PROGRESS, ctx["name"], action_id, receiver_department),
    )
    responsible_people = [
        name for name in _responsible_people(action)
        if people.get(name) != receiver_department
    ] + [ctx["name"]]
    responsible_people = list(dict.fromkeys(responsible_people))
    db.execute_update(
        "UPDATE action_item SET responsible_person_id=%s,responsible_person_ids=%s,"
        "last_progress_at=COALESCE(last_progress_at,NOW()) WHERE id=%s",
        (
            responsible_people[0] if responsible_people else ctx["name"],
            json.dumps(responsible_people, ensure_ascii=False), action_id,
        ),
    )
    rollup = _refresh_action_rollup(action_id)
    if receiver_department and receiver_department in _responsible_departments(action):
        _close_department_receiver_reminders(
            action_id, receiver_department, people, jobs,
        )
    db.execute_update(
        "UPDATE action_reminder SET read_at=COALESCE(read_at,NOW()) "
        "WHERE action_item_id=%s AND reminder_target=%s",
        (action_id, ctx["name"]),
    )
    role_note = "责任人" if _is_responsible_person(action, ctx["name"]) else "责任科室主任/副主任"
    _event(ctx["name"], "接收", f"{receiver_department}：{role_note}已接收行动项", action_id=action_id,
           meeting_id=action.get("source_meeting_id"))
    return {
        "success": True,
        "message": (
            "本科室已接收，所有责任科室均已接收，行动项进入进行中"
            if rollup.get("current_status") == STATUS_IN_PROGRESS
            else "本科室已接收，等待其他责任科室接收"
        ),
    }


@router.post("/{action_id}/assign")
def assign_responsible_person(action_id: int, req: AssignRequest):
    """责任科室负责人在任务接收前完成初次分配，并保留完整变更和审批记录。"""
    ensure_tables()
    ctx = _role_context(req.current_user)
    action = _action_row(action_id)
    _assert_action_visible(ctx, action)
    allowed_departments = _assignment_departments(ctx, action)
    if not allowed_departments:
        raise HTTPException(status_code=403, detail="仅责任科室负责人、综合室或部门领导可分配")
    if action["current_status"] != STATUS_PENDING_RECEIVE:
        raise HTTPException(status_code=400, detail="仅待接收行动项可直接分配；执行中的调整请提交变更申请")
    person = (req.responsible_person_id or "").strip()
    rows = db.execute_query(
        "SELECT name,lsys,jb FROM yggl WHERE name=%s AND COALESCE(zaizhi,0)=0 LIMIT 1",
        (person,),
    )
    if not rows:
        raise HTTPException(status_code=400, detail="所选责任人不存在或已离职")
    department = (rows[0].get("lsys") or "").strip()
    selected_job = (rows[0].get("jb") or "").strip()
    if department not in allowed_departments:
        if len(allowed_departments) == 1 and allowed_departments[0] == ctx.get("dept"):
            raise HTTPException(status_code=403, detail="科室主任或副主任只能分配本科室在职人员")
        raise HTTPException(status_code=403, detail="只能在该行动项的责任科室范围内分配在职人员")
    if ctx["department_leader"] and _job_director_priority(selected_job) is None:
        raise HTTPException(
            status_code=403,
            detail="部门领导分配行动项时，只能选择责任科室的主任、主任责或副主任",
        )
    executions = _ensure_department_executions(action)
    department_execution = next(
        (row for row in executions if row.get("department") == department),
        None,
    )
    if department_execution and department_execution.get("execution_status") != STATUS_PENDING_RECEIVE:
        raise HTTPException(status_code=400, detail="该科室已接收行动项，不能再执行初次分配")
    departments, people, jobs, supervisors, default_responsibles = _action_directory_maps()
    department, person, supervisor = _resolve_assignment_rules(
        department, person, action.get("supervisor_id"), departments, people, jobs,
        supervisors, default_responsibles,
    )
    old_departments = _responsible_departments(action)
    old_people = _responsible_people(action)
    new_departments = list(old_departments)
    if department and department not in new_departments:
        new_departments.append(department)
    new_people = [
        name for name in old_people
        if people.get(name) != department
    ]
    if person:
        new_people.append(person)
    new_people = list(dict.fromkeys(new_people))
    supervisor = _resolve_scope_supervisor(
        new_departments, new_people, supervisor, jobs, supervisors, None,
        "\n".join((
            action.get("title") or "", action.get("content") or "",
            action.get("source_quote") or "",
        )),
    )
    primary_department = new_departments[0] if new_departments else department
    primary_person = new_people[0] if new_people else person
    old_person = (action.get("responsible_person_id") or "").strip()
    old_department = (action.get("responsible_department_id") or "").strip()
    old_supervisor = (action.get("supervisor_id") or "").strip()
    before = {
        "responsible_department_id": old_department or None,
        "responsible_person_id": old_person or None,
        "responsible_department_ids": old_departments,
        "responsible_person_ids": old_people,
        "supervisor_id": old_supervisor or None,
    }
    after = {
        "responsible_department_id": primary_department,
        "responsible_person_id": primary_person,
        "responsible_department_ids": new_departments,
        "responsible_person_ids": new_people,
        "supervisor_id": supervisor,
    }
    change_id = db.execute_insert(
        "INSERT INTO action_change "
        "(action_item_id,change_type,before_content,after_content,change_reason,applicant,"
        "approval_status,approver,approval_opinion,approved_at) "
        "VALUES (%s,'负责人变更',%s,%s,'责任科室负责人初次分配',%s,'已通过',%s,"
        "'待接收阶段授权分配',NOW())",
        (
            action_id, json.dumps(before, ensure_ascii=False),
            json.dumps(after, ensure_ascii=False), ctx["name"], ctx["name"],
        ),
    )
    db.execute_update(
        "INSERT INTO action_approval "
        "(business_type,business_id,action_item_id,approver,approval_result,approval_opinion) "
        "VALUES ('责任人分配',%s,%s,%s,'通过','待接收阶段授权分配')",
        (change_id, action_id, ctx["name"]),
    )
    db.execute_update(
        "UPDATE action_item SET responsible_department_id=%s,responsible_person_id=%s,"
        "responsible_department_ids=%s,responsible_person_ids=%s,supervisor_id=%s WHERE id=%s",
        (
            primary_department, primary_person,
            json.dumps(new_departments, ensure_ascii=False),
            json.dumps(new_people, ensure_ascii=False),
            supervisor, action_id,
        ),
    )
    _ensure_department_executions(action, people)
    db.execute_update(
        "UPDATE action_department_execution SET responsible_person=%s,"
        "execution_status=%s,received_by=NULL,received_at=NULL "
        "WHERE action_item_id=%s AND department=%s",
        (person, STATUS_PENDING_RECEIVE, action_id, department),
    )
    _refresh_action_rollup(action_id)
    _close_department_receiver_reminders(
        action_id, department, people, jobs,
    )
    _event(
        ctx["name"], "责任人分配",
        f"责任人由 {old_person or '待确认'} 调整为 {person}，责任科室同步为 {department}",
        action_id=action_id, meeting_id=action.get("source_meeting_id"),
        data={"changeId": change_id, "before": before, "after": after},
    )
    updated = _action_row(action_id)
    _create_reminder(
        updated, person, "责任人分配", "您有一条新的行动项待接收",
        f"assign:{action_id}:{person}:{change_id}",
    )
    if old_person:
        _create_reminder(
            updated, old_person, "负责人变更", f"行动项责任人已变更为 {person}",
            f"assign-old:{action_id}:{old_person}:{change_id}",
        )
    return {"success": True, "message": "责任人已分配并记录变更历史"}


@router.post("/{action_id}/progress")
async def add_progress(
    action_id: int, current_user: str = Form(...),
    progress_percent: int = Form(...), current_progress: str = Form(...),
    completed_work: str = Form(""), existing_problems: str = Form(""),
    next_plan: str = Form(""), expected_completion_date: str = Form(""),
    delay_risk: bool = Form(False), files: List[UploadFile] = File(default=[]),
):
    ensure_tables()
    ctx = _role_context(current_user)
    action = _action_row(action_id)
    _departments, people, _jobs, _supervisors, _defaults = _action_directory_maps()
    reporter_department = str(people.get(ctx["name"]) or ctx.get("dept") or "").strip()
    executions = _ensure_department_executions(action, people)
    execution = next(
        (
            row for row in executions
            if row.get("department") == reporter_department
            and row.get("responsible_person") == ctx["name"]
        ),
        None,
    )
    if not execution:
        raise HTTPException(status_code=403, detail="仅本科室已接收的实际责任人可填报进展")
    if execution.get("execution_status") not in (STATUS_IN_PROGRESS, STATUS_RETURNED):
        raise HTTPException(status_code=400, detail="本科室当前状态不可填报进展")
    if not 0 <= progress_percent <= 100:
        raise HTTPException(status_code=400, detail="进度必须为 0-100")
    if not current_progress.strip():
        raise HTTPException(status_code=400, detail="本期进展不能为空")
    progress_id = db.execute_insert(
        "INSERT INTO action_progress "
        "(action_item_id,responsible_department,progress_percent,current_progress,completed_work,existing_problems,"
        "next_plan,expected_completion_date,delay_risk,attachments,reporter) "
        "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,'[]',%s)",
        (action_id, reporter_department, progress_percent, current_progress.strip(), completed_work.strip(),
         existing_problems.strip(), next_plan.strip(), expected_completion_date or None,
         1 if delay_risk else 0, ctx["name"]),
    )
    attachments = []
    for file in files:
        attachments.append(await _save_upload(
            file, EVIDENCE_DIR, "progress", progress_id, ctx["name"],
            action_id=action_id, allowed=ALLOWED_EVIDENCE_EXT,
        ))
    db.execute_update(
        "UPDATE action_progress SET attachments=%s WHERE id=%s",
        (json.dumps(attachments, ensure_ascii=False), progress_id),
    )
    db.execute_update(
        "UPDATE action_department_execution SET progress_percent=%s,execution_status=%s "
        "WHERE action_item_id=%s AND department=%s",
        (progress_percent, STATUS_IN_PROGRESS, action_id, reporter_department),
    )
    db.execute_update(
        "UPDATE action_item SET risk_status=%s,last_progress_at=NOW() WHERE id=%s",
        ("存在延期风险" if delay_risk else "", action_id),
    )
    rollup = _refresh_action_rollup(action_id)
    _event(ctx["name"], "进展填报", f"{reporter_department}填报进度 {progress_percent}%：{current_progress.strip()}",
           action_id=action_id, meeting_id=action.get("source_meeting_id"),
           data={"progressId": progress_id, "attachments": attachments})
    return {
        "success": True, "id": progress_id,
        "message": f"本科室进展已保存，行动项总进度为 {rollup.get('current_progress', 0)}%",
    }


@router.post("/{action_id}/completion")
async def apply_completion(
    action_id: int, current_user: str = Form(...),
    completion_description: str = Form(...), actual_completion_date: str = Form(...),
    completion_results: str = Form(""), remaining_issues: str = Form(""),
    files: List[UploadFile] = File(default=[]),
):
    ensure_tables()
    ctx = _role_context(current_user)
    action = _action_row(action_id)
    _departments, people, _jobs, _supervisors, _defaults = _action_directory_maps()
    applicant_department = str(people.get(ctx["name"]) or ctx.get("dept") or "").strip()
    executions = _ensure_department_executions(action, people)
    execution = next(
        (
            row for row in executions
            if row.get("department") == applicant_department
            and row.get("responsible_person") == ctx["name"]
        ),
        None,
    )
    if not execution:
        raise HTTPException(status_code=403, detail="仅本科室已接收的实际责任人可申请完工")
    if execution.get("execution_status") not in (STATUS_IN_PROGRESS, STATUS_RETURNED):
        raise HTTPException(status_code=400, detail="本科室当前状态不可申请完工")
    if not completion_description.strip() or not actual_completion_date:
        raise HTTPException(status_code=400, detail="完成情况和实际完成时间不能为空")
    application_id = db.execute_insert(
        "INSERT INTO completion_application "
        "(action_item_id,responsible_department,completion_description,actual_completion_date,completion_results,"
        "evidence_materials,remaining_issues,applicant) VALUES (%s,%s,%s,%s,%s,'[]',%s,%s)",
        (action_id, applicant_department, completion_description.strip(), actual_completion_date,
         completion_results.strip(), remaining_issues.strip(), ctx["name"]),
    )
    attachments = []
    for file in files:
        attachments.append(await _save_upload(
            file, EVIDENCE_DIR, "completion", application_id, ctx["name"],
            action_id=action_id, allowed=ALLOWED_EVIDENCE_EXT,
        ))
    db.execute_update(
        "UPDATE completion_application SET evidence_materials=%s WHERE id=%s",
        (json.dumps(attachments, ensure_ascii=False), application_id),
    )
    db.execute_update(
        "UPDATE action_department_execution SET execution_status=%s,progress_percent=100 "
        "WHERE action_item_id=%s AND department=%s",
        (STATUS_PENDING_COMPLETION, action_id, applicant_department),
    )
    rollup = _refresh_action_rollup(action_id)
    _event(ctx["name"], "完工申请", f"{applicant_department}责任人提交完工申请", action_id=action_id,
           meeting_id=action.get("source_meeting_id"),
           data={"applicationId": application_id, "attachments": attachments})
    updated = _action_row(action_id)
    _create_reminder(
        updated, updated.get("supervisor_id") or "", "完工审批",
        "您有一条行动项完工申请待审批",
        f"completion:{application_id}:{updated.get('supervisor_id')}",
    )
    return {
        "success": True, "id": application_id,
        "message": f"本科室完工申请已提交，行动项总进度为 {rollup.get('current_progress', 0)}%",
    }


@router.post("/completions/{application_id}/approve")
def approve_completion(application_id: int, req: ApprovalRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    rows = db.execute_query(
        "SELECT c.*,a.supervisor_id,a.source_meeting_id,a.title,a.required_completion_date,"
        "a.current_status "
        "FROM completion_application c JOIN action_item a ON a.id=c.action_item_id "
        "WHERE c.id=%s LIMIT 1", (application_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="完工申请不存在")
    item = rows[0]
    _departments, people, _jobs, _supervisors, _defaults = _action_directory_maps()
    application_department = (
        (item.get("responsible_department") or "").strip()
        or (people.get(item.get("applicant") or "") or "").strip()
        or (_responsible_departments(item)[0] if _responsible_departments(item) else "")
    )
    if application_department and not item.get("responsible_department"):
        db.execute_update(
            "UPDATE completion_application SET responsible_department=%s WHERE id=%s",
            (application_department, application_id),
        )
    if not (
        ctx["admin"] or ctx["department_leader"]
        or (item.get("supervisor_id") or "") == ctx["name"]
    ):
        raise HTTPException(status_code=403, detail="仅主管领导可审批该完工申请")
    if item["approval_status"] != "待审批":
        raise HTTPException(status_code=400, detail="该申请已审批")
    execution_rows = db.execute_query(
        "SELECT * FROM action_department_execution "
        "WHERE action_item_id=%s AND department=%s LIMIT 1",
        (item["action_item_id"], application_department),
    )
    if not execution_rows or execution_rows[0].get("execution_status") != STATUS_PENDING_COMPLETION:
        raise HTTPException(status_code=409, detail="该科室执行状态已变化，请刷新后重新处理")
    approved = req.result == "通过"
    app_status = "已通过" if approved else req.result
    db.execute_update(
        "UPDATE completion_application SET approval_status=%s WHERE id=%s",
        (app_status, application_id),
    )
    db.execute_update(
        "INSERT INTO action_approval "
        "(business_type,business_id,action_item_id,approver,approval_result,approval_opinion) "
        "VALUES ('完工申请',%s,%s,%s,%s,%s)",
        (application_id, item["action_item_id"], ctx["name"], req.result, req.opinion.strip()),
    )
    if approved:
        db.execute_update(
            "UPDATE action_department_execution SET execution_status=%s,progress_percent=100,"
            "completed_at=NOW() WHERE action_item_id=%s AND department=%s",
            (
                STATUS_COMPLETED, item["action_item_id"],
                application_department,
            ),
        )
    else:
        db.execute_update(
            "UPDATE action_department_execution SET execution_status=%s "
            "WHERE action_item_id=%s AND department=%s",
            (
                STATUS_RETURNED, item["action_item_id"],
                application_department,
            ),
        )
    rollup = _refresh_action_rollup(item["action_item_id"])
    _event(ctx["name"], "完工审批", f"{application_department or '责任科室'}"
           f"{req.result}：{req.opinion.strip()}",
           action_id=item["action_item_id"], meeting_id=item.get("source_meeting_id"),
           data={"applicationId": application_id})
    if not approved:
        action = _action_row(item["action_item_id"])
        execution_person = execution_rows[0].get("responsible_person") or ""
        _create_reminder(
            action, execution_person, "审批退回",
            f"完工申请被{req.result}：{req.opinion.strip() or '请继续整改'}",
            f"completion-return:{application_id}:{execution_person}",
        )
    return {
        "success": True,
        "message": f"本科室审批已完成，行动项总进度为 {rollup.get('current_progress', 0)}%",
    }


@router.post("/{action_id}/changes")
def apply_change(action_id: int, req: ChangeRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    action = _action_row(action_id)
    _assert_action_visible(ctx, action)
    allowed = (
        ctx["minutes_admin"] or ctx["department_leader"]
        or _is_responsible_person(action, ctx["name"])
        or (ctx["dept_manager"] and _is_responsible_department(action, ctx["dept"]))
    )
    if not allowed:
        raise HTTPException(status_code=403, detail="无权申请该行动项变更")
    if action["current_status"] in (STATUS_COMPLETED, STATUS_CANCELLED):
        raise HTTPException(status_code=400, detail="已完成或已取消行动项不可变更")
    if not req.reason.strip():
        raise HTTPException(status_code=400, detail="变更原因不能为空")
    before = {
        "required_completion_date": _date_text(action.get("required_completion_date")),
        "responsible_person_id": action.get("responsible_person_id"),
        "responsible_department_id": action.get("responsible_department_id"),
        "title": action.get("title"), "content": action.get("content"),
        "current_status": action.get("current_status"),
    }
    change_id = db.execute_insert(
        "INSERT INTO action_change "
        "(action_item_id,change_type,before_content,after_content,change_reason,applicant) "
        "VALUES (%s,%s,%s,%s,%s,%s)",
        (action_id, req.change_type, json.dumps(before, ensure_ascii=False, default=str),
         json.dumps(req.after_content, ensure_ascii=False, default=str),
         req.reason.strip(), ctx["name"]),
    )
    _event(ctx["name"], "变更申请", f"申请{req.change_type}：{req.reason.strip()}",
           action_id=action_id, meeting_id=action.get("source_meeting_id"),
           data={"changeId": change_id, "after": req.after_content})
    _create_reminder(
        action, action.get("supervisor_id") or "", "变更审批",
        f"您有一条{req.change_type}申请待审批",
        f"change:{change_id}:{action.get('supervisor_id')}",
    )
    return {"success": True, "id": change_id, "message": "变更申请已提交"}


@router.post("/changes/{change_id}/approve")
def approve_change(change_id: int, req: ApprovalRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    rows = db.execute_query(
        "SELECT c.*,a.supervisor_id,a.source_meeting_id,a.title,a.required_completion_date,"
        "a.current_status "
        "FROM action_change c JOIN action_item a ON a.id=c.action_item_id WHERE c.id=%s LIMIT 1",
        (change_id,),
    )
    if not rows:
        raise HTTPException(status_code=404, detail="变更申请不存在")
    change = rows[0]
    if not (
        ctx["admin"] or ctx["department_leader"]
        or (change.get("supervisor_id") or "") == ctx["name"]
    ):
        raise HTTPException(status_code=403, detail="仅主管领导可审批该变更")
    if change["approval_status"] != "待审批":
        raise HTTPException(status_code=400, detail="该变更已审批")
    if change["current_status"] in (STATUS_COMPLETED, STATUS_CANCELLED):
        raise HTTPException(status_code=409, detail="行动项已完成或取消，不能再审批该变更")
    approved = req.result == "通过"
    status = "已通过" if approved else "已退回"
    after = _json(change.get("after_content"), {})
    updates: dict[str, Any] = {}
    if approved:
        change_type = change["change_type"]
        if change_type == "延期":
            deadline = (after.get("required_completion_date") or "").strip()
            try:
                date.fromisoformat(deadline)
            except (TypeError, ValueError) as exc:
                raise HTTPException(status_code=400, detail="延期后的完成时间必须为 YYYY-MM-DD") from exc
            updates["required_completion_date"] = deadline
        elif change_type == "负责人变更":
            person = (after.get("responsible_person_id") or "").strip()
            action_now = _action_row(change["action_item_id"])
            departments, people, jobs, supervisors, default_responsibles = (
                _action_directory_maps()
            )
            if person not in people:
                raise HTTPException(status_code=400, detail="新责任人不存在或已离职")
            department, person, supervisor = _resolve_assignment_rules(
                None, person, action_now.get("supervisor_id"), departments, people, jobs,
                supervisors, default_responsibles,
            )
            supervisor = _resolve_scope_supervisor(
                [department] if department else [], [person] if person else [],
                supervisor, jobs, supervisors, None,
                "\n".join((
                    action_now.get("title") or "", action_now.get("content") or "",
                    action_now.get("source_quote") or "",
                )),
            )
            updates["responsible_department_id"] = department
            updates["responsible_person_id"] = person
            updates["responsible_department_ids"] = json.dumps(
                [department] if department else [], ensure_ascii=False,
            )
            updates["responsible_person_ids"] = json.dumps(
                [person] if person else [], ensure_ascii=False,
            )
            updates["supervisor_id"] = supervisor
            updates["current_status"] = STATUS_PENDING_RECEIVE
        elif change_type == "责任科室变更":
            department = _normalize_department_name(
                after.get("responsible_department_id")
            )
            action_now = _action_row(change["action_item_id"])
            departments, people, jobs, supervisors, default_responsibles = (
                _action_directory_maps()
            )
            if department not in departments:
                raise HTTPException(status_code=400, detail="新责任科室不存在")
            department, person, supervisor = _resolve_assignment_rules(
                department, None, action_now.get("supervisor_id"), departments, people,
                jobs, supervisors, default_responsibles,
            )
            responsible_people = _department_responsible_people(department, people, jobs)
            supervisor = _resolve_scope_supervisor(
                [department] if department else [],
                responsible_people or ([person] if person else []),
                supervisor, jobs, supervisors, None,
                "\n".join((
                    action_now.get("title") or "", action_now.get("content") or "",
                    action_now.get("source_quote") or "",
                )),
            )
            updates["responsible_department_id"] = department
            updates["responsible_person_id"] = person
            updates["responsible_department_ids"] = json.dumps(
                [department] if department else [], ensure_ascii=False,
            )
            updates["responsible_person_ids"] = json.dumps(
                responsible_people or ([person] if person else []), ensure_ascii=False,
            )
            updates["supervisor_id"] = supervisor
            updates["current_status"] = STATUS_PENDING_RECEIVE
        elif change_type == "内容调整":
            title = after.get("title")
            content = after.get("content")
            if title is None and content is None:
                raise HTTPException(status_code=400, detail="内容调整必须提供标题或内容")
            if title is not None:
                if not str(title).strip():
                    raise HTTPException(status_code=400, detail="行动项标题不能为空")
                updates["title"] = str(title).strip()[:500]
            if content is not None:
                if not str(content).strip():
                    raise HTTPException(status_code=400, detail="行动项内容不能为空")
                updates["content"] = str(content).strip()
        elif change_type == "取消":
            updates["current_status"] = STATUS_CANCELLED
    db.execute_update(
        "UPDATE action_change SET approval_status=%s,approver=%s,approval_opinion=%s,"
        "approved_at=NOW() WHERE id=%s",
        (status, ctx["name"], req.opinion.strip(), change_id),
    )
    db.execute_update(
        "INSERT INTO action_approval "
        "(business_type,business_id,action_item_id,approver,approval_result,approval_opinion) "
        "VALUES ('行动项变更',%s,%s,%s,%s,%s)",
        (change_id, change["action_item_id"], ctx["name"], req.result, req.opinion.strip()),
    )
    if approved:
        if updates:
            fields = ",".join(f"{key}=%s" for key in updates)
            db.execute_update(
                f"UPDATE action_item SET {fields} WHERE id=%s",
                tuple(list(updates.values()) + [change["action_item_id"]]),
            )
    _event(ctx["name"], "变更审批", f"{req.result}{change['change_type']}申请：{req.opinion.strip()}",
           action_id=change["action_item_id"], meeting_id=change.get("source_meeting_id"),
           data={"changeId": change_id})
    action = _action_row(change["action_item_id"])
    _create_reminder(
        action, change.get("applicant") or "", "变更结果",
        f"{change['change_type']}申请审批结果：{status}",
        f"change-result:{change_id}:{change.get('applicant')}",
    )
    if approved and change["change_type"] in ("负责人变更", "责任科室变更"):
        _create_responsible_reminders(
            action, change["change_type"],
            "您被指定为该行动项的新责任人，请及时接收",
            f"change-owner:{change_id}",
        )
    return {"success": True, "message": "变更审批已完成"}


@router.get("/approvals/pending")
def pending_approvals(current_user: str = Query(...)):
    """返回当前用户有权审批的完工申请和行动项变更。"""
    ensure_tables()
    ctx = _role_context(current_user)
    visible, visible_params = _visible_action_sql(ctx, "a")
    eligibility = "1=1" if (ctx["admin"] or ctx["department_leader"]) else "a.supervisor_id=%s"
    params = list(visible_params)
    if eligibility != "1=1":
        params.append(ctx["name"])
    completions = db.execute_query(
        "SELECT c.id business_id,'完工申请' business_type,c.action_item_id,a.action_number,"
        "a.title,a.responsible_department_id,a.responsible_person_id,"
        "c.responsible_department,c.applicant,"
        "c.applied_at,c.completion_description summary "
        "FROM completion_application c JOIN action_item a ON a.id=c.action_item_id "
        f"WHERE c.approval_status='待审批' AND {visible} AND {eligibility} "
        "ORDER BY c.applied_at DESC",
        tuple(params),
    )
    params = list(visible_params)
    if eligibility != "1=1":
        params.append(ctx["name"])
    changes = db.execute_query(
        "SELECT c.id business_id,CONCAT('变更-',c.change_type) business_type,c.action_item_id,"
        "a.action_number,a.title,a.responsible_department_id,a.responsible_person_id,"
        "c.applicant,c.applied_at,c.change_reason summary "
        "FROM action_change c JOIN action_item a ON a.id=c.action_item_id "
        f"WHERE c.approval_status='待审批' AND {visible} AND {eligibility} "
        "ORDER BY c.applied_at DESC",
        tuple(params),
    )
    items = list(completions) + list(changes)
    for item in items:
        item["applied_at"] = _dt_text(item.get("applied_at"))
    items.sort(key=lambda x: x.get("applied_at") or "", reverse=True)
    return {"success": True, "items": items, "total": len(items)}


@router.post("/{action_id}/remind")
def manual_remind(
    action_id: int, current_user: str = Form(...), note: str = Form("请及时更新行动项进展"),
):
    ensure_tables()
    ctx = _role_context(current_user)
    if not (ctx["minutes_admin"] or ctx["department_leader"]):
        raise HTTPException(status_code=403, detail="仅综合室或部门领导可催办")
    action = _action_row(action_id)
    executions = _ensure_department_executions(action)
    targets = list(dict.fromkeys(
        row.get("responsible_person")
        for row in executions
        if row.get("responsible_person")
        and row.get("execution_status") != STATUS_COMPLETED
    ))
    if not targets:
        return {
            "success": True, "sent": False,
            "message": "责任人尚未识别，本次催办已自动跳过",
        }
    reminder_batch = f"manual:{action_id}:{uuid.uuid4().hex}"
    sent = sum(
        int(_create_reminder(
            action, target, "人工催办",
            note.strip() or "请及时更新行动项进展",
            f"{reminder_batch}:{target}",
        ))
        for target in targets
    )
    if not sent:
        return {
            "success": True, "sent": False,
            "message": "未识别到有效提醒对象，本次催办已自动跳过",
        }
    try:
        _event(ctx["name"], "人工催办", note.strip(), action_id=action_id,
               meeting_id=action.get("source_meeting_id"))
    except Exception as exc:
        logger.warning("人工催办审计记录失败，但不影响提醒结果 action=%s: %s", action_id, exc)
    return {
        "success": True, "sent": True,
        "message": f"催办提醒已发送给 {sent} 名责任人",
    }


@router.get("/reminders/my")
def my_reminders(
    current_user: str = Query(...), unread_only: bool = Query(True),
    limit: int = Query(50, ge=1, le=200),
):
    ensure_tables()
    ctx = _role_context(current_user)
    where = "reminder_target=%s"
    if unread_only:
        where += " AND read_at IS NULL"
    rows = db.execute_query(
        f"SELECT r.*,a.title,a.action_number,a.current_status FROM action_reminder r "
        f"JOIN action_item a ON a.id=r.action_item_id WHERE {where} "
        "ORDER BY r.reminder_time DESC LIMIT %s",
        (ctx["name"], limit),
    )
    items = []
    for row in rows:
        item = dict(row)
        item["reminder_time"] = _dt_text(item.get("reminder_time"))
        item["read_at"] = _dt_text(item.get("read_at"))
        items.append(item)
    return {"success": True, "items": items}


@router.post("/reminders/{reminder_id}/read")
def read_reminder(reminder_id: int, req: ReceiveRequest):
    ensure_tables()
    ctx = _role_context(req.current_user)
    affected = db.execute_update(
        "UPDATE action_reminder SET read_at=NOW() WHERE id=%s AND reminder_target=%s",
        (reminder_id, ctx["name"]),
    )
    if affected <= 0:
        raise HTTPException(status_code=404, detail="提醒不存在")
    return {"success": True}


def run_action_reminder_scan() -> dict:
    """Generate idempotent due-date and stale-update reminders."""
    ensure_tables()
    rows = db.execute_query(
        "SELECT * FROM action_item WHERE current_status IN "
        "('待接收','进行中','待完工审批','退回整改')"
    )
    created = 0
    today = date.today()
    for action in rows:
        executions = _ensure_department_executions(action)
        targets = list(dict.fromkeys(
            execution.get("responsible_person")
            for execution in executions
            if execution.get("responsible_person")
            and execution.get("execution_status") != STATUS_COMPLETED
        ))
        if not targets:
            continue
        deadline = action.get("required_completion_date")
        if isinstance(deadline, datetime):
            deadline = deadline.date()
        if deadline and not isinstance(deadline, date):
            try:
                deadline = date.fromisoformat(str(deadline)[:10])
            except Exception:
                deadline = None
        latest = action.get("last_progress_at") or action.get("published_at")
        if latest:
            if isinstance(latest, str):
                try:
                    latest = datetime.fromisoformat(latest[:19])
                except Exception:
                    latest = None
        for target in targets:
            if deadline:
                delta = (deadline - today).days
                if delta in (7, 3, 0):
                    created += _create_reminder(
                        action, target, f"截止前{delta}天" if delta else "截止当天",
                        f"行动项将在 {deadline:%Y-%m-%d} 截止，请及时推进",
                        f"due:{action['id']}:{target}:{deadline}:{delta}",
                    )
                elif delta < 0 and abs(delta) % 3 == 0:
                    created += _create_reminder(
                        action, target, "逾期提醒",
                        f"行动项已逾期 {abs(delta)} 天，请尽快更新或提交变更申请",
                        f"overdue:{action['id']}:{target}:{today}",
                    )
            if latest and (datetime.now() - latest).days >= 7:
                week_key = today.isocalendar()
                created += _create_reminder(
                    action, target, "长期未更新",
                    "行动项已超过 7 天未更新进展，请及时填报",
                    f"stale:{action['id']}:{target}:{week_key.year}-{week_key.week}",
                )
    return {"scanned": len(rows), "created": int(created)}


@router.post("/reminders/scan")
def scan_reminders(current_user: str = Query(...)):
    ctx = _role_context(current_user)
    if not ctx["admin"]:
        raise HTTPException(status_code=403, detail="仅系统管理员可手工触发提醒扫描")
    return {"success": True, **run_action_reminder_scan()}


async def action_reminder_background_loop() -> None:
    import asyncio
    while True:
        try:
            await asyncio.to_thread(run_action_reminder_scan)
        except Exception as exc:
            logger.exception("行动项提醒扫描失败: %s", exc)
        await asyncio.sleep(3600)


@router.get("/attachments/{attachment_id}")
def download_attachment(attachment_id: int, current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    rows = db.execute_query("SELECT * FROM action_attachment WHERE id=%s LIMIT 1", (attachment_id,))
    if not rows:
        raise HTTPException(status_code=404, detail="附件不存在")
    attachment = rows[0]
    if attachment.get("action_item_id"):
        action = _action_row(int(attachment["action_item_id"]))
        _assert_action_visible(ctx, action)
    elif attachment["business_type"] == "meeting_minutes":
        _user(current_user)
    path = (BASE_DIR / attachment["stored_path"]).resolve()
    if BASE_DIR.resolve() not in path.parents or not path.is_file():
        raise HTTPException(status_code=404, detail="附件文件不存在")
    return FileResponse(
        str(path), filename=attachment["original_name"],
        media_type=attachment.get("content_type") or "application/octet-stream",
    )


@router.get("/{action_id}")
def get_action_detail(action_id: int, current_user: str = Query(...)):
    ensure_tables()
    ctx = _role_context(current_user)
    action = _action_row(action_id)
    _assert_action_visible(ctx, action)
    _departments, people_map, _jobs, _supervisors, _defaults = _action_directory_maps()
    department_executions = _ensure_department_executions(action, people_map)
    _refresh_action_rollup(action_id)
    action = _action_row(action_id)
    assignment_departments = _assignment_departments(ctx, action)
    pending_execution_departments = {
        row.get("department")
        for row in department_executions
        if row.get("execution_status") == STATUS_PENDING_RECEIVE
    }
    assignment_departments = [
        department for department in assignment_departments
        if department in pending_execution_departments
    ]
    assignment_people = []
    if assignment_departments and action.get("current_status") == STATUS_PENDING_RECEIVE:
        _departments, people, jobs, _supervisors, _defaults = _action_directory_maps()
        assignment_people = sorted(
            (
                {
                    "name": name,
                    "department": department,
                    "job": jobs.get(name, ""),
                }
                for name, department in people.items()
                if department in assignment_departments
                and (
                    not ctx["department_leader"]
                    or _job_director_priority(jobs.get(name, "")) is not None
                )
            ),
            key=lambda person: (
                assignment_departments.index(person["department"]),
                _job_director_priority(person["job"])
                if _job_director_priority(person["job"]) is not None else 99,
                person["name"],
            ),
        )
    item = _serialize_action(action)
    minutes = []
    if action.get("source_meeting_id"):
        minutes = db.execute_query(
            "SELECT id,meeting_name,meeting_type,meeting_date,minutes_number,meeting_subject "
            "FROM meeting_minutes WHERE id=%s LIMIT 1", (action["source_meeting_id"],)
        )
    item["meeting"] = minutes[0] if minutes else None
    if item["meeting"]:
        item["meeting"]["meeting_date"] = _date_text(item["meeting"].get("meeting_date"))
    progress = db.execute_query(
        "SELECT * FROM action_progress WHERE action_item_id=%s ORDER BY reported_at DESC,id DESC",
        (action_id,),
    )
    for row in progress:
        row["attachments"] = _json(row.get("attachments"))
        row["expected_completion_date"] = _date_text(row.get("expected_completion_date"))
        row["reported_at"] = _dt_text(row.get("reported_at"))
    completions = db.execute_query(
        "SELECT * FROM completion_application WHERE action_item_id=%s ORDER BY applied_at DESC,id DESC",
        (action_id,),
    )
    for row in completions:
        row["evidence_materials"] = _json(row.get("evidence_materials"))
        row["actual_completion_date"] = _date_text(row.get("actual_completion_date"))
        row["applied_at"] = _dt_text(row.get("applied_at"))
    approvals = db.execute_query(
        "SELECT * FROM action_approval WHERE action_item_id=%s ORDER BY approved_at DESC,id DESC",
        (action_id,),
    )
    for row in approvals:
        row["approved_at"] = _dt_text(row.get("approved_at"))
    for row in department_executions:
        row["received_at"] = _dt_text(row.get("received_at"))
        row["completed_at"] = _dt_text(row.get("completed_at"))
        row["updated_at"] = _dt_text(row.get("updated_at"))
    changes = db.execute_query(
        "SELECT * FROM action_change WHERE action_item_id=%s ORDER BY applied_at DESC,id DESC",
        (action_id,),
    )
    for row in changes:
        row["before_content"] = _json(row.get("before_content"), {})
        row["after_content"] = _json(row.get("after_content"), {})
        row["applied_at"] = _dt_text(row.get("applied_at"))
        row["approved_at"] = _dt_text(row.get("approved_at"))
    reminders = db.execute_query(
        "SELECT * FROM action_reminder WHERE action_item_id=%s ORDER BY reminder_time DESC,id DESC",
        (action_id,),
    )
    for row in reminders:
        row["reminder_time"] = _dt_text(row.get("reminder_time"))
    events = db.execute_query(
        "SELECT * FROM action_event_log WHERE action_item_id=%s ORDER BY created_at DESC,id DESC",
        (action_id,),
    )
    events = _group_timeline_events(events)
    attachments = db.execute_query(
        "SELECT id,business_type,business_id,original_name,file_size,uploader,uploaded_at "
        "FROM action_attachment WHERE action_item_id=%s ORDER BY uploaded_at DESC,id DESC",
        (action_id,),
    )
    for row in attachments:
        row["uploaded_at"] = _dt_text(row.get("uploaded_at"))
    return {
        "success": True, "item": item, "progress": progress, "completions": completions,
        "approvals": approvals, "changes": changes, "reminders": reminders,
        "events": events, "attachments": attachments,
        "departmentExecutions": department_executions, "permissions": _permissions(ctx),
        "user": ctx,
        "assignmentScope": {
            "departments": assignment_departments,
            "people": assignment_people,
        },
    }
