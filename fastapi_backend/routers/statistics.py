# -*- coding: utf-8 -*-
"""
科室统计 API - 请假/加班/公出按科室汇总
- 请假: qj 表, lsys, 仅已通过 qjzt=4
- 加班: jiaban 表, lsys, 仅已通过 jiabanzt=4
- 公出: gcsqb 表, lsysjm, 仅已批准 bldzt>=2 and szrzt>=2
领导人看板扩展：满勤率、科室横向对比、全员排序
- 统计与筛选中排除：名字末尾为1、科室(lsys)末尾为1（视为已离职人员/组织）
- 多数科室汇总接口在全员口径下排除「部办」；工作强度 /leader/work-intensity 全员口径包含「部办」并在按科室列表中展示
"""
from fastapi import APIRouter, HTTPException, Query, Body
from fastapi.responses import Response
from io import BytesIO
from urllib.parse import quote
from decimal import Decimal, ROUND_HALF_UP

# 「部办」：多数科室汇总的全员口径仍排除；绩效/满勤口径纳入部办非经理人员
LEADER_EXCLUDE_LSYS = "部办"
# 不参与任何考勤/统计的虚拟科室
OTHER_DEPT_NAMES = ("其他部门员工", "其他部门成员")
# SQL 片段：用于 WHERE 条件中排除虚拟科室（拼接在已有 != 部办 之后）
_EXCL_OTHER = "AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') "
_EXCL_OTHER_YGGL = "AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') "
_EXCL_BUBAN_MANAGERS = (
    "AND NOT (TRIM(lsys) = '部办' AND "
    "TRIM(COALESCE(jb,'')) IN ('经理','副经理')) "
)
_EXCL_BUBAN_MANAGERS_YGGL = (
    "AND NOT (TRIM(yggl.lsys) = '部办' AND "
    "TRIM(COALESCE(yggl.jb,'')) IN ('经理','副经理')) "
)
# 工作强度：公出仅计境内/境外；gclx 空与考勤模块一致视同「境内公出」
_WI_TRIP_SQL_EXCLUDE_CITY = (
    "AND COALESCE(NULLIF(TRIM(gcsqb.gclx), ''), '境内公出') != '市内公出'"
)
from typing import Optional, List, Tuple, Dict, Any
from datetime import datetime, date
from database import db
import logging
import re
from collections import defaultdict
from utils.helpers import time_to_decimal, format_time
from utils.holiday_loader import load_holidays_dict

try:
    from routers.suggestions import (
        collect_valid_times_with_marks,
        build_intervals_from_marks,
        _load_holiday_festival_map,
        _is_incentive_festival,
        is_workday,
        calc_suggestion_style_overtime_for_record,
    )
except Exception:  # pragma: no cover - import fallback for unusual startup order
    collect_valid_times_with_marks = None
    build_intervals_from_marks = None
    _load_holiday_festival_map = None
    _is_incentive_festival = None
    is_workday = None
    calc_suggestion_style_overtime_for_record = None

logger = logging.getLogger(__name__)


def _can_access_holiday_duty_attendance(name: Optional[str]) -> bool:
    """
    假期值班出勤核查：部长/副部长、系统管理员 admin1、综合技术室主任/副主任（含主任责等主任职级）、人事管理员 admin2。
    与考勤纪律审查页、排班管理入口的可用范围一致。
    """
    n = (name or "").strip()
    if not n:
        return False
    try:
        wc = db.execute_query("SELECT admin1, admin2 FROM webconfig WHERE id = 1 LIMIT 1")
        if wc:
            row = wc[0]
            a1 = (row.get("admin1") or "").strip()
            a2 = (row.get("admin2") or "").strip()
            if a1 and n == a1:
                return True
            if a2 and n == a2:
                return True
    except Exception:
        pass
    try:
        from routers.approvers import _get_user_info, _jb_match, is_zonghe_tech_director
    except Exception:
        return False
    user = _get_user_info(n)
    if not user:
        return False
    jb = (user.get("jb") or "").strip()
    if _jb_match(jb, "部长") or _jb_match(jb, "副部长"):
        return True
    if is_zonghe_tech_director(user):
        return True
    return False


def _can_access_leader_overtime_stats(name: Optional[str]) -> bool:
    """
    领导加班统计权限：
    - 部长/副部长
    - 系统管理员 admin1、人事管理员 admin2
    （不含综合技术室主任/副主任，与驾驶舱其它入口区分）
    """
    n = (name or "").strip()
    if not n:
        return False
    try:
        wc = db.execute_query("SELECT admin1, admin2 FROM webconfig WHERE id = 1 LIMIT 1")
        if wc:
            row = wc[0]
            a1 = (row.get("admin1") or "").strip()
            a2 = (row.get("admin2") or "").strip()
            if a1 and n == a1:
                return True
            if a2 and n == a2:
                return True
    except Exception:
        pass
    try:
        from routers.approvers import _get_user_info, _jb_match
    except Exception:
        return False
    user = _get_user_info(n)
    if not user:
        return False
    jb = (user.get("jb") or "").strip()
    if _jb_match(jb, "部长") or _jb_match(jb, "副部长"):
        return True
    return False


INCENTIVE_FESTIVALS = {"春节", "国庆节", "高温防暑休假"}


def _load_holiday_festival_map(year: int) -> Dict[str, str]:
    """
    加载某年假期的 日期 -> 节日名称(festival) 映射。
    若 holiday 表中无 festival 或读取失败，则返回空字典。
    """
    try:
        from utils.holiday_loader import load_holidays_for_year

        rows = load_holidays_for_year(str(year))
        mapping: Dict[str, str] = {}
        for r in rows:
            date_str = (r.get("date") or "").strip()
            fest = (r.get("festival") or "").strip()
            if date_str:
                mapping[date_str] = fest
        return mapping
    except Exception:
        return {}


def _aggregate_overtime_with_incentive(
    rows: List[Dict],
    holiday_festival_map: Dict[str, str],
    zhibanfei: float,
):
    """
    对原始加班记录按「人+日期」聚合，并按节日激励规则计算：
    - 春节/国庆节/高温防暑休假 这三类节日当天：若当日加班时长(已扣午休) >= 8 小时，则固定奖励 200 元；
      超过 8 小时的部分按 zhibanfei 元/小时额外计算。
    - 其他日期或不足 8 小时的节日，其他绩效激励按 zhibanfei 元/小时计算。
    返回:
    - per_month: { "YYYY-MM": {"hours": 总小时数, "pay": 总金额} }
    - per_employee: { name: {"hours": 总小时数, "pay": 总金额} }
    """
    # 先按 (name, date_str) 聚合每天的小时数
    per_day: Dict[tuple, float] = defaultdict(float)
    for r in rows or []:
        name = (r.get("emp_name") or r.get("name") or "").strip()
        if not name:
            continue
        timedate = r.get("timedate")
        if timedate is None:
            continue
        date_str = str(timedate)[:10]
        if len(date_str) < 10:
            continue
        try:
            hours = float(r.get("hours") if r.get("hours") is not None else r.get("jbf") or 0)
        except (TypeError, ValueError):
            hours = 0.0
        if hours <= 0:
            continue
        per_day[(name, date_str)] += hours

    per_month: Dict[str, Dict[str, float]] = defaultdict(lambda: {"hours": 0.0, "pay": 0.0})
    per_employee: Dict[str, Dict[str, float]] = defaultdict(lambda: {"hours": 0.0, "pay": 0.0})

    for (name, date_str), day_hours in per_day.items():
        month_key = date_str[:7]
        festival = holiday_festival_map.get(date_str, "")
        is_incentive = festival in INCENTIVE_FESTIVALS

        incentive_pay = 0.0
        normal_hours = 0.0

        if is_incentive and day_hours >= 8.0:
            incentive_pay = 200.0
            normal_hours = day_hours - 8.0
        else:
            normal_hours = day_hours

        day_pay = incentive_pay + normal_hours * zhibanfei

        # 员工维度：小时数依然展示真实加班小时；金额为激励 + 普通小时费
        per_employee[name]["hours"] += day_hours
        per_employee[name]["pay"] += day_pay

        # 月维度：同样累计真实小时和金额
        per_month[month_key]["hours"] += day_hours
        per_month[month_key]["pay"] += day_pay

    return per_month, per_employee


def _parse_date(v) -> Optional[date]:
    """将 DB 返回的 datetime/str 转为 date"""
    if v is None:
        return None
    if hasattr(v, "date"):
        return v.date()
    s = str(v)[:10]
    if len(s) == 10 and s[4] == "-" and s[7] == "-":
        try:
            return datetime.strptime(s, "%Y-%m-%d").date()
        except ValueError:
            pass
    return None


def _merge_intervals(intervals: List[Tuple[date, date]]) -> List[Tuple[date, date]]:
    """将多个 [start, end] 闭区间做并集去重，返回合并后的区间列表"""
    if not intervals:
        return []
    sorted_list = sorted([(s, e) for s, e in intervals if s and e])
    if not sorted_list:
        return []
    merged = []
    cur_s, cur_e = sorted_list[0]
    for s, e in sorted_list[1:]:
        if s <= cur_e:
            cur_e = max(cur_e, e)
        else:
            merged.append((cur_s, cur_e))
            cur_s, cur_e = s, e
    merged.append((cur_s, cur_e))
    return merged


def _merge_intervals_days(intervals: List[Tuple[date, date]]) -> float:
    """将多个 [start, end] 区间做并集后计算总天数（去重）。"""
    return sum((e - s).days + 1 for s, e in _merge_intervals(intervals))


def _merge_intervals_split_workdays(intervals: List[Tuple[date, date]]) -> Tuple[float, float]:
    """
    将公出区间合并后，分别统计其中的工作日天数和非工作日（节假日+周末）天数。
    返回 (total_days, holiday_days)，其中 holiday_days 是公出期间的非工作日天数。
    跨年时按各日所属年份加载节假日字典。
    """
    merged = _merge_intervals(intervals)
    if not merged:
        return 0.0, 0.0

    from datetime import timedelta

    holiday_cache: Dict[int, dict] = {}

    def _holidays_for_year(y: int) -> dict:
        if y not in holiday_cache:
            try:
                from utils.holiday_loader import load_holidays_dict
                holiday_cache[y] = load_holidays_dict(str(y))
            except Exception:
                holiday_cache[y] = {}
        return holiday_cache[y]

    total = 0
    holiday_count = 0
    for s, e in merged:
        d = s
        while d <= e:
            total += 1
            holidays = _holidays_for_year(d.year)
            date_str = d.strftime("%Y-%m-%d")
            is_weekend = d.weekday() in [5, 6]
            is_holiday = False
            is_workday_override = False

            if date_str in holidays:
                t = holidays[date_str] or ""
                if "假" in t or "休" in t:
                    is_holiday = True
                if "班" in t:
                    is_workday_override = True

            if is_workday_override:
                pass  # 调休上班日，算工作日
            elif is_weekend or is_holiday:
                holiday_count += 1

            d += timedelta(days=1)

    return float(total), float(holiday_count)


def _qj_leave_date_bounds(row: Dict) -> Tuple[Optional[date], Optional[date]]:
    """请假记录的起止日期（闭区间）。优先 timefrom/timeto，缺省回退 timefromdate。"""
    start = _parse_date(row.get("timefrom")) or _parse_date(row.get("timefromdate"))
    end = _parse_date(row.get("timeto")) or start
    if start and end and end < start:
        start, end = end, start
    return start, end


def _allocate_leave_tian_to_period(
    tian: float, leave_start: date, leave_end: date, period_start: date, period_end: date
) -> float:
    """
    将整条请假的 tian 按「与统计区间相交的日历天数 / 请假整段日历天数」比例分摊。
    解决跨月/跨年请假只在「开始月」计满全部天数的问题。
    """
    try:
        tian_f = float(tian or 0)
    except (TypeError, ValueError):
        tian_f = 0.0
    if tian_f <= 0 or not leave_start or not leave_end:
        return 0.0
    total_span = (leave_end - leave_start).days + 1
    if total_span <= 0:
        return 0.0
    ov_s = max(leave_start, period_start)
    ov_e = min(leave_end, period_end)
    if ov_e < ov_s:
        return 0.0
    ov_days = (ov_e - ov_s).days + 1
    return tian_f * float(ov_days) / float(total_span)


def _leave_overlap_fraction(
    leave_start: date, leave_end: date, period_start: date, period_end: date
) -> float:
    """请假区间与统计区间日历重叠比例 (0~1]，用于将整笔换休小时分摊到当期。"""
    if not leave_start or not leave_end:
        return 0.0
    total_span = (leave_end - leave_start).days + 1
    if total_span <= 0:
        return 0.0
    ov_s = max(leave_start, period_start)
    ov_e = min(leave_end, period_end)
    if ov_e < ov_s:
        return 0.0
    ov_days = (ov_e - ov_s).days + 1
    return float(ov_days) / float(total_span)


def _qj_hx_row_total_hours(row: Dict) -> float:
    """单条换休类请假的总小时：优先 xiaoshi，否则 tian×8（与申请端一致）。"""
    try:
        xs_raw = row.get("xiaoshi")
        if xs_raw is not None and str(xs_raw).strip() != "":
            xs = float(str(xs_raw).strip().replace(",", "."))
            if xs > 0:
                return xs
    except (TypeError, ValueError):
        pass
    try:
        tian = float(row.get("tian") or 0)
    except (TypeError, ValueError):
        tian = 0.0
    return max(0.0, tian * 8.0)


def _leave_overlap_sql_bounds() -> str:
    """WHERE 片段：请假区间与 [ps, pe] 闭区间有交集（ps/pe 为 'YYYY-MM-DD' 字符串）。"""
    return """
            AND COALESCE(DATE(timefrom), DATE(timefromdate)) <= %s
            AND COALESCE(DATE(timeto), DATE(timefrom), DATE(timefromdate)) >= %s
        """


router = APIRouter(tags=["统计"])


def _count_workdays_in_month(year: int, month: int) -> int:
    """计算某月应出勤工作日数（考虑假期与调休）"""
    try:
        from utils.holiday_loader import load_holidays_dict
        holidays = load_holidays_dict(str(year))
    except Exception:
        holidays = {}
    count = 0
    try:
        import calendar
        _, last = calendar.monthrange(year, month)
        for day in range(1, last + 1):
            d = datetime(year, month, day)
            date_str = d.strftime("%Y-%m-%d")
            weekday = d.weekday()
            is_weekend = weekday in [5, 6]
            is_holiday = False
            if date_str in holidays:
                t = holidays[date_str] or ""
                if "假" in t or "休" in t:
                    is_holiday = True
            if date_str in holidays and "班" in holidays[date_str]:
                is_weekend = False
                is_holiday = False
            if not is_weekend and not is_holiday:
                count += 1
    except Exception as e:
        logger.warning(f"计算工作日失败: {e}, 使用当月天数估算")
        import calendar
        _, last = calendar.monthrange(year, month)
        count = min(last, 22)
    return count


def _count_workdays_in_month_until(year: int, month: int, end_day: int) -> int:
    """计算某月从 1 号到 end_day 的应出勤工作日数（含调休）"""
    try:
        import calendar
        _, last = calendar.monthrange(year, month)
        end_day = max(1, min(end_day, last))
    except Exception:
        end_day = max(1, end_day)
    try:
        from utils.holiday_loader import load_holidays_dict
        holidays = load_holidays_dict(str(year))
    except Exception:
        holidays = {}

    count = 0
    for day in range(1, end_day + 1):
        d = datetime(year, month, day)
        date_str = d.strftime("%Y-%m-%d")
        weekday = d.weekday()
        is_weekend = weekday in [5, 6]
        is_holiday = False
        if date_str in holidays:
            t = holidays[date_str] or ""
            if "假" in t or "休" in t:
                is_holiday = True
        if date_str in holidays and "班" in holidays[date_str]:
            is_weekend = False
            is_holiday = False
        if not is_weekend and not is_holiday:
            count += 1
    return count


def _count_workdays_between(d_start: date, d_end: date) -> int:
    """计算闭区间 [d_start, d_end] 内应出勤工作日数（含调休），与按月统计规则一致。"""
    if d_end < d_start:
        return 0
    try:
        from utils.holiday_loader import load_holidays_dict
    except Exception:
        load_holidays_dict = None  # type: ignore

    holiday_cache: Dict[int, dict] = {}

    def _holidays_for_year(y: int) -> dict:
        if y not in holiday_cache:
            if load_holidays_dict:
                try:
                    holiday_cache[y] = load_holidays_dict(str(y))
                except Exception:
                    holiday_cache[y] = {}
            else:
                holiday_cache[y] = {}
        return holiday_cache[y]

    count = 0
    d = d_start
    from datetime import timedelta
    while d <= d_end:
        holidays = _holidays_for_year(d.year)
        date_str = d.strftime("%Y-%m-%d")
        weekday = d.weekday()
        is_weekend = weekday in [5, 6]
        is_holiday = False
        if date_str in holidays:
            t = holidays[date_str] or ""
            if "假" in t or "休" in t:
                is_holiday = True
        if date_str in holidays and holidays[date_str] and "班" in holidays[date_str]:
            is_weekend = False
            is_holiday = False
        if not is_weekend and not is_holiday:
            count += 1
        d += timedelta(days=1)
    return count


def _get_lsysjm_list(lsys: str) -> List[str]:
    """根据 lsys 获取对应的 lsysjm 列表（用于公出表）"""
    if not lsys:
        return []
    rows = db.execute_query(
        "SELECT DISTINCT lsysjm FROM yggl WHERE lsys = %s AND lsysjm IS NOT NULL AND lsysjm != '' AND (COALESCE(zaizhi,0)=0)",
        (lsys,)
    )
    result = [r["lsysjm"].strip() for r in rows if r.get("lsysjm")]
    if not result and lsys:
        result = [lsys]  # 若无映射则用 lsys 本身
    return result


# ==================== 部办打卡加班统计（领导高亮） ====================

LEADER_OVERTIME_JB_RE = r"经理助理|副经理|经理|副部长|部长"
LEADER_OVERTIME_BASELINE_YEAR = 2025
LEADER_OVERTIME_BASELINE_DAYS: Tuple[float, ...] = (
    23.53, 20.80, 20.51, 16.67, 15.09, 14.89, 14.89, 14.41, 14.27, 14.22,
    13.45, 13.15, 12.98, 12.80, 12.45, 12.43, 12.32, 12.20, 11.89, 11.59,
    11.44, 11.27, 11.04, 11.02, 10.75, 10.64, 10.51, 10.34, 10.33, 10.21,
    10.15, 10.03, 9.75, 9.69, 9.60, 9.59, 9.41, 9.32, 9.31, 9.26,
    9.10, 9.10, 9.06, 9.00, 8.81, 8.75, 8.74, 8.66, 8.65, 8.64,
    8.62, 8.60, 8.59, 8.47, 8.33, 8.15, 7.86, 7.85, 7.78, 7.69,
    7.67, 7.59, 7.51, 7.43, 7.41, 7.41, 7.39, 7.25, 6.99, 6.94,
    6.85, 6.71, 6.58, 6.53, 6.52, 6.49, 6.48, 6.47, 6.41, 6.39,
    6.38, 6.27, 6.23, 6.21, 6.08, 5.98, 5.83, 5.77, 5.73, 5.68,
    5.55, 5.43, 5.38, 5.32, 5.23, 5.12, 4.97, 4.93, 4.89, 4.84,
    4.37, 4.29, 4.28, 4.20, 4.05, 3.37, 3.22, 2.64, 2.31, 2.23,
    2.01, 1.90, 1.71, 1.63, 1.54, 1.36, 1.35, 1.29, 0.88, 0.82,
    0.75, 0.00, 0.00, 0.00, 0.00, 0.00,
)


def _round_01_half_up(value: float) -> float:
    return float(Decimal(str(value or 0)).quantize(Decimal("0.1"), rounding=ROUND_HALF_UP))


def _round_02_half_up(value: float) -> float:
    return float(Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP))


def _ensure_leader_overtime_baseline() -> None:
    db.execute_update(
        """
        CREATE TABLE IF NOT EXISTS leader_overtime_rank_baseline (
            id INT AUTO_INCREMENT PRIMARY KEY,
            baseline_year INT NOT NULL,
            rank_no INT NOT NULL,
            monthly_avg_days DECIMAL(6,2) NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY uk_leader_overtime_rank (baseline_year, rank_no)
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COMMENT='中层领导干部月均加班天数排名基准'
        """,
        (),
    )
    rows = db.execute_query(
        "SELECT COUNT(*) AS cnt FROM leader_overtime_rank_baseline WHERE baseline_year = %s",
        (LEADER_OVERTIME_BASELINE_YEAR,),
    )
    count = int((rows[0] or {}).get("cnt") or 0) if rows else 0
    if count == len(LEADER_OVERTIME_BASELINE_DAYS):
        return
    params = [
        (LEADER_OVERTIME_BASELINE_YEAR, idx, value)
        for idx, value in enumerate(LEADER_OVERTIME_BASELINE_DAYS, start=1)
    ]
    db.execute_many(
        """
        INSERT INTO leader_overtime_rank_baseline (baseline_year, rank_no, monthly_avg_days)
        VALUES (%s, %s, %s)
        ON DUPLICATE KEY UPDATE monthly_avg_days = VALUES(monthly_avg_days)
        """,
        params,
    )


def _resolve_leader_overtime_baseline_year(query_year: Optional[int] = None) -> int:
    _ensure_leader_overtime_baseline()
    try:
        if query_year:
            rows = db.execute_query(
                "SELECT MAX(baseline_year) AS y FROM leader_overtime_rank_baseline WHERE baseline_year < %s",
                (int(query_year),),
            )
            year_val = (rows[0] or {}).get("y") if rows else None
            if year_val:
                return int(year_val)
        rows = db.execute_query(
            "SELECT MAX(baseline_year) AS y FROM leader_overtime_rank_baseline",
            (),
        )
        year_val = (rows[0] or {}).get("y") if rows else None
        return int(year_val or LEADER_OVERTIME_BASELINE_YEAR)
    except Exception as e:
        logger.error(f"解析领导加班排名基准年份失败: {str(e)}")
        return LEADER_OVERTIME_BASELINE_YEAR


def _get_leader_overtime_baseline(query_year: Optional[int] = None, baseline_year: Optional[int] = None) -> Tuple[int, List[float]]:
    try:
        _ensure_leader_overtime_baseline()
        target_year = int(baseline_year) if baseline_year else _resolve_leader_overtime_baseline_year(query_year)
        rows = db.execute_query(
            "SELECT monthly_avg_days FROM leader_overtime_rank_baseline "
            "WHERE baseline_year = %s ORDER BY rank_no ASC",
            (target_year,),
        )
        values = [float(r.get("monthly_avg_days") or 0) for r in rows or []]
        return target_year, values or list(LEADER_OVERTIME_BASELINE_DAYS)
    except Exception as e:
        logger.error(f"读取领导加班排名基准失败: {str(e)}")
        return LEADER_OVERTIME_BASELINE_YEAR, list(LEADER_OVERTIME_BASELINE_DAYS)


def _save_leader_overtime_baseline_values(year: int, values: List[float]) -> int:
    _ensure_leader_overtime_baseline()
    normalized = sorted([_round_02_half_up(v) for v in values], reverse=True)
    db.execute_update(
        "DELETE FROM leader_overtime_rank_baseline WHERE baseline_year = %s",
        (year,),
    )
    params = [(year, idx, value) for idx, value in enumerate(normalized, start=1)]
    return db.execute_many(
        """
        INSERT INTO leader_overtime_rank_baseline (baseline_year, rank_no, monthly_avg_days)
        VALUES (%s, %s, %s)
        """,
        params,
    )


def _estimate_leader_overtime_rank(monthly_avg_days: float, baseline_values: List[float]) -> Optional[int]:
    if not baseline_values:
        return None
    return min(sum(1 for v in baseline_values if float(v) > monthly_avg_days) + 1, len(baseline_values))


def _period_month_count(start: date, end: date, explicit_month: Optional[int]) -> float:
    if explicit_month:
        return 1.0
    import calendar
    if start.day == 1 and end.day == calendar.monthrange(end.year, end.month)[1]:
        return float((end.year - start.year) * 12 + end.month - start.month + 1)
    days = (end - start).days + 1
    return max(days / 30.0, 1 / 30.0)


def _parse_attendance_date(val: Any) -> Optional[datetime]:
    if not val:
        return None
    if isinstance(val, datetime):
        return val
    if isinstance(val, date):
        return datetime.combine(val, datetime.min.time())
    s = str(val).strip().replace("/", "-")[:10]
    try:
        return datetime.strptime(s, "%Y-%m-%d")
    except Exception:
        return None


def _leader_overtime_for_record(record: Dict, holidays: Dict[str, str], holiday_festival_map: Dict[str, str]) -> Tuple[float, List[Dict], str]:
    """打卡加班：与考勤页智能建议「加班建议」同款算法（见 suggestions.calc_suggestion_style_overtime_for_record）。"""
    if not calc_suggestion_style_overtime_for_record:
        return 0.0, [], ""
    return calc_suggestion_style_overtime_for_record(record, holidays, holiday_festival_map)


@router.get("/dept/leader-overtime")
async def get_leader_overtime_from_attendance(
    year: int = Query(..., description="年份"),
    month: Optional[int] = Query(None, ge=1, le=12, description="月份，不传为全年"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD，与 date_to 同时传入时按日期区间统计"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD，与 date_from 同时传入时按日期区间统计"),
    current_user: Optional[str] = Query(None, description="当前用户，用于权限校验"),
):
    """
    部办人员加班统计：从 attendance_records 打卡数据临时识别，领导岗位单独标记。
    识别口径与考勤页智能建议「加班建议」一致：工作日 17:00 后与 [17,24] 交集满 1 小时/段；
    休息日/假期同 analyze_restday。最小精确到 0.1 小时，四舍五入。
    """
    try:
        if not (current_user or "").strip() or not _can_access_leader_overtime_stats(current_user):
            raise HTTPException(status_code=403, detail="无权查看领导加班统计")

        import calendar
        if date_from or date_to:
            if not date_from or not date_to:
                raise HTTPException(status_code=400, detail="date_from 与 date_to 需同时传入")
            start = _parse_date(date_from)
            end = _parse_date(date_to)
            if not start or not end:
                raise HTTPException(status_code=400, detail="日期格式无效，请使用 YYYY-MM-DD")
            if end < start:
                start, end = end, start
            # 该接口按年度假期配置识别加班类型，暂不支持跨年区间
            if start.year != end.year:
                raise HTTPException(status_code=400, detail="日期区间暂不支持跨年，请在同一年内查询")
            year = start.year
            month = None
        else:
            start = date(year, int(month), 1) if month else date(year, 1, 1)
            end = date(year, int(month), calendar.monthrange(year, int(month))[1]) if month else date(year, 12, 31)
            # 查询“当年全年”时，仅统计到今天，避免把未来日期纳入统计区间
            if not month and year == date.today().year:
                end = date.today()

        members = db.execute_query(
            "SELECT TRIM(name) AS name, TRIM(jb) AS jb, TRIM(lsys) AS lsys "
            "FROM yggl WHERE TRIM(lsys) = %s AND COALESCE(zaizhi,0)=0 "
            "AND name IS NOT NULL AND TRIM(name) != '' AND RIGHT(TRIM(name),1) != '1'",
            (LEADER_EXCLUDE_LSYS,),
        )
        members = [r for r in (members or []) if r.get("name")]
        names = [r["name"] for r in members]
        if not names:
            return {
                "success": True,
                "year": year,
                "month": month,
                "totalHours": 0,
                "personCount": 0,
                "rankBaselineYear": _resolve_leader_overtime_baseline_year(year),
                "rankTotal": len(LEADER_OVERTIME_BASELINE_DAYS),
                "list": [],
                "details": [],
            }

        ph = ",".join(["%s"] * len(names))
        rows = db.execute_query(
            f"SELECT * FROM attendance_records WHERE employee_name IN ({ph}) "
            "AND attendance_date >= %s AND attendance_date <= %s "
            "ORDER BY attendance_date DESC, employee_name",
            tuple(names) + (start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d")),
        )

        holidays = load_holidays_dict(str(year))
        holiday_festival_map = _load_holiday_festival_map(year) if _load_holiday_festival_map else {}
        baseline_year, baseline_values = _get_leader_overtime_baseline(query_year=year)
        rank_total = len(baseline_values)
        period_months = _period_month_count(start, end, month)
        by_name: Dict[str, Dict] = {
            r["name"]: {
                "name": r["name"],
                "jb": r.get("jb") or "",
                "lsys": r.get("lsys") or LEADER_EXCLUDE_LSYS,
                "isLeader": bool(re.search(LEADER_OVERTIME_JB_RE, r.get("jb") or "")),
                "hoursRaw": 0.0,
                "days": 0,
            }
            for r in members
        }
        details: List[Dict] = []
        for row in rows or []:
            name = (row.get("employee_name") or "").strip()
            if name not in by_name:
                continue
            hours, segments, day_type = _leader_overtime_for_record(row, holidays, holiday_festival_map)
            if hours <= 0:
                continue
            rounded = _round_01_half_up(hours)
            by_name[name]["hoursRaw"] += hours
            by_name[name]["days"] += 1
            details.append({
                "name": name,
                "isLeader": bool(by_name[name].get("isLeader")),
                "date": str(row.get("attendance_date") or "")[:10],
                "dayType": day_type,
                "hours": rounded,
                "segments": segments,
            })

        list_data = []
        for item in by_name.values():
            h = _round_01_half_up(item.pop("hoursRaw", 0.0))
            item["hours"] = h
            monthly_avg_days = _round_02_half_up(h / 8 / period_months)
            item["monthlyAvgDays"] = monthly_avg_days
            item["estimatedRank"] = _estimate_leader_overtime_rank(monthly_avg_days, baseline_values) if item.get("isLeader") else None
            item["rankTotal"] = rank_total if item.get("isLeader") else None
            list_data.append(item)
        list_data.sort(key=lambda x: (-x["hours"], x.get("name") or ""))

        return {
            "success": True,
            "year": year,
            "month": month,
            "range": {"start": start.strftime("%Y-%m-%d"), "end": end.strftime("%Y-%m-%d")},
            "periodMonths": _round_02_half_up(period_months),
            "rankBaselineYear": baseline_year,
            "rankTotal": rank_total,
            "totalHours": _round_01_half_up(sum(i["hours"] for i in list_data)),
            "personCount": len(list_data),
            "list": list_data,
            "details": details,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"领导加班统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/dept/leader-overtime-baseline")
async def get_leader_overtime_rank_baseline(
    year: Optional[int] = Query(None, description="基准年份，不传默认上一年"),
    current_user: Optional[str] = Query(None, description="当前用户，用于权限校验"),
):
    """读取中层领导干部月均加班天数排名基准。"""
    try:
        if not (current_user or "").strip() or not _can_access_leader_overtime_stats(current_user):
            raise HTTPException(status_code=403, detail="无权查看领导加班排名基准")
        target_year = int(year or (date.today().year - 1))
        _ensure_leader_overtime_baseline()
        rows = db.execute_query(
            "SELECT monthly_avg_days FROM leader_overtime_rank_baseline "
            "WHERE baseline_year = %s ORDER BY rank_no ASC",
            (target_year,),
        )
        values = [float(r.get("monthly_avg_days") or 0) for r in rows or []]
        return {
            "success": True,
            "year": target_year,
            "count": len(values),
            "values": values,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"读取领导加班排名基准失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/dept/leader-overtime-baseline")
async def save_leader_overtime_rank_baseline(payload: Dict[str, Any] = Body(...)):
    """保存中层领导干部月均加班天数排名基准，一次覆盖指定年份。"""
    try:
        current_user = (payload.get("current_user") or "").strip()
        if not current_user or not _can_access_leader_overtime_stats(current_user):
            raise HTTPException(status_code=403, detail="无权维护领导加班排名基准")
        try:
            year = int(payload.get("year"))
        except Exception:
            raise HTTPException(status_code=400, detail="请填写有效年份")
        if year < 2000 or year > 2100:
            raise HTTPException(status_code=400, detail="年份范围无效")

        raw_values = payload.get("values") or []
        if not isinstance(raw_values, list) or not raw_values:
            raise HTTPException(status_code=400, detail="请录入月均加班天数")
        if len(raw_values) > 1000:
            raise HTTPException(status_code=400, detail="录入数量过多")

        values: List[float] = []
        for raw in raw_values:
            try:
                value = float(raw)
            except Exception:
                raise HTTPException(status_code=400, detail=f"存在无效数字：{raw}")
            if value < 0:
                raise HTTPException(status_code=400, detail="月均加班天数不能为负数")
            values.append(value)

        affected = _save_leader_overtime_baseline_values(year, values)
        if affected < 0:
            raise HTTPException(status_code=500, detail="保存失败")
        _, saved_values = _get_leader_overtime_baseline(baseline_year=year)
        return {
            "success": True,
            "year": year,
            "count": len(saved_values),
            "values": saved_values,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"保存领导加班排名基准失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 科室列表（部长/副部长可选任意科室） ====================

@router.get("/dept/lsys-list")
async def get_dept_lsys_list():
    """
    获取全部隶属科室列表（用于领导人看板：部长/副部长可下拉选择任意科室）
    返回: { success, list: ["部办", "科室A", ...] }
    """
    try:
        rows = db.execute_query(
            "SELECT DISTINCT lsys FROM yggl WHERE lsys IS NOT NULL AND lsys != '' "
            "AND RIGHT(TRIM(lsys), 1) != '1' "
            "AND TRIM(lsys) NOT IN ('其他部门员工', '其他部门成员') "
            "AND (COALESCE(zaizhi,0)=0) ORDER BY lsys",
        )
        list_data = [r["lsys"].strip() for r in rows if r.get("lsys")]
        return {"success": True, "list": list_data}
    except Exception as e:
        logger.error(f"科室列表查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 请假科室统计 ====================

@router.get("/dept/leave")
async def get_dept_leave_stats(
    lsys: Optional[str] = Query(None, description="隶属于室，不传或空为全员"),
    year: Optional[int] = None,
    month: Optional[int] = None,
    quarter: Optional[str] = None,
    hx_only: Optional[bool] = Query(False, description="为 true 时仅统计请假类型为换休、员工换休票的记录"),
):
    """
    科室请假统计（按人汇总天数）。不传 lsys 时为全员（排除部办）。
    返回: { totalDays, personCount, list: [{ name, days }] }
    仅统计 qjzt=4 已通过。
    hx_only=true 时仅 qjfs 为换休/员工换休票（TRIM 后匹配）。
    跨月/跨季请假按日历重叠比例将 tian 分摊到各统计区间（与领导人看板月份/季度/全年一致）。
    """
    try:
        import calendar

        if year is None:
            year = datetime.now().year
        all_staff = not (lsys and lsys.strip())

        if month:
            period_start = date(year, int(month), 1)
            period_end = date(year, int(month), calendar.monthrange(year, int(month))[1])
        elif quarter:
            q = str(quarter).strip()
            q_map = {"1": (1, 3), "2": (4, 6), "3": (7, 9), "4": (10, 12)}
            lo, hi = q_map.get(q, (1, 12))
            period_start = date(year, lo, 1)
            period_end = date(year, hi, calendar.monthrange(year, hi)[1])
        else:
            period_start = date(year, 1, 1)
            period_end = date(year, 12, 31)

        pe_str = period_end.strftime("%Y-%m-%d")
        ps_str = period_start.strftime("%Y-%m-%d")
        ov = _leave_overlap_sql_bounds()
        hx_filter = " AND TRIM(qjfs) IN ('换休','员工换休票')" if hx_only else ""

        if all_staff:
            query = f"""
                SELECT xm AS name, timefrom, timeto, timefromdate, CAST(tian AS DECIMAL(10,2)) AS tian
                FROM qj
                WHERE qjzt = 4 AND RIGHT(TRIM(xm), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' AND TRIM(lsys) != %s
                {hx_filter}
                {ov}
            """
            rows = db.execute_query(query, (LEADER_EXCLUDE_LSYS, pe_str, ps_str))
        else:
            query = f"""
                SELECT xm AS name, timefrom, timeto, timefromdate, CAST(tian AS DECIMAL(10,2)) AS tian
                FROM qj
                WHERE lsys = %s AND qjzt = 4 AND RIGHT(TRIM(xm), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1'
                {hx_filter}
                {ov}
            """
            rows = db.execute_query(query, (lsys, pe_str, ps_str))

        by_name: Dict[str, float] = defaultdict(float)
        for r in rows or []:
            name = (r.get("name") or "").strip()
            if not name:
                continue
            ls_d, le_d = _qj_leave_date_bounds(r)
            if not ls_d or not le_d:
                continue
            try:
                tian_f = float(r.get("tian") or 0)
            except (TypeError, ValueError):
                tian_f = 0.0
            alloc = _allocate_leave_tian_to_period(tian_f, ls_d, le_d, period_start, period_end)
            if alloc > 0:
                by_name[name] += alloc

        list_data = [{"name": n, "days": round(v, 2)} for n, v in sorted(by_name.items(), key=lambda x: -x[1])]
        total_days = sum(by_name.values())

        return {
            "success": True,
            "totalDays": round(total_days, 2),
            "personCount": len(list_data),
            "list": list_data
        }
    except Exception as e:
        logger.error(f"请假科室统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 加班科室统计 ====================


def _query_hx_hours_all(
    year: int,
    month: Optional[int] = None,
    quarter: Optional[str] = None,
    lsys: Optional[str] = None,
    all_staff: bool = True,
) -> Dict[str, float]:
    """
    统计期内每人应扣换休小时：qjzt=4 且 qjfs 为换休/员工换休票；
    与请假统计一致，按请假区间与统计期日历重叠比例分摊整笔小时；
    小时优先 xiaoshi，缺省按 tian×8。净加班 = 该人加班小时 − 该值。
    """
    import calendar

    if month:
        period_start = date(year, int(month), 1)
        period_end = date(year, int(month), calendar.monthrange(year, int(month))[1])
    elif quarter:
        q = str(quarter).strip()
        q_map = {"1": (1, 3), "2": (4, 6), "3": (7, 9), "4": (10, 12)}
        lo, hi = q_map.get(q, (1, 12))
        period_start = date(year, lo, 1)
        period_end = date(year, hi, calendar.monthrange(year, hi)[1])
    else:
        period_start = date(year, 1, 1)
        period_end = date(year, 12, 31)

    pe_str = period_end.strftime("%Y-%m-%d")
    ps_str = period_start.strftime("%Y-%m-%d")
    ov = _leave_overlap_sql_bounds()

    if all_staff:
        query = f"""
            SELECT TRIM(qj.xm) AS name, qj.timefrom, qj.timeto, qj.timefromdate,
                CAST(qj.tian AS DECIMAL(10,2)) AS tian, qj.xiaoshi
            FROM qj INNER JOIN yggl ON qj.xm = yggl.name
                AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1'
                AND TRIM(yggl.lsys) != %s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
            WHERE qj.qjzt = 4 AND TRIM(qj.qjfs) IN ('换休','员工换休票') AND RIGHT(TRIM(qj.xm), 1) != '1'
            {ov}
        """
        rows = db.execute_query(query, (LEADER_EXCLUDE_LSYS, pe_str, ps_str))
    else:
        query = f"""
            SELECT TRIM(qj.xm) AS name, qj.timefrom, qj.timeto, qj.timefromdate,
                CAST(qj.tian AS DECIMAL(10,2)) AS tian, qj.xiaoshi
            FROM qj INNER JOIN yggl ON qj.xm = yggl.name AND yggl.lsys = %s
                AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1'
                AND (COALESCE(yggl.zaizhi,0)=0)
            WHERE qj.qjzt = 4 AND TRIM(qj.qjfs) IN ('换休','员工换休票') AND RIGHT(TRIM(qj.xm), 1) != '1'
            {ov}
        """
        rows = db.execute_query(query, (lsys, pe_str, ps_str))

    by_name: Dict[str, float] = defaultdict(float)
    for r in rows or []:
        name = (r.get("name") or "").strip()
        if not name:
            continue
        ls_d, le_d = _qj_leave_date_bounds(r)
        if not ls_d or not le_d:
            continue
        h_full = _qj_hx_row_total_hours(r)
        if h_full <= 0:
            continue
        frac = _leave_overlap_fraction(ls_d, le_d, period_start, period_end)
        if frac <= 0:
            continue
        by_name[name] += h_full * frac

    return dict(by_name)


def _dept_overtime_scope_names(lsys: Optional[str]) -> List[str]:
    """与 /dept/overtime 统计范围一致的在职员工姓名列表。"""
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        rows = db.execute_query(
            "SELECT TRIM(name) AS name FROM yggl WHERE name IS NOT NULL AND name != '' "
            "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
            "AND TRIM(lsys) != %s AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') "
            "AND (COALESCE(zaizhi,0)=0)",
            (LEADER_EXCLUDE_LSYS,),
        )
    else:
        rows = db.execute_query(
            "SELECT TRIM(name) AS name FROM yggl WHERE lsys = %s AND name IS NOT NULL AND name != '' "
            "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
            "AND (COALESCE(zaizhi,0)=0)",
            (lsys.strip(),),
        )
    return [(r.get("name") or "").strip() for r in (rows or []) if (r.get("name") or "").strip()]


def _dept_overtime_period_bounds(
    year: int,
    month: Optional[int] = None,
    quarter: Optional[str] = None,
) -> Tuple[date, date]:
    """驾驶舱加班统计期 [start, end]（结束日不晚于今天）。"""
    import calendar as _cal

    today = date.today()
    if month is not None:
        d0 = date(year, month, 1)
        d1 = date(year, month, _cal.monthrange(year, month)[1])
        if year == today.year and month == today.month:
            d1 = min(d1, today)
        return d0, d1
    if quarter in ("1", "2", "3", "4"):
        qm = {"1": (1, 3), "2": (4, 6), "3": (7, 9), "4": (10, 12)}[quarter]
        d0 = date(year, qm[0], 1)
        d1 = date(year, qm[1], _cal.monthrange(year, qm[1])[1])
        if year == today.year:
            d1 = min(d1, today)
        return d0, d1
    d0 = date(year, 1, 1)
    d1 = date(year, 12, 31)
    if year == today.year:
        d1 = min(d1, today)
    return d0, d1


def _calc_auto_overtime_hours_from_attendance(
    scope_names: List[str],
    period_start: date,
    period_end: date,
) -> Tuple[float, int]:
    """
    打卡自动识别加班总时长（与智能建议「加班建议」合计一致）。
    返回 (总小时, 有识别记录人数)。
    """
    if not scope_names or period_start > period_end:
        return 0.0, 0
    if not collect_valid_times_with_marks:
        return 0.0, 0
    att_map = _leader_style_overtime_hours_from_attendance(scope_names, period_start, period_end)
    total = 0.0
    person_count = 0
    for n in scope_names:
        att_h = round(float(att_map.get(n, 0.0)), 2)
        if att_h > 0:
            person_count += 1
        total += att_h
    return round(total, 2), person_count


@router.get("/dept/overtime")
async def get_dept_overtime_stats(
    lsys: Optional[str] = Query(None, description="隶属于室，不传或空为全员"),
    year: Optional[int] = None,
    month: Optional[int] = None,
    quarter: Optional[str] = None,
    net: Optional[bool] = Query(False, description="true 时返回净加班（加班小时 - 换休请假小时）"),
):
    """
    科室加班统计（按人汇总小时）。不传 lsys 时为全员（排除部办）。
    net=true 时，每人加班小时减去该人在同期换休类请假（换休/员工换休票）应扣小时，
    与请假统计一致按日历重叠比例分摊；小时优先取 xiaoshi，缺省按 tian×8。
    返回: { totalHours, personCount, list, autoCalculatedHours, autoCalculatedPersonCount }
    已申报：jiabanzt=4 已通过；autoCalculatedHours=智能建议同款打卡加班合计
    """
    try:
        if year is None:
            year = __import__("datetime").datetime.now().year
        all_staff = not (lsys and lsys.strip())

        if all_staff:
            join_cond = "INNER JOIN yggl ON jiaban.xm = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) != %s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)"
            join_param = (LEADER_EXCLUDE_LSYS,)
        else:
            join_cond = "INNER JOIN yggl ON jiaban.xm = yggl.name AND yggl.lsys = %s AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND (COALESCE(yggl.zaizhi,0)=0)"
            join_param = (lsys,)
        if month:
            month_str = f"{year}-{month:02d}"
            query = f"""
                SELECT TRIM(jiaban.xm) AS name, COUNT(1) AS times, SUM(CAST(COALESCE(jiaban.tian1, 0) AS DECIMAL(10,2))) AS hours
                FROM jiaban {join_cond}
                WHERE jiaban.jiabanzt = 4
                AND (jiaban.timedate LIKE %s OR SUBSTRING(jiaban.timedate, 1, 7) = %s)
                GROUP BY TRIM(jiaban.xm)
                ORDER BY hours DESC
            """
            rows = db.execute_query(query, join_param + (f"{month_str}%", month_str))
        else:
            if quarter:
                if quarter == "1":
                    mon_cond = "MONTH(jiaban.timedate) BETWEEN 1 AND 3"
                elif quarter == "2":
                    mon_cond = "MONTH(jiaban.timedate) BETWEEN 4 AND 6"
                elif quarter == "3":
                    mon_cond = "MONTH(jiaban.timedate) BETWEEN 7 AND 9"
                else:
                    mon_cond = "MONTH(jiaban.timedate) BETWEEN 10 AND 12"
                query = f"""
                    SELECT TRIM(jiaban.xm) AS name, COUNT(1) AS times, SUM(CAST(COALESCE(jiaban.tian1, 0) AS DECIMAL(10,2))) AS hours
                    FROM jiaban {join_cond}
                    WHERE jiaban.jiabanzt = 4
                    AND YEAR(jiaban.timedate) = %s AND {mon_cond}
                    GROUP BY TRIM(jiaban.xm)
                    ORDER BY hours DESC
                """
                rows = db.execute_query(query, join_param + (year,))
            else:
                query = f"""
                    SELECT TRIM(jiaban.xm) AS name, COUNT(1) AS times, SUM(CAST(COALESCE(jiaban.tian1, 0) AS DECIMAL(10,2))) AS hours
                    FROM jiaban {join_cond}
                    WHERE jiaban.jiabanzt = 4
                    AND (jiaban.timedate LIKE %s OR YEAR(jiaban.timedate) = %s)
                    GROUP BY TRIM(jiaban.xm)
                    ORDER BY hours DESC
                """
                rows = db.execute_query(query, join_param + (f"{year}%", year))

        ot_map = {}
        times_map = {}
        for r in rows:
            n = (r.get("name") or "").strip()
            if n:
                ot_map[n] = float(r.get("hours") or 0)
                times_map[n] = int(r.get("times") or 0)

        hx_map: Dict[str, float] = {}
        if net:
            hx_map = _query_hx_hours_all(year, month, quarter, lsys, all_staff)
            for n in hx_map:
                if n not in ot_map:
                    ot_map[n] = 0

        if net:
            all_names = sorted(ot_map, key=lambda k: -(ot_map[k] - hx_map.get(k, 0)))
        else:
            all_names = sorted(ot_map, key=lambda k: -ot_map[k])

        period_start, period_end = _dept_overtime_period_bounds(year, month, quarter)
        scope_names = _dept_overtime_scope_names(lsys)
        att_map: Dict[str, float] = {}
        if scope_names and period_start <= period_end and collect_valid_times_with_marks:
            att_map = _leader_style_overtime_hours_from_attendance(
                scope_names, period_start, period_end
            )
        auto_hours = round(
            sum(round(float(att_map.get(n, 0.0)), 2) for n in scope_names), 2
        )
        auto_person_count = sum(
            1 for n in scope_names if round(float(att_map.get(n, 0.0)), 2) > 0
        )

        list_data = []
        total_hours = 0
        total_times = 0
        for n in all_names:
            h = round(ot_map[n] - hx_map.get(n, 0), 2) if net else round(ot_map[n], 2)
            t = int(times_map.get(n, 0))
            list_data.append({
                "name": n,
                "hours": h,
                "times": t,
                "autoHours": round(float(att_map.get(n, 0.0)), 2),
            })
            total_hours += h
            total_times += t

        return {
            "success": True,
            "totalHours": round(total_hours, 2),
            "totalTimes": total_times,
            "personCount": len(list_data),
            "list": list_data,
            "autoCalculatedHours": auto_hours,
            "autoCalculatedPersonCount": auto_person_count,
        }
    except Exception as e:
        logger.error(f"加班科室统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def _apply_overtime_pay_scope(
    scope: Optional[str],
    current_user: Optional[str],
    scope_lsys: Optional[str],
    lsys: Optional[str],
    name: Optional[str],
) -> Tuple[Optional[str], Optional[str]]:
    """按其他绩效激励统计页权限 scope 强制修正 lsys/name。返回 (lsys, name)。"""
    if not current_user or not scope:
        return lsys, name
    if scope == "self":
        return None, (current_user or "").strip()
    if scope == "lsys" and scope_lsys:
        return (scope_lsys or "").strip(), None
    return lsys, name


def _resolve_jiaban_period_filter(
    year: Optional[int],
    month: Optional[int],
    date_from: Optional[str],
    date_to: Optional[str],
) -> Tuple[str, tuple, int, Optional[date], Optional[date]]:
    """
    解析 jiaban.timedate 筛选：年月 或 自定义闭区间 [date_from, date_to]。
    返回 (sql_and片段, 参数元组, 主年份用于缺省, range_start, range_end)。
    """
    if date_from or date_to:
        if not date_from or not date_to:
            raise HTTPException(status_code=400, detail="date_from 与 date_to 需同时传入")
        ds = _parse_date(date_from)
        de = _parse_date(date_to)
        if not ds or not de:
            raise HTTPException(status_code=400, detail="日期格式无效，请使用 YYYY-MM-DD")
        if ds > de:
            raise HTTPException(status_code=400, detail="开始日期不能晚于结束日期")
        cond = " AND DATE(jiaban.timedate) >= %s AND DATE(jiaban.timedate) <= %s"
        return cond, (ds.isoformat(), de.isoformat()), ds.year, ds, de

    y = year if year is not None else datetime.now().year
    if month is not None:
        month_key = f"{y}-{month:02d}"
        cond = (
            " AND (jiaban.timedate LIKE %s OR YEAR(jiaban.timedate) = %s)"
            " AND (MONTH(jiaban.timedate) = %s OR SUBSTRING(jiaban.timedate, 1, 7) = %s)"
        )
        return cond, (f"{y}%", y, month, month_key), y, None, None

    cond = " AND (jiaban.timedate LIKE %s OR YEAR(jiaban.timedate) = %s)"
    return cond, (f"{y}%", y), y, None, None


def _load_holiday_festival_map_span(
    year: int,
    range_start: Optional[date] = None,
    range_end: Optional[date] = None,
) -> Dict[str, str]:
    """加载节假日映射；跨年时合并各年假期表。"""
    if range_start and range_end:
        years = range(range_start.year, range_end.year + 1)
        merged: Dict[str, str] = {}
        for y in years:
            merged.update(_load_holiday_festival_map(y))
        return merged
    return _load_holiday_festival_map(year)


def _month_key_in_period(month_key: str, range_start: date, range_end: date) -> bool:
    """YYYY-MM 是否与 [range_start, range_end] 有交集。"""
    if not month_key or len(month_key) < 7:
        return False
    try:
        parts = month_key.split("-")
        y, m = int(parts[0]), int(parts[1])
        import calendar

        last_d = calendar.monthrange(y, m)[1]
        m_start = date(y, m, 1)
        m_end = date(y, m, last_d)
        return not (m_end < range_start or m_start > range_end)
    except (ValueError, IndexError):
        return False


@router.get("/dept/overtime-pay-by-month")
async def get_dept_overtime_pay_by_month(
    lsys: Optional[str] = Query(None, description="隶属于室，不传或空为全员"),
    year: Optional[int] = None,
    month: Optional[int] = Query(None, ge=1, le=12, description="筛选月份，不传为全年；与 date_from/date_to 互斥"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD"),
    name: Optional[str] = Query(None, description="仅查某人（如普通员工查本人）"),
    current_user: Optional[str] = Query(None, description="当前登录用户，与 scope 配合做权限过滤"),
    scope: Optional[str] = Query(None, description="可见范围：self=本人, lsys=本室, all=全部门"),
    scope_lsys: Optional[str] = Query(None, description="scope=lsys 时本室名称"),
):
    """
    其他绩效激励按月份统计。仅统计 jiaban 审核完成(jiabanzt=4)、换休票为否(hx 非「是」)，
    激励规则：
    - 若某天是假期表中节日为 春节/国庆节/高温防暑休假，且当天加班时长(已扣午休) >= 8 小时，则固定奖励 200 元，
      超出 8 小时部分按 zhibanfei 元/小时额外计算；
    - 其他日期或不足 8 小时部分，按 webconfig.zhibanfei（默认 15 元/小时）计算；
    支持 name=某人 仅查本人；month=1~12 仅查该月。
    当传入 current_user+scope 时按权限强制过滤：self 仅本人，lsys 仅本室，all 不限制。
    返回: { success, zhibanfei, list: [{ month, monthLabel, hours, pay }] }
    """
    try:
        lsys, name = _apply_overtime_pay_scope(scope, current_user, scope_lsys, lsys, name)
        date_cond, date_params, year, range_start, range_end = _resolve_jiaban_period_filter(
            year, month, date_from, date_to
        )
        zhibanfei = 15.0
        try:
            wc = db.execute_query("SELECT zhibanfei FROM webconfig WHERE id = 1 LIMIT 1")
            if wc and wc[0].get("zhibanfei") is not None:
                zhibanfei = float(wc[0]["zhibanfei"])
        except Exception:
            pass

        only_person = name and name.strip()

        if only_person:
            join_cond = "INNER JOIN yggl ON jiaban.xm = yggl.name AND jiaban.xm = %s AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND (COALESCE(yggl.zaizhi,0)=0)"
            join_param = (name.strip(),)
        else:
            all_staff = not (lsys and lsys.strip())
            if all_staff:
                join_cond = (
                    "INNER JOIN yggl ON jiaban.xm = yggl.name "
                    "AND RIGHT(TRIM(yggl.name), 1) != '1' "
                    "AND RIGHT(TRIM(yggl.lsys), 1) != '1' "
                    f"{_EXCL_OTHER_YGGL}{_EXCL_BUBAN_MANAGERS_YGGL}"
                    "AND (COALESCE(yggl.zaizhi,0)=0)"
                )
                join_param = ()
            else:
                join_cond = (
                    "INNER JOIN yggl ON jiaban.xm = yggl.name AND yggl.lsys = %s "
                    "AND RIGHT(TRIM(yggl.name), 1) != '1' "
                    "AND RIGHT(TRIM(yggl.lsys), 1) != '1' "
                    f"{_EXCL_BUBAN_MANAGERS_YGGL}"
                    "AND (COALESCE(yggl.zaizhi,0)=0)"
                )
                join_param = (lsys,)

        # 拉取原始加班记录（逐条），后续在 Python 中按人+日期聚合并应用激励规则
        query = f"""
            SELECT jiaban.xm AS emp_name,
                   jiaban.timedate,
                   CAST(COALESCE(jiaban.jbf, 0) AS DECIMAL(10,2)) AS hours
            FROM jiaban {join_cond}
            WHERE jiaban.jiabanzt = 4
              AND (jiaban.hx IS NULL OR TRIM(jiaban.hx) != '是'){date_cond}
        """
        rows = db.execute_query(query, join_param + date_params)

        holiday_map = _load_holiday_festival_map_span(year, range_start, range_end)
        per_month, _ = _aggregate_overtime_with_incentive(rows, holiday_map, zhibanfei)

        list_data = []
        # 若指定 month（非自定义区间），仅返回该月；自定义区间返回区间内各月
        for month_key, agg in sorted(per_month.items()):
            if month is not None and range_start is None and month_key != f"{year}-{month:02d}":
                continue
            if range_start and range_end and not _month_key_in_period(month_key, range_start, range_end):
                continue
            hours = round(agg["hours"], 2)
            pay = round(agg["pay"], 2)
            if month_key and len(month_key) == 7:
                month_label = f"{int(month_key.split('-')[1])}月"
            else:
                month_label = month_key or "-"
            list_data.append({"month": month_key, "monthLabel": month_label, "hours": hours, "pay": pay})

        return {"success": True, "zhibanfei": zhibanfei, "list": list_data}
    except Exception as e:
        logger.error(f"其他绩效激励按月考勤失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/dept/overtime-pay-by-employee")
async def get_dept_overtime_pay_by_employee(
    lsys: Optional[str] = Query(None, description="隶属于室，不传或空返回空列表（传 name 时可为空）"),
    year: Optional[int] = None,
    month: Optional[int] = Query(None, ge=1, le=12, description="筛选月份，不传为全年；与 date_from/date_to 互斥"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD"),
    name: Optional[str] = Query(None, description="仅查某人（如普通员工查本人）"),
    current_user: Optional[str] = Query(None, description="当前登录用户，与 scope 配合做权限过滤"),
    scope: Optional[str] = Query(None, description="可见范围：self=本人, lsys=本室, all=全部门"),
    scope_lsys: Optional[str] = Query(None, description="scope=lsys 时本室名称"),
):
    """
    科室员工其他绩效激励明细：指定科室、年份，按人汇总（审核通过且换休票为否）。
    激励规则与 /dept/overtime-pay-by-month 相同。
    支持 name=某人 仅查本人；month=1~12 仅查该月。
    当传入 current_user+scope 时按权限强制过滤。
    返回: { success, zhibanfei, list: [{ name, hours, pay }] }
    """
    try:
        lsys, name = _apply_overtime_pay_scope(scope, current_user, scope_lsys, lsys, name)
        date_cond, date_params, year, range_start, range_end = _resolve_jiaban_period_filter(
            year, month, date_from, date_to
        )
        zhibanfei = 15.0
        try:
            wc = db.execute_query("SELECT zhibanfei FROM webconfig WHERE id = 1 LIMIT 1")
            if wc and wc[0].get("zhibanfei") is not None:
                zhibanfei = float(wc[0]["zhibanfei"])
        except Exception:
            pass

        only_person = name and name.strip()

        if only_person:
            join_cond = "INNER JOIN yggl ON jiaban.xm = yggl.name AND jiaban.xm = %s AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND (COALESCE(yggl.zaizhi,0)=0)"
            params = (name.strip(),) + date_params
        else:
            if not lsys or not lsys.strip():
                return {"success": True, "zhibanfei": zhibanfei, "list": []}
            join_cond = (
                "INNER JOIN yggl ON jiaban.xm = yggl.name AND yggl.lsys = %s "
                "AND RIGHT(TRIM(yggl.name), 1) != '1' "
                "AND RIGHT(TRIM(yggl.lsys), 1) != '1' "
                f"{_EXCL_BUBAN_MANAGERS_YGGL}"
                "AND (COALESCE(yggl.zaizhi,0)=0)"
            )
            params = (lsys.strip(),) + date_params

        query = f"""
            SELECT jiaban.xm AS emp_name,
                   jiaban.timedate,
                   CAST(COALESCE(jiaban.jbf, 0) AS DECIMAL(10,2)) AS hours
            FROM jiaban {join_cond}
            WHERE jiaban.jiabanzt = 4
              AND (jiaban.hx IS NULL OR TRIM(jiaban.hx) != '是'){date_cond}
        """
        rows = db.execute_query(query, params)

        holiday_map = _load_holiday_festival_map_span(year, range_start, range_end)
        _, per_employee = _aggregate_overtime_with_incentive(rows, holiday_map, zhibanfei)

        list_data = []
        for emp_name, agg in per_employee.items():
            total_hours = round(agg["hours"], 2)
            pay = round(agg["pay"], 2)
            list_data.append({"name": emp_name, "hours": total_hours, "pay": pay})

        # 按加班小时降序、姓名排序
        list_data.sort(key=lambda x: (-x["hours"], x["name"]))

        return {"success": True, "zhibanfei": zhibanfei, "list": list_data}
    except Exception as e:
        logger.error(f"其他绩效激励按员工统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/dept/overtime-pay-export")
async def get_overtime_pay_export(
    year: Optional[int] = Query(None, description="年份（与 month 配合；自定义区间时可省略）"),
    month: Optional[int] = Query(None, ge=1, le=12, description="月份；与 date_from/date_to 互斥"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD"),
    current_user: Optional[str] = Query(None, description="当前登录用户，与 scope 配合做权限过滤"),
    scope: Optional[str] = Query(None, description="可见范围：self=本人, lsys=本室, all=全部门"),
    scope_lsys: Optional[str] = Query(None, description="scope=lsys 时本室名称"),
):
    """
    导出其他绩效激励工资报表数据：全员 + 各科室。
    指定 month 或 date_from+date_to 其一；以 yggl 名单为准，无记录者激励为 0。
    返回: { success, zhibanfei, all: [{ name, pay }], byDept: [{ lsys, list: [{ name, pay }] }] }
    """
    try:
        if not ((date_from and date_to) or (year is not None and month is not None)):
            raise HTTPException(
                status_code=400,
                detail="请指定 year+month，或同时传入 date_from 与 date_to",
            )
        date_cond, date_params, year, range_start, range_end = _resolve_jiaban_period_filter(
            year, month, date_from, date_to
        )

        zhibanfei = 15.0
        try:
            wc = db.execute_query("SELECT zhibanfei FROM webconfig WHERE id = 1 LIMIT 1")
            if wc and wc[0].get("zhibanfei") is not None:
                zhibanfei = float(wc[0]["zhibanfei"])
        except Exception:
            pass

        q_rows = f"""
            SELECT jiaban.xm AS emp_name,
                   jiaban.timedate,
                   CAST(COALESCE(jiaban.jbf, 0) AS DECIMAL(10,2)) AS hours
            FROM jiaban
            INNER JOIN yggl ON jiaban.xm = yggl.name
            WHERE jiaban.jiabanzt = 4
              AND (jiaban.hx IS NULL OR TRIM(jiaban.hx) != '是'){date_cond}
              AND RIGHT(TRIM(yggl.name), 1) != '1'
              AND RIGHT(TRIM(yggl.lsys), 1) != '1'
              AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员')
              AND NOT (TRIM(yggl.lsys) = '部办' AND TRIM(COALESCE(yggl.jb,'')) IN ('经理','副经理'))
              AND (COALESCE(yggl.zaizhi,0)=0)
        """
        rows = db.execute_query(q_rows, date_params)

        holiday_map = _load_holiday_festival_map_span(year, range_start, range_end)
        _, per_employee = _aggregate_overtime_with_incentive(rows, holiday_map, zhibanfei)

        # 先准备全员名单（含部办非经理人员），再按 per_employee 中的 pay 填值，保证人全
        yggl_rows = db.execute_query(
            "SELECT name, lsys FROM yggl WHERE lsys IS NOT NULL AND lsys != '' AND RIGHT(TRIM(lsys), 1) != '1' "
            f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
            "AND RIGHT(TRIM(name), 1) != '1' AND (COALESCE(zaizhi,0)=0)",
        )

        list_all = []
        for r in (yggl_rows or []):
            emp_name = (r.get("name") or "").strip()
            if not emp_name:
                continue
            agg = per_employee.get(emp_name, {"pay": 0.0, "hours": 0.0})
            pay = round(agg["pay"], 2)
            list_all.append({"name": emp_name, "pay": pay})

        # 科室列表（与 lsys-list 一致，部办仅含非经理人员）
        lsys_list = sorted({(r.get("lsys") or "").strip() for r in (yggl_rows or []) if r.get("lsys")})

        # 各科室：从全员名单中按 lsys 划分，同样用 per_employee 中的 pay
        by_dept = []
        for lsys in lsys_list:
            dept_list = []
            for r in (yggl_rows or []):
                emp_name = (r.get("name") or "").strip()
                emp_lsys = (r.get("lsys") or "").strip()
                if not emp_name or emp_lsys != lsys:
                    continue
                agg = per_employee.get(emp_name, {"pay": 0.0, "hours": 0.0})
                pay = round(agg["pay"], 2)
                dept_list.append({"name": emp_name, "pay": pay})
            by_dept.append({"lsys": lsys, "list": dept_list})

        # 按权限 scope 过滤导出结果
        if current_user and scope == "self":
            list_all = [x for x in list_all if (x.get("name") or "").strip() == (current_user or "").strip()]
            by_dept = []
        elif scope == "lsys" and scope_lsys:
            lsys_val = (scope_lsys or "").strip()
            by_dept = [x for x in by_dept if (x.get("lsys") or "").strip() == lsys_val]
            dept_names = {n.get("name") for d in by_dept for n in (d.get("list") or [])}
            list_all = [x for x in list_all if (x.get("name") or "") in dept_names]

        return {"success": True, "zhibanfei": zhibanfei, "all": list_all, "byDept": by_dept}
    except Exception as e:
        logger.error(f"其他绩效激励按月导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/dept/overtime-hours-export")
async def get_overtime_hours_export(
    year: Optional[int] = Query(None, description="年份；自定义区间时可仅依赖 date_from"),
    month: Optional[int] = Query(None, ge=1, le=12, description="月份，不传为全年；与 date_from/date_to 互斥"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD"),
    current_user: Optional[str] = Query(None, description="当前登录用户，与 scope 配合做权限过滤"),
    scope: Optional[str] = Query(None, description="可见范围：self=本人, lsys=本室, all=全部门"),
    scope_lsys: Optional[str] = Query(None, description="scope=lsys 时本室名称"),
):
    """
    导出全部加班时长（含其他绩效激励与换休票两类，审核通过 jiabanzt=4）。
    以 yggl 名单为准，无记录者各项为 0。
    返回: { success, all: [{ name, totalHours, payHours, hxHours, times }], byDept: [...] }
    """
    try:
        if not year and not (date_from and date_to):
            raise HTTPException(
                status_code=400,
                detail="请传入 year，或同时传入 date_from 与 date_to",
            )
        date_cond, date_params, year, range_start, range_end = _resolve_jiaban_period_filter(
            year, month, date_from, date_to
        )

        q_rows = f"""
            SELECT jiaban.xm AS emp_name,
                   CAST(COALESCE(NULLIF(TRIM(jiaban.tian1), ''), jiaban.jbf, 0) AS DECIMAL(10,2)) AS hours,
                   TRIM(COALESCE(jiaban.hx, '')) AS hx_flag
            FROM jiaban
            INNER JOIN yggl ON jiaban.xm = yggl.name
            WHERE jiaban.jiabanzt = 4
              {date_cond}
              AND RIGHT(TRIM(yggl.name), 1) != '1'
              AND RIGHT(TRIM(yggl.lsys), 1) != '1'
              AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员')
              AND NOT (TRIM(yggl.lsys) = '部办' AND TRIM(COALESCE(yggl.jb,'')) IN ('经理','副经理'))
              AND (COALESCE(yggl.zaizhi,0)=0)
        """
        rows = db.execute_query(q_rows, date_params)

        per_employee: Dict[str, Dict[str, float]] = defaultdict(
            lambda: {"totalHours": 0.0, "payHours": 0.0, "hxHours": 0.0, "times": 0}
        )
        for r in rows or []:
            emp_name = (r.get("emp_name") or "").strip()
            if not emp_name:
                continue
            try:
                h = float(r.get("hours") or 0)
            except (TypeError, ValueError):
                h = 0.0
            if h <= 0:
                continue
            agg = per_employee[emp_name]
            agg["totalHours"] += h
            agg["times"] += 1
            if (r.get("hx_flag") or "").strip() == "是":
                agg["hxHours"] += h
            else:
                agg["payHours"] += h

        yggl_rows = db.execute_query(
            "SELECT name, lsys FROM yggl WHERE lsys IS NOT NULL AND lsys != '' AND RIGHT(TRIM(lsys), 1) != '1' "
            f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
            "AND RIGHT(TRIM(name), 1) != '1' AND (COALESCE(zaizhi,0)=0)",
        )

        def _row_from_agg(emp_name: str) -> Dict[str, Any]:
            agg = per_employee.get(emp_name, {})
            return {
                "name": emp_name,
                "totalHours": round(float(agg.get("totalHours") or 0), 2),
                "payHours": round(float(agg.get("payHours") or 0), 2),
                "hxHours": round(float(agg.get("hxHours") or 0), 2),
                "times": int(agg.get("times") or 0),
            }

        list_all = []
        for r in yggl_rows or []:
            emp_name = (r.get("name") or "").strip()
            if emp_name:
                list_all.append(_row_from_agg(emp_name))

        lsys_list = sorted({(r.get("lsys") or "").strip() for r in (yggl_rows or []) if r.get("lsys")})
        by_dept = []
        for lsys in lsys_list:
            dept_list = []
            for r in yggl_rows or []:
                emp_name = (r.get("name") or "").strip()
                emp_lsys = (r.get("lsys") or "").strip()
                if emp_name and emp_lsys == lsys:
                    dept_list.append(_row_from_agg(emp_name))
            by_dept.append({"lsys": lsys, "list": dept_list})

        if current_user and scope == "self":
            list_all = [x for x in list_all if (x.get("name") or "").strip() == (current_user or "").strip()]
            by_dept = []
        elif scope == "lsys" and scope_lsys:
            lsys_val = (scope_lsys or "").strip()
            by_dept = [x for x in by_dept if (x.get("lsys") or "").strip() == lsys_val]
            dept_names = {n.get("name") for d in by_dept for n in (d.get("list") or [])}
            list_all = [x for x in list_all if (x.get("name") or "") in dept_names]

        return {"success": True, "all": list_all, "byDept": by_dept}
    except Exception as e:
        logger.error(f"全部加班时长导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 公出科室统计 ====================

@router.get("/dept/business-trip")
async def get_dept_business_trip_stats(
    lsys: Optional[str] = Query(None, description="隶属于室，不传或空为全员"),
    year: Optional[int] = None,
    month: Optional[int] = None,
    quarter: Optional[str] = None
):
    """
    科室公出统计（按人汇总人天）。不传 lsys 时为全员（排除部办）。
    与领导人看板「全体员工排序-公出」同一逻辑：仅已批准(bldzt=2, szrzt=2)，按区间并集计天数，不重复累加。
    """
    try:
        import calendar
        if year is None:
            year = __import__("datetime").datetime.now().year
        all_staff = not (lsys and lsys.strip())

        # 与 rankings 公出一致：拉取原始记录，仅已批准，再按区间并集算天数
        if all_staff:
            join_cond = """
                gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name
                AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) != %s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
            """
            join_param: tuple = (LEADER_EXCLUDE_LSYS,)
        else:
            join_cond = """
                gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name
                AND yggl.lsys = %s AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND (COALESCE(yggl.zaizhi,0)=0)
            """
            join_param = (lsys,)

        month_start = date(year, month, 1) if month else None
        month_end = date(year, month, calendar.monthrange(year, month)[1]) if month else None
        year_start = date(year, 1, 1)
        year_end = date(year, 12, 31)
        q_start, q_end = None, None
        if quarter:
            if quarter == "1":
                q_start, q_end = date(year, 1, 1), date(year, 3, 31)
            elif quarter == "2":
                q_start, q_end = date(year, 4, 1), date(year, 6, 30)
            elif quarter == "3":
                q_start, q_end = date(year, 7, 1), date(year, 9, 30)
            else:
                q_start, q_end = date(year, 10, 1), date(year, 12, 31)

        if month:
            trip_raw_query = f"""
                SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
                FROM {join_cond}
                WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                  AND COALESCE(gcsqb.gcsj, gcsqb.yjcfsj) <= %s AND COALESCE(gcsqb.sjfhtime, gcsqb.yjfhsj) >= %s
            """
            trip_rows = db.execute_query(trip_raw_query, join_param + (month_end.strftime("%Y-%m-%d"), month_start.strftime("%Y-%m-%d")))
        elif quarter and q_start and q_end:
            month_str_s = q_start.strftime("%Y-%m-%d")
            month_str_e = q_end.strftime("%Y-%m-%d")
            trip_raw_query = f"""
                SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
                FROM {join_cond}
                WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                  AND COALESCE(gcsqb.gcsj, gcsqb.yjcfsj) <= %s AND COALESCE(gcsqb.sjfhtime, gcsqb.yjfhsj) >= %s
            """
            trip_rows = db.execute_query(trip_raw_query, join_param + (month_str_e, month_str_s))
        else:
            month_str = f"{year}%"
            trip_raw_query = f"""
                SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
                FROM {join_cond}
                WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                  AND YEAR(COALESCE(gcsqb.gcsj, gcsqb.yjcfsj)) = %s
            """
            trip_rows = db.execute_query(trip_raw_query, join_param + (year,))

        by_person: dict = defaultdict(list)
        for row in trip_rows:
            gcr = (row.get("gcr") or "").strip()
            start_d = _parse_date(row.get("gcsj") or row.get("yjcfsj"))
            end_d = _parse_date(row.get("sjfhtime") or row.get("yjfhsj"))
            if start_d and end_d and end_d >= start_d:
                by_person[gcr].append((start_d, end_d))

        list_data = []
        today = date.today()
        effective_year_end = min(year_end, today)
        for gcr, intervals in by_person.items():
            if not intervals:
                continue
            if month and month_start and month_end:
                effective_month_end = min(month_end, today)
                clipped = [(max(s, month_start), min(e, effective_month_end)) for s, e in intervals if s <= effective_month_end and e >= month_start]
            elif quarter and q_start and q_end:
                effective_q_end = min(q_end, today)
                clipped = [(max(s, q_start), min(e, effective_q_end)) for s, e in intervals if s <= effective_q_end and e >= q_start]
            else:
                clipped = [(max(s, year_start), min(e, effective_year_end)) for s, e in intervals if s <= effective_year_end and e >= year_start]
            days = _merge_intervals_days(clipped)
            if days > 0:
                list_data.append({"name": gcr, "days": round(days, 2)})
        list_data.sort(key=lambda x: -x["days"])
        total_days = sum(d["days"] for d in list_data)

        return {
            "success": True,
            "totalDays": round(total_days, 2),
            "personCount": len(list_data),
            "list": list_data
        }
    except Exception as e:
        logger.error(f"公出科室统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 满勤判定辅助 ====================

def _to_comparable_dt(val) -> Optional[str]:
    """将 datetime / date / str 转为 19 位 'YYYY-MM-DD HH:MM:SS' 用于闭区间比较。"""
    if val is None:
        return None
    if hasattr(val, "strftime"):
        out = val.strftime("%Y-%m-%d %H:%M:%S") if hasattr(val, "hour") else val.strftime("%Y-%m-%d") + " 00:00:00"
        return out[:19]
    s = str(val).strip()
    if "." in s:
        s = s.split(".")[0]
    if len(s) >= 19:
        return s[:19]
    if len(s) == 10:
        return s + " 00:00:00"
    if len(s) == 16 and s[10] == " " and ":" in s[11:]:
        return s + ":00"
    return s


def _interval_covered_by(s_start: str, s_end: str, rows: list, get_start_end) -> bool:
    """建议区间 [s_start, s_end] 是否被 rows 中某条记录的区间包含。"""
    for r in rows:
        r_start, r_end = get_start_end(r)
        r_start = _to_comparable_dt(r_start)
        r_end = _to_comparable_dt(r_end)
        if r_start and r_end and r_start <= s_start and s_end <= r_end:
            return True
    return False


def _compute_full_attendance_info(names: List[str], year: int, month: Optional[int] = None) -> Tuple[set, Dict[str, int]]:
    """
    计算不满勤的人员姓名集合，以及满勤人员各自的公出天数。

    返回 (abnormal_set, gc_days_by_name):
      - abnormal_set: 不满勤人员姓名集合
      - gc_days_by_name: {姓名: 公出天数}，仅包含满勤且有公出的人员
    """
    empty = (set(), {})
    if not names:
        return empty

    if month:
        sugg_rows = db.execute_query(
            """SELECT employee_name, start_time, end_time
               FROM attendance_suggestions
               WHERE status = 1 AND year = %s AND month = %s""",
            (year, month)
        )
    else:
        sugg_rows = db.execute_query(
            """SELECT employee_name, start_time, end_time
               FROM attendance_suggestions
               WHERE status = 1 AND year = %s""",
            (year,)
        )

    if not sugg_rows:
        return empty

    per_person: Dict[str, list] = defaultdict(list)
    for r in sugg_rows:
        n = (r.get("employee_name") or "").strip()
        if n:
            per_person[n].append(r)

    name_set = set(names)
    relevant = {n: items for n, items in per_person.items() if n in name_set}
    if not relevant:
        return empty

    if month:
        month_start = f"{year}-{month:02d}-01"
        month_end = f"{year + 1}-01-01" if month == 12 else f"{year}-{month + 1:02d}-01"
    else:
        month_start = f"{year}-01-01"
        month_end = f"{year + 1}-01-01"

    rel_names = list(relevant.keys())
    ph = ",".join(["%s"] * len(rel_names))

    try:
        qj_rows = db.execute_query(
            f"SELECT xm, timefrom, timeto FROM qj "
            f"WHERE xm IN ({ph}) AND qjzt = 4 AND timefrom < %s AND timeto >= %s",
            tuple(rel_names) + (month_end, month_start),
        )
    except Exception:
        qj_rows = []

    try:
        gc_rows = db.execute_query(
            f"SELECT gcr AS xm, yjcfsj, yjfhsj, gcsj, sjfhtime FROM gcsqb "
            f"WHERE gcr IN ({ph}) AND bldzt = 2 AND szrzt = 2 "
            f"AND (yjcfsj IS NOT NULL OR yjfhsj IS NOT NULL) "
            f"AND COALESCE(yjcfsj, gcsj) < %s AND COALESCE(yjfhsj, sjfhtime, yjcfsj, gcsj) >= %s",
            tuple(rel_names) + (month_end, month_start),
        )
    except Exception:
        gc_rows = []

    qj_by_name: Dict[str, list] = defaultdict(list)
    for r in qj_rows:
        n = (r.get("xm") or "").strip()
        if n:
            qj_by_name[n].append(r)

    gc_by_name: Dict[str, list] = defaultdict(list)
    for r in gc_rows:
        n = (r.get("xm") or "").strip()
        if n:
            gc_by_name[n].append(r)

    abnormal = set()
    gc_dates_by_name: Dict[str, set] = defaultdict(set)
    for n, items in relevant.items():
        is_abnormal = False
        for s in items:
            s_start = _to_comparable_dt(s.get("start_time"))
            s_end = _to_comparable_dt(s.get("end_time"))
            if not s_start or not s_end:
                is_abnormal = True
                break
            covered_by_gc = _interval_covered_by(
                s_start, s_end, gc_by_name.get(n, []),
                lambda r: (r.get("yjcfsj") or r.get("gcsj"), r.get("yjfhsj") or r.get("sjfhtime")),
            )
            if covered_by_gc:
                gc_dates_by_name[n].add(s_start[:10])
                continue
            is_abnormal = True
            break
        if is_abnormal:
            abnormal.add(n)

    gc_days_by_name = {n: len(dates) for n, dates in gc_dates_by_name.items() if n not in abnormal}
    return abnormal, gc_days_by_name


def _compute_abnormal_set(names: List[str], year: int, month: Optional[int] = None) -> set:
    """计算不满勤的人员姓名集合（简便包装）。"""
    abnormal, _ = _compute_full_attendance_info(names, year, month)
    return abnormal


# ==================== 个人满勤查询 ====================

@router.get("/person/full-attendance")
async def get_person_full_attendance(
    name: str = Query(..., description="员工姓名"),
    year: int = Query(..., description="年份"),
    month: int = Query(..., ge=1, le=12, description="月份"),
):
    """
    查询指定员工某月是否满勤。
    满勤 = 当月 status=1 的考勤异常全部由已通过公出覆盖（或无异常）。
    返回: { success, isFull: bool }
    """
    try:
        name = (name or "").strip()
        if not name:
            return {"success": False, "isFull": False}
        abnormal = _compute_abnormal_set([name], year, month)
        return {"success": True, "isFull": name not in abnormal}
    except Exception as e:
        logger.error(f"个人满勤查询失败: {str(e)}")
        return {"success": False, "isFull": False}


# ==================== 领导人看板扩展 API ====================

@router.get("/leader/full-attendance")
async def get_leader_full_attendance(
    year: int = Query(..., description="年份"),
    month: int = Query(..., description="月份"),
    lsys: Optional[str] = Query(None, description="隶属科室，不传则全员")
):
    """
    满勤率：指定月份全员或指定科室的满勤率。
    满勤 = 当月 status=1 的考勤异常全部由已通过公出覆盖（或无异常）；有请假覆盖或未处理则不满勤。
    返回: workdays(当月应出勤工作日，仅作参考), totalPeople, fullCount, rate, fullNames(满勤人员姓名),
    byDept(仅当未传lsys时，各科室明细)
    """
    try:
        workdays = _count_workdays_in_month(year, month)
        month_str = f"{year}-{month:02d}"

        if lsys:
            rows = db.execute_query(
                "SELECT name FROM yggl WHERE lsys = %s AND name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
                (lsys,)
            )
            names = [r["name"].strip() for r in rows if r.get("name")]
        else:
            rows = db.execute_query(
                "SELECT name, lsys FROM yggl WHERE name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
            )
            names = [r["name"].strip() for r in rows if r.get("name")]

        if not names:
            return {
                "success": True,
                "workdays": workdays,
                "totalPeople": 0,
                "fullCount": 0,
                "rate": 0,
                "fullNames": [],
                "byDept": []
            }

        abnormal_set = _compute_abnormal_set(names, year, month)
        full_count = sum(1 for n in names if n not in abnormal_set)
        total = len(names)
        rate = round(full_count / total, 4) if total else 0
        full_names_all = sorted(n for n in names if n not in abnormal_set)

        result = {
            "success": True,
            "workdays": workdays,
            "totalPeople": total,
            "fullCount": full_count,
            "rate": rate,
            "fullNames": full_names_all,
            "byDept": []
        }

        if lsys is None and rows and len(rows) > 0 and "lsys" in (rows[0] or {}):
            dept_names = {}
            for r in rows:
                n = (r.get("name") or "").strip()
                d = (r.get("lsys") or "").strip()
                if not n:
                    continue
                dept_names.setdefault(d, []).append(n)
            by_dept = []
            for d, nlist in dept_names.items():
                fc = sum(1 for n in nlist if n not in abnormal_set)
                tot = len(nlist)
                full_list = sorted(n for n in nlist if n not in abnormal_set)
                by_dept.append({
                    "lsys": d,
                    "totalPeople": tot,
                    "fullCount": fc,
                    "rate": round(fc / tot, 4) if tot else 0,
                    "fullNames": full_list,
                })
            result["byDept"] = by_dept

        return result
    except Exception as e:
        logger.error(f"满勤率查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/leader/full-attendance-export")
async def get_leader_full_attendance_export(
    year: int = Query(..., description="年份"),
    month: Optional[int] = Query(None, description="月份，不传则全年"),
    lsys: Optional[str] = Query(None, description="隶属科室，不传则全员")
):
    """
    满勤名单导出：与领导人看板满勤统计同一逻辑（异常全部由公出覆盖仍算满勤，有请假或未处理则不满勤）。
    返回 byDept 中每项含 fullNames（满勤人员姓名列表），用于 Excel 等导出。
    """
    try:
        if month:
            workdays = _count_workdays_in_month(year, month)
        else:
            workdays = sum(_count_workdays_in_month(year, m) for m in range(1, 13))

        if lsys:
            rows = db.execute_query(
                "SELECT name FROM yggl WHERE lsys = %s AND name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
                (lsys,)
            )
            names = [r["name"].strip() for r in rows if r.get("name")]
        else:
            rows = db.execute_query(
                "SELECT name, lsys FROM yggl WHERE name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
            )
            names = [r["name"].strip() for r in rows if r.get("name")]

        if not names:
            return {"success": True, "workdays": workdays, "totalPeople": 0, "fullCount": 0, "rate": 0, "byDept": []}

        abnormal_set, gc_days_map = _compute_full_attendance_info(names, year, month)
        full_count = sum(1 for n in names if n not in abnormal_set)
        total = len(names)
        rate = round(full_count / total, 4) if total else 0

        def _build_details(name_list):
            details = []
            for n in sorted(name_list):
                gc = gc_days_map.get(n, 0)
                details.append({"name": n, "attendDays": workdays - gc, "businessDays": gc})
            return details

        by_dept = []
        if lsys:
            full_list = [n for n in names if n not in abnormal_set]
            by_dept.append({
                "lsys": lsys.strip(),
                "totalPeople": total,
                "fullCount": full_count,
                "rate": rate,
                "fullNames": sorted(full_list),
                "fullDetails": _build_details(full_list),
            })
        elif rows and len(rows) > 0 and "lsys" in (rows[0] or {}):
            dept_names = {}
            for r in rows:
                n = (r.get("name") or "").strip()
                d = (r.get("lsys") or "").strip()
                if not n:
                    continue
                dept_names.setdefault(d, []).append(n)
            for d, nlist in sorted(dept_names.items()):
                fc = sum(1 for n in nlist if n not in abnormal_set)
                tot = len(nlist)
                full_list = [n for n in nlist if n not in abnormal_set]
                by_dept.append({
                    "lsys": d,
                    "totalPeople": tot,
                    "fullCount": fc,
                    "rate": round(fc / tot, 4) if tot else 0,
                    "fullNames": sorted(full_list),
                    "fullDetails": _build_details(full_list),
                })

        return {
            "success": True,
            "workdays": workdays,
            "totalPeople": total,
            "fullCount": full_count,
            "rate": rate,
            "byDept": by_dept
        }
    except Exception as e:
        logger.error(f"满勤名单导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/leader/full-attendance-year")
async def get_leader_full_attendance_year(
    year: int = Query(..., description="年份"),
    lsys: Optional[str] = Query(None, description="隶属科室，不传则全员")
):
    """
    满勤率（全年）：指定年份全员或指定科室的全年满勤率。
    全年满勤 = 该年度内 status=1 异常全部由已通过公出覆盖（或无异常）。
    返回: totalPeople, fullCount, rate, fullNames(满勤人员姓名), byDept(仅当未传lsys时)，无 workdays。
    """
    try:
        year_prefix = f"{year}-"
        if lsys:
            rows = db.execute_query(
                "SELECT name FROM yggl WHERE lsys = %s AND name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
                (lsys,)
            )
            names = [r["name"].strip() for r in rows if r.get("name")]
        else:
            rows = db.execute_query(
                "SELECT name, lsys FROM yggl WHERE name IS NOT NULL AND name != '' "
                "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                "AND (COALESCE(zaizhi,0)=0)",
            )
            names = [r["name"].strip() for r in rows if r.get("name")]

        if not names:
            return {
                "success": True,
                "totalPeople": 0,
                "fullCount": 0,
                "rate": 0,
                "fullNames": [],
                "byDept": []
            }

        abnormal_set = _compute_abnormal_set(names, year)
        full_count = sum(1 for n in names if n not in abnormal_set)
        total = len(names)
        rate = round(full_count / total, 4) if total else 0
        full_names_all = sorted(n for n in names if n not in abnormal_set)

        result = {
            "success": True,
            "totalPeople": total,
            "fullCount": full_count,
            "rate": rate,
            "fullNames": full_names_all,
            "byDept": []
        }

        if lsys is None and rows and len(rows) > 0 and "lsys" in (rows[0] or {}):
            dept_names = {}
            for r in rows:
                n = (r.get("name") or "").strip()
                d = (r.get("lsys") or "").strip()
                if not n:
                    continue
                dept_names.setdefault(d, []).append(n)
            by_dept = []
            for d, nlist in dept_names.items():
                fc = sum(1 for n in nlist if n not in abnormal_set)
                tot = len(nlist)
                full_list = sorted(n for n in nlist if n not in abnormal_set)
                by_dept.append({
                    "lsys": d,
                    "totalPeople": tot,
                    "fullCount": fc,
                    "rate": round(fc / tot, 4) if tot else 0,
                    "fullNames": full_list,
                })
            result["byDept"] = by_dept

        return result
    except Exception as e:
        logger.error(f"全年满勤率查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/leader/full-attendance-by-month")
async def get_leader_full_attendance_by_month(
    year: int = Query(..., description="年份"),
    lsys: Optional[str] = Query(None, description="隶属科室，不传则全员")
):
    """
    按月考勤满勤人数：横轴月份，纵轴满勤人数，可筛选科室。
    满勤 = 当月异常全部由公出覆盖或无异常。返回 12 个月每月的 fullCount、totalPeople。
    返回: list[{ month, monthLabel, fullCount, totalPeople }]
    """
    try:
        list_data = []
        for month in range(1, 13):
            month_str = f"{year}-{month:02d}"
            if lsys:
                rows = db.execute_query(
                    "SELECT name FROM yggl WHERE lsys = %s AND name IS NOT NULL AND name != '' "
                    "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                    f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                    "AND (COALESCE(zaizhi,0)=0)",
                    (lsys,)
                )
                names = [r["name"].strip() for r in rows if r.get("name")]
            else:
                rows = db.execute_query(
                    "SELECT name FROM yggl WHERE name IS NOT NULL AND name != '' "
                    "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
                    f"{_EXCL_OTHER}{_EXCL_BUBAN_MANAGERS}"
                    "AND (COALESCE(zaizhi,0)=0)",
                )
                names = [r["name"].strip() for r in rows if r.get("name")]

            if not names:
                list_data.append({
                    "month": month,
                    "monthLabel": f"{month}月",
                    "fullCount": 0,
                    "totalPeople": 0
                })
                continue

            abnormal_set = _compute_abnormal_set(names, year, month)
            full_count = sum(1 for n in names if n not in abnormal_set)
            total = len(names)
            list_data.append({
                "month": month,
                "monthLabel": f"{month}月",
                "fullCount": full_count,
                "totalPeople": total
            })
        return {"success": True, "list": list_data}
    except Exception as e:
        logger.error(f"按月满勤人数查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/leader/dept-comparison")
async def get_leader_dept_comparison(
    year: int = Query(..., description="年份"),
    month: Optional[int] = Query(None, description="月份，不传则全年")
):
    """
    科室横向对比：各科室加班、请假、公出总数及人均。
    返回: list[{ lsys, personCount, overtimeTotal, leaveTotal, tripTotal, overtimePerCapita, leavePerCapita, tripPerCapita }]
    """
    try:
        import calendar

        month_str = f"{year}-{month:02d}%" if month else f"{year}%"
        month_cond_overtime = "AND (timedate LIKE %s OR SUBSTRING(timedate, 1, 7) = %s)" if month else "AND (timedate LIKE %s OR YEAR(timedate) = %s)"
        params_overtime = (month_str, month_str) if month else (month_str, year)

        if month:
            leave_period_start = date(year, month, 1)
            leave_period_end = date(year, month, calendar.monthrange(year, month)[1])
        else:
            leave_period_start = date(year, 1, 1)
            leave_period_end = date(year, 12, 31)
        leave_pe_str = leave_period_end.strftime("%Y-%m-%d")
        leave_ps_str = leave_period_start.strftime("%Y-%m-%d")
        leave_ov = _leave_overlap_sql_bounds()

        person_rows = db.execute_query(
            "SELECT lsys, COUNT(*) AS cnt FROM yggl WHERE lsys IS NOT NULL AND lsys != '' AND RIGHT(TRIM(lsys), 1) != '1' AND RIGHT(TRIM(name), 1) != '1' AND TRIM(lsys) != %s AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0) GROUP BY lsys ORDER BY lsys",
            (LEADER_EXCLUDE_LSYS,)
        )
        person_by_lsys = {r["lsys"].strip(): int(r.get("cnt") or 0) for r in person_rows if r.get("lsys")}

        leave_raw = f"""
            SELECT lsys, timefrom, timeto, timefromdate, CAST(tian AS DECIMAL(10,2)) AS tian
            FROM qj WHERE qjzt = 4 AND RIGHT(TRIM(xm), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' AND TRIM(lsys) != %s AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员')
            {leave_ov}
        """
        leave_rows = db.execute_query(leave_raw, (LEADER_EXCLUDE_LSYS, leave_pe_str, leave_ps_str))
        leave_by_lsys_acc: Dict[str, float] = defaultdict(float)
        for row in leave_rows or []:
            l = (row.get("lsys") or "").strip()
            if not l:
                continue
            ls_d, le_d = _qj_leave_date_bounds(row)
            if not ls_d or not le_d:
                continue
            try:
                tian_f = float(row.get("tian") or 0)
            except (TypeError, ValueError):
                tian_f = 0.0
            alloc = _allocate_leave_tian_to_period(tian_f, ls_d, le_d, leave_period_start, leave_period_end)
            if alloc > 0:
                leave_by_lsys_acc[l] += alloc
        leave_by_lsys = {k: round(v, 2) for k, v in leave_by_lsys_acc.items()}

        overtime_query = f"""
            SELECT yggl.lsys, SUM(CAST(COALESCE(jiaban.tian1, 0) AS DECIMAL(10,2))) AS total
            FROM jiaban INNER JOIN yggl ON jiaban.xm = yggl.name AND yggl.lsys IS NOT NULL AND yggl.lsys != '' AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
            WHERE jiaban.jiabanzt = 4 {month_cond_overtime}
            GROUP BY yggl.lsys
        """
        ot_rows = db.execute_query(overtime_query, params_overtime)
        overtime_by_lsys = {r["lsys"].strip(): round(float(r.get("total") or 0), 2) for r in ot_rows if r.get("lsys")}

        # 公出与「全体员工排序」一致：仅已批准(bldzt=2, szrzt=2)，按区间并集计天数后按科室汇总
        month_start = date(year, month, 1) if month else None
        month_end = date(year, month, calendar.monthrange(year, month)[1]) if month else None
        year_start = date(year, 1, 1)
        year_end = date(year, 12, 31)
        if month:
            trip_raw_query = """
                SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj, yggl.lsys
                FROM gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) != %s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
                WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                  AND COALESCE(gcsqb.gcsj, gcsqb.yjcfsj) <= %s AND COALESCE(gcsqb.sjfhtime, gcsqb.yjfhsj) >= %s
            """
            trip_rows = db.execute_query(trip_raw_query, (LEADER_EXCLUDE_LSYS, month_end.strftime("%Y-%m-%d"), month_start.strftime("%Y-%m-%d")))
        else:
            trip_raw_query = """
                SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj, yggl.lsys
                FROM gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) != %s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
                WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                  AND YEAR(COALESCE(gcsqb.gcsj, gcsqb.yjcfsj)) = %s
            """
            trip_rows = db.execute_query(trip_raw_query, (LEADER_EXCLUDE_LSYS, year))
        by_person_trip: dict = defaultdict(list)
        for row in trip_rows:
            gcr = (row.get("gcr") or "").strip()
            lsys = (row.get("lsys") or "").strip()
            if not lsys:
                continue
            start_d = _parse_date(row.get("gcsj") or row.get("yjcfsj"))
            end_d = _parse_date(row.get("sjfhtime") or row.get("yjfhsj"))
            if start_d and end_d and end_d >= start_d:
                by_person_trip[(gcr, lsys)].append((start_d, end_d))
        trip_by_lsys = defaultdict(float)
        today = date.today()
        effective_year_end = min(year_end, today)
        for (gcr, lsys), intervals in by_person_trip.items():
            if not intervals:
                continue
            if month and month_start and month_end:
                effective_month_end = min(month_end, today)
                clipped = [(max(s, month_start), min(e, effective_month_end)) for s, e in intervals if s <= effective_month_end and e >= month_start]
            else:
                clipped = [(max(s, year_start), min(e, effective_year_end)) for s, e in intervals if s <= effective_year_end and e >= year_start]
            days = _merge_intervals_days(clipped)
            if days > 0:
                trip_by_lsys[lsys] += round(days, 2)
        trip_by_lsys = dict(trip_by_lsys)

        # 按科室汇总换休请假小时（与 /dept/overtime net 一致：区间重叠比例 × 小时）
        hx_raw = f"""
            SELECT yggl.lsys, qj.timefrom, qj.timeto, qj.timefromdate, CAST(qj.tian AS DECIMAL(10,2)) AS tian, qj.xiaoshi
            FROM qj INNER JOIN yggl ON qj.xm = yggl.name AND yggl.lsys IS NOT NULL AND yggl.lsys != ''
                AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
            WHERE qj.qjzt = 4 AND TRIM(qj.qjfs) IN ('换休','员工换休票') AND RIGHT(TRIM(qj.xm), 1) != '1'
                AND TRIM(yggl.lsys) != %s
            {leave_ov}
        """
        hx_rows = db.execute_query(hx_raw, (LEADER_EXCLUDE_LSYS, leave_pe_str, leave_ps_str))
        hx_by_lsys_acc: Dict[str, float] = defaultdict(float)
        for row in hx_rows or []:
            lsys_k = (row.get("lsys") or "").strip()
            if not lsys_k or lsys_k == LEADER_EXCLUDE_LSYS:
                continue
            ls_d, le_d = _qj_leave_date_bounds(row)
            if not ls_d or not le_d:
                continue
            h_full = _qj_hx_row_total_hours(row)
            if h_full <= 0:
                continue
            frac = _leave_overlap_fraction(ls_d, le_d, leave_period_start, leave_period_end)
            if frac <= 0:
                continue
            hx_by_lsys_acc[lsys_k] += h_full * frac
        hx_by_lsys = {k: round(v, 2) for k, v in hx_by_lsys_acc.items()}

        if month:
            workdays = _count_workdays_in_month(year, month)
        else:
            workdays = sum(_count_workdays_in_month(year, m) for m in range(1, 13))

        all_lsys = sorted(person_by_lsys.keys())
        list_data = []
        for l in all_lsys:
            pc = person_by_lsys.get(l, 0)
            ot = overtime_by_lsys.get(l, 0)
            lv = leave_by_lsys.get(l, 0)
            tr = trip_by_lsys.get(l, 0)
            hx = hx_by_lsys.get(l, 0)
            net_ot = round(ot - hx, 2)
            effective_pd = workdays * pc - tr
            list_data.append({
                "lsys": l,
                "personCount": pc,
                "workdays": workdays,
                "overtimeTotal": ot,
                "netOvertimeTotal": net_ot,
                "leaveTotal": lv,
                "tripTotal": tr,
                "overtimePerCapita": round(ot / pc, 2) if pc else 0,
                "netOvertimePerCapita": round(net_ot / pc, 2) if pc else 0,
                "overtimePerWorkday": round(ot / effective_pd, 2) if effective_pd > 0 else 0,
                "netOvertimePerWorkday": round(net_ot / effective_pd, 2) if effective_pd > 0 else 0,
                "leavePerCapita": round(lv / pc, 2) if pc else 0,
                "tripPerCapita": round(tr / pc, 2) if pc else 0
            })
        return {"success": True, "list": list_data}
    except Exception as e:
        logger.error(f"科室对比查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/leader/rankings")
async def get_leader_rankings(
    year: int = Query(..., description="年份"),
    month: Optional[int] = Query(None, description="月份，不传则全年"),
    type_: str = Query("overtime", alias="type", description="overtime|leave|trip")
):
    """
    全体员工排序：按加班小时/请假天数/公出天数排序。
    返回: list[{ rank, name, lsys, value, unit }]
    """
    try:
        import calendar

        month_str = f"{year}-{month:02d}%" if month else f"{year}%"
        month_param = (month_str, month_str) if month else (month_str, year)

        rows = []
        unit = ""

        if type_ == "overtime":
            query = """
                SELECT jiaban.xm AS name, yggl.lsys,
                    SUM(CAST(COALESCE(jiaban.tian1, 0) AS DECIMAL(10,2))) AS value
                FROM jiaban INNER JOIN yggl ON jiaban.xm = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
                WHERE jiaban.jiabanzt = 4 AND RIGHT(TRIM(jiaban.xm), 1) != '1' AND (jiaban.timedate LIKE %s OR YEAR(jiaban.timedate) = %s)
                GROUP BY jiaban.xm, yggl.lsys ORDER BY value DESC
            """
            unit = "小时"
            rows = db.execute_query(query, month_param)
        elif type_ == "leave":
            if month:
                lp_start = date(year, month, 1)
                lp_end = date(year, month, calendar.monthrange(year, month)[1])
            else:
                lp_start = date(year, 1, 1)
                lp_end = date(year, 12, 31)
            pe_str = lp_end.strftime("%Y-%m-%d")
            ps_str = lp_start.strftime("%Y-%m-%d")
            lov = _leave_overlap_sql_bounds()
            leave_raw = f"""
                SELECT qj.xm AS name, qj.lsys, timefrom, timeto, timefromdate, CAST(qj.tian AS DECIMAL(10,2)) AS tian
                FROM qj WHERE qj.qjzt = 4 AND RIGHT(TRIM(qj.xm), 1) != '1' AND RIGHT(TRIM(qj.lsys), 1) != '1' AND TRIM(qj.lsys) NOT IN ('其他部门员工','其他部门成员')
                {lov}
            """
            lrows = db.execute_query(leave_raw, (pe_str, ps_str))
            acc_lv: Dict[Tuple[str, str], float] = defaultdict(float)
            for row in lrows or []:
                nm = (row.get("name") or "").strip()
                ls = (row.get("lsys") or "").strip()
                if not nm:
                    continue
                ls_d, le_d = _qj_leave_date_bounds(row)
                if not ls_d or not le_d:
                    continue
                try:
                    tian_f = float(row.get("tian") or 0)
                except (TypeError, ValueError):
                    tian_f = 0.0
                alloc = _allocate_leave_tian_to_period(tian_f, ls_d, le_d, lp_start, lp_end)
                if alloc > 0:
                    acc_lv[(nm, ls)] += alloc
            rows = [
                {"name": k[0], "lsys": k[1], "value": round(v, 2)}
                for k, v in sorted(acc_lv.items(), key=lambda x: -x[1])
                if v > 0
            ]
            unit = "天"
        elif type_ == "trip":
            # 公出天数按区间并集计算，避免重复申报导致超过 365 天
            month_start = date(year, month, 1) if month else None
            month_end = date(year, month, calendar.monthrange(year, month)[1]) if month else None
            if month:
                trip_raw_query = """
                    SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj, yggl.lsys
                    FROM gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
                    WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                      AND COALESCE(gcsqb.gcsj, gcsqb.yjcfsj) <= %s AND COALESCE(gcsqb.sjfhtime, gcsqb.yjfhsj) >= %s
                """
                trip_rows = db.execute_query(trip_raw_query, (month_end.strftime("%Y-%m-%d"), month_start.strftime("%Y-%m-%d")))
            else:
                trip_raw_query = """
                    SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj, yggl.lsys
                    FROM gcsqb INNER JOIN yggl ON gcsqb.gcr = yggl.name AND RIGHT(TRIM(yggl.name), 1) != '1' AND RIGHT(TRIM(yggl.lsys), 1) != '1' AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)
                    WHERE RIGHT(TRIM(gcsqb.gcr), 1) != '1' AND (gcsqb.bldzt = 2 AND gcsqb.szrzt = 2)
                      AND YEAR(COALESCE(gcsqb.gcsj, gcsqb.yjcfsj)) = %s
                """
                trip_rows = db.execute_query(trip_raw_query, (year,))
            # 按 (gcr, lsys) 分组，每人收集 [start, end] 区间后做并集再算天数
            by_person: dict = defaultdict(list)
            for row in trip_rows:
                gcr = (row.get("gcr") or "").strip()
                lsys = (row.get("lsys") or "").strip()
                start_d = _parse_date(row.get("gcsj") or row.get("yjcfsj"))
                end_d = _parse_date(row.get("sjfhtime") or row.get("yjfhsj"))
                if start_d and end_d and end_d >= start_d:
                    by_person[(gcr, lsys)].append((start_d, end_d))
            list_trip = []
            year_start = date(year, 1, 1)
            year_end = date(year, 12, 31)
            today = date.today()
            effective_year_end = min(year_end, today)
            for (gcr, lsys), intervals in by_person.items():
                if not intervals:
                    continue
                if month and month_start and month_end:
                    effective_month_end = min(month_end, today)
                    clipped = [(max(s, month_start), min(e, effective_month_end)) for s, e in intervals if s <= effective_month_end and e >= month_start]
                    days = _merge_intervals_days(clipped)
                else:
                    clipped = [(max(s, year_start), min(e, effective_year_end)) for s, e in intervals if s <= effective_year_end and e >= year_start]
                    days = _merge_intervals_days(clipped)
                if days > 0:
                    list_trip.append({"name": gcr, "lsys": lsys, "value": round(days, 2)})
            list_trip.sort(key=lambda x: -x["value"])
            rows = list_trip
            unit = "天"
        else:
            raise HTTPException(status_code=400, detail="type 须为 overtime|leave|trip")

        list_data = []
        for i, r in enumerate(rows, 1):
            list_data.append({
                "rank": i,
                "name": (r.get("name") or "").strip(),
                "lsys": (r.get("lsys") or "").strip(),
                "value": round(float(r.get("value") or 0), 2),
                "unit": unit
            })
        return {"success": True, "list": list_data, "unit": unit}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"全员排序查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 考勤纪律审查 ====================

def _parse_time_str(val) -> Optional[str]:
    """将 DB 返回的时间字段转为 HH:MM:SS 字符串，兼容 timedelta / datetime / str。"""
    if val is None:
        return None
    if hasattr(val, "total_seconds"):
        total = int(val.total_seconds())
        h, rem = divmod(total, 3600)
        m, s = divmod(rem, 60)
        return f"{h:02d}:{m:02d}:{s:02d}"
    if hasattr(val, "strftime"):
        return val.strftime("%H:%M:%S")
    s = str(val).strip()
    if not s:
        return None
    if " " in s:
        s = s.split()[-1]
    parts = s.split(":")
    if len(parts) >= 3:
        return f"{parts[0].zfill(2)}:{parts[1].zfill(2)}:{parts[2].zfill(2)}"
    if len(parts) == 2:
        return f"{parts[0].zfill(2)}:{parts[1].zfill(2)}:00"
    return None


def _time_in_range(t_str: str, lo: str, hi: str) -> bool:
    """判断 HH:MM:SS 格式的时间是否在 [lo, hi] 闭区间内（进一原则：超过整分钟的秒数归入下一分钟）。"""
    if len(lo) == 5:
        lo = lo + ":00"
    if len(hi) == 5:
        hi = hi + ":00"
    return lo <= t_str <= hi


def _get_last_time(row: dict) -> Optional[str]:
    """取一行考勤记录中最后一个有效打卡时间（time_10 → time_1 倒序查找）。"""
    for i in range(10, 0, -1):
        val = row.get(f"time_{i}")
        t = _parse_time_str(val)
        if t:
            return t
    return None


def _get_first_time(row: dict) -> Optional[str]:
    """Return the first valid punch time in a merged attendance row."""
    for i in range(1, 11):
        t = _parse_time_str(row.get(f"time_{i}"))
        if t:
            return t
    return None


def _count_valid_times(row: dict) -> int:
    return sum(1 for i in range(1, 11) if _parse_time_str(row.get(f"time_{i}")))


def _rate(numerator: int, denominator: int) -> float:
    if denominator <= 0:
        return 0.0
    return round(numerator / denominator, 4)


def _load_non_workday_set(start_date: str, end_date: str) -> set:
    """
    返回 start_date ~ end_date 之间所有非工作日的日期字符串集合。
    非工作日 = 周末（周六日）+ 法定假日（holiday 表 type 含'假'或'休'），
    但 holiday 表 type 含'班'的调休上班日排除在外（视为工作日）。
    """
    from utils.holiday_loader import load_holidays_dict
    import calendar as _cal

    years = set()
    try:
        years.add(int(start_date[:4]))
        years.add(int(end_date[:4]))
    except (ValueError, IndexError):
        pass

    holidays: dict = {}
    for y in years:
        holidays.update(load_holidays_dict(str(y)))

    non_work = set()
    try:
        from datetime import timedelta
        cur = datetime.strptime(start_date, "%Y-%m-%d")
        end = datetime.strptime(end_date, "%Y-%m-%d")
        while cur <= end:
            ds = cur.strftime("%Y-%m-%d")
            is_weekend = cur.weekday() in (5, 6)
            ht = holidays.get(ds, "")
            is_holiday_off = ("假" in ht or "休" in ht)
            is_makeup_work = ("班" in ht)
            if is_makeup_work:
                pass
            elif is_weekend or is_holiday_off:
                non_work.add(ds)
            cur += timedelta(days=1)
    except Exception as e:
        logger.warning(f"构建非工作日集合失败: {e}")
    return non_work


_DISCIPLINE_MINUTE_WHITELIST = frozenset((2, 3, 4, 5, 10, 20, 30, 60))


def _normalize_discipline_minutes(v, default: int = 2) -> int:
    """踩点阈值仅允许白名单分钟数，非法值回退为 default。"""
    try:
        n = int(v)
    except (TypeError, ValueError):
        return default
    return n if n in _DISCIPLINE_MINUTE_WHITELIST else default


@router.get("/discipline/clock-in-stats")
async def get_clock_in_discipline_stats(
    year: int = Query(...),
    month: Optional[int] = Query(None),
    lsys: Optional[str] = Query(None),
    dimension: str = Query("person", description="聚合维度: person / month / dept"),
    clock_in_minutes: int = Query(2, description="踩点上班阈值：8:00 前 N 分钟，可选 2/3/4/5/10/20/30/60"),
    clock_out_minutes: int = Query(2, description="踩点下班阈值：17:00 后 N 分钟，可选 2/3/4/5/10/20/30/60"),
    exclude_holidays: bool = Query(False, description="是否排除节假日（周末+法定假日）"),
):
    """
    打卡纪律大数据检测。
    统计 attendance_records 中踩点上班与踩点下班的情况。
    clock_in_minutes: 8:00前N分钟为踩点上班区间，默认2（即7:58-8:00）
    clock_out_minutes: 17:00后N分钟为踩点下班区间，默认2（即17:00-17:02）
    """
    try:
        ci_min = _normalize_discipline_minutes(clock_in_minutes)
        co_min = _normalize_discipline_minutes(clock_out_minutes)

        ci_start_hour, ci_start_min = divmod(60 * 8 - ci_min, 60)
        ci_lo = f"{ci_start_hour:02d}:{ci_start_min:02d}"
        ci_hi = "08:00"
        co_lo = "17:00"
        co_end_hour, co_end_min = divmod(60 * 17 + co_min, 60)
        co_hi = f"{co_end_hour:02d}:{co_end_min:02d}"

        date_lo = f"{year}-01-01"
        date_hi = f"{year}-12-31"
        if month:
            date_lo = f"{year}-{month:02d}-01"
            if month == 12:
                date_hi = f"{year}-12-31"
            else:
                date_hi = f"{year}-{month + 1:02d}-01"

        sql = (
            "SELECT a.employee_name, a.department, a.attendance_date, "
            "a.time_1, a.time_2, a.time_3, a.time_4, a.time_5, "
            "a.time_6, a.time_7, a.time_8, a.time_9, a.time_10 "
            "FROM attendance_records a "
            "WHERE a.attendance_date >= %s AND a.attendance_date < %s "
        )
        params: list = [date_lo, date_hi]
        if lsys:
            sql += "AND a.department = %s "
            params.append(lsys)
        sql += "ORDER BY a.attendance_date, a.employee_name"

        rows = db.execute_query(sql, tuple(params))

        valid_names_rows = db.execute_query(
            "SELECT name, lsys FROM yggl WHERE name IS NOT NULL AND name != '' "
            "AND RIGHT(TRIM(name), 1) != '1' AND RIGHT(TRIM(lsys), 1) != '1' "
            "AND TRIM(lsys) != %s AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0)",
            (LEADER_EXCLUDE_LSYS,),
        )
        valid_names = {r["name"].strip() for r in valid_names_rows if r.get("name")}
        name_dept_map = {}
        for r in valid_names_rows:
            n = (r.get("name") or "").strip()
            if n:
                name_dept_map[n] = (r.get("lsys") or "").strip()

        non_workdays = _load_non_workday_set(date_lo, date_hi) if exclude_holidays else set()

        clock_in_data = []
        clock_out_data = []

        for row in rows or []:
            name = (row.get("employee_name") or "").strip()
            if not name or name not in valid_names:
                continue
            dept = name_dept_map.get(name, (row.get("department") or "").strip())
            date_str = str(row.get("attendance_date") or "")[:10]
            if date_str in non_workdays:
                continue
            month_key = date_str[:7]

            first_time = _parse_time_str(row.get("time_1"))
            last_time = _get_last_time(row)

            if first_time and _time_in_range(first_time, ci_lo, ci_hi):
                clock_in_data.append({"name": name, "dept": dept, "date": date_str, "month": month_key, "time": first_time})

            if last_time and _time_in_range(last_time, co_lo, co_hi):
                clock_out_data.append({"name": name, "dept": dept, "date": date_str, "month": month_key, "time": last_time})

        def _aggregate(data_list, dim):
            agg = defaultdict(lambda: {"count": 0, "dates": []})
            for item in data_list:
                if dim == "person":
                    key = item["name"]
                elif dim == "month":
                    key = item["month"]
                else:
                    key = item["dept"]
                agg[key]["count"] += 1
                if len(agg[key]["dates"]) < 50:
                    agg[key]["dates"].append({"date": item["date"], "name": item["name"], "time": item["time"]})
            result = []
            for k, v in agg.items():
                entry = {"key": k, "count": v["count"], "dates": v["dates"]}
                if dim == "person":
                    entry["dept"] = name_dept_map.get(k, "")
                result.append(entry)
            result.sort(key=lambda x: -x["count"])
            return result

        return {
            "success": True,
            "clockIn": _aggregate(clock_in_data, dimension),
            "clockOut": _aggregate(clock_out_data, dimension),
            "clockInTotal": len(clock_in_data),
            "clockOutTotal": len(clock_out_data),
            "clockInRange": f"{ci_lo}-{ci_hi}",
            "clockOutRange": f"{co_lo}-{co_hi}",
        }
    except Exception as e:
        logger.error(f"打卡纪律统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/discipline/person-scatter")
async def get_person_scatter(
    name: str = Query(..., description="员工姓名"),
    start_date: str = Query(..., description="起始日期 YYYY-MM-DD"),
    end_date: str = Query(..., description="结束日期 YYYY-MM-DD"),
    exclude_holidays: bool = Query(False, description="是否排除节假日"),
):
    """
    获取某员工在日期区间内的逐日上班/下班打卡时间，用于散点图展示。
    返回 { success, name, data: [{ date, clockIn, clockOut }] }
    clockIn / clockOut 为小时浮点数（如 7.97 = 7:58:12），null 表示当天无记录。
    """
    try:
        sql = (
            "SELECT attendance_date, "
            "time_1, time_2, time_3, time_4, time_5, "
            "time_6, time_7, time_8, time_9, time_10 "
            "FROM attendance_records "
            "WHERE employee_name = %s AND attendance_date >= %s AND attendance_date <= %s "
            "ORDER BY attendance_date"
        )
        rows = db.execute_query(sql, (name, start_date, end_date))
        non_workdays = _load_non_workday_set(start_date, end_date) if exclude_holidays else set()

        def _to_hours(val):
            t = _parse_time_str(val)
            if not t:
                return None
            parts = t.split(":")
            h = int(parts[0])
            m = int(parts[1])
            s = int(parts[2]) if len(parts) > 2 else 0
            return round(h + m / 60 + s / 3600, 4)

        data = []
        for row in rows or []:
            d = str(row.get("attendance_date") or "")[:10]
            if len(d) < 10:
                continue
            if d in non_workdays:
                continue
            ci = _to_hours(row.get("time_1"))
            last_raw = None
            for i in range(10, 0, -1):
                v = row.get(f"time_{i}")
                if v is not None and _parse_time_str(v):
                    last_raw = v
                    break
            co = _to_hours(last_raw)
            if ci is None and co is None:
                continue
            data.append({"date": d, "clockIn": ci, "clockOut": co})

        return {"success": True, "name": name, "data": data}
    except Exception as e:
        logger.error(f"个人散点图数据查询失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
#  工作强度统计：口径 A=加班/在岗；口径 B=（加班−请假h）/在岗
#  加班=智能建议同款打卡识别；公出仅境内/境外（不含市内；gclx 空视同境内）
# ============================================================

@router.get("/discipline/holiday-duty-attendance")
async def get_holiday_duty_attendance(
    name: str = Query(..., description="当前登录用户姓名，用于鉴权"),
    start_date: str = Query(..., description="Start date, YYYY-MM-DD"),
    end_date: str = Query(..., description="End date, YYYY-MM-DD"),
    lsys: Optional[str] = Query(None, description="Department filter"),
):
    """
    Check duty attendance by shift schedule.

    Day shift: 08:00-17:00. Night shift: 17:00-22:00.
    The first punch is used as arrival time and the last punch as leave time.
    For night shift with only one same-day punch, leave time is treated as 24:00.
    """
    try:
        if not _can_access_holiday_duty_attendance(name):
            raise HTTPException(status_code=403, detail="无假期值班出勤核查权限")
        try:
            d0 = datetime.strptime(str(start_date)[:10], "%Y-%m-%d").date()
            d1 = datetime.strptime(str(end_date)[:10], "%Y-%m-%d").date()
        except ValueError:
            raise HTTPException(status_code=400, detail="日期格式应为 YYYY-MM-DD")
        if d1 < d0:
            raise HTTPException(status_code=400, detail="结束日期不能早于开始日期")
        if (d1 - d0).days > 92:
            raise HTTPException(status_code=400, detail="单次最多查询 93 天")

        expected_ranges = {
            "白班": ("08:00:00", "17:00:00"),
            "夜班": ("17:00:00", "22:00:00"),
        }

        sql = """
            SELECT
                s.department,
                s.employee_name,
                s.shift_date,
                s.shift_type,
                a.time_1, a.time_2, a.time_3, a.time_4, a.time_5,
                a.time_6, a.time_7, a.time_8, a.time_9, a.time_10
            FROM shift_schedule s
            INNER JOIN yggl y
                ON TRIM(y.name) = TRIM(s.employee_name)
               AND TRIM(y.lsys) = TRIM(s.department)
               AND COALESCE(y.zaizhi, 0) = 0
               AND RIGHT(TRIM(y.name), 1) != '1'
               AND RIGHT(TRIM(y.lsys), 1) != '1'
               AND TRIM(y.lsys) != %s
               AND TRIM(y.lsys) NOT IN ('其他部门员工','其他部门成员')
            LEFT JOIN attendance_records a
                ON TRIM(a.employee_name) = TRIM(s.employee_name)
               AND DATE(a.attendance_date) = s.shift_date
            WHERE s.shift_date >= %s
              AND s.shift_date <= %s
              AND s.shift_type IN (%s, %s)
        """
        params: list = [
            LEADER_EXCLUDE_LSYS,
            d0.strftime("%Y-%m-%d"),
            d1.strftime("%Y-%m-%d"),
            "白班",
            "夜班",
        ]
        if lsys:
            sql += " AND s.department = %s"
            params.append(lsys.strip())
        sql += " ORDER BY s.shift_date, s.department, s.employee_name"

        rows = db.execute_query(sql, tuple(params))

        staff_sql = """
            SELECT TRIM(lsys) AS dept, COUNT(DISTINCT TRIM(name)) AS cnt
            FROM yggl
            WHERE name IS NOT NULL
              AND TRIM(name) != ''
              AND lsys IS NOT NULL
              AND TRIM(lsys) != ''
              AND COALESCE(zaizhi, 0) = 0
              AND RIGHT(TRIM(name), 1) != '1'
              AND RIGHT(TRIM(lsys), 1) != '1'
              AND TRIM(lsys) != %s
              AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员')
        """
        staff_params: list = [LEADER_EXCLUDE_LSYS]
        if lsys:
            staff_sql += " AND TRIM(lsys) = %s"
            staff_params.append(lsys.strip())
        staff_sql += " GROUP BY TRIM(lsys)"
        staff_rows = db.execute_query(staff_sql, tuple(staff_params))
        dept_member_counts = {
            (r.get("dept") or "").strip(): int(r.get("cnt") or 0)
            for r in (staff_rows or [])
            if (r.get("dept") or "").strip()
        }
        scope_member_count = sum(dept_member_counts.values())

        detail_rows = []
        by_date = defaultdict(lambda: {
            "date": "",
            "scheduled": 0,
            "attended": 0,
            "normal": 0,
            "late": 0,
            "earlyLeave": 0,
            "absent": 0,
            "_attendedPeople": set(),
            "_attendedPeopleWithTrip": set(),
            "_domesticTripPeople": set(),
        })
        by_dept = defaultdict(lambda: {
            "dept": "",
            "scheduled": 0,
            "attended": 0,
            "normal": 0,
            "late": 0,
            "earlyLeave": 0,
            "absent": 0,
            "_attendedPeople": set(),
            "_attendedPeopleWithTrip": set(),
            "_domesticTripPeople": set(),
        })

        totals = {
            "scheduled": 0,
            "attended": 0,
            "normal": 0,
            "late": 0,
            "earlyLeave": 0,
            "absent": 0,
        }
        scheduled_people = set()
        attended_people = set()
        attended_people_with_trip = set()
        domestic_trip_people = set()

        for row in rows or []:
            shift_type = (row.get("shift_type") or "").strip()
            if shift_type not in expected_ranges:
                continue

            dept = (row.get("department") or "").strip()
            name = (row.get("employee_name") or "").strip()
            shift_date = str(row.get("shift_date") or "")[:10]
            expected_start, expected_end = expected_ranges[shift_type]

            first_time = _get_first_time(row)
            last_time = _get_last_time(row)
            punch_count = _count_valid_times(row)

            is_absent = not first_time
            if shift_type == "夜班" and first_time and punch_count <= 1:
                last_time = "24:00:00"

            is_late = bool(first_time and first_time > expected_start)
            is_early = bool(first_time and last_time and last_time < expected_end)
            is_attended = bool(first_time)
            is_normal = is_attended and not is_late and not is_early

            if is_absent:
                status = "absent"
                status_text = "缺勤"
            elif is_late and is_early:
                status = "late_early"
                status_text = "迟到、早退"
            elif is_late:
                status = "late"
                status_text = "迟到"
            elif is_early:
                status = "early_leave"
                status_text = "早退"
            else:
                status = "normal"
                status_text = "正常"

            totals["scheduled"] += 1
            totals["attended"] += 1 if is_attended else 0
            totals["normal"] += 1 if is_normal else 0
            totals["late"] += 1 if is_late else 0
            totals["earlyLeave"] += 1 if is_early else 0
            totals["absent"] += 1 if is_absent else 0

            scheduled_people.add(name)
            if is_attended:
                attended_people.add(name)
                attended_people_with_trip.add(name)

            for bucket in (by_date[shift_date], by_dept[dept]):
                if "date" in bucket:
                    bucket["date"] = shift_date
                if "dept" in bucket:
                    bucket["dept"] = dept
                bucket["scheduled"] += 1
                bucket["attended"] += 1 if is_attended else 0
                bucket["normal"] += 1 if is_normal else 0
                bucket["late"] += 1 if is_late else 0
                bucket["earlyLeave"] += 1 if is_early else 0
                bucket["absent"] += 1 if is_absent else 0
                if is_attended:
                    bucket["_attendedPeople"].add(name)
                    bucket["_attendedPeopleWithTrip"].add(name)

            detail_rows.append({
                "date": shift_date,
                "dept": dept,
                "name": name,
                "shiftType": shift_type,
                "expectedStart": expected_start[:5],
                "expectedEnd": expected_end[:5],
                "firstIn": first_time[:5] if first_time else "",
                "lastOut": last_time[:5] if last_time else "",
                "punchCount": punch_count,
                "late": is_late,
                "earlyLeave": is_early,
                "absent": is_absent,
                "status": status,
                "statusText": status_text,
            })

        trip_sql = """
            SELECT
                TRIM(g.gcr) AS name,
                TRIM(y.lsys) AS dept,
                COALESCE(g.gcsj, g.yjcfsj) AS trip_start,
                COALESCE(g.sjfhtime, g.yjfhsj) AS trip_end
            FROM gcsqb g
            INNER JOIN yggl y
                ON TRIM(g.gcr) = TRIM(y.name)
               AND COALESCE(y.zaizhi, 0) = 0
               AND RIGHT(TRIM(y.name), 1) != '1'
               AND RIGHT(TRIM(y.lsys), 1) != '1'
               AND TRIM(y.lsys) != %s
               AND TRIM(y.lsys) NOT IN ('其他部门员工','其他部门成员')
            WHERE RIGHT(TRIM(g.gcr), 1) != '1'
              AND g.bldzt = 2
              AND g.szrzt = 2
              AND COALESCE(NULLIF(TRIM(g.gclx), ''), '境内公出') = '境内公出'
              AND COALESCE(g.gcsj, g.yjcfsj) <= %s
              AND COALESCE(g.sjfhtime, g.yjfhsj) >= %s
        """
        trip_params: list = [
            LEADER_EXCLUDE_LSYS,
            d1.strftime("%Y-%m-%d 23:59:59"),
            d0.strftime("%Y-%m-%d 00:00:00"),
        ]
        if lsys:
            trip_sql += " AND TRIM(y.lsys) = %s"
            trip_params.append(lsys.strip())

        try:
            trip_rows = db.execute_query(trip_sql, tuple(trip_params))
        except Exception as e:
            logger.warning("查询境内公出人数失败: %s", e)
            trip_rows = []

        from datetime import timedelta
        for tr in trip_rows or []:
            trip_name = (tr.get("name") or "").strip()
            trip_dept = (tr.get("dept") or "").strip()
            trip_start = _parse_date(tr.get("trip_start"))
            trip_end = _parse_date(tr.get("trip_end"))
            if not trip_name or not trip_dept or not trip_start or not trip_end:
                continue
            cur = max(trip_start, d0)
            end_cur = min(trip_end, d1)
            while cur <= end_cur:
                ds = cur.strftime("%Y-%m-%d")
                if ds in by_date:
                    by_date[ds]["_domesticTripPeople"].add(trip_name)
                    by_date[ds]["_attendedPeopleWithTrip"].add(trip_name)
                if trip_dept in by_dept:
                    by_dept[trip_dept]["_domesticTripPeople"].add(trip_name)
                    by_dept[trip_dept]["_attendedPeopleWithTrip"].add(trip_name)
                domestic_trip_people.add(trip_name)
                attended_people_with_trip.add(trip_name)
                cur += timedelta(days=1)

        def finalize_bucket(item: dict) -> dict:
            scheduled = item.get("scheduled") or 0
            attended = item.get("attended") or 0
            attended_people_count = len(item.pop("_attendedPeople", set()))
            attended_people_with_trip_count = len(item.pop("_attendedPeopleWithTrip", set()))
            domestic_trip_people_count = len(item.pop("_domesticTripPeople", set()))
            dept_key = (item.get("dept") or "").strip()
            member_total = dept_member_counts.get(dept_key, 0) if dept_key else scope_member_count
            item["attendanceRate"] = _rate(attended, scheduled)
            item["absentRate"] = _rate(item.get("absent") or 0, scheduled)
            item["lateRate"] = _rate(item.get("late") or 0, scheduled)
            item["earlyLeaveRate"] = _rate(item.get("earlyLeave") or 0, scheduled)
            item["memberTotal"] = member_total
            item["attendedPeople"] = attended_people_count
            item["memberAttendanceRate"] = _rate(attended_people_count, member_total)
            item["domesticTripPeople"] = domestic_trip_people_count
            item["attendedPeopleWithTrip"] = attended_people_with_trip_count
            item["memberAttendanceRateWithTrip"] = _rate(attended_people_with_trip_count, member_total)
            return item

        date_rows = [finalize_bucket(v) for _k, v in sorted(by_date.items(), key=lambda kv: kv[0])]
        dept_rows = [finalize_bucket(v) for _k, v in sorted(by_dept.items(), key=lambda kv: kv[0])]
        summary = finalize_bucket({
            **totals,
            "scheduledPeople": len(scheduled_people),
            "attendancePersonRate": _rate(len(attended_people), len(scheduled_people)),
            "domesticTripPeople": len(domestic_trip_people),
            "_attendedPeople": set(attended_people),
            "_attendedPeopleWithTrip": set(attended_people_with_trip),
            "_domesticTripPeople": set(domestic_trip_people),
        })

        return {
            "success": True,
            "startDate": d0.strftime("%Y-%m-%d"),
            "endDate": d1.strftime("%Y-%m-%d"),
            "summary": summary,
            "byDate": date_rows,
            "byDept": dept_rows,
            "details": detail_rows,
            "rules": {
                "dayShift": "08:00-17:00",
                "nightShift": "17:00-22:00",
                "nightMissingLeaveAs": "24:00",
            },
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"值班出勤核查失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/discipline/holiday-duty-attendance/export")
async def export_holiday_duty_attendance(
    name: str = Query(..., description="当前登录用户姓名，用于鉴权"),
    start_date: str = Query(..., description="Start date, YYYY-MM-DD"),
    end_date: str = Query(..., description="End date, YYYY-MM-DD"),
    lsys: Optional[str] = Query(None, description="Department filter"),
):
    """Export holiday duty attendance check tables to Excel."""
    try:
        if not _can_access_holiday_duty_attendance(name):
            raise HTTPException(status_code=403, detail="无假期值班出勤核查权限")
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
    except Exception:
        raise HTTPException(status_code=500, detail="服务端未安装 openpyxl，无法导出")

    data = await get_holiday_duty_attendance(name=name, start_date=start_date, end_date=end_date, lsys=lsys)

    def pct(rate) -> str:
        try:
            return f"{float(rate or 0) * 100:.1f}%"
        except (TypeError, ValueError):
            return "0.0%"

    wb = Workbook()
    header_fill = PatternFill("solid", fgColor="E5E7EB")
    title_fill = PatternFill("solid", fgColor="DBEAFE")
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    title_font = Font(name="Microsoft YaHei", size=14, bold=True, color="1F2937")
    header_font = Font(name="Microsoft YaHei", size=10, bold=True, color="111827")
    body_font = Font(name="Microsoft YaHei", size=10, color="111827")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def style_range(ws):
        for row in ws.iter_rows():
            for cell in row:
                cell.border = border
                cell.alignment = center
                cell.font = body_font
        for cell in ws[1]:
            cell.font = title_font
            cell.fill = title_fill
        for cell in ws[3]:
            cell.font = header_font
            cell.fill = header_fill
        for col in range(1, ws.max_column + 1):
            max_len = 10
            for row in range(1, ws.max_row + 1):
                v = ws.cell(row=row, column=col).value
                if v is not None:
                    max_len = max(max_len, min(len(str(v)) + 2, 36))
            ws.column_dimensions[get_column_letter(col)].width = max_len
        ws.freeze_panes = "A4"

    scope_name = (lsys or "全部科室").strip()
    title = f"{scope_name} {data.get('startDate')} 至 {data.get('endDate')} 假期值班出勤核查"

    ws = wb.active
    ws.title = "按日期汇总"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
    ws.cell(1, 1, title)
    summary = data.get("summary") or {}
    ws.append([])
    ws.append(["日期", "应出勤", "已出勤", "出勤率", "出勤人员占比", "出勤人员占比（含境内公出）", "迟到", "早退", "缺勤"])
    for item in data.get("byDate") or []:
        ws.append([
            item.get("date", ""),
            item.get("scheduled", 0),
            item.get("attended", 0),
            pct(item.get("attendanceRate")),
            f"{item.get('attendedPeople', 0)} / {item.get('memberTotal', 0)} ({pct(item.get('memberAttendanceRate'))})",
            f"{item.get('attendedPeopleWithTrip', 0)} / {item.get('memberTotal', 0)} ({pct(item.get('memberAttendanceRateWithTrip'))})",
            item.get("late", 0),
            item.get("earlyLeave", 0),
            item.get("absent", 0),
        ])
    ws.append([])
    ws.append([
        "合计",
        summary.get("scheduled", 0),
        summary.get("attended", 0),
        pct(summary.get("attendanceRate")),
        f"{summary.get('attendedPeople', 0)} / {summary.get('memberTotal', 0)} ({pct(summary.get('memberAttendanceRate'))})",
        f"{summary.get('attendedPeopleWithTrip', 0)} / {summary.get('memberTotal', 0)} ({pct(summary.get('memberAttendanceRateWithTrip'))})",
        summary.get("late", 0),
        summary.get("earlyLeave", 0),
        summary.get("absent", 0),
    ])
    style_range(ws)

    ws_dept = wb.create_sheet("科室汇总")
    ws_dept.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
    ws_dept.cell(1, 1, title)
    ws_dept.append([])
    ws_dept.append(["科室", "应出勤", "已出勤", "出勤率", "出勤人员占比", "出勤人员占比（含境内公出）", "迟到", "早退", "缺勤"])
    for item in data.get("byDept") or []:
        ws_dept.append([
            item.get("dept", ""),
            item.get("scheduled", 0),
            item.get("attended", 0),
            pct(item.get("attendanceRate")),
            f"{item.get('attendedPeople', 0)} / {item.get('memberTotal', 0)} ({pct(item.get('memberAttendanceRate'))})",
            f"{item.get('attendedPeopleWithTrip', 0)} / {item.get('memberTotal', 0)} ({pct(item.get('memberAttendanceRateWithTrip'))})",
            item.get("late", 0),
            item.get("earlyLeave", 0),
            item.get("absent", 0),
        ])
    style_range(ws_dept)

    ws_detail = wb.create_sheet("异常明细")
    ws_detail.merge_cells(start_row=1, start_column=1, end_row=1, end_column=10)
    ws_detail.cell(1, 1, title)
    ws_detail.append([])
    ws_detail.append(["日期", "科室", "姓名", "班次", "应到", "应离", "首入", "末出", "打卡次数", "状态"])
    for item in data.get("details") or []:
        if item.get("status") == "normal":
            continue
        ws_detail.append([
            item.get("date", ""),
            item.get("dept", ""),
            item.get("name", ""),
            item.get("shiftType", ""),
            item.get("expectedStart", ""),
            item.get("expectedEnd", ""),
            item.get("firstIn", ""),
            item.get("lastOut", ""),
            item.get("punchCount", 0),
            item.get("statusText", ""),
        ])
    style_range(ws_detail)
    for row in ws_detail.iter_rows(min_row=4, min_col=2, max_col=3):
        for cell in row:
            cell.alignment = left

    bio = BytesIO()
    wb.save(bio)
    fname = f"{scope_name}_{data.get('startDate')}_{data.get('endDate')}_假期值班出勤核查.xlsx"
    return Response(
        content=bio.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"},
    )


def _get_overtime_by_person(year: int, month: Optional[int], lsys: Optional[str]) -> dict:
    """按人汇总加班小时 {name: hours}，仅已通过(jiabanzt=4)"""
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        join_cond = ("INNER JOIN yggl ON jiaban.xm = yggl.name "
                     "AND RIGHT(TRIM(yggl.name),1)!='1' AND RIGHT(TRIM(yggl.lsys),1)!='1' "
                     f"AND TRIM(yggl.lsys)!=%s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)")
        join_param = (LEADER_EXCLUDE_LSYS,)
    else:
        join_cond = ("INNER JOIN yggl ON jiaban.xm = yggl.name "
                     "AND yggl.lsys=%s AND RIGHT(TRIM(yggl.name),1)!='1' "
                     "AND RIGHT(TRIM(yggl.lsys),1)!='1' AND (COALESCE(yggl.zaizhi,0)=0)")
        join_param = (lsys,)

    if month:
        ms = f"{year}-{month:02d}"
        sql = f"""
            SELECT TRIM(jiaban.xm) AS name,
                   SUM(CAST(COALESCE(jiaban.tian1,0) AS DECIMAL(10,2))) AS hours
            FROM jiaban {join_cond}
            WHERE jiaban.jiabanzt=4
              AND (jiaban.timedate LIKE %s OR SUBSTRING(jiaban.timedate,1,7)=%s)
            GROUP BY TRIM(jiaban.xm)
        """
        rows = db.execute_query(sql, join_param + (f"{ms}%", ms))
    else:
        sql = f"""
            SELECT TRIM(jiaban.xm) AS name,
                   SUM(CAST(COALESCE(jiaban.tian1,0) AS DECIMAL(10,2))) AS hours
            FROM jiaban {join_cond}
            WHERE jiaban.jiabanzt=4
              AND (jiaban.timedate LIKE %s OR YEAR(jiaban.timedate)=%s)
            GROUP BY TRIM(jiaban.xm)
        """
        rows = db.execute_query(sql, join_param + (f"{year}%", year))

    result = {}
    for r in (rows or []):
        n = (r.get("name") or "").strip()
        if n:
            result[n] = float(r.get("hours") or 0)
    return result


def _get_trip_days_by_person(year: int, month: Optional[int], lsys: Optional[str]) -> dict:
    """按人汇总公出天数与节假日公出天数 {name: {tripDays, holidayTripDays}}，区间并集去重，仅已批准；不含市内公出（gclx 空视同境内）。"""
    import calendar as _cal
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        jc = ("gcsqb INNER JOIN yggl ON gcsqb.gcr=yggl.name "
              "AND RIGHT(TRIM(yggl.name),1)!='1' AND RIGHT(TRIM(yggl.lsys),1)!='1' "
              "AND TRIM(yggl.lsys)!=%s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)")
        jp = (LEADER_EXCLUDE_LSYS,)
    else:
        jc = ("gcsqb INNER JOIN yggl ON gcsqb.gcr=yggl.name "
              "AND yggl.lsys=%s AND RIGHT(TRIM(yggl.name),1)!='1' "
              "AND RIGHT(TRIM(yggl.lsys),1)!='1' AND (COALESCE(yggl.zaizhi,0)=0)")
        jp = (lsys,)

    if month:
        ms = date(year, month, 1)
        me = date(year, month, _cal.monthrange(year, month)[1])
        sql = f"""
            SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
            FROM {jc}
            WHERE RIGHT(TRIM(gcsqb.gcr),1)!='1'
              AND gcsqb.bldzt=2 AND gcsqb.szrzt=2
              {_WI_TRIP_SQL_EXCLUDE_CITY}
              AND COALESCE(gcsqb.gcsj,gcsqb.yjcfsj)<=%s
              AND COALESCE(gcsqb.sjfhtime,gcsqb.yjfhsj)>=%s
        """
        rows = db.execute_query(sql, jp + (me.strftime("%Y-%m-%d"), ms.strftime("%Y-%m-%d")))
        clip_start, clip_end = ms, min(me, date.today())
    else:
        sql = f"""
            SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
            FROM {jc}
            WHERE RIGHT(TRIM(gcsqb.gcr),1)!='1'
              AND gcsqb.bldzt=2 AND gcsqb.szrzt=2
              {_WI_TRIP_SQL_EXCLUDE_CITY}
              AND YEAR(COALESCE(gcsqb.gcsj,gcsqb.yjcfsj))=%s
        """
        rows = db.execute_query(sql, jp + (year,))
        clip_start = date(year, 1, 1)
        clip_end = min(date(year, 12, 31), date.today())

    by_person: dict = defaultdict(list)
    for row in (rows or []):
        gcr = (row.get("gcr") or "").strip()
        s = _parse_date(row.get("gcsj") or row.get("yjcfsj"))
        e = _parse_date(row.get("sjfhtime") or row.get("yjfhsj"))
        if s and e and e >= s:
            by_person[gcr].append((max(s, clip_start), min(e, clip_end)))

    result = {}
    for gcr, intervals in by_person.items():
        total_days, holiday_days = _merge_intervals_split_workdays(intervals)
        if total_days > 0:
            result[gcr] = {
                "tripDays": round(total_days, 2),
                "holidayTripDays": round(holiday_days, 2),
            }
    return result


def _get_overtime_by_person_range(d0: date, d1: date, lsys: Optional[str]) -> dict:
    """按人汇总加班小时 {name: hours}，日期区间 [d0,d1]（闭区间），仅已通过(jiabanzt=4)"""
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        join_cond = ("INNER JOIN yggl ON jiaban.xm = yggl.name "
                     "AND RIGHT(TRIM(yggl.name),1)!='1' AND RIGHT(TRIM(yggl.lsys),1)!='1' "
                     f"AND TRIM(yggl.lsys)!=%s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)")
        join_param = (LEADER_EXCLUDE_LSYS,)
    else:
        join_cond = ("INNER JOIN yggl ON jiaban.xm = yggl.name "
                     "AND yggl.lsys=%s AND RIGHT(TRIM(yggl.name),1)!='1' "
                     "AND RIGHT(TRIM(yggl.lsys),1)!='1' AND (COALESCE(yggl.zaizhi,0)=0)")
        join_param = (lsys,)

    sql = f"""
        SELECT TRIM(jiaban.xm) AS name,
               SUM(CAST(COALESCE(jiaban.tian1,0) AS DECIMAL(10,2))) AS hours
        FROM jiaban {join_cond}
        WHERE jiaban.jiabanzt=4
          AND DATE(jiaban.timedate) >= %s AND DATE(jiaban.timedate) <= %s
        GROUP BY TRIM(jiaban.xm)
    """
    rows = db.execute_query(sql, join_param + (d0.strftime("%Y-%m-%d"), d1.strftime("%Y-%m-%d")))

    result = {}
    for r in (rows or []):
        n = (r.get("name") or "").strip()
        if n:
            result[n] = float(r.get("hours") or 0)
    return result


def _get_trip_days_by_person_range(d0: date, d1: date, lsys: Optional[str]) -> dict:
    """按人汇总公出天数与节假日公出天数，区间与按月逻辑一致，裁剪到 [d0, min(d1,today)]；不含市内公出（gclx 空视同境内）。"""
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        jc = ("gcsqb INNER JOIN yggl ON gcsqb.gcr=yggl.name "
              "AND RIGHT(TRIM(yggl.name),1)!='1' AND RIGHT(TRIM(yggl.lsys),1)!='1' "
              "AND TRIM(yggl.lsys)!=%s AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(yggl.zaizhi,0)=0)")
        jp = (LEADER_EXCLUDE_LSYS,)
    else:
        jc = ("gcsqb INNER JOIN yggl ON gcsqb.gcr=yggl.name "
              "AND yggl.lsys=%s AND RIGHT(TRIM(yggl.name),1)!='1' "
              "AND RIGHT(TRIM(yggl.lsys),1)!='1' AND (COALESCE(yggl.zaizhi,0)=0)")
        jp = (lsys,)

    sql = f"""
        SELECT gcsqb.gcr, gcsqb.gcsj, gcsqb.sjfhtime, gcsqb.yjfhsj, gcsqb.yjcfsj
        FROM {jc}
        WHERE RIGHT(TRIM(gcsqb.gcr),1)!='1'
          AND gcsqb.bldzt=2 AND gcsqb.szrzt=2
          {_WI_TRIP_SQL_EXCLUDE_CITY}
          AND COALESCE(gcsqb.gcsj,gcsqb.yjcfsj)<=%s
          AND COALESCE(gcsqb.sjfhtime,gcsqb.yjfhsj)>=%s
    """
    rows = db.execute_query(sql, jp + (d1.strftime("%Y-%m-%d"), d0.strftime("%Y-%m-%d")))
    clip_start = d0
    clip_end = min(d1, date.today())

    by_person: dict = defaultdict(list)
    for row in (rows or []):
        gcr = (row.get("gcr") or "").strip()
        s = _parse_date(row.get("gcsj") or row.get("yjcfsj"))
        e = _parse_date(row.get("sjfhtime") or row.get("yjfhsj"))
        if s and e and e >= s:
            by_person[gcr].append((max(s, clip_start), min(e, clip_end)))

    result = {}
    for gcr, intervals in by_person.items():
        total_days, holiday_days = _merge_intervals_split_workdays(intervals)
        if total_days > 0:
            result[gcr] = {
                "tripDays": round(total_days, 2),
                "holidayTripDays": round(holiday_days, 2),
            }
    return result


def _get_leave_days_by_person_period(period_start: date, period_end: date, lsys: Optional[str]) -> Dict[str, float]:
    """
    qjzt=4 已通过请假，与 /dept/leave 相同的天数重叠分摊逻辑；统计期闭区间 [period_start, period_end]。
    员工范围与工作强度一致（yggl 在职、含部办）。返回 { 姓名: 分摊后天数 }。
    """
    pe_str = period_end.strftime("%Y-%m-%d")
    ps_str = period_start.strftime("%Y-%m-%d")
    ov = _leave_overlap_sql_bounds()
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        query = f"""
            SELECT TRIM(qj.xm) AS name, qj.timefrom, qj.timeto, qj.timefromdate, CAST(qj.tian AS DECIMAL(10,2)) AS tian
            FROM qj INNER JOIN yggl ON TRIM(qj.xm) = TRIM(yggl.name)
            WHERE qj.qjzt = 4
              AND RIGHT(TRIM(yggl.name),1) != '1' AND RIGHT(TRIM(yggl.lsys),1) != '1'
              AND TRIM(yggl.lsys) NOT IN ('其他部门员工','其他部门成员')
              AND (COALESCE(yggl.zaizhi,0) = 0)
            {ov}
        """
        rows = db.execute_query(query, (pe_str, ps_str))
    else:
        query = f"""
            SELECT TRIM(qj.xm) AS name, qj.timefrom, qj.timeto, qj.timefromdate, CAST(qj.tian AS DECIMAL(10,2)) AS tian
            FROM qj INNER JOIN yggl ON TRIM(qj.xm) = TRIM(yggl.name) AND TRIM(yggl.lsys) = %s
            WHERE qj.qjzt = 4
              AND RIGHT(TRIM(yggl.name),1) != '1' AND RIGHT(TRIM(yggl.lsys),1) != '1'
              AND (COALESCE(yggl.zaizhi,0) = 0)
            {ov}
        """
        rows = db.execute_query(query, (lsys.strip(), pe_str, ps_str))

    by_name: Dict[str, float] = defaultdict(float)
    for r in rows or []:
        name = (r.get("name") or "").strip()
        if not name:
            continue
        ls_d, le_d = _qj_leave_date_bounds(r)
        if not ls_d or not le_d:
            continue
        try:
            tian_f = float(r.get("tian") or 0)
        except (TypeError, ValueError):
            tian_f = 0.0
        alloc = _allocate_leave_tian_to_period(tian_f, ls_d, le_d, period_start, period_end)
        if alloc > 0:
            by_name[name] += alloc
    return dict(by_name)


def _get_staff_with_dept(lsys: Optional[str]) -> list:
    """返回 [{name, lsys, jb}]，在职员工。全员（不传 lsys）时包含「部办」，以便工作强度按科室展示部办卡片。"""
    all_staff = not (lsys and lsys.strip())
    if all_staff:
        rows = db.execute_query(
            "SELECT name, lsys, jb FROM yggl WHERE name IS NOT NULL AND name!='' "
            "AND RIGHT(TRIM(name),1)!='1' AND RIGHT(TRIM(lsys),1)!='1' "
            "AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0)",
        )
    else:
        rows = db.execute_query(
            "SELECT name, lsys, jb FROM yggl WHERE lsys=%s AND name IS NOT NULL AND name!='' "
            "AND RIGHT(TRIM(name),1)!='1' AND RIGHT(TRIM(lsys),1)!='1' "
            "AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') AND (COALESCE(zaizhi,0)=0)",
            (lsys,),
        )
    return [{
        "name": (r["name"] or "").strip(),
        "lsys": (r.get("lsys") or "").strip(),
        "jb": (r.get("jb") or "").strip(),
    } for r in (rows or []) if (r.get("name") or "").strip()]


def _leader_style_overtime_hours_from_attendance(names: List[str], start: date, end: date) -> Dict[str, float]:
    """
    与智能建议「加班建议」同款：从 attendance_records 识别加班时长（工作日 17:00 后；休息日/假期同 analyze_restday）。
    返回 { name: raw_hours }，与领导加班页一致：先累加原始小时，最后再统一 round。
    """
    if not names:
        return {}
    if not collect_valid_times_with_marks or not build_intervals_from_marks or not is_workday:
        return {n: 0.0 for n in names}

    name_set = set(names)
    ph = ",".join(["%s"] * len(names))
    rows = db.execute_query(
        f"SELECT * FROM attendance_records WHERE employee_name IN ({ph}) "
        "AND attendance_date >= %s AND attendance_date <= %s "
        "ORDER BY attendance_date ASC, employee_name",
        tuple(names) + (start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d")),
    )

    holidays_by_year: Dict[int, Dict[str, str]] = {}
    festival_by_year: Dict[int, Dict[str, str]] = {}
    for y in range(start.year, end.year + 1):
        holidays_by_year[y] = load_holidays_dict(str(y))
        festival_by_year[y] = _load_holiday_festival_map(y) if _load_holiday_festival_map else {}

    totals: Dict[str, float] = defaultdict(float)
    for row in rows or []:
        name = (row.get("employee_name") or "").strip()
        if name not in name_set:
            continue
        date_obj = _parse_attendance_date(row.get("attendance_date"))
        if not date_obj:
            continue
        y = date_obj.year
        holidays = holidays_by_year.get(y) or {}
        festival_map = festival_by_year.get(y) or {}
        if calc_suggestion_style_overtime_for_record:
            hours, _segments, _day_type = calc_suggestion_style_overtime_for_record(
                row, holidays, festival_map
            )
        else:
            hours, _segments, _day_type = _leader_overtime_for_record(row, holidays, festival_map)
        if hours > 0:
            totals[name] += float(hours)

    return {n: float(totals.get(n, 0.0)) for n in names}


@router.get("/leader/work-intensity")
async def get_work_intensity(
    year: int = Query(..., description="年份（日期区间模式仍需传入以保持兼容）"),
    month: Optional[int] = Query(None, description="月份，不传则全年；与 date_from/date_to 互斥"),
    lsys: Optional[str] = Query(None, description="科室，不传则全员"),
    date_from: Optional[str] = Query(None, description="起始日期 YYYY-MM-DD"),
    date_to: Optional[str] = Query(None, description="结束日期 YYYY-MM-DD，与 date_from 同时传则按闭区间统计"),
    intensity_formula: str = Query(
        "a",
        description="工作强度口径：a=加班÷在岗；b=（加班−请假小时）÷在岗。加班=智能建议同款打卡识别；请假小时=统计期内已通过请假按天重叠分摊×8；公出仅境内/境外（不含市内）",
    ),
):
    """
    工作强度统计：
    口径 A = 加班时长 / 实际在岗时长
    口径 B =（加班时长 − 请假时间）/ 实际在岗时长；请假时间为统计期内已通过请假（qjzt=4）按天重叠分摊后×8 小时，与请假汇总卡片一致。
    加班时长：全员按打卡数据识别（与考勤智能建议「加班建议」合计一致，非 jiaban 申报汇总）。
    实际在岗时长 = 应出勤时长 - 公出时长 + 公出期间节假日时长
    公出时长仅含境内/境外公出（不含市内公出；gclx 空视同境内公出）。
    时长单位统一为小时（天数×8）。
    返回：全员、各科室、每个人的 intensity 及 intensityFormula。
    传入 date_from + date_to 时按自定义日期区间统计（忽略 month），区间结束日晚于今天时按今天截断。
    """
    try:
        HOURS_PER_DAY = 8

        today = date.today()
        range_meta: Dict = {}

        if date_from or date_to:
            if not date_from or not date_to:
                raise HTTPException(status_code=400, detail="date_from 与 date_to 需同时传入")
            ds = _parse_date(date_from)
            de = _parse_date(date_to)
            if not ds or not de:
                raise HTTPException(status_code=400, detail="日期格式无效，请使用 YYYY-MM-DD")
            if de < ds:
                ds, de = de, ds
            d_end_eff = min(de, today)
            if ds > d_end_eff:
                workdays = 0
            else:
                workdays = _count_workdays_between(ds, d_end_eff)
            expected_hours = workdays * HOURS_PER_DAY
            staff = _get_staff_with_dept(lsys)
            trip_map = _get_trip_days_by_person_range(ds, de, lsys)
            range_meta = {
                "rangeMode": True,
                "dateFrom": ds.isoformat(),
                "dateTo": de.isoformat(),
                "effectiveDateTo": d_end_eff.isoformat(),
            }
            ot_att_start, ot_att_end = ds, d_end_eff
            leave_ps, leave_pe = ds, d_end_eff
        elif month:
            # 当统计“当前年月”时，应出勤只统计到今天，避免按整月放大分母
            if year == today.year and month == today.month:
                workdays = _count_workdays_in_month_until(year, month, today.day)
            else:
                workdays = _count_workdays_in_month(year, month)
            expected_hours = workdays * HOURS_PER_DAY
            staff = _get_staff_with_dept(lsys)
            trip_map = _get_trip_days_by_person(year, month, lsys)
            range_meta = {"rangeMode": False}
            import calendar as _cal
            month_first = date(year, month, 1)
            month_last = date(year, month, _cal.monthrange(year, month)[1])
            ot_att_start = month_first
            ot_att_end = month_last
            if year == today.year and month == today.month:
                ot_att_end = min(month_last, today)
            leave_ps, leave_pe = month_first, month_last
        else:
            # 统计全年时：当年仅统计到今天，历史年份统计整年
            if year == today.year:
                workdays = sum(_count_workdays_in_month(year, m) for m in range(1, today.month))
                workdays += _count_workdays_in_month_until(year, today.month, today.day)
            else:
                workdays = sum(_count_workdays_in_month(year, m) for m in range(1, 13))
            expected_hours = workdays * HOURS_PER_DAY
            staff = _get_staff_with_dept(lsys)
            trip_map = _get_trip_days_by_person(year, None, lsys)
            range_meta = {"rangeMode": False}
            year_first = date(year, 1, 1)
            year_last = date(year, 12, 31)
            ot_att_start = year_first
            ot_att_end = year_last if year < today.year else min(year_last, today)
            leave_ps = year_first
            leave_pe = year_last if year < today.year else min(year_last, today)

        all_names = sorted({s["name"] for s in staff})
        if all_names:
            att_ot = _leader_style_overtime_hours_from_attendance(all_names, ot_att_start, ot_att_end)
            ot_map = {n: round(float(att_ot.get(n, 0.0)), 2) for n in all_names}
        else:
            ot_map = {}

        formula = (intensity_formula or "a").strip().lower()
        if formula not in ("a", "b"):
            formula = "a"
        leave_days_map: Dict[str, float] = (
            _get_leave_days_by_person_period(leave_ps, leave_pe, lsys) if formula == "b" else {}
        )

        person_list = []
        dept_agg = defaultdict(
            lambda: {"ot": 0.0, "leave_h": 0.0, "trip_days": 0.0, "trip_holiday_days": 0.0, "count": 0}
        )

        for s in staff:
            name = s["name"]
            dept = s["lsys"]
            jb = s.get("jb") or ""
            ot = ot_map.get(name, 0)
            trip_d = float((trip_map.get(name) or {}).get("tripDays", 0))
            trip_holiday_d = float((trip_map.get(name) or {}).get("holidayTripDays", 0))
            trip_h = trip_d * HOURS_PER_DAY
            trip_holiday_h = trip_holiday_d * HOURS_PER_DAY
            actual_h = expected_hours - trip_h + trip_holiday_h
            leave_h = float(leave_days_map.get(name, 0.0)) * HOURS_PER_DAY if formula == "b" else 0.0
            ot_net = ot - leave_h if formula == "b" else ot
            intensity = round(ot_net / actual_h, 4) if actual_h > 0 else 0

            row_p = {
                "name": name,
                "lsys": dept,
                "jb": jb,
                "overtimeHours": round(ot, 2),
                "tripDays": round(trip_d, 2),
                "tripHolidayDays": round(trip_holiday_d, 2),
                "actualHours": round(actual_h, 2),
                "intensity": intensity,
            }
            if formula == "b":
                row_p["leaveHours"] = round(leave_h, 2)
            person_list.append(row_p)

            da = dept_agg[dept]
            da["ot"] += ot
            if formula == "b":
                da["leave_h"] += leave_h
            da["trip_days"] += trip_d
            da["trip_holiday_days"] += trip_holiday_d
            da["count"] += 1

        person_list.sort(key=lambda x: -x["intensity"])

        total_ot = sum(p["overtimeHours"] for p in person_list)
        total_leave_h = sum(float(p.get("leaveHours") or 0) for p in person_list) if formula == "b" else 0.0
        total_trip_d = sum(p["tripDays"] for p in person_list)
        total_trip_holiday_d = sum(p.get("tripHolidayDays", 0) for p in person_list)
        total_trip_h = total_trip_d * HOURS_PER_DAY
        total_trip_holiday_h = total_trip_holiday_d * HOURS_PER_DAY
        total_actual = expected_hours * len(staff) - total_trip_h + total_trip_holiday_h
        if formula == "b":
            overall_intensity = round((total_ot - total_leave_h) / total_actual, 4) if total_actual > 0 else 0
        else:
            overall_intensity = round(total_ot / total_actual, 4) if total_actual > 0 else 0

        dept_list = []
        for dept_name, da in dept_agg.items():
            dept_actual = expected_hours * da["count"] - da["trip_days"] * HOURS_PER_DAY + da["trip_holiday_days"] * HOURS_PER_DAY
            if formula == "b":
                dept_ot_net = da["ot"] - da["leave_h"]
                dept_i = round(dept_ot_net / dept_actual, 4) if dept_actual > 0 else 0
            else:
                dept_i = round(da["ot"] / dept_actual, 4) if dept_actual > 0 else 0
            row_d = {
                "lsys": dept_name,
                "personCount": da["count"],
                "overtimeHours": round(da["ot"], 2),
                "tripDays": round(da["trip_days"], 2),
                "tripHolidayDays": round(da["trip_holiday_days"], 2),
                "actualHours": round(dept_actual, 2),
                "intensity": dept_i,
            }
            if formula == "b":
                row_d["leaveHours"] = round(da["leave_h"], 2)
            dept_list.append(row_d)
        dept_list.sort(key=lambda x: -x["intensity"])

        return {
            "success": True,
            "workdays": workdays,
            "expectedHoursPerPerson": expected_hours,
            "totalPeople": len(staff),
            "totalOvertimeHours": round(total_ot, 2),
            "totalActualHours": round(total_actual, 2),
            "totalLeaveHours": round(total_leave_h, 2) if formula == "b" else None,
            "overallIntensity": overall_intensity,
            "intensityFormula": formula,
            "overtimeCalcMethod": "suggestion",
            "overtimeCalcNote": "加班时长与考勤页智能建议「加班建议」合计一致（非 jiaban 申报汇总）",
            "byDept": dept_list,
            "byPerson": person_list,
            **range_meta,
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"工作强度统计失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
#  考勤表 Word 导出
# ============================================================

import math
from io import BytesIO
from fastapi.responses import StreamingResponse

_LEAVE_TYPE_COLS = [
    "事假", "病假", "工伤", "生育", "探亲", "丧假", "婚假", "旷工", "职工假",
]
_HX_TYPES = {"换休", "员工换休票", "书面请假换休", "书面申请领取"}

_LEAVE_TYPE_MAP = {
    "事假": "事假",
    "书面请假": "事假",
    "书面申请": "事假",
    "病假": "病假",
    "工伤": "工伤",
    "产假": "生育",
    "哺乳假": "生育",
    "护理假": "生育",
    "探亲假": "探亲",
    "丧假": "丧假",
    "婚假": "婚假",
    "旷工": "旷工",
    "带薪休假": "职工假",
    "带薪休年假": "职工假",
    "带薪年休假": "职工假",
    "换休": "换休",
    "员工换休票": "换休",
    "书面请假换休": "换休",
    "书面申请领取": "换休",
    "异常打卡": None,
}


def _ceil_quarter(val: float) -> float:
    """向上取整到 0.25 的整数倍。"""
    if val <= 0:
        return 0.0
    return math.ceil(val * 4) / 4


def _fmt_days(val: float) -> str:
    """格式化天数：整数不带小数点，否则去除尾部多余零。如 1.0→'1', 0.5→'0.5', 0.25→'0.25'"""
    if val == int(val):
        return str(int(val))
    s = f"{val:.2f}".rstrip("0").rstrip(".")
    return s


@router.get("/leader/attendance-report-export")
async def attendance_report_export(
    year: int = Query(...),
    month: int = Query(..., ge=1, le=12),
    lsys: Optional[str] = Query(None, description="科室，不传则按全部科室分别生成"),
):
    """
    根据模版生成 Word 考勤表。
    返回：每科室一份 sheet 数据，前端据此填充 Word；
    或直接生成 .docx 文件流返回下载。
    """
    import calendar
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from copy import deepcopy

    try:
        _, last_day = calendar.monthrange(year, month)
        workdays = _count_workdays_in_month(year, month)

        # 确定科室列表
        if lsys and lsys.strip():
            dept_list = [lsys.strip()]
        else:
            rows = db.execute_query(
                "SELECT DISTINCT TRIM(lsys) AS lsys FROM yggl "
                "WHERE lsys IS NOT NULL AND lsys!='' AND RIGHT(TRIM(lsys),1)!='1' "
                "AND TRIM(lsys)!=%s AND TRIM(lsys) NOT IN ('其他部门员工','其他部门成员') "
                "AND (COALESCE(zaizhi,0)=0) ORDER BY lsys",
                (LEADER_EXCLUDE_LSYS,)
            )
            dept_list = [r["lsys"] for r in (rows or []) if r.get("lsys")]

        period_start = date(year, month, 1)
        period_end = date(year, month, last_day)
        pe_str = period_end.strftime("%Y-%m-%d")
        ps_str = period_start.strftime("%Y-%m-%d")
        ov = _leave_overlap_sql_bounds()

        # 加载模版
        import os
        template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
                                     "智能制造技术室XXX年XX月份考勤表模版.docx")
        if not os.path.exists(template_path):
            raise HTTPException(status_code=500, detail="考勤表模版文件不存在")

        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        FONT_CJK = "宋体"
        FONT_LATIN = "Times New Roman"

        def _set_run_font(run, size=None, bold=False):
            """统一 run 字体：中文宋体，数字/英文 Times New Roman"""
            run.font.name = FONT_LATIN
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:ascii'), FONT_LATIN)
            rfonts.set(qn('w:hAnsi'), FONT_LATIN)
            rfonts.set(qn('w:eastAsia'), FONT_CJK)
            if size:
                run.font.size = size
            run.font.bold = bold

        def _set_cell_font(cell, size=Pt(9)):
            """遍历单元格所有段落所有 run，统一字体"""
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, size=size)

        base_doc = Document(template_path)
        out_doc = Document()
        # 继承模版页面设置
        out_sec = out_doc.sections[0]
        tmpl_sec = base_doc.sections[0]
        out_sec.page_width = tmpl_sec.page_width
        out_sec.page_height = tmpl_sec.page_height
        out_sec.left_margin = tmpl_sec.left_margin
        out_sec.right_margin = tmpl_sec.right_margin
        out_sec.top_margin = tmpl_sec.top_margin
        out_sec.bottom_margin = tmpl_sec.bottom_margin
        out_sec.orientation = tmpl_sec.orientation

        is_first_dept = True

        for dept_name in dept_list:
            # 获取该科室在职人员
            staff_rows = db.execute_query(
                "SELECT TRIM(name) AS name FROM yggl WHERE lsys=%s AND name IS NOT NULL AND name!='' "
                "AND RIGHT(TRIM(name),1)!='1' AND (COALESCE(zaizhi,0)=0) ORDER BY name",
                (dept_name,)
            )
            staff_names = [r["name"] for r in (staff_rows or []) if r.get("name")]
            if not staff_names:
                continue

            # 查该科室请假数据
            leave_sql = f"""
                SELECT TRIM(xm) AS name, TRIM(qjfs) AS qjfs,
                       timefrom, timeto, timefromdate,
                       CAST(tian AS DECIMAL(10,2)) AS tian
                FROM qj
                WHERE lsys=%s AND qjzt=4
                  AND RIGHT(TRIM(xm),1)!='1'
                  {ov}
                ORDER BY xm
            """
            leave_rows = db.execute_query(leave_sql, (dept_name, pe_str, ps_str))

            # 按人按模板列类型汇总（通过 _LEAVE_TYPE_MAP 映射原始 qjfs）
            person_leave: Dict[str, Dict[str, float]] = defaultdict(lambda: defaultdict(float))
            for r in (leave_rows or []):
                nm = (r.get("name") or "").strip()
                qjfs = (r.get("qjfs") or "").strip()
                if not nm or not qjfs:
                    continue
                mapped = _LEAVE_TYPE_MAP.get(qjfs)
                if mapped is None:
                    continue
                ls_d, le_d = _qj_leave_date_bounds(r)
                if not ls_d or not le_d:
                    continue
                try:
                    tian_f = float(r.get("tian") or 0)
                except (TypeError, ValueError):
                    tian_f = 0.0
                alloc = _allocate_leave_tian_to_period(tian_f, ls_d, le_d, period_start, period_end)
                if alloc > 0:
                    person_leave[nm][mapped] += alloc

            # 查该科室加班数据
            ot_map = _get_overtime_by_person(year, month, dept_name)

            # 查该科室排班数据：值班(白班)天数、夜班天数
            shift_sql = """
                SELECT employee_name,
                       SUM(CASE WHEN shift_type='白班' THEN 1 ELSE 0 END) AS day_count,
                       SUM(CASE WHEN shift_type='夜班' THEN 1 ELSE 0 END) AS night_count
                FROM shift_schedule
                WHERE department=%s AND year=%s AND month=%s
                GROUP BY employee_name
            """
            shift_rows = db.execute_query(shift_sql, (dept_name, year, month))
            shift_day_map: Dict[str, int] = {}
            shift_night_map: Dict[str, int] = {}
            for r in (shift_rows or []):
                en = (r.get("employee_name") or "").strip()
                if en:
                    dc = int(r.get("day_count") or 0)
                    nc = int(r.get("night_count") or 0)
                    if dc > 0:
                        shift_day_map[en] = dc
                    if nc > 0:
                        shift_night_map[en] = nc

            # 从模版复制表格结构
            tmpl_table1 = base_doc.tables[1]

            if not is_first_dept:
                # 分页
                run = out_doc.add_paragraph().add_run()
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                run._element.append(br)

            # ---- 标题：居中大号加粗段落 ----
            title_p = out_doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(f"{year}年{month}月份考勤表")
            _set_run_font(title_run, size=Pt(16), bold=True)
            title_p.space_after = Pt(6)

            # ---- Table 1: 主体考勤表 ----
            new_tbl1 = deepcopy(tmpl_table1._tbl)
            out_doc.element.body.append(new_tbl1)
            tbl1_obj = out_doc.tables[-1]

            def _cell_set_text(cell, text):
                """清除单元格所有段落内容，在第一段写入文本并设字体。"""
                for p in cell.paragraphs:
                    p.clear()
                run = cell.paragraphs[0].add_run(text)
                _set_run_font(run, size=Pt(9))

            # 修改表头 Row0 中的科室和月/日/天信息
            hdr_cells = tbl1_obj.rows[0].cells
            _cell_set_text(hdr_cells[2], dept_name)
            _cell_set_text(hdr_cells[8], f"{month}月")
            _cell_set_text(hdr_cells[14], f" {month}月")
            _cell_set_text(hdr_cells[16], f"{last_day}日计")
            _cell_set_text(hdr_cells[19], f"{workdays}天）")

            # 统一表头 Row0 和 Row1 中原有 run 的字体
            for ri in range(min(2, len(tbl1_obj.rows))):
                for cell in tbl1_obj.rows[ri].cells:
                    _set_cell_font(cell, size=Pt(9))

            # 删除模版数据行（保留 row0=表头区域, row1=列标题）
            data_rows_to_remove = list(tbl1_obj.rows[2:])
            for row in data_rows_to_remove:
                tbl1_obj._tbl.remove(row._tr)

            # 为每位员工添加一行
            template_data_row_tr = tmpl_table1.rows[2]._tr  # 用作格式模版
            from docx.table import _Row

            for idx, name in enumerate(staff_names, 1):
                new_tr = deepcopy(template_data_row_tr)
                tbl1_obj._tbl.append(new_tr)
                row_obj = _Row(new_tr, tbl1_obj)
                cells = row_obj.cells

                def _data_cell(cell, text):
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(text)
                    _set_run_font(run, size=Pt(9))

                # Col0: 编号
                _data_cell(cells[0], str(idx))
                # Col1 (gridSpan=2): 姓名
                _data_cell(cells[1], name)
                # Col3: 基本工资（留空）
                _data_cell(cells[3], "")

                # 获取该人的请假数据（已映射到模板列名）
                pl = person_leave.get(name, {})

                # 判断出勤：除"换休"列外有任何假别即非全勤
                non_hx_leave = sum(v for k, v in pl.items() if k != "换休")
                # Col4: 出勤
                _data_cell(cells[4], "全勤" if non_hx_leave <= 0 else "")

                # 各假别列映射 (col_index, leave_type)
                leave_col_map = [
                    (5, "事假"),
                    (6, "病假"),
                    (7, "工伤"),
                    (9, "生育"),
                    (11, "探亲"),
                    (12, "丧假"),
                    (13, "婚假"),
                    (15, "旷工"),
                ]
                for ci, lt in leave_col_map:
                    val = pl.get(lt, 0)
                    _data_cell(cells[ci], _fmt_days(_ceil_quarter(val)) if val > 0 else "")

                # Col17: 值班（白班天数）
                day_shifts = shift_day_map.get(name, 0)
                _data_cell(cells[17], str(day_shifts) if day_shifts > 0 else "")

                # Col18 (gridSpan=2): 加班 - 小时转天
                ot_hours = ot_map.get(name, 0)
                if ot_hours > 0:
                    ot_days = _ceil_quarter(ot_hours / 8.0)
                    _data_cell(cells[18], _fmt_days(ot_days))
                else:
                    _data_cell(cells[18], "")

                # Col20: 夜班天数
                night_shifts = shift_night_map.get(name, 0)
                _data_cell(cells[20], str(night_shifts) if night_shifts > 0 else "")

                # Col21 (gridSpan=2): 职工假
                zgj = pl.get("职工假", 0)
                _data_cell(cells[21], _fmt_days(_ceil_quarter(zgj)) if zgj > 0 else "")

                # Col23: 换休（已映射合并为 "换休" 键）
                hx_total = pl.get("换休", 0)
                _data_cell(cells[23], _fmt_days(_ceil_quarter(hx_total)) if hx_total > 0 else "0")

            # 底部签名段落
            footer_p = out_doc.add_paragraph()
            footer_p.space_before = Pt(12)
            footer_run = footer_p.add_run("主    管                                   考 勤 员:")
            _set_run_font(footer_run, size=Pt(11))

            is_first_dept = False

        # 输出文件
        buf = BytesIO()
        out_doc.save(buf)
        buf.seek(0)
        month_str = f"{month:02d}"
        fname = f"考勤表_{year}年{month}月"
        if lsys and lsys.strip():
            fname += f"_{lsys.strip()}"
        fname += ".docx"

        from urllib.parse import quote
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"考勤表导出失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
